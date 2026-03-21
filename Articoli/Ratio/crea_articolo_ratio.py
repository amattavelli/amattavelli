#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = "/home/user/amattavelli/Articoli/Ratio/03.2026 - Agenti AI - La nuova frontiera per studi e imprese.docx"

doc = Document()

# --- Impostazioni pagina ---
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# --- Stili ---
styles = doc.styles

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph_with_style(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

# === RUBRICA / CATEGORIA ===
p_rubrica = add_paragraph_with_style(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4)
r = p_rubrica.add_run("INNOVAZIONE DIGITALE  |  INTELLIGENZA ARTIFICIALE")
set_font(r, size=8, bold=True, color=(120, 120, 120))

# === TITOLO ===
p_titolo = add_paragraph_with_style(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=4)
r = p_titolo.add_run("Agenti AI:\nla nuova frontiera per studi professionali e imprese")
set_font(r, size=22, bold=True, color=(30, 30, 30))

# === SOTTOTITOLO ===
p_sub = add_paragraph_with_style(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=10)
r = p_sub.add_run(
    "Oltre i chatbot: i sistemi multi-agente stanno ridisegnando processi, "
    "competenze e modelli di business nel mondo dei professionisti e delle PMI."
)
set_font(r, size=12, italic=True, color=(60, 60, 60))

# === AUTORE ===
p_autore = add_paragraph_with_style(doc, "", space_before=0, space_after=14)
r = p_autore.add_run("di Alberto Mattavelli  |  Consulente in innovazione digitale e AI")
set_font(r, size=9, italic=True, color=(100, 100, 100))

# Linea separatrice
doc.add_paragraph("─" * 80)

# === ABSTRACT ===
p_abs_label = add_paragraph_with_style(doc, "", space_before=10, space_after=2)
r = p_abs_label.add_run("IN SINTESI")
set_font(r, size=8, bold=True, color=(0, 90, 160))

p_abs = add_paragraph_with_style(doc, "", space_before=0, space_after=16)
p_abs.paragraph_format.left_indent  = Cm(0.5)
p_abs.paragraph_format.right_indent = Cm(0.5)
r = p_abs.add_run(
    "Gli agenti di intelligenza artificiale rappresentano un cambio di paradigma "
    "rispetto ai tradizionali strumenti AI generativi. Non si limitano a rispondere "
    "a domande: pianificano, ragionano, delegano subtask e interagiscono con sistemi "
    "esterni per raggiungere obiettivi complessi in modo autonomo. Per studi "
    "professionali e aziende, la posta in gioco non è solo l'efficienza operativa: "
    "è la ridefinizione stessa del valore della consulenza."
)
set_font(r, size=10, italic=True, color=(40, 40, 40))

# === CORPO ARTICOLO ===

def h2(doc, testo):
    p = add_paragraph_with_style(doc, "", space_before=16, space_after=6)
    r = p.add_run(testo)
    set_font(r, size=13, bold=True, color=(0, 90, 160))
    return p

def body(doc, testo):
    p = add_paragraph_with_style(doc, "", space_before=0, space_after=8)
    p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(testo)
    set_font(r, size=11)
    return p

def quote(doc, testo):
    p = add_paragraph_with_style(doc, "", space_before=12, space_after=12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    r = p.add_run(f"\u201c{testo}\u201d")
    set_font(r, size=12, italic=True, color=(0, 90, 160))
    return p

def elenco(doc, items, bold_prefix=None):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.8)
        if bold_prefix and item.startswith(bold_prefix):
            pass
        if ":" in item:
            parts = item.split(":", 1)
            r1 = p.add_run(parts[0] + ":")
            set_font(r1, size=11, bold=True)
            r2 = p.add_run(parts[1])
            set_font(r2, size=11)
        else:
            r = p.add_run(item)
            set_font(r, size=11)

def nota(doc, testo):
    p = add_paragraph_with_style(doc, "", space_before=4, space_after=4)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("► " + testo)
    set_font(r, size=9, color=(80, 80, 80))

# ---- SEZIONE 1 ----
h2(doc, "1. Dal prompt all'agente: un salto qualitativo")

body(doc,
    "Fino a poco tempo fa, quando si parlava di intelligenza artificiale in azienda o in uno "
    "studio professionale, si intendeva essenzialmente un assistente testuale avanzato: si poneva "
    "una domanda, si otteneva una risposta. Utile, ma intrinsecamente passivo. Gli agenti AI "
    "segnano una discontinuità radicale."
)
body(doc,
    "Un agente AI è un sistema software capace di perseguire obiettivi in modo autonomo: "
    "scompone un compito in passi, utilizza strumenti esterni (database, API, software gestionali, "
    "browser), valuta i risultati intermedi e riadatta la strategia. In altri termini, non risponde "
    "soltanto — agisce."
)
body(doc,
    "Nel 2025 i principali fornitori di piattaforme — OpenAI, Google, Anthropic, Microsoft e "
    "una miriade di startup — hanno rilasciato framework per la costruzione di agenti: AutoGen, "
    "CrewAI, LangGraph, Claude Agent SDK, Google ADK. Nel 2026 questi strumenti sono già maturi "
    "e presenti nei contesti produttivi di medie e grandi organizzazioni. La domanda per studi e "
    "imprese non è più «se» adottarli, ma «come» farlo in modo consapevole."
)

quote(doc,
    "L'agente AI non risponde soltanto: pianifica, decide e agisce, delegando subtask "
    "ad altri agenti specializzati e interagendo con i sistemi informativi aziendali."
)

# ---- SEZIONE 2 ----
h2(doc, "2. Architettura multi-agente: cosa cambia nella pratica")

body(doc,
    "Nei sistemi più evoluti non esiste un singolo agente tuttofare, bensì una rete di agenti "
    "specializzati orchestrati da un agente coordinatore. Un esempio concreto in ambito "
    "contabile-fiscale potrebbe essere:"
)
elenco(doc, [
    "Agente di estrazione dati: raccoglie e normalizza informazioni da fatture, estratti conto e documenti gestionali.",
    "Agente di analisi: verifica coerenze, segnala anomalie, confronta KPI con benchmark di settore.",
    "Agente normativo: consulta le fonti aggiornate (circolari Agenzia delle Entrate, Gazzetta Ufficiale, OIC) e verifica la conformità.",
    "Agente di reporting: redige bozze di relazioni, note integrative o comunicazioni al cliente.",
    "Agente orchestratore: coordina i precedenti, gestisce le eccezioni e coinvolge il professionista umano nelle decisioni ad alto valore o ambigue.",
])
body(doc,
    "Il professionista interviene nelle fasi che richiedono giudizio, responsabilità e relazione: "
    "non scompare dalla catena del valore, ma si concentra sul livello più alto di essa."
)

# ---- SEZIONE 3 ----
h2(doc, "3. Applicazioni concrete per studi professionali")

body(doc,
    "Per commercialisti, consulenti del lavoro, avvocati d'impresa e revisori, le applicazioni "
    "più immediate riguardano quattro aree:"
)
elenco(doc, [
    "Analisi di bilancio automatizzata: lettura, riclassificazione e commento di bilanci in pochi minuti, con segnalazione automatica degli indici critici.",
    "Due diligence documentale: analisi massiva di contratti, visure, atti notarili per l'individuazione di clausole rilevanti o rischi latenti.",
    "Monitoraggio fiscale continuativo: aggiornamento in tempo reale sulle scadenze, sulle novità normative e sugli impatti per ciascun cliente.",
    "Redazione assistita di atti e comunicazioni: bozze di istanze, ricorsi, relazioni di stima, lettere di incarico — riviste e validate dal professionista.",
])
body(doc,
    "Il guadagno di produttività stimato da diversi studi di settore oscilla tra il 30% e il 50% "
    "per le attività di back-office, con punte superiori nelle fasi di raccolta e normalizzazione dei dati."
)

nota(doc,
    "Attenzione: il risparmio di tempo ha valore solo se riorientato verso attività a maggiore "
    "valore aggiunto — consulenza strategica, presidio della relazione, sviluppo commerciale."
)

# ---- SEZIONE 4 ----
h2(doc, "4. Imprese: dall'automazione di processo all'intelligenza operativa")

body(doc,
    "Per le aziende — in particolare le PMI che ancora faticano a strutturare processi digitali "
    "di base — gli agenti AI offrono una leva straordinaria, a condizione di affrontare il tema "
    "con metodo."
)
body(doc,
    "Le aree di maggiore impatto operativo includono:"
)
elenco(doc, [
    "Ciclo attivo e passivo: riconciliazione automatica, gestione eccezioni, comunicazioni ai fornitori e ai clienti.",
    "Controllo di gestione in tempo reale: monitoraggio continuo di margini, scostamenti, cash flow previsto.",
    "Gestione del personale: onboarding documentale, monitoraggio scadenze contrattuali, supporto alle richieste HR.",
    "Intelligenza commerciale: analisi del mercato, monitoraggio dei concorrenti, scoring dei lead, redazione di offerte.",
])
body(doc,
    "Un aspetto spesso sottovalutato è l'impatto sulla qualità delle decisioni. Gli agenti non "
    "sostituiscono il management: forniscono un'analisi più rapida, più completa e meno soggetta "
    "a bias cognitivi, liberando l'attenzione dei responsabili per le scelte strategiche."
)

quote(doc,
    "Non si tratta di automatizzare il passato, ma di riprogettare i processi a partire "
    "dalle nuove capacità disponibili."
)

# ---- SEZIONE 5 ----
h2(doc, "5. Rischi e responsabilità: il perimetro che non va dimenticato")

body(doc,
    "L'entusiasmo per le potenzialità degli agenti AI non deve far abbassare la guardia sui "
    "rischi reali, che in un contesto professionale e aziendale assumono rilievo concreto:"
)
elenco(doc, [
    "Allucinazioni e confidenza eccessiva: i modelli linguistici possono generare informazioni plausibili ma errate. In ambito normativo o contrattuale, l'errore può avere conseguenze rilevanti.",
    "Sicurezza dei dati: gli agenti accedono a sistemi e dati sensibili. È indispensabile definire perimetri di accesso, log di audit e policy di data governance.",
    "Responsabilità professionale: l'output dell'agente non scarica il professionista dalla responsabilità. La validazione umana rimane un obbligo deontologico e spesso legale.",
    "Dipendenza dai fornitori: i modelli AI sono servizi cloud soggetti a variazioni di pricing, discontinuità e modifiche nei termini contrattuali.",
    "Bias nei dati di addestramento: i modelli riflettono i dati su cui sono stati addestrati, con possibili distorsioni in contesti specifici o di nicchia.",
])
body(doc,
    "L'AI Act europeo (in vigore dal 2026) introduce obblighi di trasparenza e conformità "
    "per i sistemi AI ad alto rischio. Studi professionali e aziende devono presidiare anche "
    "questo fronte normativo, che si intreccia con il GDPR e con le responsabilità dei "
    "data controller."
)

# ---- SEZIONE 6 ----
h2(doc, "6. Come iniziare: un approccio pragmatico")

body(doc,
    "Non esiste un percorso universale, ma è possibile indicare una traiettoria razionale per "
    "chi vuole avvicinarsi agli agenti AI in modo strutturato e sostenibile."
)
elenco(doc, [
    "Mappare i processi: identificare le attività ripetitive, ad alto volume e a bassa discrezionalità — sono i candidati ideali per la prima automazione.",
    "Definire metriche di successo: prima di implementare, stabilire cosa si vuole misurare (tempo risparmiato, tasso di errore, soddisfazione del cliente).",
    "Partire in piccolo, imparare in fretta: un pilota circoscritto permette di acquisire competenze, identificare criticità e costruire fiducia interna prima di scalare.",
    "Formare le persone: il cambiamento organizzativo è il fattore critico. La tecnologia è necessaria ma non sufficiente.",
    "Presidiare governance e compliance: definire chi è responsabile dell'output degli agenti, come vengono registrate le azioni, come si gestiscono gli errori.",
])
body(doc,
    "Il ruolo del consulente — che si tratti del commercialista, del consulente del lavoro "
    "o del revisore — è sempre più quello di accompagnare il cliente in questo percorso: "
    "non come esperto di tecnologia in senso stretto, ma come interprete di business capace "
    "di tradurre le potenzialità dell'AI in valore misurabile."
)

# ---- CONCLUSIONI ----
h2(doc, "Conclusioni: il vantaggio competitivo si costruisce ora")

body(doc,
    "Gli agenti AI non sono una promessa futura: sono una realtà operativa del 2026. "
    "Le organizzazioni che stanno acquisendo competenze oggi — che siano studi da dieci "
    "professionisti o aziende da cento dipendenti — si stanno costruendo un vantaggio "
    "difficile da recuperare per chi ritarda."
)
body(doc,
    "Il messaggio per i professionisti è chiaro: la domanda non è se l'AI cambierà il "
    "vostro lavoro, ma in che misura voi guiderete quel cambiamento. Chi adotta "
    "consapevolezza, metodo e una solida cultura del rischio sarà in grado di trasformare "
    "questa discontinuità tecnologica in un'opportunità concreta di differenziazione."
)
body(doc,
    "La consulenza di valore, nel 2026, si misura anche dalla capacità di aiutare i "
    "propri clienti a navigare intelligentemente nell'ecosistema dell'intelligenza artificiale."
)

# Linea finale
doc.add_paragraph("─" * 80)

# === NOTE A PIÈ DI ARTICOLO ===
p_note_label = add_paragraph_with_style(doc, "", space_before=8, space_after=4)
r = p_note_label.add_run("RIFERIMENTI E APPROFONDIMENTI")
set_font(r, size=8, bold=True, color=(100, 100, 100))

ref_items = [
    "European AI Act — Regolamento (UE) 2024/1689, applicazione progressiva 2024-2026",
    "Anthropic, «Building effective agents», blog tecnico, dicembre 2024",
    "McKinsey Global Institute, «The economic potential of generative AI», aggiornamento 2025",
    "CNDCEC — Commissione AI, «Linee guida per l'uso dell'intelligenza artificiale negli studi professionali», 2025",
    "OpenAI, Google DeepMind, Anthropic — documentazione tecnica framework agentici, 2025-2026",
]
for ref in ref_items:
    p_ref = add_paragraph_with_style(doc, "", space_before=1, space_after=2)
    p_ref.paragraph_format.left_indent = Cm(0.5)
    r = p_ref.add_run("• " + ref)
    set_font(r, size=8, color=(80, 80, 80))

# === BIO AUTORE ===
doc.add_paragraph("")
p_bio_box = add_paragraph_with_style(doc, "", space_before=12, space_after=4)
p_bio_box.paragraph_format.left_indent  = Cm(0.5)
p_bio_box.paragraph_format.right_indent = Cm(0.5)
r = p_bio_box.add_run("L'AUTORE  |  ")
set_font(r, size=9, bold=True, color=(0, 90, 160))
r2 = p_bio_box.add_run(
    "Alberto Mattavelli è consulente in innovazione digitale e intelligenza artificiale, "
    "con focus su studi professionali e PMI. Autore di articoli e formatore su temi di "
    "digital transformation, AI applicata alla consulenza e governance dei dati."
)
set_font(r2, size=9, color=(60, 60, 60))

# === SALVA ===
doc.save(OUTPUT_PATH)
print(f"Articolo salvato in: {OUTPUT_PATH}")
