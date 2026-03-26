#!/usr/bin/env python3
"""
Crea 2026-015 La memoria dello studio.docx
copiando lo stile di 2026-013 AI e junior.docx
"""
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from copy import deepcopy

SRC = Path("Articoli/RATIO/2026-013 AI e junior.docx")
DST = Path("Articoli/RATIO/2026-015 La memoria dello studio.docx")

shutil.copy2(SRC, DST)

doc = Document(DST)

# Contenuto articolo 2026-015
TITOLO_LABEL = "Titolo"
TITOLO_CONTENT = "La memoria dello studio: l\u2019AI che non dimentica mai"
SOTTOTITOLO_LABEL = "Sottotitolo"
SOTTOTITOLO_CONTENT = (
    "Ogni studio accumula anni di prassi non scritte. "
    "Il problema \u00e8 che quella conoscenza vive nelle teste delle persone \u2014 "
    "e quando se ne vanno, se ne va con loro. "
    "L\u2019AI pu\u00f2 diventare la memoria collettiva dello studio."
)
APPROFONDIMENTO_LABEL = "Approfondimento"
APPROFONDIMENTO_PARAS = [
    (
        "Il collaboratore pi\u00f9 bravo dello studio ha appena comunicato che se ne va. "
        "Buona opportunit\u00e0 di carriera, non si poteva fare altrimenti. "
        "Peccato che nella sua testa ci fossero tre anni di prassi non scritte: "
        "come preferisce ricevere la documentazione il cliente Rossi, "
        "perch\u00e9 con il cliente Bianchi bisogna sempre fare il punto prima di inviare "
        "qualsiasi comunicazione, quale template di analisi funziona meglio per le PMI "
        "del settore manifatturiero."
    ),
    (
        "Questo \u00e8 il problema della conoscenza tacita negli studi professionali: "
        "esiste, vale moltissimo, ma non \u00e8 da nessuna parte."
    ),
    (
        "L\u2019intelligenza artificiale offre oggi uno strumento concreto per affrontare "
        "questa fragilit\u00e0, a patto di usarla in modo sistematico e non episodico."
    ),
    (
        "Il primo livello \u00e8 quello delle istruzioni personalizzate. "
        "I principali assistenti AI \u2014 ChatGPT, Claude, Gemini \u2014 consentono di "
        "configurare istruzioni persistenti che definiscono il contesto professionale, "
        "il tono preferito, le aree di specializzazione e le prassi dello studio. "
        "Non si tratta di un\u2019operazione tecnica: \u00e8 un atto di formalizzazione "
        "della propria metodologia. Scrivere \u201cnel nostro studio utilizziamo questo "
        "schema di riclassificazione del bilancio\u201d o \u201cquando analizziamo un\u2019azienda "
        "in difficolt\u00e0 partiamo sempre da questi tre indicatori\u201d significa codificare "
        "un sapere che altrimenti rimarrebbe implicito."
    ),
    (
        "Il secondo livello riguarda i progetti strutturati. "
        "Molti strumenti AI permettono di creare spazi dedicati a specifici clienti "
        "o tipologie di lavoro, nei quali accumulare documenti, analisi e note nel tempo. "
        "Un progetto dedicato a un cliente diventa progressivamente pi\u00f9 ricco e "
        "contestualizzato: l\u2019AI risponde con la logica del caso specifico, non in astratto."
    ),
    (
        "Il terzo livello, forse il pi\u00f9 sottovalutato, \u00e8 quello della documentazione "
        "attiva. Alla fine di ogni analisi complessa, si pu\u00f2 chiedere all\u2019AI di produrre "
        "una sintesi metodologica: non solo \u201ccosa abbiamo trovato\u201d ma \u201ccome abbiamo "
        "ragionato\u201d. Quel documento diventa un asset dello studio, riutilizzabile, "
        "aggiornabile, trasferibile al prossimo collaboratore \u2014 senza che questi debba "
        "ricominciare da zero o aspettare che qualcuno si ricordi di spiegargli le cose."
    ),
    (
        "Il risultato non \u00e8 un sistema infallibile. "
        "L\u2019AI non sostituisce il giudizio professionale, n\u00e9 la relazione con il cliente. "
        "Ma rende meno fragile quella conoscenza che oggi dipende troppo dalle persone "
        "e troppo poco dai processi."
    ),
    (
        "La domanda da porsi non \u00e8 \u201cl\u2019AI pu\u00f2 diventare la memoria del nostro studio?\u201d "
        "ma \u201cstiamo facendo qualcosa per non perdere la memoria del nostro studio?\u201d "
        "L\u2019AI \u00e8 uno degli strumenti disponibili. Probabilmente il pi\u00f9 accessibile. "
        "Resta per\u00f2 inutile se non lo usiamo con metodo \u2014 che \u00e8 esattamente quello "
        "che il collaboratore pi\u00f9 bravo stava cercando di insegnarci, prima di salutarci."
    ),
]

# Mapping: paragraphs 0-12 nella struttura attesa
EXPECTED = [
    TITOLO_LABEL,           # 0
    TITOLO_CONTENT,         # 1
    SOTTOTITOLO_LABEL,      # 2
    SOTTOTITOLO_CONTENT,    # 3
    APPROFONDIMENTO_LABEL,  # 4
] + APPROFONDIMENTO_PARAS   # 5-12


def set_para_text(para, new_text):
    """Svuota tutti i runs e aggiunge il testo come nuovo run, preservando pPr."""
    p = para._p
    # Rimuovi tutti i run (w:r) esistenti
    for r in p.findall(qn('w:r')):
        p.remove(r)
    # Rimuovi bookmarkStart/End e proofErr che potrebbero interferire
    for tag in ['w:bookmarkStart', 'w:bookmarkEnd', 'w:proofErr']:
        for elem in p.findall(qn(tag)):
            p.remove(elem)
    # Aggiungi il nuovo testo come run pulito
    # Copia l'rPr dal primo run del template se esiste
    new_run = etree.SubElement(p, qn('w:r'))
    t = etree.SubElement(new_run, qn('w:t'))
    t.text = new_text
    if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def set_label_para_text(para, new_text, font_name='Aptos Display'):
    """Come set_para_text ma aggiunge anche rPr con il font specificato."""
    p = para._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    for tag in ['w:bookmarkStart', 'w:bookmarkEnd', 'w:proofErr']:
        for elem in p.findall(qn(tag)):
            p.remove(elem)
    new_run = etree.SubElement(p, qn('w:r'))
    # Crea rPr con font
    rPr = etree.SubElement(new_run, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsiaTheme'), 'minorEastAsia')
    rFonts.set(qn('w:cstheme'), 'minorBidi')
    t = etree.SubElement(new_run, qn('w:t'))
    t.text = new_text


# Verifica che i paragrafi 0-4 abbiano la struttura attesa
label_paras = [0, 2, 4]
print("Struttura 2026-013 (prima della modifica):")
for i in range(13):
    print(f"  Para {i}: '{doc.paragraphs[i].text[:60]}'")

print("\nApplico il nuovo contenuto...")

# Sostituisci label paragraphs (0, 2, 4)
set_label_para_text(doc.paragraphs[0], TITOLO_LABEL)
set_label_para_text(doc.paragraphs[2], SOTTOTITOLO_LABEL)
set_label_para_text(doc.paragraphs[4], APPROFONDIMENTO_LABEL)

# Sostituisci content paragraphs
set_para_text(doc.paragraphs[1], TITOLO_CONTENT)
set_para_text(doc.paragraphs[3], SOTTOTITOLO_CONTENT)

for i, text in enumerate(APPROFONDIMENTO_PARAS):
    set_para_text(doc.paragraphs[5 + i], text)

doc.save(DST)
print(f"Salvato: {DST}")

# Verifica
doc2 = Document(DST)
print("\nVerifica paragrafi:")
for i in range(13):
    print(f"  Para {i}: '{doc2.paragraphs[i].text[:80]}'")
