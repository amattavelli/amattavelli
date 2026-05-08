"""
Articolo 4 — Polizze RC professionale e AI
Ratio/articoli/2026-05-08_polizze-rc-errori-ai-professionisti.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-05-08_polizze-rc-errori-ai-professionisti.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Maggio 2026  |  Responsabilità e Gestione del Rischio")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("L'errore che la polizza non copre:\ncome l'AI sta cambiando la RC professionale")
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Le polizze di responsabilità civile professionale non sono state scritte per un "
    "mondo in cui il software suggerisce la detrazione sbagliata. "
    "Alcune compagnie stanno aggiornando le condizioni. La maggior parte ancora no."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  8 maggio 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Un commercialista usa un software tributario con funzioni AI per predisporre "
    "il modello 730 di un cliente. Il software suggerisce l'inserimento di una detrazione "
    "che in realtà non spetta: un errore di interpretazione normativa generato dall'algoritmo. "
    "Il professionista non verifica nel dettaglio, il cliente presenta la dichiarazione, "
    "l'Agenzia delle Entrate contesta e irroga sanzioni. Chi paga? La risposta, nella "
    "maggior parte dei casi, è: il commercialista. La domanda successiva, che molti non "
    "si sono ancora posti, è: la sua polizza RC professionale copre questo tipo di errore?"
)

para(doc,
    "Le polizze di responsabilità civile professionale per commercialisti, avvocati e "
    "consulenti del lavoro sono state progettate per un mondo in cui l'errore è umano. "
    "Il professionista valuta, decide, firma. Se sbaglia, la colpa è sua e la polizza lo "
    "copre. Con l'introduzione di sistemi AI nel flusso di lavoro, la catena causale si "
    "allunga: c'è uno strumento che elabora, suggerisce o decide, e poi c'è il professionista "
    "che accetta o non verifica abbastanza. Questo schema, che è diventato la normalità "
    "in molti studi nel corso del 2025, non è contemplato nella maggior parte dei contratti "
    "assicurativi in essere."
)

heading(doc, "Come il mercato assicurativo sta reagendo")

para(doc,
    "Il tema è entrato nell'agenda del mercato assicurativo italiano nel 2025 e nel 2026 "
    "si è concretizzato in alcune modifiche nelle polizze collettive degli ordini professionali. "
    "UIA International, associazione internazionale di avvocati, ha pubblicato una proposta "
    "di adeguamento delle coperture RC che esplicitamente richiama il rischio di danni "
    "causati da sistemi AI usati nello svolgimento dell'incarico. Alcune compagnie "
    "assicurative hanno iniziato a inserire clausole specifiche: alcune escludono i danni "
    "derivanti da output AI non verificati, altre li includono solo se il professionista "
    "dimostra di aver adottato procedure di controllo documentate."
)

para(doc,
    "Questa biforcazione è il punto cruciale per chi usa l'AI nel lavoro quotidiano. "
    "Non si tratta di sapere se la tua polizza copre l'AI in astratto: si tratta di sapere "
    "se copre il tuo processo di lavoro concreto. Se usi Genya, Datev, Wolters Kluwer o "
    "qualsiasi altro software tributario con funzioni AI integrate, e non hai mai riletto "
    "le condizioni della tua polizza aggiornate al 2025-2026, non sai cosa è coperto "
    "nel caso in cui l'AI generi un errore che supera il tuo controllo."
)

heading(doc, "Il quadro normativo che cambia le responsabilità")

para(doc,
    "La Legge 132/2025 ha chiarito che la responsabilità civile e deontologica del "
    "professionista non si trasferisce all'AI. Se il software sbaglia, il professionista "
    "risponde. Questa chiarezza normativa, però, lascia aperta una questione pratica: "
    "il professionista che dimostra di aver usato lo strumento seguendo le istruzioni del "
    "fornitore, di aver eseguito una verifica ragionevole e di aver documentato il processo "
    "ha un profilo di rischio diverso da chi non ha fatto nessuna di queste cose. "
    "Le compagnie assicurative stanno cominciando a fare questa distinzione nelle nuove "
    "condizioni di polizza."
)

para(doc,
    "Il concetto che sta emergendo nel mercato assicurativo è quello della diligenza "
    "aumentata: chi usa l'AI ha l'obbligo di controllare l'output in misura proporzionata "
    "al rischio dell'attività. Per un parere fiscale su un'operazione di ristrutturazione "
    "aziendale, la diligenza richiesta è molto maggiore che per la bozza di una lettera "
    "standard. Questo significa che la stessa polizza può coprire o non coprire lo stesso "
    "errore a seconda di come il professionista ha lavorato. Il fatto che il software "
    "abbia generato l'errore non sposta la responsabilità: sposta solo la complessità "
    "della prova."
)

heading(doc, "Cosa fare prima di rinnovare la polizza")

para(doc,
    "Per uno studio professionale, la conseguenza operativa è concreta. Aggiornare la "
    "propria polizza RC non basta: bisogna poter dimostrare, con documentazione interna, "
    "di aver adottato un processo di supervisione degli output AI. Non serve un protocollo "
    "sofisticato: serve una procedura scritta che dica quali tipi di output AI richiedono "
    "verifica, chi la effettua e come si traccia. Questo documento, che potrebbe sembrare "
    "una formalità burocratica, diventa la prova della diligenza professionale in caso "
    "di contestazione davanti all'ordine o in sede civile."
)

para(doc,
    "C'è un effetto collaterale di questo scenario che riguarda la relazione con i clienti. "
    "Se un cliente usa strumenti AI per preparare documentazione da fornire allo studio, "
    "per esempio bilanci gestionali o estratti conto elaborati da software con funzioni AI, "
    "il professionista che accetta quella documentazione senza una verifica adeguata assume "
    "un rischio che la sua polizza potrebbe non coprire. La catena dell'AI non riguarda "
    "solo quello che succede dentro lo studio: riguarda anche quello che arriva dall'esterno."
)

heading(doc, "La verifica che si rimanda sempre")

para(doc,
    "Il mercato assicurativo si adatta sempre con un ritardo rispetto alla tecnologia. "
    "Chi ha già rivisto la propria polizza alla luce dell'AI ha un vantaggio: sa cosa è "
    "coperto e cosa no. Chi non l'ha fatto, lavora con uno strumento di gestione del rischio "
    "costruito per un mondo che non esiste più. Chiedere all'assicuratore o al broker una "
    "lettura esplicita delle condizioni di copertura in presenza di AI nel flusso di lavoro "
    "non è un esercizio di cautela eccessiva: è la verifica di base che ogni professionista "
    "dovrebbe fare prima di rinnovare il contratto. Il rinnovo annuale, che molti trattano "
    "come un adempimento automatico, è il momento giusto per porre questa domanda. "
    "Se il rinnovo è già passato, il momento giusto è adesso."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "Legge 132/2025 — Disciplina organica sull'intelligenza artificiale in Italia",
    "UIA International — Proposta di adeguamento coperture RC professionale e AI (2026)",
    "KTS Finance — \"Assicurazione Professionale Commercialisti e IA: 3 Rischi del 730/2026\" (2026)",
    "EC News — \"AI e responsabilità del professionista: i nuovi confini del rischio\" (2026)",
    "Il Fatto Quotidiano — \"I professionisti che usano l'Intelligenza Artificiale hanno l'obbligo di "
    "informare i clienti e sono responsabili di eventuali errori\" (28 marzo 2026)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
