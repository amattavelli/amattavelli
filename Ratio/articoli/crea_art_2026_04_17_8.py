#!/usr/bin/env python3
"""
Articolo 4: Quanto vale davvero l'AI per le PMI italiane: ROI, numeri e aspettative reali
File: Ratio/articoli/2026-04-17_roi-ai-pmi-italiane-numeri-reali.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/home/user/amattavelli/Ratio/articoli/2026-04-17_roi-ai-pmi-italiane-numeri-reali.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(3)
sec.right_margin = Cm(3)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot)
    pPr.append(pBdr)


def para(doc, text, size=11, italic=False, color=None, sa=8, sb=0,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


# --- TESTATA ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.all_caps = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Economia e Gestione d\u2019Impresa")
r.font.name = 'Calibri'
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# --- TITOLO ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "Quanto vale davvero l\u2019AI per le PMI italiane: "
    "ROI, numeri e aspettative reali"
)
r.font.name = 'Calibri'
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Le promesse sull\u2019intelligenza artificiale sono spesso ottimistiche. "
    "I risultati reali nelle PMI italiane lo sono di meno, ma restano significativi. "
    "Una lettura dei dati disponibili per aiutare imprenditori e consulenti "
    "a ragionare con i numeri in mano."
)
r.font.name = 'Calibri'
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  \u2022  17 aprile 2026")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x60)

# --- CORPO ---
para(doc,
    "Circola un numero che viene citato spesso nei convegni sull\u2019AI: +37% di "
    "produttivit\u00e0. Lo si attribuisce, secondo i contesti, a ricerche di McKinsey, "
    "Gartner o ai fornitori stessi di software AI. Il problema di questo numero "
    "\u00e8 che raramente viene contestualizzato: +37% rispetto a cosa, in quale settore, "
    "con quale definizione di produttivit\u00e0, in quali condizioni di adozione. "
    "Questo non significa che sia falso. Significa che per un imprenditore o "
    "un consulente che deve prendere decisioni concrete, non \u00e8 sufficiente."
)

para(doc,
    "Proviamo allora a guardare i dati con un po\u2019 pi\u00f9 di granularit\u00e0, concentrandoci "
    "sul contesto italiano e sulle PMI. Secondo l\u2019elaborazione Istat del primo "
    "trimestre 2026, il 16,4% delle aziende italiane con pi\u00f9 di dieci dipendenti "
    "ha adottato soluzioni AI nel 2025. Ma la quota delle PMI con meno di cinquanta "
    "dipendenti che utilizza AI in modo sistematico \u00e8 stimata intorno al 7-8%. "
    "Il gap tra adozione dichiarata e utilizzo effettivo \u00e8 ampio, e dipende da "
    "fattori che hanno poco a che fare con la tecnologia: competenze interne, "
    "cultura aziendale, capacit\u00e0 di misurare i risultati."
)

h2(doc, "I settori dove il ROI \u00e8 pi\u00f9 documentato")

para(doc,
    "I dati di ritorno sull\u2019investimento pi\u00f9 affidabili per le PMI italiane "
    "provengono da tre aree funzionali. La prima \u00e8 il customer service: "
    "le aziende che hanno implementato chatbot e agenti AI per la gestione "
    "delle richieste in ingresso riportano una riduzione del carico sugli operatori "
    "umani tra il 40% e il 65%, con tempi di risposta medi scesi da ore a minuti. "
    "Il beneficio \u00e8 misurabile direttamente: meno ore di lavoro umano per pratica "
    "gestita, maggiore soddisfazione del cliente misurata con NPS."
)

para(doc,
    "La seconda area \u00e8 la gestione documentale e la contabilit\u00e0. Le aziende che "
    "hanno adottato software con AI integrata per la categorizzazione delle fatture, "
    "la riconciliazione bancaria e il controllo degli adempimenti riportano riduzioni "
    "del tempo dedicato a queste operazioni nell\u2019ordine del 30-50%. Il dato \u00e8 coerente "
    "con le stime dei principali fornitori di software gestionale italiani e con "
    "le valutazioni della Fondazione Nazionale di Ricerca dei Commercialisti. "
    "La terza area \u00e8 il marketing: l\u2019AI per la generazione di contenuti, "
    "la segmentazione della clientela e la personalizzazione delle comunicazioni "
    "mostra ritorni pi\u00f9 variabili, ma con picchi significativi nelle aziende "
    "che avevano gap strutturali nella produzione di contenuti."
)

h2(doc, "Dove il ROI \u00e8 pi\u00f9 lento del previsto")

para(doc,
    "Esiste una categoria di investimenti AI dove il ritorno \u00e8 reale ma richiede "
    "pi\u00f9 tempo del previsto: i sistemi di analisi predittiva e business intelligence. "
    "Molte PMI italiane hanno acquistato strumenti di questo tipo convinte di "
    "ottenere insight utili in poche settimane. La realt\u00e0 \u00e8 che questi sistemi "
    "richiedono dati storici strutturati, integrazione con le fonti esistenti, "
    "e un periodo di calibrazione che pu\u00f2 durare mesi. "
    "Senza un\u2019architettura dati di base sufficientemente ordinata, "
    "il modello AI produce output inaffidabili, e il ROI si allontana."
)

para(doc,
    "Un secondo ostacolo documentato \u00e8 il cosiddetto \u201cAI adoption gap\u201d: "
    "il divario tra la disponibilit\u00e0 dello strumento e la capacit\u00e0 del personale "
    "di usarlo efficacemente. Secondo l\u2019OCSE, il 46% dei lavoratori italiani "
    "necessita di aggiornamento delle competenze per restare competitivo. "
    "Nelle PMI, dove la formazione \u00e8 spesso trattata come un costo discrezionale, "
    "lo strumento AI acquistato viene sottoutilizzato perch\u00e9 nessuno in azienda "
    "sa sfruttarne le funzioni pi\u00f9 avanzate. Il risultato \u00e8 un ROI apparentemente "
    "basso che in realt\u00e0 \u00e8 un problema di implementazione, non di tecnologia."
)

h2(doc, "Come calcolare il ROI prima di investire")

para(doc,
    "La valutazione corretta di un investimento AI parte dall\u2019identificazione "
    "del processo da automatizzare, non dalla tecnologia. Il punto di partenza "
    "utile \u00e8 una stima del costo attuale del processo: quante ore di lavoro "
    "richiede, con quale frequenza, con quale margine di errore. A partire da "
    "questa base, \u00e8 possibile stimare il risparmio atteso dall\u2019automazione, "
    "confrontarlo con il costo del software (licenza, implementazione, formazione, "
    "manutenzione) e calcolare il tempo di ritorno dell\u2019investimento."
)

para(doc,
    "Un calcolo di questo tipo porta spesso a conclusioni diverse da quelle "
    "che si ottengono leggendo i white paper dei fornitori. Pu\u00f2 rivelare che "
    "un investimento apparentemente costoso ha un payback molto rapido, "
    "perch\u00e9 il processo che automatizza \u00e8 un collo di bottiglia critico. "
    "Oppure pu\u00f2 mostrare che uno strumento economico non vale la pena perch\u00e9 "
    "il processo che automatizza \u00e8 gi\u00e0 efficiente. La chiarezza sui numeri "
    "\u00e8 il prerequisito per qualsiasi decisione razionale sull\u2019AI in azienda."
)

h2(doc, "Il ruolo del consulente nella valutazione")

para(doc,
    "Il commercialista e il consulente aziendale hanno un ruolo specifico "
    "in questa fase di adozione dell\u2019AI nelle PMI italiane: quello di "
    "aiutare l\u2019imprenditore a fare domande giuste. Non \u201cquale AI devo comprare\u201d, "
    "ma \u201cquale problema voglio risolvere, quanto mi costa oggi, e quanto "
    "sono disposto a investire per risolverlo\u201d. Non \u201cl\u2019AI aumenter\u00e0 "
    "la mia produttivit\u00e0\u201d, ma \u201cin quale funzione, misurata come, "
    "rispetto a quale baseline\u201d."
)

para(doc,
    "Le aziende italiane che hanno ottenuto i risultati migliori dall\u2019AI "
    "sono quelle che hanno avviato l\u2019adozione con un progetto pilota circoscritto, "
    "misurato i risultati con KPI definiti prima dell\u2019implementazione, "
    "e scalato solo dopo aver validato i dati. \u00c8 un approccio meno spettacolare "
    "di quello che si legge nei comunicati stampa, ma \u00e8 quello che funziona. "
    "E il professionista che sa accompagnare un cliente in questo percorso "
    "offre un valore che nessun software pu\u00f2 sostituire."
)

sep(doc)

# --- FONTI ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Fonti e riferimenti")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Istat \u2014 Indagine sull\u2019adozione dell\u2019AI nelle imprese italiane, Q1 2026",
    "Fondazione Nazionale di Ricerca dei Commercialisti \u2014 Indagine sull\u2019uso "
    "dell\u2019AI negli studi professionali, luglio-settembre 2025",
    "OCSE \u2014 Competenze e mercato del lavoro nell\u2019era dell\u2019intelligenza artificiale (2025)",
    "AI4Business \u2014 AI 2026: l\u2019anno dell\u2019adozione sistemica e delle scelte "
    "irreversibili nelle aziende",
    "Kinetikon \u2014 Agenti IA nelle PMI italiane: adozione e casi d\u2019uso (2026)",
    "Cosmonet \u2014 AI Agent Italia 2026: la rivoluzione degli assistenti virtuali "
    "per le aziende",
    "Economy Magazine \u2014 La grande nuova frontiera dell\u2019AI per le aziende (2026)",
]
for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"\u2022 {s}")
    r.font.name = 'Calibri'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte")
r.font.name = 'Calibri'
r.font.size = Pt(8)
r.italic = True
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

doc.save(OUTPUT)
print(f"Salvato: {OUTPUT}")
