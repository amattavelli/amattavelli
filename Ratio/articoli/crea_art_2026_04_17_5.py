#!/usr/bin/env python3
"""
Articolo 1: AI Agentica: quando il software agisce senza che tu lo chieda
File: Ratio/articoli/2026-04-17_ai-agentica-imprese-italiane-2026.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/home/user/amattavelli/Ratio/articoli/2026-04-17_ai-agentica-imprese-italiane-2026.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(3)
sec.right_margin = Cm(3)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot)
    pPr.append(pBdr)


def para(doc, text, size=11, italic=False, color=None, sa=8, sb=0,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


# --- TESTATA ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.all_caps = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Innovazione e Strategia")
r.font.name = 'Calibri'
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# --- TITOLO ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "AI Agentica: quando il software agisce senza che tu lo chieda"
)
r.font.name = 'Calibri'
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Nel 2026 la frontiera non \u00e8 pi\u00f9 l\u2019AI che risponde alle domande: "
    "\u00e8 l\u2019AI che pianifica, decide e agisce. Le aziende italiane iniziano "
    "a fare i conti con un cambiamento molto pi\u00f9 profondo del previsto."
)
r.font.name = 'Calibri'
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  \u2022  17 aprile 2026")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# --- CORPO ---
para(doc,
    "Per anni l\u2019intelligenza artificiale nelle aziende ha significato una cosa sola: "
    "un sistema che risponde. Fai una domanda, ottieni una risposta. Carichi un documento, "
    "ricevi un riassunto. L\u2019AI generativa, quella dei modelli linguistici grandi che hanno "
    "dominato il dibattito dal 2022 in poi, funzionava fondamentalmente cos\u00ec: riceveva "
    "un input e restituiva un output. Utile, certamente. Ma ancora lontana dall\u2019agire "
    "autonomo. Nel 2026 questo modello sta cambiando, e la transizione in corso ha un nome "
    "preciso: AI agentica."
)

para(doc,
    "Un agente AI non si limita a rispondere. Riceve un obiettivo, lo scompone in sotto-task, "
    "decide quali strumenti usare, esegue azioni nel mondo digitale, verifica i risultati e "
    "itera. Pu\u00f2 inviare email, accedere a database, compilare moduli, aggiornare sistemi "
    "gestionali, attivare processi su altri software. Non ha bisogno che un utente approvi "
    "ogni singolo passaggio: pu\u00f2 lavorare in autonomia per ore, scalando un obiettivo "
    "che prima avrebbe richiesto il lavoro coordinato di pi\u00f9 persone."
)

para(doc,
    "Il termine \u201cagente\u201d non \u00e8 nuovo nell\u2019informatica, ma la sua applicazione pratica "
    "ai modelli linguistici \u00e8 diventata concreta solo negli ultimi diciotto mesi. Piattaforme "
    "come Claude di Anthropic, GPT-4o di OpenAI e Gemini di Google hanno introdotto funzioni "
    "di orchestrazione che permettono ai modelli di usare strumenti esterni, mantenere memoria "
    "tra sessioni e collaborare tra loro. Il risultato \u00e8 una categoria di software che si "
    "comporta, almeno parzialmente, come un collaboratore digitale che sa cosa fare "
    "senza aspettare istruzioni passo dopo passo."
)

h2(doc, "Cosa sta succedendo nelle aziende italiane")

para(doc,
    "Secondo i dati Istat elaborati nel primo trimestre 2026, il 16,4% delle aziende italiane "
    "con pi\u00f9 di dieci dipendenti ha adottato soluzioni di AI nel 2025, un dato raddoppiato "
    "rispetto all\u20198,2% dell\u2019anno precedente. La crescita \u00e8 reale, ma nasconde una "
    "distinzione importante: la maggioranza delle implementazioni riguarda ancora AI assistiva, "
    "cio\u00e8 strumenti che supportano l\u2019operatore senza sostituirne le decisioni. "
    "L\u2019AI agentica \u00e8 presente in una quota molto pi\u00f9 ridotta, stimata intorno al 7-8% "
    "delle PMI, concentrata soprattutto nelle funzioni di customer service, gestione ordini "
    "e monitoraggio finanziario."
)

para(doc,
    "I settori che mostrano l\u2019adozione pi\u00f9 rapida sono quelli con processi ad alto volume "
    "e bassa variabilit\u00e0: logistica, e-commerce, gestione documentale, fatturazione attiva. "
    "Un agente che verifica le fatture in entrata, le riconcilia con i movimenti bancari e "
    "segnala le anomalie pu\u00f2 comprimere a pochi minuti un lavoro che normalmente impegna "
    "un operatore per ore. In ambito customer service, i dati di settore indicano che gli "
    "agenti AI sono in grado di risolvere autonomamente tra il 60% e l\u201985% delle richieste "
    "in ingresso, scalando all\u2019operatore umano solo i casi che richiedono giudizio "
    "contestuale o empatia relazionale."
)

h2(doc, "Il nodo della supervisione umana")

para(doc,
    "L\u2019autonomia degli agenti AI apre una questione che le aziende non possono ignorare: "
    "chi \u00e8 responsabile di ci\u00f2 che l\u2019agente fa? Se un sistema automatico invia una "
    "comunicazione errata a un cliente, approva un pagamento non autorizzato o cancella "
    "un record nel gestionale, la responsabilit\u00e0 ricade comunque sull\u2019azienda che "
    "ha attivato il sistema. L\u2019AI Act europeo, con le scadenze operative che si "
    "avvicinano ad agosto 2026, richiede esplicitamente che i sistemi AI ad alto rischio "
    "mantengano un livello di supervisione umana documentato e verificabile. "
    "Per gli agenti AI che operano in aree sensibili, come il credito, la salute o "
    "la gestione del personale, questo requisito diventa stringente."
)

para(doc,
    "La risposta operativa che molte aziende stanno adottando \u00e8 il modello \u201chuman-in-the-loop\u201d: "
    "l\u2019agente lavora in autonomia su task a basso rischio, ma richiede approvazione esplicita "
    "per azioni con impatto elevato. Un sistema di questo tipo pu\u00f2 gestire decine di "
    "operazioni al giorno senza intervento umano, chiedendo conferma solo nei casi in cui "
    "il valore in gioco o la complessit\u00e0 superano una soglia predefinita. \u00c8 un equilibrio "
    "che richiede progettazione attenta, ma che rende l\u2019automazione sostenibile anche in "
    "contesti regolamentati."
)

h2(doc, "Cosa cambia per il professionista e per il manager")

para(doc,
    "L\u2019arrivo degli agenti AI non elimina il lavoro umano: lo ridistribuisce. "
    "Le attivit\u00e0 ripetitive, a basso valore aggiunto e alta prevedibilit\u00e0 vengono "
    "assorbite dagli agenti. Le attivit\u00e0 che richiedono giudizio, relazione, creativit\u00e0 "
    "e responsabilit\u00e0 restano saldamente nelle mani delle persone. "
    "Il rischio non \u00e8 la sostituzione, ma la marginalizzazione di chi non impara "
    "a lavorare con questi sistemi. Il professionista che sa definire obiettivi chiari "
    "per un agente AI, verificarne l\u2019output e integrarne il lavoro nel proprio processo "
    "diventa esponenzialmente pi\u00f9 produttivo di chi ancora affronta ogni task manualmente."
)

para(doc,
    "Per i manager, la sfida \u00e8 di governance. Introdurre un agente AI in azienda significa "
    "decidere quali processi automatizzare, con quali vincoli, con quale livello di "
    "supervisione e con quale piano di fallback quando il sistema commette errori. "
    "Significa anche formare il personale a convivere con un sistema che agisce, "
    "non solo suggerisce. Le aziende che stanno gestendo meglio questa transizione "
    "sono quelle che hanno iniziato con piloti circoscritti, misurato i risultati "
    "con metriche chiare e allargato progressivamente il perimetro di autonomia "
    "solo dove i dati giustificavano la fiducia."
)

sep(doc)

# --- FONTI ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Fonti e riferimenti")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Istat \u2014 Indagine sull\u2019adozione dell\u2019AI nelle imprese italiane, Q1 2026",
    "Corriere Nazionale \u2014 La nuova era dell\u2019efficienza: come l\u2019AI agentica sta trasformando "
    "il lavoro e l\u2019impresa in Italia nel 2026 (16 aprile 2026)",
    "Kinetikon \u2014 Agenti IA nelle PMI italiane: adozione e casi d\u2019uso (2026)",
    "AI4Business \u2014 AI 2026: l\u2019anno dell\u2019adozione sistemica e delle scelte irreversibili "
    "nelle aziende",
    "Alessandria Today \u2014 2026: il passaggio dall\u2019intelligenza artificiale conversazionale "
    "agli agenti autonomi (gennaio 2026)",
    "Security Open Lab \u2014 Agenti AI e nuove regolamentazioni: il quadro normativo italiano "
    "ed europeo 2026 per le imprese",
]
for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"\u2022 {s}")
    r.font.name = 'Calibri'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte")
r.font.name = 'Calibri'
r.font.size = Pt(8)
r.italic = True
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

doc.save(OUTPUT)
print(f"Salvato: {OUTPUT}")
