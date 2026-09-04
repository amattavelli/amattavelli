"""
Quattro articoli Ratio -- 4 settembre 2026

1. 2026-09-04_simest-200-milioni-ai-pmi-dal-21-settembre.docx
2. 2026-09-04_doppia-compliance-ai-act-legge-132-imprese-italiane.docx
3. 2026-09-04_decreti-attuativi-ai-act-agosto-2026-imprese.docx
4. 2026-09-04_agenti-ai-processi-aziendali-dal-chatbot-all-operatore.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# SIMEST: 200 milioni per l'AI nelle PMI dal 21 settembre
# ============================================================

doc = new_doc()
testata(doc, "Settembre 2026", "Finanza Agevolata e Incentivi AI")
titolo(
    doc,
    "SIMEST apre il bando AI: 200 milioni per le PMI dal 21 settembre.",
    "Dal 21 settembre 2026 le PMI con vocazione internazionale possono accedere "
    "a 200 milioni di euro di finanziamenti agevolati per investimenti in intelligenza artificiale. "
    "La misura SIMEST copre sviluppo, adozione e integrazione di soluzioni AI. "
    "Chi ha i requisiti non dovrebbe aspettare.",
    "A cura della Redazione Ratio  •  4 settembre 2026"
)

para(doc,
    "Nel 1962, il governo italiano varò la Cassa per il Mezzogiorno "
    "con una dotazione senza precedenti per l'epoca. "
    "Molte imprese meridionali non ne seppero mai l'esistenza, "
    "o la scoprirono quando i fondi erano già esauriti. "
    "La storia degli incentivi pubblici in Italia "
    "si ripete con una certa regolarità: "
    "le risorse ci sono, l'accesso è complicato, "
    "l'informazione arriva in ritardo. "
    "Dal 21 settembre 2026, SIMEST — la società del Gruppo CDP "
    "che supporta l'internazionalizzazione delle imprese italiane — "
    "apre le domande per una nuova misura dedicata agli investimenti "
    "in intelligenza artificiale e tecnologie quantistiche, "
    "con una dotazione di 200 milioni di euro. "
    "Si tratta di una delle misure più concrete e operative "
    "attualmente disponibili per le PMI italiane che vogliono adottare l'AI "
    "su processi reali, non solo sperimentarla."
)

para(doc,
    "La misura si rivolge alle PMI con almeno il 3% del fatturato realizzato con l'estero "
    "— una soglia significativamente più bassa rispetto al 10% normalmente richiesto "
    "dal Fondo rotativo 394/1981 per altre agevolazioni SIMEST. "
    "Possono accedere anche i subfornitori delle imprese esportatrici, "
    "purché destinino almeno il 10% del proprio fatturato a clienti che esportano. "
    "In pratica, la platea di potenziali beneficiari è molto più ampia "
    "di quanto il nome 'misura per esportatori' possa suggerire: "
    "include produttori manifatturieri con componenti di export, "
    "studi di servizi che lavorano per clienti internazionali, "
    "fornitori di filiere orientate ai mercati esteri."
)

heading(doc, "Le condizioni economiche della misura")

para(doc,
    "Il finanziamento SIMEST AI ha una struttura agevolata su tre livelli. "
    "Il primo elemento è il tasso agevolato: il finanziamento viene erogato "
    "a condizioni inferiori a quelle di mercato, con una durata fino a otto anni "
    "e un periodo di preammortamento fino a due anni — "
    "il che significa che l'impresa può iniziare a restituire il capitale "
    "solo dopo aver avuto il tempo di generare i ritorni dall'investimento. "
    "Il secondo elemento è l'anticipo: le imprese ammesse "
    "possono ricevere fino al 50% dell'importo finanziato in anticipo, "
    "prima che le spese siano sostenute. "
    "Questo risolve uno dei principali problemi di liquidità "
    "che frenano le PMI nell'accesso ai bandi tradizionali. "
    "Il terzo elemento è il contributo a fondo perduto: "
    "fino al 10% dell'intervento agevolativo, "
    "nei limiti del regime de minimis (200.000 euro in tre anni). "
    "Per una PMI con un investimento AI da 300.000 euro, "
    "questo significa potenzialmente 30.000 euro non rimborsabili."
)

heading(doc, "Cosa finanzia la misura")

para(doc,
    "La misura copre lo sviluppo, l'adozione e l'integrazione di soluzioni AI "
    "nei processi aziendali. In termini pratici, questo include "
    "l'acquisto o lo sviluppo di software con componenti AI, "
    "l'integrazione di strumenti AI nei gestionali esistenti, "
    "i costi di consulenza per l'impostazione dei progetti, "
    "la formazione del personale sugli strumenti adottati "
    "e i costi di hardware necessari per l'elaborazione dei dati. "
    "Sono esclusi gli abbonamenti a servizi AI generici "
    "— il ChatGPT Team da 30 euro al mese non è un investimento "
    "ammissibile in questa misura — "
    "mentre rientrano i progetti con un investimento strutturato "
    "e un obiettivo di integrazione nei processi aziendali. "
    "La distinzione è tra 'uso dell'AI' e 'investimento nell'AI': "
    "la misura finanzia la seconda categoria."
)

heading(doc, "Come prepararsi prima del 21 settembre")

para(doc,
    "Le domande aprono il 21 settembre attraverso il portale SIMEST. "
    "Per arrivare pronti, le PMI interessate devono completare alcune verifiche "
    "nei giorni che restano. "
    "Prima: verificare il requisito di export "
    "— estrarre dal bilancio 2025 la quota di fatturato estera "
    "o identificare i clienti esportatori che giustificano la partecipazione come subfornitori. "
    "Seconda: identificare chiaramente il progetto AI "
    "— cosa si vuole fare, su quale processo, "
    "con quale fornitore o soluzione, con quale investimento stimato. "
    "Terza: raccogliere la documentazione aziendale standard "
    "— visura camerale aggiornata, bilanci degli ultimi due esercizi, "
    "documento antimafia se richiesto. "
    "Chi arriva al 21 settembre con la documentazione pronta "
    "ha un vantaggio concreto: "
    "i fondi sono contingentati e le domande vengono processate nell'ordine in cui arrivano."
)

para(doc,
    "La Cassa per il Mezzogiorno chiuse nel 1984 dopo vent'anni di attività. "
    "Molte delle risorse distribuite produssero poco perché "
    "arrivarono alle imprese sbagliate, nel momento sbagliato, "
    "senza un progetto chiaro dietro. "
    "SIMEST AI può funzionare diversamente: "
    "le risorse ci sono, la finestra è aperta dal 21 settembre, "
    "e il progetto lo deve portare l'imprenditore. "
    "Il professionista che lo aiuta a farlo arrivare pronto "
    "vale quanto il finanziamento stesso."
)

riferimenti(doc, [
    "MySolution — 'IA nelle PMI: dal 21 settembre il nuovo finanziamento SIMEST' (agosto 2026)",
    "Industria Italiana — 'Simest mette 200 milioni sull'AI: finanziamenti alle Pmi fino a otto anni' (2026)",
    "Tom's Hardware — 'Fino a 200 milioni per l'AI nelle PMI esportatrici, domande dal 21 settembre' (2026)",
    "Edotto.com — 'Fondo Simest per l'AI: 200 milioni alle Pmi dal 21 settembre' (2026)",
    "Incentivimpresa.it — 'Bandi Intelligenza Artificiale 2026: Finanziamenti AI PMI'",
    "Regolamento CE n. 394/1981 — Fondo rotativo per l'internazionalizzazione delle imprese",
    "Agenzia delle Entrate — Regime de minimis: soglia 200.000 euro in tre anni",
])
doc.save(BASE + "2026-09-04_simest-200-milioni-ai-pmi-dal-21-settembre.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# La doppia compliance che inchioda le imprese: AI Act + Legge 132
# ============================================================

doc = new_doc()
testata(doc, "Settembre 2026", "Normativa AI e Compliance")
titolo(
    doc,
    "Due regimi, un'azienda. La doppia compliance che nessuno ha spiegato bene.",
    "Le imprese italiane non devono rispettare solo l'AI Act europeo. "
    "Devono rispettare anche la Legge 23 settembre 2025 n. 132, "
    "che aggiunge obblighi nazionali distinti: sulle relazioni di lavoro, "
    "sui deepfake, sul diritto d'autore. "
    "Due regimi, due autorità, due sistemi di sanzioni. "
    "Molte aziende non lo sanno ancora.",
    "A cura della Redazione Ratio  •  4 settembre 2026"
)

para(doc,
    "Nel 1993, quando l'Italia recepì la Direttiva IVA comunitaria, "
    "il legislatore nazionale aggiunse una serie di adempimenti "
    "che la Direttiva non richiedeva: fatturazioni specifiche, "
    "registri supplementari, comunicazioni periodiche. "
    "I commercialisti che seguivano solo il testo europeo "
    "si trovarono in difetto. "
    "Trent'anni dopo, il copione si ripete con l'intelligenza artificiale. "
    "Il Regolamento UE 2024/1689 — l'AI Act — "
    "stabilisce un quadro europeo di regole sull'AI. "
    "Ma l'Italia, come consentito dal Regolamento, "
    "ha aggiunto un livello normativo nazionale: "
    "la Legge 23 settembre 2025, n. 132, "
    "entrata in vigore il 10 ottobre 2025 "
    "e ora completata dai due decreti attuativi approvati "
    "dal Consiglio dei Ministri il 4 agosto 2026. "
    "Per le imprese italiane, la compliance AI non è un binario: è una doppia corsia."
)

para(doc,
    "L'AI Act europeo fissa le regole sui sistemi AI classificati per livello di rischio: "
    "vietati, ad alto rischio, a rischio limitato, a rischio minimo. "
    "Stabilisce obblighi di trasparenza, AI literacy, governance dei modelli di uso generale. "
    "Designa la Commissione europea e l'AI Office come autorità di vigilanza sovranazionali. "
    "La Legge 132/2025 agisce su un piano diverso: "
    "integra l'AI Act con disposizioni nazionali "
    "che il Regolamento europeo lasciava alla discrezionalità degli Stati membri. "
    "E aggiunge materie che l'AI Act non toccava: "
    "la disciplina penale dei deepfake lesivi, "
    "la protezione del diritto d'autore per le opere create con assistenza AI, "
    "gli obblighi informativi delle imprese verso i lavoratori "
    "che usano sistemi AI nelle proprie mansioni."
)

heading(doc, "I tre obblighi nuovi della Legge 132 che non vengono dall'AI Act")

para(doc,
    "Il primo obbligo riguarda i lavoratori. "
    "La Legge 132 impone alle imprese che utilizzano sistemi AI "
    "nei processi di lavoro di informare preventivamente i lavoratori "
    "sull'uso di tali sistemi, sui dati elaborati, "
    "sulle modalità di supervisione e sulle eventuali conseguenze "
    "sulle condizioni di lavoro. "
    "Questo obbligo non deriva dall'AI Act — "
    "che disciplina i sistemi AI usati per prendere decisioni sui lavoratori "
    "solo nella categoria ad alto rischio — "
    "ma dalla Legge 132, che lo estende a un perimetro più ampio. "
    "Il secondo obbligo riguarda i deepfake. "
    "La Legge 132 ha introdotto nel Codice Penale il reato "
    "di diffusione illecita di contenuti generati da AI "
    "che ledono l'immagine o la dignità di persone reali. "
    "Non è un obbligo diretto per le imprese, "
    "ma genera responsabilità per le aziende "
    "che producono o diffondono contenuti AI senza i controlli adeguati. "
    "Il terzo obbligo riguarda il diritto d'autore. "
    "La Legge 132 stabilisce che le opere create con assistenza AI significativa "
    "devono indicare questa circostanza: "
    "chi pubblica contenuti aziendali generati da AI "
    "senza disclosure adeguata può incorrere in violazioni "
    "dei diritti di terzi o in contestazioni sulla paternità dell'opera."
)

heading(doc, "Le due autorità che vigilano")

para(doc,
    "La doppia compliance produce anche una doppia vigilanza. "
    "Per l'AI Act europeo, in Italia le autorità designate "
    "sono l'Agenzia per l'Italia Digitale (AgID) "
    "— come autorità di notifica — "
    "e l'Agenzia per la Cybersicurezza Nazionale (ACN) "
    "— come autorità di sorveglianza del mercato, "
    "con poteri di ispezione e sanzione. "
    "Per la Legge 132, si aggiungono i decreti attuativi "
    "approvati il 4 agosto 2026, "
    "che definiscono le competenze dell'ACN anche nel contesto nazionale "
    "e le regole sull'uso dell'AI nella formazione. "
    "Le sanzioni previste dall'AI Act arrivano fino al 7% del fatturato mondiale "
    "o 35 milioni di euro per i sistemi ad alto rischio; "
    "fino al 3% o 15 milioni per le violazioni degli obblighi di trasparenza. "
    "La Legge 132 aggiunge sanzioni specifiche "
    "sul versante dei rapporti di lavoro e della responsabilità penale per i deepfake."
)

heading(doc, "Come impostare la compliance duale")

para(doc,
    "Per un'impresa o uno studio professionale, "
    "affrontare la doppia compliance in modo ordinato richiede "
    "un approccio a due livelli separati ma coordinati. "
    "Il primo livello è la mappatura dei sistemi AI in uso: "
    "identificare ogni strumento con componenti AI, "
    "classificarlo per categoria di rischio AI Act "
    "e verificare se il suo utilizzo rientra "
    "negli obblighi della Legge 132 sui rapporti di lavoro. "
    "Il secondo livello è la documentazione: "
    "per l'AI Act, le disclosure ai clienti e la formazione del personale; "
    "per la Legge 132, le informative ai lavoratori "
    "e le policy sui contenuti AI pubblicati. "
    "Il professionista che affianca l'impresa "
    "deve conoscere entrambi i livelli: "
    "consigliare solo l'AI Act lascia l'impresa esposta "
    "agli obblighi nazionali, e viceversa."
)

para(doc,
    "Nel 1993, i commercialisti che avevano letto solo la Direttiva IVA "
    "e non il decreto legislativo di recepimento "
    "si trovarono a rispiegare ai clienti "
    "perché c'erano adempimenti in più. "
    "La situazione con l'AI è identica, "
    "con la differenza che le sanzioni sono un ordine di grandezza superiori "
    "e i sistemi sono già in produzione. "
    "Meglio leggere entrambe le norme adesso."
)

riferimenti(doc, [
    "Tom's Hardware — 'La doppia compliance che inchioda le imprese italiane all'AI Act' (2026)",
    "Legge 23 settembre 2025, n. 132 — Disposizioni in materia di intelligenza artificiale",
    "Regolamento UE 2024/1689 (AI Act) — in vigore dal 2 agosto 2026",
    "Federprivacy — 'Il Consiglio dei Ministri approva i decreti di adeguamento all'AI Act' (agosto 2026)",
    "Diritto Mercato Tecnologia — 'AI Act: il Governo approva definitivamente i decreti di adeguamento' (2026)",
    "FISCOeTASSE.com — 'Regolamento intelligenza artificiale in Italia 2026: approvati i decreti' (2026)",
    "Quotidianopiù — 'AI Act: imprese nel labirinto della compliance' (2026)",
])
doc.save(BASE + "2026-09-04_doppia-compliance-ai-act-legge-132-imprese-italiane.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# I decreti attuativi AI Act di agosto 2026: cosa cambia davvero
# ============================================================

doc = new_doc()
testata(doc, "Settembre 2026", "Normativa AI e Compliance")
titolo(
    doc,
    "I decreti AI di agosto. Cosa cambia, chi vigila, cosa fare.",
    "Il 4 agosto 2026 il Consiglio dei Ministri ha approvato "
    "i due decreti legislativi che completano il quadro normativo italiano sull'AI. "
    "Uno riguarda i poteri delle autorità nazionali. "
    "L'altro l'uso dell'AI nella formazione. "
    "Per le imprese, cambia chi può fare ispezioni e con quali poteri.",
    "A cura della Redazione Ratio  •  4 settembre 2026"
)

para(doc,
    "Nel 1974, quando l'Italia recepì la Direttiva europea sulla sicurezza sul lavoro, "
    "molte imprese non modificarono nulla per i due anni successivi. "
    "La norma era in vigore, ma le autorità di vigilanza "
    "non avevano ancora gli strumenti operativi per applicarla. "
    "Poi, quando gli ispettorati furono potenziati, "
    "arrivarono le prime sanzioni. "
    "Chi aveva aspettato l'ispezione per adeguarsi "
    "si trovò nella peggiore posizione possibile: "
    "nel mezzo di una verifica. "
    "Il 4 agosto 2026, il Consiglio dei Ministri ha approvato "
    "in via definitiva i due decreti legislativi "
    "di adeguamento dell'ordinamento italiano all'AI Act. "
    "Erano stati approvati in esame preliminare il 10 giugno 2026 "
    "e trasmessi al Parlamento per i pareri delle commissioni competenti. "
    "Con l'approvazione definitiva, "
    "il quadro normativo italiano sull'AI è ora completo. "
    "E le autorità hanno i poteri per usarlo."
)

para(doc,
    "I due decreti attuativi della Legge 132/2025 "
    "intervengono su due aree distinte. "
    "Il primo decreto è quello sulle autorità nazionali competenti: "
    "definisce le funzioni dell'Agenzia per la Cybersicurezza Nazionale (ACN) "
    "come autorità di vigilanza del mercato AI in Italia, "
    "e dell'Agenzia per l'Italia Digitale (AgID) "
    "come autorità di notifica per i sistemi ad alto rischio. "
    "Il secondo decreto riguarda l'uso dell'AI nella formazione: "
    "stabilisce i criteri per l'utilizzo di sistemi AI "
    "nei percorsi educativi e formativi, "
    "con obblighi di trasparenza e limiti specifici "
    "per i sistemi che incidono sulla valutazione degli studenti."
)

heading(doc, "I nuovi poteri dell'ACN")

para(doc,
    "Il decreto sulle autorità nazionali attribuisce all'ACN "
    "poteri di vigilanza, ispezione e sanzione "
    "che rendono operativa la compliance AI in Italia. "
    "In concreto, l'ACN può richiedere alle imprese "
    "la documentazione sui sistemi AI in uso, "
    "accedere alle loro sedi per effettuare verifiche, "
    "ordinare la sospensione di sistemi AI non conformi "
    "e applicare le sanzioni previste dall'AI Act "
    "— fino al 3% del fatturato mondiale per le violazioni di trasparenza, "
    "fino al 7% per i sistemi ad alto rischio non conformi. "
    "I poteri dell'ACN si aggiungono a quelli della Commissione europea "
    "e dell'AI Office, che mantengono la competenza "
    "sui fornitori di modelli di uso generale "
    "e sui casi più gravi di violazione sistematica. "
    "Per le PMI, l'autorità con cui avranno più probabilità di interfacciarsi "
    "è l'ACN: è quella che conosce il territorio italiano "
    "e che ha i poteri di ispezione diretta."
)

heading(doc, "Cosa verificherà l'ACN nelle imprese")

para(doc,
    "Sulla base dei poteri attribuiti dal decreto, "
    "un'ispezione ACN su un'impresa che usa sistemi AI "
    "si concentrerà su tre aree principali. "
    "La prima è la mappatura dei sistemi: "
    "l'impresa sa quali sistemi AI usa? "
    "Li ha classificati per categoria di rischio? "
    "Ha un registro aggiornato? "
    "La seconda è la trasparenza: "
    "i clienti, i lavoratori e gli interlocutori dell'impresa "
    "sono informati dell'uso dell'AI nelle interazioni con loro? "
    "Esiste una disclosure leggibile e accessibile? "
    "La terza è la formazione: "
    "il personale che usa sistemi AI "
    "ha ricevuto formazione adeguata? "
    "Esiste documentazione che lo attesti? "
    "Chi arriva a un'ispezione senza risposta a queste tre domande "
    "è esposto alle sanzioni, anche se i sistemi in uso "
    "non sono tecnicamente 'ad alto rischio'."
)

heading(doc, "Il calendario degli adempimenti aggiornato")

para(doc,
    "Con i decreti di agosto, il quadro degli adempimenti "
    "si può riassumere in tre scadenze operative. "
    "Prima scadenza — già scaduta il 2 agosto 2026: "
    "trasparenza per i sistemi AI a contatto con persone, "
    "AI literacy del personale, "
    "governance dei modelli di uso generale. "
    "Chi non è ancora in regola su questi punti "
    "deve adeguarsi immediatamente. "
    "Seconda scadenza — entro il 10 ottobre 2026: "
    "completare gli adempimenti della Legge 132/2025 "
    "sui rapporti di lavoro e sull'informativa ai lavoratori "
    "sull'uso di sistemi AI nelle loro mansioni. "
    "Terza scadenza — entro il 2 dicembre 2027: "
    "valutazione della conformità per i sistemi ad alto rischio "
    "indicati nell'Allegato III dell'AI Act. "
    "Le prime due scadenze sono già nell'orizzonte immediato. "
    "La terza dà tempo di prepararsi, "
    "non di non prepararsi."
)

para(doc,
    "Le imprese che nel 1974 aspettarono l'ispezione "
    "per adeguarsi alla normativa sulla sicurezza sul lavoro "
    "non fecero una scelta sbagliata solo eticamente: "
    "fecero una scelta sbagliata anche economicamente, "
    "perché i costi dell'adeguamento forzato "
    "sono sempre maggiori di quelli dell'adeguamento programmato. "
    "Con l'AI Act, la finestra per l'adeguamento programmato "
    "è aperta adesso. "
    "Tra qualche mese potrebbe esserlo di meno."
)

riferimenti(doc, [
    "Federprivacy — 'Il Consiglio dei Ministri approva i decreti di adeguamento all'AI Act' (4 agosto 2026)",
    "Diritto Mercato Tecnologia — 'AI Act: il Governo approva definitivamente i decreti di adeguamento' (2026)",
    "FISCOeTASSE.com — 'Regolamento intelligenza artificiale in Italia 2026' (agosto 2026)",
    "Legge 23 settembre 2025, n. 132 — Disposizioni in materia di intelligenza artificiale",
    "Regolamento UE 2024/1689 (AI Act), Allegati I e III",
    "Regolamento UE 2026/1744 (Digital Omnibus on AI) — modifiche alle scadenze AI Act",
    "Agenda Digitale — 'AI Act: come cambia il riconoscimento facciale con il decreto italiano' (2026)",
    "PMI.it — 'AI Act dal 2 agosto 2026, obblighi in vigore e rinvii' (2026)",
])
doc.save(BASE + "2026-09-04_decreti-attuativi-ai-act-agosto-2026-imprese.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Agenti AI nei processi aziendali: dal chatbot all'operatore
# ============================================================

doc = new_doc()
testata(doc, "Settembre 2026", "Strumenti AI per Professionisti")
titolo(
    doc,
    "Dal chatbot all'operatore. Come cambia l'AI nei processi aziendali.",
    "Nel 2026, l'AI non risponde più solo alle domande: agisce. "
    "Gli agenti autonomi leggono email, aprono software, compilano moduli, "
    "producono report e prenotano risorse senza intervento umano su ogni passaggio. "
    "Per le PMI italiane, la domanda non è se adottarli: "
    "è come farlo senza perdere il controllo.",
    "A cura della Redazione Ratio  •  4 settembre 2026"
)

para(doc,
    "Nel 1913, Ford introdusse la catena di montaggio a Highland Park. "
    "Non fu un'invenzione: fu un'organizzazione diversa di ciò che già esisteva. "
    "I lavoratori erano già lì; le macchine anche. "
    "Ciò che cambiò fu la sequenza: "
    "ogni persona faceva una cosa sola, in un ordine preciso, "
    "senza fermarsi a pensare al processo complessivo. "
    "Il risultato fu un aumento di produttività dell'ottocento per cento "
    "nel giro di tre anni. "
    "Gli agenti AI nel 2026 funzionano secondo la stessa logica. "
    "Non sono nuove capacità cognitive: "
    "sono una nuova organizzazione di capacità già esistenti. "
    "Un agente AI prende un obiettivo — 'aggiorna il CRM con i dati "
    "delle email ricevute questa settimana' — "
    "e lo scompone in una sequenza di azioni: "
    "legge le email, estrae i dati rilevanti, "
    "apre il CRM, cerca il cliente corrispondente, "
    "aggiorna i campi. "
    "Ogni passo è semplice; la sequenza automatizzata cambia la produttività."
)

para(doc,
    "La differenza tra un agente AI e un chatbot tradizionale "
    "è la capacità di agire, non solo di rispondere. "
    "Un chatbot riceve una domanda e produce una risposta testuale. "
    "Un agente riceve un obiettivo e produce un risultato: "
    "un documento aggiornato, un'email inviata, un record modificato, "
    "un report generato. "
    "I principali sistemi agentici disponibili nel settembre 2026 "
    "— Claude Opus 5 con computer use, GPT-5.6 con Operator, "
    "Gemini 3.7 con Project Mariner — "
    "possono operare su interfacce web, "
    "leggere e scrivere file, interagire con gestionali via API "
    "o direttamente sull'interfaccia grafica, "
    "e produrre output strutturati che si integrano nei workflow esistenti."
)

heading(doc, "I sette processi più adatti agli agenti AI nelle PMI italiane")

para(doc,
    "Sulla base dell'adozione nelle imprese italiane nel 2026, "
    "i processi che producono i risultati migliori con gli agenti AI "
    "hanno alcune caratteristiche comuni: "
    "alto volume di operazioni ripetitive, "
    "dati strutturati o semi-strutturati in ingresso, "
    "output verificabile da un operatore umano prima dell'azione finale. "
    "I sette casi d'uso più diffusi nelle PMI italiane sono: "
    "la gestione della posta in arrivo (classificazione, prioritizzazione, bozze di risposta); "
    "il pre-qualifica dei lead (lettura dei form, incrocio con il CRM, scoring automatico); "
    "il follow-up su offerte (invio automatico di promemoria sulle trattative aperte); "
    "la gestione dei ticket di assistenza (risposta alle richieste frequenti, "
    "instradamento di quelle complesse); "
    "la reportistica settimanale (aggregazione da fonti multiple, "
    "produzione del documento); "
    "la riconciliazione contabile di primo livello "
    "(incrocio tra fatture, movimenti bancari e registrazioni contabili); "
    "e la sintesi di documenti tecnici o contrattuali "
    "(estratti strutturati da contratti, bilanci, verbali). "
    "Nessuno di questi casi richiede un agente completamente autonomo: "
    "tutti prevedono un punto di verifica umana prima delle azioni più sensibili."
)

heading(doc, "Il perimetro di autonomia: la scelta che conta di più")

para(doc,
    "La variabile che determina se un agente AI produce valore o problemi "
    "non è la capacità tecnica del modello: "
    "è la chiarezza con cui è stato definito il suo perimetro di autonomia. "
    "Un perimetro di autonomia ben definito risponde a tre domande: "
    "quali azioni l'agente può compiere in autonomia senza chiedere conferma? "
    "Quali azioni richiedono un'approvazione umana prima di essere eseguite? "
    "Cosa succede quando l'agente incontra un caso non previsto? "
    "Le aziende che hanno avuto esperienze negative con gli agenti AI "
    "hanno quasi sempre saltato la definizione esplicita di queste regole, "
    "lasciando al modello la discrezionalità di decidere autonomamente "
    "fin dove spingersi. "
    "Un agente che lavora su un alto volume di operazioni "
    "replica un errore sistematico centinaia di volte "
    "prima che venga rilevato. "
    "La governance non è un freno alla produttività: "
    "è ciò che rende la produttività sostenibile nel tempo."
)

heading(doc, "Il professionista come punto di connessione")

para(doc,
    "Per molte PMI italiane, la figura che può più concretamente aiutare "
    "nell'adozione degli agenti AI non è un consulente tecnologico esterno: "
    "è il professionista di fiducia già inserito nei processi aziendali. "
    "Il commercialista che conosce il ciclo di fatturazione, "
    "i flussi di cassa, le scadenze fiscali dell'impresa "
    "è nella posizione ideale per identificare "
    "quali processi si prestano all'automazione agentiva "
    "e quali no. "
    "Non deve diventare un esperto di AI: "
    "deve saper rispondere a domande come 'questo processo è abbastanza strutturato "
    "da poter essere affidato a un agente?', "
    "'qual è il punto di supervisione umana in questo flusso?', "
    "'come documentiamo l'uso di questo agente per l'AI Act?'. "
    "Queste domande valgono consulenza. "
    "E le PMI che le pongono al professionista già presente nel loro ecosistema "
    "hanno più probabilità di rispondere bene di quelle "
    "che cercano un nuovo specialista dall'esterno."
)

para(doc,
    "Ford non inventò l'automobile: la rese accessibile. "
    "Non inventò il lavoro in fabbrica: lo organizò meglio. "
    "Gli agenti AI del 2026 non inventano nuove capacità per le imprese: "
    "organizzano meglio quelle che già esistono. "
    "La catena di montaggio cambiò l'economia mondiale "
    "non perché era sofisticata, ma perché era semplice, "
    "scalabile e replicabile. "
    "Il processo ben definito è ancora la unità fondamentale. "
    "L'agente AI è il nuovo modo di eseguirlo."
)

riferimenti(doc, [
    "AI4Business — 'Agenti e workflow: dall'AI conversazionale all'AI operativa' (2026)",
    "Impesud — 'Agentic AI in Italia 2026: Dalla Teoria all'Azione' (2026)",
    "AiChain Solutions — 'Automazioni AI Aziende: i 7 Processi (2026)'",
    "Best Tech Partner — 'Workflow Automation AI: Guida Strategica agli Agenti Intelligenti 2026'",
    "EM-EasyMobile — 'AI nelle PMI: come adottarla senza perdere il controllo' (2 settembre 2026)",
    "Osservatorio Artificial Intelligence, Politecnico di Milano — Rapporto mercato AI 2025-2026",
    "Tom's Hardware — 'Agenti AI e nuove regolamentazioni: il quadro normativo italiano ed europeo 2026'",
    "Corriere Comunicazioni — 'Agenti AI e regolamentazione 2026: cosa cambia per le imprese'",
])
doc.save(BASE + "2026-09-04_agenti-ai-processi-aziendali-dal-chatbot-all-operatore.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
