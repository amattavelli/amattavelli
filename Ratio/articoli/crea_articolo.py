"""
Crea l'articolo: "Gli AI Agent al servizio dei professionisti"
Destinazione: Ratio/articoli/
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/AI_Agent_Professionisti_PMI_2025.docx"

doc = Document()

# --- Impostazioni pagina ---
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# --- Stili base ---
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

# Helper per aggiungere uno stile paragrafo se non esiste
def get_or_create_style(doc, name, base_style_name='Normal'):
    try:
        return doc.styles[name]
    except KeyError:
        s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = doc.styles[base_style_name]
        return s

# --- Funzioni helper ---
def add_heading(doc, text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def add_paragraph(doc, text, bold=False, italic=False, size=11,
                  color=None, space_after=8, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_boxed_paragraph(doc, text, bg_color="E8F0FE"):
    """Paragrafo con sfondo colorato (evidenziazione)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.italic = True
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ============================================================
# TESTATA
# ============================================================

# Rivista / sezione
p_rivista = doc.add_paragraph()
p_rivista.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = p_rivista.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
run_r.font.name = 'Calibri'
run_r.font.size = Pt(9)
run_r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run_r.font.bold = True
run_r.font.all_caps = True

add_separator(doc)

# Data e numero
p_data = doc.add_paragraph()
p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_d = p_data.add_run(f"Marzo 2025  |  Speciale Intelligenza Artificiale")
run_d.font.name = 'Calibri'
run_d.font.size = Pt(8.5)
run_d.font.italic = True
run_d.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ============================================================
# TITOLO
# ============================================================
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_after = Pt(6)
run_t = p_title.add_run(
    "Gli AI Agent al servizio dei professionisti"
)
run_t.font.name = 'Calibri'
run_t.font.size = Pt(26)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# Sottotitolo
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_after = Pt(10)
run_s = p_sub.add_run(
    "La nuova frontiera dell'intelligenza artificiale autonoma per studi professionali e imprese italiane"
)
run_s.font.name = 'Calibri'
run_s.font.size = Pt(14)
run_s.font.italic = True
run_s.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_separator(doc)

# Autore e data
p_autore = doc.add_paragraph()
p_autore.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_autore.paragraph_format.space_after = Pt(16)
run_a = p_autore.add_run("A cura della Redazione Ratio  •  24 marzo 2025")
run_a.font.name = 'Calibri'
run_a.font.size = Pt(9)
run_a.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ============================================================
# SOMMARIO / ABSTRACT
# ============================================================
p_abs = doc.add_paragraph()
p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_abs.paragraph_format.left_indent = Cm(0.5)
p_abs.paragraph_format.right_indent = Cm(0.5)
p_abs.paragraph_format.space_after = Pt(16)
run_ab = p_abs.add_run(
    "Dal 2 febbraio 2025 sono entrati in vigore i primi obblighi dell'AI Act europeo. "
    "Ma la vera rivoluzione per gli studi professionali e le PMI italiane non è normativa: è operativa. "
    "Gli AI Agent — sistemi di intelligenza artificiale capaci di ragionare, pianificare e agire in autonomia "
    "su compiti complessi — stanno ridisegnando i processi di lavoro di commercialisti, consulenti, avvocati e manager. "
    "In questo articolo analizziamo cosa sono, cosa possono fare concretamente e come affrontare la transizione con metodo."
)
run_ab.font.name = 'Calibri'
run_ab.font.size = Pt(11)
run_ab.font.italic = True
run_ab.font.color.rgb = RGBColor(0x24, 0x24, 0x24)

# ============================================================
# SEZIONE 1 — Dal Copilot all'Agent
# ============================================================
add_heading(doc, "1. Dal Copilot all'Agent: il salto di qualità dell'IA", level=2)

add_paragraph(doc,
    "Per gran parte del 2023 e del 2024 il dibattito sull'intelligenza artificiale in ambito professionale "
    "si è concentrato sugli strumenti di assistenza generativa: ChatGPT, Copilot di Microsoft, Gemini di Google. "
    "Strumenti preziosi, ma con un limite strutturale: l'utente deve guidare ogni singolo passo, "
    "formulare la domanda giusta, interpretare la risposta, decidere il passo successivo."
)

add_paragraph(doc,
    "Il 2025 segna il passaggio a una fase diversa. Gli "
    + "AI Agent"
    + " — o agenti di intelligenza artificiale — sono sistemi in grado di: "
    "ricevere un obiettivo complesso in linguaggio naturale, scomporlo autonomamente in sotto-compiti, "
    "utilizzare strumenti esterni (database, API, file, browser, gestionali) e portare a termine l'incarico "
    "con supervisione umana minima o nulla. Non rispondono: agiscono."
)

add_boxed_paragraph(doc,
    "Esempio pratico — Un agente configurato per il monitoraggio dei crediti può: accedere al gestionale, "
    "identificare le fatture scadute, verificare la situazione del cliente via PEC o portale, "
    "redigere la lettera di sollecito personalizzata e archiviarla nel fascicolo — il tutto in pochi minuti, "
    "senza intervento umano."
)

add_paragraph(doc,
    "Secondo il Politecnico di Milano (Osservatorio Artificial Intelligence, rapporto 2024), "
    "il mercato italiano dell'AI ha raggiunto 760 milioni di euro, con una crescita del 52% sull'anno precedente. "
    "Ma i dati più interessanti per i professionisti riguardano la composizione: "
    "la quota destinata all'automazione intelligente dei processi (RPA + AI generativa) "
    "ha superato per la prima volta il 40% del totale. "
    "Il segnale è chiaro: le imprese stanno passando dall'esplorazione all'implementazione."
)

# ============================================================
# SEZIONE 2 — Come funzionano
# ============================================================
add_heading(doc, "2. Come funzionano gli AI Agent: l'architettura essenziale", level=2)

add_paragraph(doc,
    "Comprendere il funzionamento degli agenti non è un esercizio accademico: "
    "è il presupposto per valutarne i rischi, le opportunità e i limiti nel contesto del proprio studio o della propria impresa."
)

add_paragraph(doc,
    "Un AI Agent è composto da quattro elementi fondamentali:", bold=False
)

add_bullet(doc, "Modello linguistico (LLM): il \"cervello\" che comprende il linguaggio, ragiona e pianifica. "
           "I principali sono GPT-4o (OpenAI), Claude 3.7 (Anthropic), Gemini 2.0 (Google).")
add_bullet(doc, "Memoria: la capacità di ricordare il contesto di una sessione o, in versioni più evolute, "
           "di archiviare informazioni persistenti tra sessioni diverse.")
add_bullet(doc, "Strumenti (tools): le \"mani\" dell'agente — API, ricerche web, esecuzione di codice, "
           "lettura/scrittura di file, connessione a gestionali.")
add_bullet(doc, "Loop di ragionamento (ReAct o simili): il ciclo con cui l'agente pianifica, "
           "agisce, osserva il risultato e itera fino al completamento del compito.")

add_paragraph(doc,
    "La combinazione di questi elementi consente a un agente ben configurato di gestire attività "
    "che richiederebbero ore di lavoro manuale qualificato. Il professionista non deve scrivere codice: "
    "deve definire l'obiettivo, i confini d'azione e le regole di escalation verso l'umano."
)

# ============================================================
# SEZIONE 3 — Applicazioni concrete per professionisti
# ============================================================
add_heading(doc, "3. Applicazioni concrete per studi professionali e PMI", level=2)

add_paragraph(doc,
    "Vediamo i principali ambiti di applicazione già operativi in Italia, "
    "ordinati per maturità tecnologica e facilità di adozione."
)

# 3.1
p31 = doc.add_paragraph()
p31.paragraph_format.space_after = Pt(4)
run_31 = p31.add_run("3.1  Analisi documentale e due diligence")
run_31.font.name = 'Calibri'
run_31.font.size = Pt(12)
run_31.font.bold = True
run_31.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_paragraph(doc,
    "L'analisi di contratti, bilanci, visure e documentazione legale è uno dei campi dove gli agenti "
    "offrono il ritorno più immediato. Un agente può analizzare centinaia di pagine di documenti, "
    "estrarre clausole rilevanti, confrontarle con un modello di riferimento e produrre una sintesi strutturata "
    "con i punti di attenzione. Studi legali e commercialisti che lavorano su operazioni straordinarie "
    "(M&A, fusioni, cessioni) stanno già riducendo i tempi di due diligence del 40-60%."
)

# 3.2
p32 = doc.add_paragraph()
p32.paragraph_format.space_after = Pt(4)
run_32 = p32.add_run("3.2  Redazione e revisione di documenti fiscali e contabili")
run_32.font.name = 'Calibri'
run_32.font.size = Pt(12)
run_32.font.bold = True
run_32.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_paragraph(doc,
    "Dai comunicati IVA alle relazioni di accompagnamento al bilancio, dalla bozza di ricorso tributario "
    "alle risposte a interpelli, gli agenti possono predisporre documenti professionali partendo "
    "da dati strutturati (estratti contabili, codici normativi, giurisprudenza). "
    "La qualità del testo prodotto richiede revisione umana, ma il risparmio di tempo nella fase di bozza "
    "è dell'ordine del 60-70%. Alcuni studi hanno già integrato questi flussi nei propri software gestionali."
)

# 3.3
p33 = doc.add_paragraph()
p33.paragraph_format.space_after = Pt(4)
run_33 = p33.add_run("3.3  Supporto al controllo di gestione e all'analisi finanziaria")
run_33.font.name = 'Calibri'
run_33.font.size = Pt(12)
run_33.font.bold = True
run_33.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_paragraph(doc,
    "Collegati ai dati di bilancio o ai file Excel del controller, gli agenti possono calcolare automaticamente "
    "indici di bilancio, confrontarli con i benchmark di settore, identificare anomalie "
    "e formulare commenti narrativi per i report direzionali. "
    "Non si tratta di sostituire il controller: si tratta di liberarlo dalle elaborazioni meccaniche "
    "per concentrarlo sull'interpretazione e sulla relazione con il management."
)

# 3.4
p34 = doc.add_paragraph()
p34.paragraph_format.space_after = Pt(4)
run_34 = p34.add_run("3.4  Gestione intelligente delle comunicazioni con clienti e banche")
run_34.font.name = 'Calibri'
run_34.font.size = Pt(12)
run_34.font.bold = True
run_34.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_paragraph(doc,
    "Agenti configurati sul profilo del cliente possono predisporre memorie descrittive per pratiche bancarie, "
    "rispondere a richieste di documentazione, gestire scadenze e follow-up via email. "
    "Nelle PMI con un piccolo team amministrativo, questo riduce significativamente il carico "
    "su attività a basso valore aggiunto e migliora la qualità e tempestività delle comunicazioni."
)

# ============================================================
# SEZIONE 4 — AI Act: cosa cambia da febbraio 2025
# ============================================================
add_heading(doc, "4. AI Act: il nuovo quadro normativo europeo e gli obblighi per le imprese", level=2)

add_paragraph(doc,
    "Dal 1° agosto 2024 è in vigore il Regolamento (UE) 2024/1689, noto come AI Act. "
    "Ma è il 2 febbraio 2025 la data che ogni professionista deve avere ben presente: "
    "da quel giorno sono operative le prime disposizioni vincolanti, in particolare "
    "il divieto assoluto di utilizzo delle pratiche di IA considerate inaccettabili "
    "(manipolazione subliminale, social scoring, riconoscimento facciale di massa in spazi pubblici)."
)

add_paragraph(doc,
    "Per studi professionali e imprese le scadenze più rilevanti sono:"
)

add_bullet(doc, "Febbraio 2025 (già in vigore): divieto dei sistemi IA a rischio inaccettabile.")
add_bullet(doc, "Agosto 2025: obbligo di formazione in materia di IA per tutti coloro che utilizzano "
           "o sviluppano sistemi di IA in ambito professionale (art. 4 AI Act — \"AI literacy\").")
add_bullet(doc, "Agosto 2026: entrata in piena applicazione delle norme sui sistemi IA ad alto rischio "
           "(inclusi i sistemi usati in ambito creditizio, selezione del personale, istruzione).")

add_paragraph(doc,
    "In Italia, le autorità nazionali competenti per la vigilanza sull'AI Act sono "
    "l'Agenzia per l'Italia Digitale (AGID) e l'Agenzia per la Cybersicurezza Nazionale (ACN). "
    "Il Garante Privacy mantiene le proprie competenze per i profili relativi al trattamento dei dati personali."
)

add_boxed_paragraph(doc,
    "Attenzione pratica — L'obbligo di \"AI literacy\" (art. 4) è già in vigore. "
    "Ogni studio o impresa che utilizza sistemi di IA deve assicurare che il personale coinvolto "
    "abbia una conoscenza sufficiente delle caratteristiche, dei limiti e dei rischi di questi strumenti. "
    "Non è richiesta una certificazione formale, ma è opportuno documentare le iniziative formative intraprese."
)

# ============================================================
# SEZIONE 5 — Rischi e buone pratiche
# ============================================================
add_heading(doc, "5. Rischi, limiti e buone pratiche per un'adozione responsabile", level=2)

add_paragraph(doc,
    "L'entusiasmo giustificato intorno agli AI Agent non deve far dimenticare i rischi concreti, "
    "soprattutto in contesti professionali dove la responsabilità verso il cliente è diretta."
)

# Tabella rischi/mitigazioni
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Rischio'
hdr_cells[1].text = 'Buona pratica di mitigazione'
for cell in hdr_cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr()
    # Colore sfondo intestazione
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497D')
    tcPr.append(shd)

risks = [
    ("Allucinazioni e informazioni errate",
     "Verificare sempre le fonti. Non usare output dell'agente senza revisione umana per atti con valenza legale o fiscale."),
    ("Violazione della riservatezza dei dati",
     "Usare soluzioni con dati ospitati in Europa (GDPR). Evitare di caricare dati sensibili su servizi cloud non conformi."),
    ("Dipendenza eccessiva dallo strumento",
     "Mantenere le competenze professionali di base. L'agente supporta il professionista, non lo sostituisce."),
    ("Responsabilità professionale",
     "La firma del professionista implica la responsabilità del contenuto. L'output dell'agente non esonera da obblighi deontologici."),
    ("Obsolescenza rapida degli strumenti",
     "Privilegiare architetture modulari. Evitare dipendenze da un solo fornitore (lock-in)."),
]

for risk, practice in risks:
    row_cells = table.add_row().cells
    row_cells[0].text = risk
    row_cells[1].text = practice
    for cell in row_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

doc.add_paragraph()

# ============================================================
# SEZIONE 6 — Come iniziare
# ============================================================
add_heading(doc, "6. Come iniziare: un percorso in quattro passi", level=2)

add_paragraph(doc,
    "Per uno studio professionale o una PMI italiana che voglia avvicinarsi agli AI Agent "
    "in modo strutturato e responsabile, suggeriamo il seguente approccio graduale."
)

steps = [
    ("Passo 1 — Mappare i processi ripetitivi",
     "Identificare le attività che assorbono più tempo e sono ad alto tasso di standardizzazione: "
     "compilazione di moduli, redazione di bozze, raccolta e sintesi di dati. "
     "Queste sono le candidate ideali per la prima automazione agente."),
    ("Passo 2 — Scegliere la piattaforma giusta",
     "Valutare soluzioni enterprise già integrate con i gestionali in uso "
     "(ad esempio Microsoft Copilot for M365, soluzioni verticali per studi commercialisti, "
     "oppure piattaforme no-code come n8n o Make). Privilegiare fornitori con data center in UE."),
    ("Passo 3 — Pilotare su un processo non critico",
     "Avviare un progetto pilota su un processo a basso rischio (es. redazione di comunicazioni standard). "
     "Misurare tempi, qualità dell'output, soddisfazione del team. Correggere prima di scalare."),
    ("Passo 4 — Formare il team e aggiornare le policy interne",
     "Adempiere all'obbligo di AI literacy (art. 4 AI Act) con percorsi formativi documentati. "
     "Aggiornare le policy interne su uso degli strumenti IA, trattamento dei dati, "
     "revisione degli output e responsabilità."),
]

for title, body in steps:
    p_step = doc.add_paragraph()
    p_step.paragraph_format.space_after = Pt(4)
    p_step.paragraph_format.space_before = Pt(8)
    run_step = p_step.add_run(title)
    run_step.font.name = 'Calibri'
    run_step.font.size = Pt(11)
    run_step.font.bold = True
    run_step.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    add_paragraph(doc, body)

# ============================================================
# CONCLUSIONI
# ============================================================
add_heading(doc, "Conclusioni", level=2)

add_paragraph(doc,
    "Gli AI Agent non sono fantascienza né una moda passeggera: sono strumenti operativi "
    "che stanno già cambiando il modo in cui i professionisti e le imprese più avanzate lavorano. "
    "Il vantaggio competitivo non appartiene a chi adotta prima, ma a chi adotta meglio: "
    "con metodo, con una chiara visione dei rischi e con la consapevolezza che la tecnologia "
    "potenzia il giudizio professionale, ma non lo sostituisce."
)

add_paragraph(doc,
    "Per i commercialisti, i consulenti aziendali e i manager italiani, "
    "la domanda non è più \"se\" dotarsi di questi strumenti, ma \"come\" farlo in modo responsabile, "
    "efficace e conforme al nuovo quadro normativo europeo. "
    "Chi inizia a costruire questa competenza oggi avrà un vantaggio significativo "
    "nei prossimi dodici mesi."
)

add_separator(doc)

# Note/Fonti
p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(10)
run_n = p_note.add_run("Fonti e riferimenti")
run_n.font.name = 'Calibri'
run_n.font.size = Pt(9)
run_n.font.bold = True
run_n.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Regolamento (UE) 2024/1689 — AI Act (GU UE, 12 luglio 2024)",
    "Osservatorio Artificial Intelligence — Politecnico di Milano, Rapporto 2024",
    "AGID / ACN — Comunicazione sulle autorità nazionali competenti AI Act (gennaio 2025)",
    "Garante per la Protezione dei Dati Personali — Linee guida sull'uso dell'IA nei servizi professionali",
    "Microsoft, Anthropic, Google — Documentazione tecnica su Copilot, Claude e Gemini (2025)",
]
for s in sources:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_after = Pt(2)
    run_s = p_s.add_run(f"• {s}")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(8.5)
    run_s.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# Piè di pagina
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(16)
run_f = p_footer.add_run(
    "© 2025 Mattavelli Amodeo — Commercialisti Associati  •  Riproduzione consentita con citazione della fonte"
)
run_f.font.name = 'Calibri'
run_f.font.size = Pt(8)
run_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_f.font.italic = True

# ============================================================
# SALVATAGGIO
# ============================================================
doc.save(OUTPUT_FILE)
print(f"Articolo salvato in: {OUTPUT_FILE}")
