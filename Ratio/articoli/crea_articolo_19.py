"""
Articolo 19: Dall'assistente all'agente — come cambia il lavoro del professionista nell'era agentica
Ratio/articoli/2026-04-10_dallagente-allassistente-lavoro-professionista.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-10_dall-assistente-all-agente-lavoro-professionista.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_paragraph(doc, text, size=11, italic=False, color=None,
                  space_after=8, space_before=0,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


# ---- TESTATA ----
p_rivista = doc.add_paragraph()
p_rivista.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = p_rivista.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
run_r.font.name = 'Calibri'
run_r.font.size = Pt(9)
run_r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run_r.font.bold = True
run_r.font.all_caps = True

add_separator(doc)

p_data = doc.add_paragraph()
p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_d = p_data.add_run("Aprile 2026  |  Riflessione e Prospettiva")
run_d.font.name = 'Calibri'
run_d.font.size = Pt(8.5)
run_d.italic = True
run_d.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_after = Pt(6)
run_t = p_title.add_run(
    "Dall\u2019assistente all\u2019agente: come cambia davvero il lavoro del professionista"
)
run_t.font.name = 'Calibri'
run_t.font.size = Pt(24)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_after = Pt(10)
run_s = p_sub.add_run(
    "Fino a ieri l\u2019AI rispondeva alle domande. Da oggi l\u2019AI esegue compiti. "
    "Il passaggio dall\u2019assistente all\u2019agente autonomo non \u00e8 un aggiornamento tecnico: "
    "\u00e8 un cambiamento nel rapporto tra il professionista e il suo lavoro."
)
run_s.font.name = 'Calibri'
run_s.font.size = Pt(13)
run_s.italic = True
run_s.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_separator(doc)

p_autore = doc.add_paragraph()
p_autore.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_autore.paragraph_format.space_after = Pt(16)
run_a = p_autore.add_run("A cura della Redazione Ratio  \u2022  10 aprile 2026")
run_a.font.name = 'Calibri'
run_a.font.size = Pt(9)
run_a.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- CORPO ----

add_paragraph(doc,
    "C\u2019\u00e8 una differenza concettuale che ancora non \u00e8 entrata nel vocabolario quotidiano "
    "di molti professionisti, e vale la pena nominarla con chiarezza. Un assistente AI \u2014 "
    "ChatGPT, Claude, Gemini usati nella modalit\u00e0 pi\u00f9 comune \u2014 risponde a ci\u00f2 "
    "che gli si chiede. Gli dai un input, ti restituisce un output. Il ciclo finisce l\u00ec. "
    "Un agente AI fa qualcosa di diverso: riceve un obiettivo e lo persegue in autonomia, "
    "compiendo una sequenza di azioni, prendendo decisioni intermedie, usando strumenti esterni, "
    "fino a produrre un risultato finale. Non risponde: esegue."
)

add_paragraph(doc,
    "Questa distinzione non \u00e8 tecnica. Ha conseguenze dirette su come il professionista "
    "organizza il proprio lavoro, su quali compiti delega, su dove concentra la propria attenzione "
    "e su cosa diventa il valore aggiunto della sua presenza. Il passaggio dall\u2019assistente "
    "all\u2019agente \u00e8 il vero punto di svolta del 2026 \u2014 non un singolo lancio di prodotto, "
    "ma una trasformazione nella natura dello strumento."
)

add_heading2(doc, "Il professionista come direttore di agenti")

add_paragraph(doc,
    "L\u2019immagine che circola sempre pi\u00f9 nelle discussioni sull\u2019AI professionale \u00e8 "
    "quella del \u201cdirettore di agenti\u201d: il professionista che non esegue i compiti ma "
    "li assegna, li supervisiona, ne verifica i risultati e interviene dove il giudizio umano "
    "\u00e8 necessario. \u00c8 un\u2019immagine suggestiva, ma richiede qualche precisazione per "
    "non diventare fuorviante."
)

add_paragraph(doc,
    "Dirigere un agente AI non \u00e8 come fare il manager di un team umano. Richiede la capacit\u00e0 "
    "di definire l\u2019obiettivo con precisione sufficiente perch\u00e9 l\u2019agente possa "
    "operare in autonomia \u2014 e questa \u00e8 una competenza non banale. Un agente a cui viene "
    "detto \u201canalizza questo bilancio\u201d produrr\u00e0 qualcosa di generico. Un agente a "
    "cui viene detto \u201canalizza questo bilancio, identifica i tre principali rischi di "
    "liquidit\u00e0 per i prossimi sei mesi, e predisponi una bozza di comunicazione al cda\u201d "
    "produrr\u00e0 qualcosa di usabile. La differenza \u00e8 nella qualit\u00e0 del briefing, "
    "non nella tecnologia."
)

add_heading2(doc, "I compiti che cambiano, quelli che restano")

add_paragraph(doc,
    "L\u2019AI agentica \u00e8 particolarmente efficace su tre categorie di lavoro. La prima \u00e8 "
    "quella dei compiti ripetitivi ad alto volume: monitoraggio delle scadenze, riconciliazione "
    "dei dati contabili, estrazione di informazioni da documenti strutturati, preparazione di "
    "reportistica standardizzata. Sono attivit\u00e0 che consumano tempo ma non richiedono giudizio "
    "professionale \u2014 e che un agente configur\u00e0 bene pu\u00f2 eseguire con affidabilit\u00e0 "
    "crescente. La seconda categoria \u00e8 quella della ricerca e sintesi: raccogliere informazioni "
    "da fonti diverse, costruire un quadro d\u2019insieme, preparare il materiale per una decisione. "
    "La terza \u00e8 quella della comunicazione strutturata: bozze di email, avvisi ai clienti, "
    "sintesi di documenti complessi in formato leggibile."
)

add_paragraph(doc,
    "Restano invece saldamente in mano al professionista umano le attivit\u00e0 che richiedono "
    "giudizio su situazioni non standardizzate, la gestione della relazione con il cliente nei "
    "momenti critici, la responsabilit\u00e0 professionale per le scelte adottate e la capacit\u00e0 "
    "di leggere contesti che vanno oltre i dati disponibili. Nessun agente AI pu\u00f2 sostituire "
    "il commercialista che capisce, parlando con un imprenditore, che il vero problema non \u00e8 "
    "quello dichiarato ma qualcosa di pi\u00f9 profondo. Quello \u00e8 il lavoro che resta, "
    "e che acquista valore proprio perch\u00e9 il resto viene automatizzato."
)

add_heading2(doc, "Il rischio dell\u2019automazione inconsapevole")

add_paragraph(doc,
    "C\u2019\u00e8 un rischio specifico nell\u2019era degli agenti che vale la pena nominare. "
    "Quando i processi di basso livello vengono automatizzati, si libera tempo \u2014 ma quel tempo "
    "non si trasforma automaticamente in lavoro di maggior valore. Pu\u00f2 tradursi in maggiore "
    "volume di lavoro a basso valore, se si risponde all\u2019efficienza guadagnata semplicemente "
    "aumentando il numero di pratiche gestite. Il salto qualitativo avviene solo se il tempo "
    "liberato viene reinvestito consciamente: in attivit\u00e0 di consulenza pi\u00f9 profonda, "
    "in sviluppo della relazione con i clienti strategici, in apprendimento e aggiornamento "
    "professionale."
)

add_paragraph(doc,
    "L\u2019automazione inconsapevole \u2014 quella che avviene senza una scelta esplicita su come "
    "usare il tempo recuperato \u2014 rischia di generare efficienza senza valore. Il professionista "
    "che ci guadagna davvero dall\u2019AI agentica non \u00e8 quello che fa di pi\u00f9: "
    "\u00e8 quello che decide cosa smettere di fare e cosa fare di pi\u00f9."
)

add_heading2(doc, "La domanda che vale la pena farsi ora")

add_paragraph(doc,
    "Il mercato dell\u2019AI agentica cresce al 67% annuo. I fornitori di software gestionale "
    "italiani stanno integrando queste capacit\u00e0 nei loro prodotti. Entro diciotto mesi, "
    "la maggior parte degli studi professionali avr\u00e0 accesso a strumenti agentici dentro "
    "i sistemi gi\u00e0 in uso, senza dover scegliere piattaforme separate. La domanda non \u00e8 "
    "se questo accadr\u00e0: \u00e8 gi\u00e0 in corso. La domanda \u00e8 chi arriva preparato."
)

add_paragraph(doc,
    "Prepararsi non significa diventare esperti di tecnologia. Significa avere gi\u00e0 risposto "
    "a tre domande: quali attivit\u00e0 del nostro studio sono candidate all\u2019automazione? "
    "Come ridefiniremo i processi di lavoro quando queste attivit\u00e0 saranno automatizzate? "
    "E soprattutto: cosa vogliamo fare con il tempo che si liberer\u00e0? Chi ha gi\u00e0 queste "
    "risposte \u00e8 in grado di usare l\u2019AI agentica come un moltiplicatore. Chi non le ha "
    "trover\u00e0 nello strumento soltanto un\u2019altra cosa da gestire."
)

add_separator(doc)

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(10)
run_n = p_note.add_run("Fonti e riferimenti")
run_n.font.name = 'Calibri'
run_n.font.size = Pt(9)
run_n.font.bold = True
run_n.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Digitech.news \u2014 Cresce l\u2019AI agentica: entro il 2030 sar\u00e0 il 31% dell\u2019intero mercato (aprile 2026)",
    "Manager.it \u2014 Agenti AI autonomi: cosa sono e come cambiano il lavoro nel 2026",
    "Jenova.ai \u2014 Assistente AI professionale: come gli agenti AI trasformeranno la produttivit\u00e0 nel 2026",
    "Cosmonet.info \u2014 AI Agent Italia 2026: la rivoluzione degli assistenti virtuali per le aziende",
    "Economy Magazine \u2014 La grande nuova frontiera dell\u2019AI per le aziende (2026)",
    "Osservatorio Intelligenza Artificiale \u2014 Politecnico di Milano, Report 2026",
]
for s in sources:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_after = Pt(2)
    run_s = p_s.add_run(f"\u2022 {s}")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(8.5)
    run_s.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(16)
run_f = p_footer.add_run(
    "\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte"
)
run_f.font.name = 'Calibri'
run_f.font.size = Pt(8)
run_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_f.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
