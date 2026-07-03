"""
Quattro articoli Ratio — 3 luglio 2026

1. 2026-07-03_ai-act-articolo-50-trenta-giorni-al-2-agosto.docx
2. 2026-07-03_modello-231-e-ai-rischio-governance-ignorato.docx
3. 2026-07-03_abbonamenti-ai-in-azienda-come-si-deducono.docx
4. 2026-07-03_ai-che-risponde-al-cliente-dove-conviene.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# AI Act art. 50: trenta giorni al 2 agosto
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Normativa AI")
titolo(
    doc,
    "L'AI Act al 2 agosto:\ntrenta giorni per non farsi trovare impreparati.",
    "Dal 2 agosto 2026 scattano gli obblighi di trasparenza dell'articolo 50 "
    "dell'AI Act. Chi usa chatbot, genera contenuti con l'AI o tratta immagini "
    "sintetiche ha un mese per adeguarsi. Le sanzioni arrivano fino al 3% del "
    "fatturato. Ma la compliance e' piu' semplice di quanto sembri, se si sa "
    "dove guardare.",
    "A cura della Redazione Ratio  •  3 luglio 2026"
)

para(doc,
    "Un'azienda di servizi del Nord Italia ha inserito sul proprio sito, a inizio "
    "anno, un assistente virtuale che risponde alle domande dei clienti sui "
    "prodotti, sugli orari e sulle procedure di reso. Il chatbot funziona bene, "
    "riduce il carico del servizio clienti e i responsabili sono soddisfatti. "
    "Nessuno si e' posto la domanda se il cliente sappia che sta parlando con "
    "un sistema automatico. Dal 2 agosto non sara' piu' una questione di buona "
    "pratica: sara' un obbligo di legge, e ignorarlo espone a sanzioni che non "
    "si recuperano facilmente."
)

para(doc,
    "L'articolo 50 del Regolamento UE 2024/1689, meglio noto come AI Act, "
    "definisce gli obblighi di trasparenza per alcune categorie specifiche di "
    "sistemi di intelligenza artificiale. Non si tratta degli obblighi piu' "
    "complessi del regolamento, quelli che riguardano i sistemi ad alto rischio "
    "e che richiedono mesi di documentazione tecnica. Si tratta di obblighi "
    "relativamente diretti, che si applicano a un numero molto piu' ampio di "
    "aziende, compresi studi professionali, PMI e qualsiasi organizzazione che "
    "utilizzi strumenti AI nei rapporti con l'esterno. Il limite di un mese "
    "non e' una data simbolica: le autorita' nazionali di vigilanza avranno "
    "i poteri sanzionatori pieni dal giorno stesso."
)

heading(doc, "Cosa prescrive concretamente l'articolo 50")

para(doc,
    "La norma si articola su quattro categorie di obblighi. La prima riguarda "
    "i sistemi conversazionali: qualsiasi chatbot, assistente virtuale o sistema "
    "che interagisce in modo automatico con persone fisiche deve rendersi "
    "riconoscibile come tale. L'utente deve sapere, prima o durante "
    "l'interazione, che sta dialogando con un'intelligenza artificiale e non "
    "con un essere umano. L'eccezione vale solo quando il contesto rende "
    "l'automatizzazione del tutto evidente, come nel caso di un risponditore "
    "telefonico che nessuno scambia per una persona reale."
)

para(doc,
    "La seconda categoria riguarda i contenuti sintetici: testi, immagini, audio "
    "e video generati o manipolati da sistemi AI devono essere marcati con "
    "un'etichetta tecnica leggibile automaticamente che ne consenta il "
    "riconoscimento. Non si tratta necessariamente di aggiungere una scritta "
    "visibile in ogni immagine: si tratta di incorporare metadati che i sistemi "
    "di rilevamento possano leggere. La terza riguarda i deepfake, cioe' i "
    "contenuti che rappresentano persone reali in situazioni o con dichiarazioni "
    "false: devono essere dichiarati esplicitamente come tali, sia con marcatura "
    "tecnica sia con avviso leggibile dall'utente. La quarta, piu' specifica, "
    "impone un'informativa preventiva alle persone fisiche quando vengono "
    "analizzate da sistemi di riconoscimento delle emozioni o di "
    "categorizzazione biometrica."
)

heading(doc, "Chi si deve adeguare e in quanto tempo")

para(doc,
    "Gli obblighi si applicano sia ai fornitori dei sistemi AI, cioe' a chi li "
    "costruisce e li commercializza, sia ai deployer, cioe' alle aziende che "
    "li utilizzano nell'esercizio della propria attivita'. Un'impresa che "
    "acquista un chatbot da un fornitore e lo installa sul proprio sito e' un "
    "deployer: ha l'obbligo di assicurarsi che il sistema rispetti i requisiti "
    "di trasparenza, indipendentemente da quanto ha fatto il fornitore. Se il "
    "fornitore non ha previsto la dichiarazione automatica, tocca al deployer "
    "aggiungerla. Questo significa che la verifica va fatta ora, non dopo "
    "il 2 agosto."
)

para(doc,
    "Per uno studio professionale che usa un chatbot sul sito, la compliance si "
    "riduce quasi sempre a due passi: verificare che il sistema identifichi se "
    "stesso all'inizio dell'interazione (basta un messaggio di benvenuto che "
    "dica 'Sono un assistente automatico') e controllare che il fornitore del "
    "software abbia aggiornato le proprie condizioni di servizio per includere "
    "la conformita' all'AI Act. Per uno studio che genera contenuti con strumenti "
    "AI a uso interno, senza pubblicarli come prodotti aziendali, l'obbligo "
    "di marcatura non si attiva: si applica ai contenuti distribuiti al pubblico "
    "o a terzi, non all'uso privato degli strumenti."
)

para(doc,
    "Le sanzioni per la violazione degli obblighi dell'articolo 50 arrivano fino "
    "a 15 milioni di euro o al 3% del fatturato annuo globale, con applicazione "
    "dell'importo maggiore. Per una PMI con un fatturato di cinque milioni, il "
    "3% significa 150.000 euro. La proporzionalita' si applica in linea teorica, "
    "ma il dato strutturale non cambia: trattarsi come un caso troppo piccolo "
    "per essere nel mirino non e' una strategia di gestione del rischio, "
    "e' un pregiudizio. Chi si adegua prima del 2 agosto non ha niente da "
    "temere. Chi aspetta di vedere cosa fanno gli altri si espone a una "
    "scadenza che non si sposta."
)

riferimenti(doc, [
    "Regolamento UE 2024/1689 (AI Act), articolo 50 — EUR-Lex",
    "AgendaDigitale.eu — 'AI Act, dal 2 agosto scatta la trasparenza obbligatoria: cosa cambia (e per chi)'",
    "AgendaDigitale.eu — 'Obblighi di trasparenza AI Act: cosa devono fare le aziende dal 2 agosto 2026'",
    "UniverseIT.blog — 'AI Act 2 agosto 2026: cosa cambia per le imprese'",
    "TrueScreen.io — 'Articolo 50 AI Act: obblighi etichettatura dal 2 agosto 2026'",
    "StudioLegaleStefanelli.it — 'Trasparenza dell'AI: cosa impone l'art. 50 dell'AI Act e perche conta'",
    "Certifico.com — 'AI Act / obblighi per le imprese dal 02 agosto 2026'",
    "AIPIA.it — 'Enforcement AI Act: sanzioni e poteri dal 2 agosto 2026'",
])
doc.save(BASE + "2026-07-03_ai-act-articolo-50-trenta-giorni-al-2-agosto.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# Modello 231 e AI: il rischio di governance ignorato
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Compliance e Governance")
titolo(
    doc,
    "Il Modello 231 che non parla di AI:\num rischio che quasi nessuno ha ancora calcolato.",
    "La Legge 132/2025 ha introdotto nuove fattispecie penali legate all'uso "
    "di intelligenza artificiale. Chi ha aggiornato il Modello 231 dell'azienda "
    "negli ultimi anni ha quasi certamente un documento che non ne fa menzione. "
    "Per i commercialisti e i consulenti che assistono i clienti sulla compliance "
    "231, questo e' un punto cieco da affrontare ora.",
    "A cura della Redazione Ratio  •  3 luglio 2026"
)

para(doc,
    "Un membro dell'Organismo di Vigilanza di una societa' manifatturiera con "
    "circa sessanta dipendenti ha riletto il Modello di Organizzazione, Gestione "
    "e Controllo della propria azienda lo scorso aprile, in preparazione a una "
    "verifica annuale. L'ultimo aggiornamento risaliva al 2023. Il documento "
    "copriva i rischi classici del D.Lgs. 231/2001: corruzione, reati tributari, "
    "sicurezza sul lavoro, reati ambientali. Dell'intelligenza artificiale non "
    "c'era traccia, neppure come voce dell'analisi dei rischi. Il problema e' "
    "che nel frattempo l'azienda aveva introdotto un sistema di valutazione "
    "automatica dei fornitori basato su AI e un chatbot per la gestione delle "
    "segnalazioni interne. Entrambi erano operativi. Entrambi erano fuori "
    "dal perimetro del Modello."
)

para(doc,
    "La situazione e' molto diffusa. La maggior parte dei Modelli 231 in "
    "circolazione nelle PMI italiane e' stata redatta o aggiornata prima "
    "che l'AI entrasse nell'uso corrente aziendale. Dal settembre 2025, "
    "con l'entrata in vigore della Legge 132/2025, il quadro normativo e' "
    "cambiato in modo rilevante: la legge ha introdotto nell'ordinamento "
    "italiano le prime norme organiche sull'intelligenza artificiale, incluse "
    "nuove aggravanti penali e fattispecie di reato connesse all'utilizzo "
    "improprio di sistemi AI. Un Modello 231 che non mappi questi nuovi rischi "
    "non e' solo incompleto: potrebbe non coprire l'azienda in caso di "
    "contestazione."
)

heading(doc, "Cosa introduce la Legge 132/2025 nel perimetro 231")

para(doc,
    "La Legge 132/2025 agisce su due livelli. Sul primo, ha introdotto nuove "
    "aggravanti per reati esistenti quando commessi con il supporto di sistemi "
    "di intelligenza artificiale: frodi, manipolazioni di mercato, "
    "appropriazione indebita di dati. Il fatto che un sistema automatizzato "
    "abbia eseguito l'operazione non esime il soggetto apicale dalla "
    "responsabilita', ma puo' aggravare il trattamento sanzionatorio. "
    "Sul secondo livello, ha esteso esplicitamente l'applicabilita' del "
    "D.Lgs. 231/2001 agli illeciti commessi attraverso l'utilizzo doloso "
    "di sistemi AI da parte di soggetti in posizione apicale o sottoposti "
    "alla loro direzione."
)

para(doc,
    "Per le aziende che usano AI in processi decisionali rilevanti, come "
    "la selezione del personale, la valutazione del merito creditizio dei "
    "clienti, la gestione automatizzata dei prezzi o l'analisi delle "
    "segnalazioni di whistleblowing, il rischio non e' solo normativo in "
    "astratto. Se un sistema AI produce una decisione discriminatoria nella "
    "selezione del personale e nessun presidio del Modello 231 prevede "
    "supervisione su quel processo, la responsabilita' dell'ente puo' "
    "scattare sulla base del fatto che l'OdV non aveva gli strumenti per "
    "rilevare il rischio. La mancanza di un protocollo specifico per l'AI "
    "diventa, in questo scenario, una lacuna del sistema di controllo."
)

heading(doc, "Cosa deve contenere un aggiornamento adeguato del MOG")

para(doc,
    "Un aggiornamento del Modello 231 che tenga conto dell'AI dovrebbe "
    "partire da una mappatura dei sistemi intelligenti effettivamente in uso "
    "in azienda: non solo i grandi progetti dichiarati, ma anche gli strumenti "
    "adottati in modo autonomo dai singoli reparti, il cosiddetto 'shadow AI'. "
    "Per ognuno di questi sistemi va valutato il processo su cui incide, il "
    "tipo di decisioni che supporta o automatizza e se rientra nelle categorie "
    "di rischio elevato individuate dall'AI Act. Questa mappatura non "
    "richiede competenze tecniche avanzate: richiede che qualcuno in azienda "
    "abbia la responsabilita' di tenerla aggiornata."
)

para(doc,
    "Il Modello aggiornato dovrebbe poi definire i protocolli di supervisione "
    "umana sui processi in cui l'AI e' coinvolta, le modalita' di "
    "documentazione delle decisioni automatizzate e i criteri per la "
    "segnalazione di anomalie all'OdV. L'Organismo di Vigilanza, da parte "
    "sua, deve verificare periodicamente non solo che il Modello esista, "
    "ma che il registro dei sistemi AI sia aggiornato e che i presidi "
    "funzionino concretamente. Una policy sull'AI scritta e non applicata "
    "non protegge l'ente."
)

para(doc,
    "Per i commercialisti e i consulenti che assistono i clienti sulla compliance "
    "231, questo e' un momento in cui il ruolo di advisor si amplia in modo "
    "naturale. Il cliente che chiede di aggiornare il Modello in vista di una "
    "verifica ha quasi certamente questa lacuna. Segnalargliela prima che "
    "diventi un problema concreto e' esattamente il tipo di valore aggiunto "
    "che distingue la consulenza dalla semplice redazione di documenti."
)

riferimenti(doc, [
    "D.Lgs. 8 giugno 2001, n. 231 — Normattiva",
    "Legge 9 settembre 2025, n. 132 (Legge sull'Intelligenza Artificiale) — Normattiva",
    "AteneoWeb.com — 'AI Act e D.Lgs. 231/2001: guida per commercialisti 2026'",
    "ValeRioCarlesimo.it — 'La compliance 231 e l'intelligenza artificiale: l'adattamento del MOG'",
    "Winple.it — 'Modello 231 e Intelligenza Artificiale: obblighi, rischi e aggiornamenti dopo la Legge 132/2025'",
    "Commercialisti.it — 'Cybersecurity e Modello 231: integrazione dei rischi informatici nella governance d'impresa' (maggio 2026)",
    "Paradigma.it — 'L'Intelligenza Artificiale applicata al sistema 231' (maggio 2026)",
    "Ransomtax.it — 'Mog 231 e intelligenza artificiale: cosa cambia con la Legge 132/2025'",
])
doc.save(BASE + "2026-07-03_modello-231-e-ai-rischio-governance-ignorato.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# Abbonamenti AI in azienda: deducibilita' e contabilizzazione
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Fiscalita' e Amministrazione")
titolo(
    doc,
    "L'abbonamento a ChatGPT in nota spese:\ncome si deduce e dove va in bilancio.",
    "Migliaia di professionisti e dipendenti pagano ogni mese un abbonamento "
    "AI con la carta personale o aziendale. Pochi sanno come trattarli "
    "fiscalmente. Il tema e' semplice ma pieno di varianti che fanno la "
    "differenza: uso personale o aziendale, abbonamento individuale o "
    "piano business, software o servizio.",
    "A cura della Redazione Ratio  •  3 luglio 2026"
)

para(doc,
    "Una consulente del lavoro con uno studio di quattro persone ha cominciato "
    "a usare Claude Pro a febbraio, pagando il canone mensile con la carta di "
    "credito personale. A maggio ha deciso di passare al piano aziendale, ha "
    "intestato l'abbonamento allo studio e ha iniziato ad addebitarlo in "
    "nota spese. Quando il suo studio ha chiuso il trimestre e ha consegnato "
    "i giustificativi al commercialista, quest'ultimo le ha fatto una domanda "
    "semplice: 'Lo usi solo tu o anche i collaboratori?'. Da quella domanda "
    "e' partita una conversazione che ha richiesto piu' tempo del previsto, "
    "perche' la risposta cambia la classificazione contabile, l'IVA e il "
    "regime di deducibilita'. La questione non e' complicata, ma richiede "
    "di guardare il caso specifico, non la categoria generica."
)

para(doc,
    "Il punto di partenza e' la natura giuridica dell'abbonamento AI. "
    "La stragrande maggioranza degli abbonamenti a strumenti come ChatGPT, "
    "Claude, Gemini o Copilot e' classificabile come acquisto di servizi "
    "digitali in abbonamento, tecnicamente 'fornitura di servizi prestati "
    "tramite mezzi elettronici' ai sensi della normativa IVA. Per le "
    "imprese e i professionisti residenti in Italia, questi acquisti da "
    "fornitori extra-UE come OpenAI o Anthropic rientrano nel meccanismo "
    "del reverse charge: l'IVA non compare nella fattura del fornitore "
    "estero, ma va autoassegnata e annotata dal committente italiano. "
    "Un errore frequente e' non applicare il reverse charge su questi "
    "pagamenti, trattandoli come se fossero esenti o fuori campo."
)

heading(doc, "La deducibilita': tre scenari diversi")

para(doc,
    "Il primo scenario e' l'abbonamento intestato all'azienda o allo studio, "
    "usato esclusivamente per l'attivita' professionale. In questo caso la "
    "deducibilita' e' piena, al 100%, come per qualsiasi altro acquisto di "
    "software o servizio strumentale all'attivita'. La classificazione "
    "contabile piu' corretta e' 'costi per servizi' o, se lo studio vuole "
    "tenere separati i costi di digitalizzazione, 'costi per software e "
    "licenze digitali'. La voce deve comunque essere documentata con la "
    "fattura del fornitore o la conferma di pagamento associata all'account "
    "aziendale. Per i piani business di OpenAI, Anthropic e Google, la "
    "fattura con i dati fiscali dell'azienda si ottiene dalla sezione "
    "fatturazione del portale: va scaricata e conservata."
)

para(doc,
    "Il secondo scenario e' l'abbonamento intestato al dipendente o al "
    "collaboratore, rimborsato dall'azienda tramite nota spese. Qui la "
    "deducibilita' e' in linea di principio possibile, ma l'inerenza va "
    "documentata in modo piu' robusto: serve una policy aziendale che autorizzi "
    "l'uso di strumenti AI personali per fini lavorativi, e il rimborso deve "
    "essere limitato alla quota di utilizzo professionale. Se il dipendente "
    "usa lo stesso abbonamento anche per scopi personali, la deducibilita' "
    "integrale e' contestabile. La soluzione piu' pulita, per le aziende con "
    "piu' utenti, e' il piano business intestato all'azienda, che consente "
    "di gestire gli accessi, avere un'unica fattura e applicare il reverse "
    "charge correttamente."
)

para(doc,
    "Il terzo scenario e' il professionista persona fisica in regime "
    "ordinario, che usa l'abbonamento AI per la propria attivita'. "
    "La deducibilita' e' al 100% se l'uso e' esclusivamente professionale, "
    "con la stessa logica del software di studio. Per i professionisti in "
    "regime forfettario, la deducibilita' analitica non si applica: i costi "
    "sono assorbiti dalla deduzione forfettaria prevista dal regime, "
    "e l'abbonamento AI non genera un beneficio fiscale aggiuntivo separato."
)

heading(doc, "Il caso dei costi di formazione AI")

para(doc,
    "Una questione che emerge spesso riguarda i corsi di formazione sull'AI: "
    "webinar, abbonamenti a piattaforme di apprendimento, partecipazione a "
    "eventi specializzati. Questi costi rientrano generalmente tra le spese "
    "di formazione del personale, deducibili al 100% se inerenti all'attivita'. "
    "Il nodo interpretativo piu' comune e' quando la formazione ha un "
    "contenuto misto, in parte professionale e in parte personale. "
    "Per i corsi certificati da enti accreditati, il problema non si pone: "
    "la finalita' professionale e' documentata dal programma. Per i corsi "
    "online acquistati individualmente, conservare il programma del corso e "
    "una nota interna che ne giustifichi la rilevanza per l'attivita' e' "
    "la misura minima per sostenere la deducibilita' in caso di verifica."
)

para(doc,
    "Sul lato IVA, una precisazione rilevante: i corsi di formazione online "
    "erogati da fornitori extra-UE, come molte piattaforme americane, rientrano "
    "anchessi nel reverse charge per i soggetti passivi IVA italiani. "
    "La scadenza di agosto per gli obblighi AI Act e il crescente controllo "
    "fiscale sulle transazioni digitali rendono questo un buon momento per "
    "fare un passaggio complessivo con il proprio commercialista su come "
    "sono stati trattati questi costi nell'ultimo anno e mezzo. Correggere "
    "adesso un'impostazione scorretta costa meno che farlo su contestazione."
)

riferimenti(doc, [
    "Agenzia delle Entrate — Circolare n. 36/E/2012 (trattamento IVA servizi elettronici)",
    "D.P.R. 26 ottobre 1972, n. 633 (IVA), articoli 7-sexies e 17 — reverse charge",
    "T.U.I.R. D.P.R. 917/1986, articoli 54 e 109 (deducibilita' costi professionisti e imprese)",
    "AgendaDigitale.eu — 'AI nelle aziende italiane: obblighi fiscali e compliance 2026'",
    "FiscoeTasse.com — 'Software in abbonamento SaaS: IVA e deducibilita' 2026'",
    "IlSole24Ore — 'Strumenti digitali e nota spese: cosa si deduce nel 2026'",
    "OpenAI Help Center — 'Billing and invoices for Teams and Enterprise'",
    "Anthropic — 'Claude for Business: billing and tax documentation'",
])
doc.save(BASE + "2026-07-03_abbonamenti-ai-in-azienda-come-si-deducono.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# AI che risponde al cliente: dove conviene e dove no
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Professione e Strumenti AI")
titolo(
    doc,
    "Quando l'AI risponde al cliente al posto tuo:\ndove funziona e dove smette di funzionare.",
    "Usare l'AI per preparare le risposte ai clienti riduce il tempo di lavoro "
    "in modo misurabile. Ma c'e' un confine tra usarla come strumento di "
    "supporto e delegarle la sostanza della consulenza. Quel confine non e' "
    "sempre dove sembra.",
    "A cura della Redazione Ratio  •  3 luglio 2026"
)

para(doc,
    "Un avvocato tributarista racconta di aver introdotto l'abitudine di usare "
    "Claude per preparare le bozze delle risposte alle email piu' tecniche dei "
    "clienti. Il processo e' semplice: incolla l'email, aggiunge il contesto "
    "del caso e chiede una bozza di risposta. La bozza arriva in trenta secondi, "
    "lui la rilegge, la corregge dove necessario e la invia. Risparmia tra "
    "i quindici e i venti minuti per ogni email complessa. Con quaranta email "
    "a settimana, il risparmio cumulato e' rilevante. A marzo, pero', ha "
    "inviato una risposta basata su una bozza AI che conteneva un riferimento "
    "normativo leggermente sbagliato, corretto nella versione del decreto "
    "successiva. Il cliente aveva accettato la risposta senza notarlo. "
    "La questione era emersa settimane dopo, durante una verifica interna. "
    "Nessuna conseguenza concreta, ma una domanda che e' rimasta: dove "
    "si trova il punto in cui il risparmio di tempo diventa un rischio?"
)

para(doc,
    "La risposta non e' nella tecnologia, ma nel tipo di contenuto che viene "
    "delegato. L'AI ha una caratteristica strutturale che non cambia "
    "indipendentemente dal modello o dalla versione: produce l'output piu' "
    "plausibile sulla base dei dati con cui e' stata addestrata. Questo "
    "significa che funziona molto bene su problemi che hanno una risposta "
    "consolidata, ampiamente documentata, coerente con la letteratura e con "
    "la prassi. Su problemi che richiedono interpretazione di normativa "
    "recente, valutazione di casi al limite, o giudizi basati su contesti "
    "specifici del cliente, la plausibilita' non coincide necessariamente "
    "con la correttezza."
)

heading(doc, "La distinzione che conta: volume contro specificita'")

para(doc,
    "La regola pratica piu' utile per i professionisti che usano l'AI nel "
    "lavoro con i clienti e' distinguere tra compiti ad alto volume e bassa "
    "specificita', dove l'AI eccelle, e compiti a basso volume e alta "
    "specificita', dove l'AI aiuta ma non puo' sostituire il giudizio "
    "professionale. Nel primo gruppo rientrano la sintesi di documenti lunghi, "
    "la strutturazione di risposte su quesiti ricorrenti, la predisposizione "
    "di checklist operative, la ricerca di precedenti su temi consolidati, "
    "la redazione di comunicazioni standard. In questi casi, la bozza AI "
    "riduce il lavoro di preparazione e libera il professionista per la "
    "verifica, che rimane necessaria ma richiede molto meno tempo."
)

para(doc,
    "Nel secondo gruppo rientrano le interpretazioni di norme recenti o "
    "oggetto di contenzioso, la valutazione di rischi specifici per il "
    "cliente, le risposte su situazioni al limite che richiedono conoscenza "
    "della storia del dossier, e qualsiasi contenuto che il cliente usera' "
    "per prendere decisioni con effetti giuridici o economici rilevanti. "
    "Su questi compiti, l'AI puo' fornire un punto di partenza utile, "
    "ma il professionista deve costruire la risposta finale con autonomia "
    "intellettuale reale, non limitarsi a correggere la forma di una bozza "
    "che non ha rivisto nella sostanza."
)

heading(doc, "L'obbligo di informativa e la Legge 132/2025")

para(doc,
    "C'e' anche una dimensione normativa che i professionisti devono tenere "
    "presente. La Legge 132/2025 ha introdotto, per chi esercita attivita' "
    "professionali regolamentate, un obbligo di informativa al cliente quando "
    "strumenti di intelligenza artificiale vengono utilizzati nell'elaborazione "
    "di pareri, perizie o consulenze. L'obbligo non impone di rivelare quale "
    "strumento si usa o come: impone di rendere il cliente consapevole che "
    "sistemi automatizzati hanno contribuito alla prestazione. La forma "
    "dell'informativa non e' ancora codificata in modo uniforme dai vari "
    "ordini professionali, ma il principio e' chiaro: il cliente ha il "
    "diritto di sapere."
)

para(doc,
    "In pratica, la maggior parte degli studi che usa l'AI in modo sistematico "
    "sta adottando una clausola standard nell'incarico professionale o nel "
    "preventivo, che specifica l'utilizzo di strumenti di supporto "
    "digitale per l'elaborazione delle informazioni. Questa clausola serve "
    "sia a rispettare l'obbligo normativo sia a gestire le aspettative: "
    "un cliente che sa che la bozza di risposta e' stata prodotta con "
    "l'AI e revisionata dal professionista ha un quadro corretto della "
    "prestazione che sta ricevendo."
)

para(doc,
    "Il punto centrale, alla fine, non riguarda la tecnologia. Un cliente "
    "paga il professionista per il giudizio, non per le ore. L'AI riduce "
    "le ore necessarie per preparare, sintetizzare e strutturare: questo "
    "e' un guadagno di efficienza reale che puo' anche giustificare tariffe "
    "piu' competitive. Ma la competenza che il cliente acquista, la "
    "capacita' di valutare la situazione specifica, identificare il rischio "
    "non evidente e scegliere la strada giusta tra opzioni equivalenti in "
    "apparenza, questa rimane integralmente a carico del professionista, "
    "con o senza AI. Quando l'AI risponde al posto del professionista, "
    "senza che il professionista abbia davvero esaminato la risposta, "
    "il cliente ha pagato per qualcosa che non ha ricevuto."
)

riferimenti(doc, [
    "Legge 9 settembre 2025, n. 132 (Legge sull'Intelligenza Artificiale), articoli sull'informativa professionale — Normattiva",
    "PaganiniBellini.it — 'Intelligenza artificiale e avvocati: quando l'errore diventa responsabilita' professionale'",
    "ItaliaOggi.it — 'AI per commercialisti, avvocati e consulenti aziendali: la specializzazione verticale'",
    "KTSFinance.com — 'Assicurazione professionale commercialisti e IA: 3 rischi del 730/2026'",
    "AiPerCommercialisti.it — 'Novita' e strumenti AI per la professione' (aggiornato giugno 2026)",
    "AtlasWorkspace.ai — 'Claude vs NotebookLM: Research Workflow Compared 2026'",
    "Deloitte Italy — 'State of AI in the Enterprise 2026' (rapporto completo)",
])
doc.save(BASE + "2026-07-03_ai-che-risponde-al-cliente-dove-conviene.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
