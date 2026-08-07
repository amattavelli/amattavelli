"""
Quattro articoli Ratio -- 7 agosto 2026

1. 2026-08-07_agenti-ai-dal-chatbot-all-agente-studio-professionale.docx
2. 2026-08-07_ai-act-articolo-50-cinque-giorni-dopo.docx
3. 2026-08-07_opus5-sonnet5-quale-scegliere-professionista.docx
4. 2026-08-07_pmi-trenta-percento-produttivita-non-ci-crede.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# Dal chatbot all'agente: cosa cambia davvero per lo studio
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "AI Agentica e Professioni")
titolo(
    doc,
    "Dal chatbot all'agente: la differenza che conta.",
    "Il mercato AI italiano ha superato 1,8 miliardi nel 2026. "
    "La parola che si sente sempre più è 'agente'. "
    "Vale la pena capire cosa significa davvero, "
    "prima di firmare un contratto o delegare un processo.",
    "A cura della Redazione Ratio  •  7 agosto 2026"
)

para(doc,
    "Nel diritto romano il procurator era la persona che agiva per conto di un'altra, "
    "con poteri definiti da un mandato e responsabilità che derivavano dall'ampiezza "
    "di quei poteri. La distinzione tra il nunzio, che trasmetteva semplicemente un "
    "messaggio, e il procurator, che poteva concludere contratti e obbligare il dominus, "
    "era giuridicamente rilevante e praticamente decisiva. Chi confondeva i due ruoli "
    "si esponeva a conseguenze che il mandante non aveva autorizzato. Gli agenti AI "
    "che stanno entrando nelle aziende e negli studi professionali italiani "
    "pongono esattamente la stessa questione, in forma tecnica e normativa insieme."
)

para(doc,
    "Il termine 'agente AI' descrive un sistema in grado di perseguire un obiettivo "
    "in autonomia, usando strumenti, prendendo decisioni intermedie e agendo su "
    "sistemi esterni senza attendere un'istruzione per ogni passaggio. È una "
    "differenza sostanziale rispetto al chatbot, che risponde a una domanda e si ferma. "
    "Un agente può navigare su un sito, compilare un modulo, inviare un'email, "
    "aggiornare un database, estrarre dati da un documento e inserirli in un gestionale. "
    "Lo fa in sequenza, sulla base di un obiettivo iniziale, senza che nessuno approvi "
    "ogni singola azione. Il mercato AI italiano ha raggiunto 1,82 miliardi di euro "
    "nel 2026, con una crescita del 51% anno su anno, e la quota più dinamica "
    "di questa crescita riguarda proprio i sistemi agentici."
)

heading(doc, "Quanto sono diffusi nelle aziende italiane")

para(doc,
    "Secondo i dati disponibili a metà 2026, il 22% delle grandi aziende italiane "
    "coordina già workflow agentici indipendenti, mentre la quota nelle PMI resta "
    "all'8%. Il divario non sorprende: gli agenti richiedono un'infrastruttura dati "
    "adeguata, la capacità di definire obiettivi operativi precisi e, soprattutto, "
    "qualcuno che sappia progettare il perimetro entro cui l'agente può muoversi. "
    "Il problema che emerge dai dati, però, non è solo la bassa adozione nelle PMI: "
    "è che il 40% delle aziende che già usa agenti AI non ha definito processi "
    "di supervisione umana sui workflow. L'agente agisce. Nessuno controlla cosa ha fatto "
    "prima che diventi definitivo."
)

heading(doc, "Il perimetro del mandato")

para(doc,
    "Per uno studio professionale o un'azienda che valuta di introdurre un agente AI, "
    "la domanda decisiva non è 'cosa sa fare questo agente' ma 'fin dove lo voglio "
    "far arrivare'. Un agente che raccoglie informazioni e prepara bozze opera in un "
    "perimetro controllabile: l'output è un documento che un professionista legge "
    "prima di usarlo. Un agente che invia comunicazioni ai clienti, approva spese, "
    "modifica record nel gestionale o interagisce con l'Agenzia delle Entrate per conto "
    "dello studio opera con poteri che hanno conseguenze reali e, in alcuni casi, "
    "giuridicamente rilevanti. La distinzione non è tecnica: è di governo."
)

heading(doc, "Cosa dice la norma")

para(doc,
    "L'AI Act, anche nella versione rivista dal Digital Omnibus, prevede che chi "
    "utilizza agenti AI in processi con impatto significativo sulle persone, "
    "sui dati o sulle decisioni aziendali debba documentare il perimetro di azione "
    "del sistema, mantenere la supervisione umana e tenere un registro delle operazioni. "
    "Dal 2 agosto 2026 le autorità di vigilanza hanno gli strumenti per ispezionare "
    "queste configurazioni. Il 48% delle aziende italiane che usano agenti AI tiene "
    "un registro; la metà non sa con certezza quanti e quali agenti operano per "
    "suo conto. Per uno studio professionale che usa questi strumenti per conto dei "
    "propri clienti, la responsabilità non si trasferisce all'agente: "
    "rimane dove è sempre rimasta."
)

para(doc,
    "Il diritto romano aveva risolto la questione con chiarezza: chi conferisce "
    "un mandato risponde delle conseguenze dell'azione del procuratore. "
    "Dodicimila anni dopo, il principio regge ancora. "
    "Il fatto che il procuratore questa volta sia un software non sposta la firma."
)

riferimenti(doc, [
    "Impesud.it – 'Agentic AI in Italia 2026: Dalla Teoria all'Azione' (2026)",
    "AI4Business.it – 'Agenti AI e nuove regolamentazioni: il quadro normativo italiano ed europeo 2026' (2026)",
    "ICT Security Magazine – 'AI governance aziendale: i rischi degli agenti autonomi' (2026)",
    "Il Giornale – 'AI agentica: quanto è diffusa nelle aziende italiane?' (2026)",
    "Regolamento UE 2024/1689 (AI Act), artt. 26-27 obblighi dei deployer",
    "Regolamento UE 2026/1744 (Digital Omnibus AI), disposizioni su PMI e proporzionalità",
])
doc.save(BASE + "2026-08-07_agenti-ai-dal-chatbot-all-agente-studio-professionale.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# AI Act art. 50: cinque giorni dopo il 2 agosto
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Normativa AI")
titolo(
    doc,
    "Cinque giorni dopo il 2 agosto: cosa è cambiato.",
    "Dal 2 agosto 2026 i chatbot devono dichiarare la propria natura artificiale. "
    "I contenuti sintetici vanno marcati. Le sanzioni arrivano al 3% del fatturato. "
    "Cinque giorni dopo, il mercato si è accorto di pochissimo.",
    "A cura della Redazione Ratio  •  7 agosto 2026"
)

para(doc,
    "Luigi Pirandello aveva capito qualcosa che le imprese europee stanno riscoprendo "
    "con qualche sorpresa: l'identità non è una proprietà dell'individuo, "
    "è una costruzione sociale che dipende dallo sguardo altrui. Enrico IV, "
    "il personaggio del suo dramma, sceglie di continuare a recitare la follia "
    "anche quando potrebbe smettere, perché l'identità che gli hanno attribuito "
    "è diventata più comoda della verità. Dal 2 agosto 2026, l'articolo 50 "
    "dell'AI Act impone a chatbot e sistemi generativi l'obbligo opposto: "
    "dichiarare la propria identità artificiale, anche quando sarebbe più "
    "comodo non farlo."
)

para(doc,
    "L'articolo 50 del Regolamento UE 2024/1689 è entrato in applicazione il "
    "2 agosto 2026, in anticipo rispetto ad altri obblighi dell'AI Act rinviati "
    "dal Digital Omnibus. I suoi contenuti principali sono tre. Primo: i sistemi AI "
    "progettati per interagire con le persone, come chatbot e assistenti vocali, devono "
    "informare l'utente fin dal primo contatto che sta parlando con un sistema "
    "artificiale, salvo il caso in cui la natura del sistema sia già evidente "
    "a un utente ragionevolmente informato. Secondo: i sistemi che generano contenuti "
    "sintetici, testi, immagini, audio, video, devono marcarli tecnicamente in modo "
    "che siano riconoscibili come generati da AI. Terzo: i fornitori di modelli "
    "che generavano contenuti sintetici prima del 2 agosto 2026 hanno tempo fino al "
    "2 dicembre 2026 per adeguarsi ai requisiti di marcatura; per i nuovi sistemi "
    "l'obbligo è immediato."
)

heading(doc, "Chi è coinvolto e con quale responsabilità")

para(doc,
    "L'articolo 50 si rivolge principalmente ai fornitori dei sistemi AI, cioè "
    "a chi sviluppa e immette sul mercato il chatbot o il generatore di contenuti. "
    "Ma coinvolge anche i deployer, ossia le aziende e i professionisti che usano "
    "questi sistemi per interagire con i propri clienti. Uno studio professionale "
    "che ha installato un assistente AI sul proprio sito per rispondere alle domande "
    "dei clienti deve verificare che il sistema sia configurato per dichiararsi "
    "non umano. Un'impresa che usa strumenti di generazione automatica di comunicazioni "
    "commerciali deve assicurarsi che queste portino una marcatura leggibile. "
    "Le sanzioni previste per le violazioni dell'articolo 50 arrivano al 3% "
    "del fatturato mondiale annuo, con il limite minimo di 15 milioni di euro "
    "per le organizzazioni più grandi."
)

heading(doc, "Cosa verifica concretamente uno studio")

para(doc,
    "Per la maggior parte degli studi professionali e delle PMI italiane, "
    "la verifica di conformità all'articolo 50 è praticabile in tempi brevi "
    "e non richiede consulenza specializzata di livello enterprise. Le domande "
    "operative sono quattro. Il chatbot sul sito o sulla piattaforma dello studio "
    "si identifica come AI nella prima interazione? I contenuti generati "
    "automaticamente, da newsletter a risposte ai clienti, recano una dichiarazione "
    "di origine? I fornitori degli strumenti utilizzati hanno comunicato come "
    "gestiscono la marcatura tecnica dei contenuti sintetici? Esiste un referente "
    "interno che ha in carico queste verifiche? Se la risposta a una di queste "
    "domande è 'non lo so', è il momento di chiederlo."
)

heading(doc, "Il nodo del 'già evidente'")

para(doc,
    "Il testo dell'articolo 50 prevede un'esenzione per i casi in cui la natura "
    "artificiale del sistema sia già evidente a un utente ragionevolmente informato. "
    "Questa formulazione apre uno spazio interpretativo che sarà riempito "
    "dalla prassi e, probabilmente, dalle prime decisioni delle autorità di vigilanza. "
    "Cosa è 'ragionevolmente evidente'? Un nome come 'Assistente virtuale' sul sito "
    "di una banca è sufficiente? Un avatar con tratti chiaramente non umani esonera "
    "dall'obbligo di comunicazione verbale? Le linee guida dell'AGCM e del Garante "
    "AI italiano, attese per i prossimi mesi, dovranno chiarire questi margini. "
    "Nel dubbio, il comportamento cautelativo è dichiarare: costa meno dell'ambiguità."
)

para(doc,
    "Pirandello conosceva bene il finale: quando l'identità fittizia viene smascherata "
    "nel momento sbagliato, le conseguenze sono difficili da gestire. "
    "L'articolo 50 non chiede ai sistemi AI di smettere di sembrare intelligenti. "
    "Chiede solo di dire chi sono. "
    "È una richiesta che, nel 2026, sembra più facile da soddisfare di quanto "
    "alcune aziende stiano dimostrando."
)

riferimenti(doc, [
    "Altalex.com – 'AI Act: che cosa cambia davvero dal 2 agosto 2026 per imprese e professionisti' (31 luglio 2026)",
    "Il Fatto Quotidiano – 'AI Act, dal 2 agosto obblighi di trasparenza per chatbot e deepfake' (2 agosto 2026)",
    "42lf.it – 'AI Act, art. 50: gli obblighi di trasparenza sull’intelligenza artificiale diventano operativi dal 2 agosto 2026'",
    "Fiscal Focus – 'AI Act, obblighi di trasparenza operativi per chatbot e contenuti generati dall’IA'",
    "CityNext.it – 'AI Act, dal 2 agosto obbligo di etichettare i contenuti creati dall’intelligenza artificiale' (31 luglio 2026)",
    "Regolamento UE 2024/1689 (AI Act), articolo 50",
])
doc.save(BASE + "2026-08-07_ai-act-articolo-50-cinque-giorni-dopo.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# Opus 5 vs Sonnet 5: quale serve a uno studio professionale
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Strumenti e Modelli AI")
titolo(
    doc,
    "Opus 5 o Sonnet 5: la domanda giusta non è quale è meglio.",
    "Anthropic ha rilasciato Sonnet 5 il 30 giugno e Opus 5 il 24 luglio 2026. "
    "Due modelli di fascia alta con prezzi distinti e capacità diverse. "
    "Per uno studio professionale, la scelta dipende da cosa si fa davvero con l'AI.",
    "A cura della Redazione Ratio  •  7 agosto 2026"
)

para(doc,
    "Ogni anno, alla fine di giugno, il Touring Club Italiano pubblica le stelle "
    "degli hotel: cinque stelle lusso, cinque stelle, quattro stelle, e così via. "
    "Il sistema funziona perché la stella misura una cosa precisa, i servizi "
    "e gli standard strutturali, non l'adeguatezza per lo scopo del viaggio. "
    "Un albergo a cinque stelle in una città sbagliata vale meno di un tre stelle "
    "a dieci minuti dal cantiere che devi visitare. I benchmark dei modelli AI "
    "funzionano con la stessa logica: misurano le prestazioni su compiti definiti, "
    "non la rispondenza al lavoro specifico che hai da fare."
)

para(doc,
    "Anthropic ha rilasciato Claude Sonnet 5 il 30 giugno 2026 e Claude Opus 5 "
    "il 24 luglio 2026, a distanza di meno di un mese. Sonnet 5 è il modello "
    "predefinito per i piani Free e Pro di Claude e viene descritto come il "
    "modello Sonnet più agentico mai rilasciato: ragiona, usa strumenti come browser "
    "e terminale, e verifica il proprio output senza che nessuno lo chieda. "
    "Opus 5 è il modello di punta, posizionato per ragionamento complesso, "
    "coding professionale e lavoro di conoscenza a livello enterprise. "
    "In termini di prezzo API, Sonnet 5 parte da 2 dollari per milione di token "
    "in input (tariffa promozionale fino al 31 agosto, poi 3 dollari); "
    "Opus 5 costa 5 dollari in input e 25 in output."
)

heading(doc, "Dove Opus 5 fa la differenza")

para(doc,
    "Opus 5 batte Sonnet 5 in modo misurabile su compiti che richiedono ragionamento "
    "multistep profondo, analisi di documenti lunghi e complessi, coding su basi di "
    "codice articolate. Su Frontier-Bench, il benchmark che misura le prestazioni "
    "su task professionali complessi, Opus 5 raddoppia il punteggio di Opus 4.8 e "
    "si avvicina a Fable 5, il modello di frontiera di Anthropic, a metà del costo. "
    "Per uno studio professionale, questo si traduce in un vantaggio misurabile su "
    "attività come l'analisi di contratti complessi con molte clausole in tensione "
    "tra loro, la revisione di bilanci consolidati con strutture articolate, "
    "l'interpretazione di circolari con rimandi incrociati a normative precedenti. "
    "Su questi compiti, la differenza di qualità rispetto a Sonnet 5 è percepibile."
)

heading(doc, "Dove Sonnet 5 è sufficiente")

para(doc,
    "Per la maggior parte dei compiti quotidiani di uno studio professionale, "
    "Sonnet 5 produce output di qualità molto vicina a Opus 5 con un costo "
    "significativamente inferiore. Redazione di lettere ai clienti, sintesi di "
    "documenti, bozze di risposte a quesiti standard, ricerca di riferimenti normativi "
    "su temi ben definiti, strutturazione di presentazioni: su questi task, "
    "la differenza tra i due modelli è difficilmente percepibile nell'output finale "
    "e Sonnet 5 è più veloce. La finestra di contesto di un milione di token, "
    "condivisa da entrambi, consente di lavorare su documenti lunghi senza tagli. "
    "Per un uso prevalentemente testuale e generativo, Sonnet 5 regge."
)

heading(doc, "La scelta che ha senso fare")

para(doc,
    "Per uno studio professionale che usa Claude tramite piano Pro o Team "
    "(dove entrambi i modelli sono disponibili senza costi aggiuntivi per messaggio), "
    "la risposta pratica è usare Sonnet 5 per impostazione predefinita e passare "
    "a Opus 5 per i compiti che richiedono il massimo: analisi legali complesse, "
    "revisione di documenti dove l'errore ha un costo alto, ragionamento su scenari "
    "con molte variabili. Per chi accede tramite API, la scelta dipende dal volume: "
    "su task in batch ad alto volume, il differenziale di costo tra i due modelli "
    "diventa rilevante e Sonnet 5 è la scelta razionale. Su task critici a basso "
    "volume, il maggior costo di Opus 5 è giustificato dalla qualità. "
    "Quello che non è razionale è scegliere il modello più potente per tutto "
    "per ragioni di principio, o scegliere quello più economico per risparmiare "
    "su compiti dove la qualità ha un valore diretto."
)

para(doc,
    "L'albergo a cinque stelle lusso è meraviglioso. "
    "Ma se il convegno è a venti chilometri, "
    "il quattro stelle con il parcheggio gratuito finisce per lavorare meglio."
)

riferimenti(doc, [
    "Anthropic – 'Introducing Claude Sonnet 5' (30 giugno 2026)",
    "Anthropic – 'Introducing Claude Opus 5' (24 luglio 2026)",
    "InfoData Il Sole 24 Ore – 'Claude Opus 5 in cinque punti: più efficiente, più veloce, ma il salto è meno netto di quanto racconti Anthropic' (27 luglio 2026)",
    "Bleap.finance – 'Claude Opus 5.0: Recensione, Benchmark, Prezzi e Guida API' (2026)",
    "DataCamp – 'Claude Opus 5 vs. Claude Sonnet 5: Choosing Which to Use' (2026)",
    "Cloudzy.com – 'Claude Sonnet 5 vs. Opus 4.8: Which Model to Actually Run' (2026)",
])
doc.save(BASE + "2026-08-07_opus5-sonnet5-quale-scegliere-professionista.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# +30% produttivita' promessa, PMI italiane ancora ferme
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "PMI e Adozione AI")
titolo(
    doc,
    "Plus trenta per cento di produttività. PMI italiane: aspetta.",
    "Lo SME-AIMIX 2026 documenta guadagni di produttività fino al 30% per le PMI "
    "che adottano l'AI. Il 58,6% delle PMI italiane che non la usa cita la mancanza "
    "di competenze interne. Il potenziale c'è. Il problema è altrove.",
    "A cura della Redazione Ratio  •  7 agosto 2026"
)

para(doc,
    "Zenone di Elea dimostrò con un argomento geometricamente ineccepibile "
    "che Achille non avrebbe mai raggiunto la tartaruga: ogni volta che l'eroe "
    "copriva la distanza che lo separava dall'animale, la tartaruga si era già "
    "spostata un po' più avanti, all'infinito. Il paradosso è falso, "
    "ovviamente, Achille supera la tartaruga, ma funziona come metafora per "
    "descrivere la posizione di molte PMI italiane rispetto all'adozione dell'AI: "
    "ogni volta che il potenziale viene documentato in modo convincente, le ragioni "
    "per non agire si spostano un po' più avanti, e il divario non si chiude."
)

para(doc,
    "Lo SME-AI Maturity Index 2026, il rapporto di Webidoo Insight Lab sulla "
    "maturità digitale delle PMI italiane ed europee, quantifica il potenziale "
    "di guadagno di produttività per le piccole e medie imprese che adottano "
    "soluzioni AI in modo strutturato: fino al 30% di miglioramento sui processi "
    "principali. Il dato è coerente con altri studi europei, che su campioni "
    "più ampi mostrano che le aziende ad alta esposizione all'AI hanno registrato "
    "una crescita media della produttività del 34% tra il 2018 e il 2025, contro "
    "il 24% dei settori a bassa esposizione. Il potenziale è documentato, "
    "misurabile e reale. Il problema è che le PMI italiane, nel 2026, "
    "lo stanno guardando da lontano."
)

heading(doc, "Le ragioni che le PMI danno")

para(doc,
    "Tra le PMI italiane che non usano AI e non prevedono di farlo nel breve, "
    "il 58,6% cita la mancanza di competenze interne, il 47,3% l'incertezza "
    "normativa e il 45,2% le difficoltà legate alla qualità e alla gestione "
    "dei dati. Il 76% non ha investito in AI e non prevede di farlo. "
    "Lo score di maturità AI delle PMI italiane nello SME-AIMIX è 8,5%, "
    "sotto la media europea. Questi numeri vanno letti con attenzione: "
    "la mancanza di competenze interne non significa che l'AI sia inaccessibile "
    "alle PMI, significa che le PMI non hanno al loro interno qualcuno in grado "
    "di valutarne l'applicazione e gestirne l'introduzione. "
    "È un problema di intermediazione professionale prima che di tecnologia."
)

heading(doc, "Dove l'AI si è già inserita")

para(doc,
    "I dati sui settori di applicazione nelle PMI che hanno adottato l'AI "
    "descrivono una concentrazione prevedibile: il 33,1% dei casi riguarda "
    "marketing e vendite, il 25,7% l'amministrazione, il 20% la ricerca e sviluppo. "
    "La contabilità, la gestione dei flussi di cassa, la preparazione di dichiarazioni "
    "e la revisione di documenti sono aree dove l'AI porta guadagni di tempo "
    "misurabili già con strumenti di livello consumer, senza infrastrutture "
    "complesse. Sono anche le aree dove il commercialista, il consulente del lavoro "
    "e il revisore hanno accesso diretto ai processi del cliente e potrebbero "
    "essere la figura che introduce lo strumento, lo configura e forma il personale. "
    "Il 30% di produttività in più non arriva dall'abbonamento: arriva "
    "da chi sa dove applicarlo."
)

heading(doc, "Il fattore che separa chi avanza da chi aspetta")

para(doc,
    "Lo SME-AIMIX 2026 identifica una correlazione chiara tra la presenza di un "
    "referente dedicato all'AI, interno o esterno all'azienda, e l'effettiva "
    "integrazione degli strumenti nei processi. Le PMI che hanno avviato "
    "percorsi strutturati di adozione quasi sempre hanno in comune "
    "una figura che ha tradotto il potenziale generico in un caso d'uso specifico: "
    "questo processo, con questo strumento, in questo modo. Le PMI che si fermano "
    "alla fase di sperimentazione con ChatGPT per redigere email hanno quasi sempre "
    "in comune l'assenza di quella figura. Il potenziale del 30% di produttività "
    "in più è condizionale: richiede qualcuno che sappia attivarlo."
)

para(doc,
    "Achille raggiunge la tartaruga perché smette di dividere lo spazio "
    "in frazioni infinite e comincia a correre. "
    "Le PMI italiane hanno ancora il paradosso in testa."
)

riferimenti(doc, [
    "MediaKey.it – 'SME-AIMIX 2026: fino al +30% di produttività con l’AI, "
    "ma le PMI italiane non sono ancora pronte' (2026)",
    "AdnKronos – 'AI: nelle PMI italiane il potenziale c’è ma la maturità resta bassa' (30 giugno 2026)",
    "AgendaDigitale.eu – 'AI nelle PMI italiane: competenze e dati frenano la svolta digitale' (2026)",
    "ESTE.it – 'AI Act, le PMI italiane arrivano impreparate alla svolta' (2026)",
    "Istat – Rapporto sull’uso delle tecnologie ICT e AI nelle imprese italiane, 2025-2026",
    "Corriere Nazionale – 'La Nuova Era dell’Efficienza: Come l’AI Agentica sta Trasformando il Lavoro e l’Impresa in Italia nel 2026' (16 aprile 2026)",
])
doc.save(BASE + "2026-08-07_pmi-trenta-percento-produttivita-non-ci-crede.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
