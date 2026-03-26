"""
Script per creare l'articolo "Gli Agenti AI" in formato .docx
Stile editoriale: coerente con la serie DRC/Excellere in Studio di Alessandro Mattavelli
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = "/home/user/amattavelli/Ratio/articoli/Agenti AI - la nuova frontiera per professionisti e PMI italiane.docx"

doc = Document()

# ─── Stili base ───────────────────────────────────────────────────────────────

def add_heading2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph(text, style="Heading 3")
    return p

def add_normal(doc, text):
    p = doc.add_paragraph(text, style="Normal")
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.style.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    # Simulate bullet with dash prefix
    p.clear()
    p.add_run(f"- {text}")
    return p

def add_numbered(doc, num, text):
    p = doc.add_paragraph(style="Normal")
    p.add_run(f"{num}. {text}")
    return p

def add_bold_intro(doc, label, text):
    """Aggiunge una riga con label in grassetto seguita da testo normale."""
    p = doc.add_paragraph(style="Normal")
    run_bold = p.add_run(label)
    run_bold.bold = True
    run_normal = p.add_run(text)
    return p

# ─── TITOLO ───────────────────────────────────────────────────────────────────

title = doc.add_paragraph(style="Normal")
title_run = title.add_run("Gli Agenti AI: la nuova frontiera per professionisti e PMI italiane")
title_run.bold = True
title_run.font.size = Pt(18)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ─── BYLINE ───────────────────────────────────────────────────────────────────

byline = doc.add_paragraph("Di Alessandro Mattavelli", style="Normal")

doc.add_paragraph("", style="Normal")  # Spazio

# ─── INTRODUZIONE ─────────────────────────────────────────────────────────────

add_normal(doc,
    "Negli ultimi mesi, il termine «agenti AI» è comparso con insistenza crescente nei convegni, nelle newsletter "
    "tecnologiche e nelle riunioni di direzione. Ma al di là del clamore mediatico, cosa significa davvero questa "
    "espressione per un commercialista, un consulente del lavoro, un responsabile amministrativo o l'imprenditore "
    "di una piccola o media impresa italiana? Significa molto più di quanto si pensi, e prima lo si comprende, "
    "prima si può scegliere consapevolmente se e come adottare questi strumenti. Il 2026 non è l'anno in cui "
    "decidere se occuparsi di agenti AI: è l'anno in cui decidere come farlo, perché chi aspetta ancora "
    "accumula uno svantaggio competitivo che diventa ogni giorno più difficile da colmare. Questo articolo "
    "vuole offrire una bussola pratica — non un elenco di promesse tecnologiche — per orientarsi in un "
    "panorama in rapida evoluzione."
)

# ─── SEZ 1: COSA SONO ─────────────────────────────────────────────────────────

add_heading2(doc, "Cosa sono gli Agenti AI — e perché non sono semplici chatbot")

add_normal(doc,
    "Quando utilizziamo un chatbot come ChatGPT, Claude o Gemini, interagiamo con un sistema che risponde "
    "a una domanda alla volta: noi scriviamo, lui risponde, noi leggiamo e decidiamo cosa fare. L'intelligenza "
    "è dentro lo schermo; il controllo e l'esecuzione restano nelle nostre mani. Un agente AI è una cosa "
    "fondamentalmente diversa."
)

add_normal(doc,
    "Un agente AI riceve un obiettivo — non una singola domanda — e poi agisce in autonomia per raggiungerlo. "
    "Osserva il contesto, pianifica i passi necessari, esegue operazioni su altri sistemi (CRM, ERP, casella "
    "di posta, gestionale contabile, piattaforma e-commerce), controlla i risultati intermedi e si corregge "
    "se qualcosa non va. Non aspetta che noi approviamo ogni singolo passaggio: lavora, e ci informa quando "
    "ha finito — o quando incontra un ostacolo che richiede la nostra attenzione."
)

add_normal(doc,
    "La differenza, in termini pratici, è enorme. Un chatbot è un amplificatore della nostra produttività "
    "individuale: ci aiuta a scrivere, a ragionare, a sintetizzare. Un agente AI è un collaboratore digitale "
    "che opera in parallelo a noi, eseguendo compiti strutturati mentre noi ci dedichiamo ad attività "
    "ad alto valore strategico. Non è fantascienza: le piattaforme di Microsoft (Copilot Studio), "
    "Salesforce (Agentforce), Google (Vertex AI Agents) e molte startup italiane e internazionali offrono "
    "già oggi strumenti di questo tipo, con livelli di complessità e costo molto diversi tra loro."
)

# ─── SEZ 2: I NUMERI ──────────────────────────────────────────────────────────

add_heading2(doc, "Il mercato AI in Italia: i numeri che contano")

add_normal(doc,
    "Prima di parlare di strategie, è utile avere un quadro realistico della situazione italiana. "
    "I dati disponibili all'inizio del 2026 raccontano una storia fatta di opportunità enormi e "
    "ritardi strutturali che non possiamo ignorare."
)

add_normal(doc,
    "Secondo l'Osservatorio Artificial Intelligence del Politecnico di Milano, il mercato AI italiano "
    "ha raggiunto i 760 milioni di euro nel 2024, con una crescita del 52% rispetto all'anno precedente, "
    "e nel 2025 ha superato il miliardo di euro. Una crescita impressionante in termini assoluti, "
    "ma che va letta con onestà: siamo ancora lontani dai volumi di investimento di Francia, Germania "
    "e Regno Unito, e la distribuzione tra grandi aziende e PMI è molto disomogenea."
)

add_normal(doc,
    "I dati ISTAT 2025 ci dicono che solo il 16% delle imprese italiane con almeno dieci addetti "
    "utilizza soluzioni di intelligenza artificiale, e oltre l'80% delle PMI si trova ancora in una "
    "fase esplorativa. Eppure il 26,7% delle piccole e medie imprese utilizza già qualche strumento AI, "
    "rispetto al 18% del 2023: la curva di adozione si sta accelerando. A livello globale, secondo "
    "Salesforce Research, il 35% delle organizzazioni dichiara già un utilizzo diffuso di agenti AI, "
    "e il 17% li ha dispiegati sull'intera azienda, con una crescita del 282% nell'adozione rispetto "
    "all'anno precedente."
)

add_normal(doc,
    "Il Microsoft AI Tour di Milano (marzo 2026) ha stimato che l'adozione pervasiva dell'AI potrebbe "
    "generare 336 miliardi di euro di valore aggiunto annuo per l'economia italiana entro il 2040. "
    "Sono numeri che vanno presi come ordini di grandezza, non come previsioni esatte, ma che "
    "confermano la dimensione strutturale — non congiunturale — di questa trasformazione."
)

# ─── SEZ 3: PROFESSIONISTI ────────────────────────────────────────────────────

add_heading2(doc, "Come gli Agenti AI cambiano il lavoro dei professionisti")

add_normal(doc,
    "Per i professionisti italiani — commercialisti, consulenti del lavoro, avvocati, revisori, "
    "consulenti aziendali — la sfida degli agenti AI è duplice: da un lato, capire come questi "
    "strumenti possono amplificare le proprie capacità; dall'altro, comprendere come cambierà "
    "la relazione con i propri clienti, che li chiedono sempre di più in anticipo."
)

add_normal(doc,
    "Vediamo alcuni ambiti concreti in cui gli agenti AI stanno già producendo cambiamenti "
    "misurabili per i professionisti:"
)

add_bullet(doc, "Analisi documentale automatizzata: un agente può esaminare decine di contratti, "
    "bilanci o dichiarazioni fiscali in parallelo, estrarre le informazioni rilevanti e "
    "produrre un report strutturato in pochi minuti. Il professionista interviene sulla "
    "valutazione strategica, non sull'estrazione dei dati.")

add_bullet(doc, "Monitoraggio normativo continuo: gli agenti possono tenere traccia delle "
    "modifiche normative (circolari, risoluzioni, decreti) e inviare al professionista "
    "un digest settimanale personalizzato per tipologia di cliente.")

add_bullet(doc, "Ricerca di agevolazioni e finanza agevolata: studi recenti indicano che "
    "le imprese che usano AI per la ricerca di bandi scoprono in media il 40% in più di "
    "opportunità rispetto alla ricerca manuale, con una qualità di matching sensibilmente "
    "superiore.")

add_bullet(doc, "Predisposizione della prima bozza di atti e comunicazioni: lettere ai "
    "clienti, risposta a questionari di due diligence, bozze di contratti standard. "
    "L'agente prepara, il professionista revisiona e firma.")

add_bullet(doc, "Gestione delle scadenze e delle to-do list di studio: integrazione con "
    "i gestionali per rilevare automaticamente le scadenze imminenti, inviare "
    "promemoria ai clienti e aggiornare il calendario dello studio.")

add_normal(doc,
    "È importante essere onesti su un punto: gli agenti AI non eliminano la necessità "
    "del giudizio professionale. Lo spostano. Il professionista del futuro prossimo non è "
    "quello che sa fare le cose meglio di un agente AI su compiti ripetitivi — è quello "
    "che sa contestualizzare, interpretare, comunicare e assumersi la responsabilità delle "
    "decisioni che contano. Le competenze non scompaiono: si spostano verso livelli di "
    "astrazione più elevati."
)

# ─── SEZ 4: PMI ───────────────────────────────────────────────────────────────

add_heading2(doc, "Applicazioni concrete per le PMI italiane")

add_normal(doc,
    "Per le piccole e medie imprese italiane, il discorso sugli agenti AI tende a "
    "bloccarsi su due obiezioni ricorrenti: «Non ho un reparto IT» e «Non ho il budget». "
    "Entrambe le obiezioni sono comprensibili, ma sempre meno valide nel 2026."
)

add_normal(doc,
    "La proliferazione di piattaforme no-code e low-code per la creazione di agenti AI "
    "ha abbassato drasticamente la barriera d'ingresso tecnica: oggi un'imprenditrice o "
    "un responsabile commerciale con buone competenze digitali di base può configurare un "
    "agente semplice senza scrivere una riga di codice. E i costi di accesso agli strumenti "
    "di base sono nell'ordine delle decine — non centinaia — di euro al mese."
)

add_normal(doc,
    "Alcune aree di applicazione particolarmente rilevanti per le PMI italiane nel breve termine:"
)

add_bullet(doc, "Customer service e gestione delle richieste: un agente può rispondere "
    "autonomamente alle domande frequenti di clienti e fornitori (stato degli ordini, "
    "disponibilità, condizioni di pagamento), liberando le risorse interne per le "
    "richieste complesse.")

add_bullet(doc, "Supporto commerciale e CRM: aggiornamento automatico del CRM dopo "
    "ogni interazione commerciale, generazione di proposte personalizzate, "
    "follow-up automatici sulle trattative aperte.")

add_bullet(doc, "Controllo della liquidità e reporting finanziario: integrazione con "
    "il gestionale contabile per produrre report settimanali sulla posizione di cassa, "
    "segnalare anomalie nei pagamenti e simulare scenari di breve termine.")

add_bullet(doc, "Recruiting e gestione delle candidature: screening iniziale dei CV, "
    "comunicazioni standardizzate ai candidati, pianificazione dei colloqui.")

add_bullet(doc, "Marketing operativo: pubblicazione programmata di contenuti sui social, "
    "A/B testing automatizzato delle campagne email, analisi delle performance "
    "e suggerimenti di ottimizzazione.")

add_normal(doc,
    "Il punto cruciale non è la singola applicazione: è la logica di accumulazione. "
    "Ogni agente che automatizza un processo ripetitivo libera tempo e attenzione che "
    "possono essere reinvestiti in attività a maggior valore. L'impatto non è lineare: "
    "è moltiplicativo, e si manifesta appieno quando gli agenti iniziano a lavorare "
    "in modo coordinato all'interno dello stesso ecosistema aziendale."
)

# ─── SEZ 5: FRAMEWORK PREA ────────────────────────────────────────────────────

add_heading2(doc, "Il framework PREA per avviare il percorso agentico")

add_normal(doc,
    "Una delle domande più frequenti che riceviamo da imprenditori e professionisti è: "
    "«Da dove cominciamo?». La risposta dipende dalla specifica realtà aziendale, "
    "ma un approccio strutturato che funziona bene nelle PMI italiane può essere "
    "sintetizzato nell'acronimo PREA:"
)

add_bold_intro(doc, "P — Processo: ",
    "identificate un processo aziendale ripetitivo, ad alto volume e a bassa variabilità. "
    "Non iniziate dall'eccezionale: iniziate dal routinario. La gestione delle richieste "
    "di preventivo, l'aggiornamento del CRM, l'invio delle conferme d'ordine sono "
    "candidati ideali.")

add_bold_intro(doc, "R — Risultato atteso: ",
    "definite in modo preciso e misurabile il risultato che volete ottenere. Non «fare "
    "prima le cose», ma «ridurre il tempo di risposta alle richieste di preventivo "
    "da 48 ore a 4 ore». Senza un obiettivo misurabile, non saprete mai se il vostro "
    "agente sta funzionando.")

add_bold_intro(doc, "E — Esperienza pilota: ",
    "avviate un progetto pilota su scala ridotta, con un perimetro limitato e un "
    "orizzonte temporale definito (8-12 settimane). Misurate i risultati, "
    "raccogliete i feedback del team, identificate i punti di attrito.")

add_bold_intro(doc, "A — Ampliamento graduale: ",
    "solo dopo aver validato il pilota, allargate il perimetro. Aggiungete nuove "
    "funzionalità, integrate altri sistemi, coinvolgete più reparti. La scalabilità "
    "dell'AI agentica è uno dei suoi punti di forza — ma funziona solo se le fondamenta "
    "del pilota sono solide.")

add_normal(doc,
    "Il 78% dei C-level intervistati in uno studio IBM concorda sul fatto che per "
    "ottenere il massimo beneficio dall'AI agentica sia necessario ripensare il modello "
    "operativo dell'impresa. Introdurre agenti AI senza una fase di analisi e ridisegno "
    "dei flussi non produce automazione intelligente: produce automazione del caos. "
    "Il framework PREA serve esattamente a evitare questa trappola."
)

# ─── SEZ 6: RISCHI E GOVERNANCE ───────────────────────────────────────────────

add_heading2(doc, "Rischi, governance e AI Act europeo")

add_normal(doc,
    "Un articolo onesto sugli agenti AI non può ignorare i rischi. Non per spaventare, "
    "ma perché la consapevolezza dei rischi è la premessa di un'adozione responsabile "
    "e duratura."
)

add_heading3(doc, "L'AI Act europeo: cosa cambia per le aziende italiane")

add_normal(doc,
    "Il Regolamento europeo sull'intelligenza artificiale (AI Act) è entrato in vigore "
    "nell'agosto 2024 e si applica progressivamente fino al 2027. Le prime obbligazioni — "
    "relative ai sistemi AI considerati «inaccettabili» — sono già operative. "
    "Le aziende italiane che utilizzano sistemi AI nei processi di assunzione del personale, "
    "nell'erogazione di credito, nella gestione di infrastrutture critiche o nel controllo "
    "dei lavoratori devono verificare con attenzione a quale categoria di rischio appartengono "
    "i loro sistemi e quali obblighi di trasparenza, documentazione e supervisione umana "
    "si applicano. Il consiglio pratico è di coinvolgere subito il proprio consulente legale "
    "e il DPO (Data Protection Officer) nella valutazione di ogni nuovo progetto agentico."
)

add_heading3(doc, "I rischi operativi degli agenti AI")

add_normal(doc,
    "Oltre al quadro normativo, esistono rischi operativi concreti che ogni impresa "
    "deve considerare prima di dispiegare agenti AI nei propri processi:"
)

add_numbered(doc, 1, "Allucinazioni e errori: gli agenti AI possono produrre informazioni "
    "errate con la stessa sicurezza con cui producono informazioni corrette. "
    "Non esiste ancora un sistema di AI che non sbagli mai. La supervisione umana "
    "sui compiti ad alto impatto non è un'opzione: è una necessità.")

add_numbered(doc, 2, "Sicurezza dei dati: un agente AI che accede ai vostri sistemi aziendali "
    "deve essere configurato con criteri rigorosi di controllo degli accessi. "
    "Chiedetevi sempre: quali dati può leggere? Quali può modificare? "
    "Chi supervisiona le sue azioni?")

add_numbered(doc, 3, "Dipendenza tecnologica: costruire processi critici su un unico "
    "fornitore di AI crea una dipendenza che può rivelarsi costosa. "
    "Valutate sempre le condizioni contrattuali, le politiche di continuità del "
    "servizio e la possibilità di migrare verso alternative.")

add_numbered(doc, 4, "Resistenza interna al cambiamento: il 49% dei top manager identifica "
    "la carenza di competenze interne come la principale barriera all'adozione dell'AI. "
    "Ma la resistenza culturale è spesso altrettanto rilevante. Investire nella "
    "formazione del team non è un costo accessorio: è una condizione abilitante.")

add_normal(doc,
    "Uno studio del MIT evidenzia che solo il 5% delle aziende è riuscito a ottenere "
    "ritorni finanziari significativi dai progetti di AI, e tra il 70% e l'80% delle "
    "iniziative agentiche non è riuscito a scalare a livello enterprise. I motivi "
    "principali: mancanza di chiarezza sugli obiettivi, sottovalutazione del change "
    "management, e scelta di iniziare dai casi d'uso più complessi anziché dai più "
    "semplici. Il framework PREA esiste proprio per evitare questi errori."
)

# ─── SEZ 7: CONCLUSIONI ───────────────────────────────────────────────────────

add_heading2(doc, "Conclusioni: non è più il momento di aspettare")

add_normal(doc,
    "La domanda che sentiamo più spesso oggi — «Dobbiamo davvero occuparci di agenti AI adesso?» — "
    "rivela un fraintendimento fondamentale sulla natura di questa trasformazione. "
    "Non si tratta di una moda tecnologica con un ciclo vitale di 18 mesi. Si tratta di un "
    "cambiamento strutturale nel modo in cui il lavoro viene organizzato, distribuito e valorizzato."
)

add_normal(doc,
    "I professionisti e le imprese che stanno costruendo oggi le loro competenze agentiche — "
    "anche con progetti pilota piccoli e circoscritti — stanno accumulando un vantaggio "
    "difficilmente recuperabile da chi aspetta. Non perché chi aspetta sarà escluso: "
    "perché chi agisce ora acquisisce l'esperienza pratica, commette gli errori utili, "
    "affina i processi e costruisce la fiducia interna che trasforma un esperimento in "
    "un vantaggio competitivo duraturo."
)

add_normal(doc,
    "Come per tutte le tecnologie trasformative, il vero rischio non è adottare gli agenti AI "
    "troppo presto. Il vero rischio è non capire che la finestra in cui il vantaggio del primo "
    "mover è massimo si sta chiudendo. Chi era già sull'AI generativa nel 2023 ha oggi "
    "due anni di vantaggio su chi ha aspettato il 2025. Chi inizia con gli agenti nel "
    "2026 avrà lo stesso vantaggio su chi aspetterà il 2028."
)

add_normal(doc,
    "Il nostro suggerimento concreto: identificate questa settimana un processo ripetitivo "
    "nel vostro studio o nella vostra azienda, chiedetevi se un agente AI potrebbe gestirlo "
    "in autonomia, e contattate un consulente specializzato per una valutazione di fattibilità. "
    "Non occorre un piano quinquennale: occorre un primo passo consapevole."
)

# ─── SPAZIO FINALE ────────────────────────────────────────────────────────────

doc.add_paragraph("", style="Normal")

note = doc.add_paragraph(style="Normal")
note_run = note.add_run(
    "Alessandro Mattavelli è consulente specializzato in intelligenza artificiale applicata "
    "ai processi aziendali e ai professionisti. È autore della serie «Excellere in Studio» "
    "e relatore in corsi di formazione su AI e innovazione digitale per PMI e studi professionali."
)
note_run.italic = True
note_run.font.size = Pt(9)

# ─── SALVATAGGIO ─────────────────────────────────────────────────────────────

doc.save(OUTPUT_PATH)
print(f"Articolo salvato in: {OUTPUT_PATH}")
