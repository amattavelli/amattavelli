"""
Articolo 2 — Claude Finance Agents di Anthropic
Ratio/articoli/2026-05-08_claude-agenti-finanza-chiusura-libri.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-05-08_claude-agenti-finanza-chiusura-libri.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Maggio 2026  |  Strumenti e Modelli AI")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Dieci agenti per chiudere i libri:\ncosa fa Claude quando lavora in finanza")
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Anthropic ha rilasciato dieci agenti AI specializzati per KYC, chiusura mensile "
    "e valutazione d'azienda. Per studi e uffici contabili, la domanda non è se usarli "
    "ma chi nel team capisce abbastanza da sapere quando smettere di fidarsi."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  8 maggio 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Se gestisci la contabilità di un'azienda o lavori in uno studio che affianca PMI "
    "nella redazione dei bilanci, c'è una fase del mese che conosci bene: la chiusura. "
    "Riconciliazioni di conti, verifica delle competenze, quadratura tra mastro e situazione "
    "bancaria, predisposizione delle scritture di assestamento. Sono ore di lavoro ripetitivo, "
    "ad alta concentrazione, con un margine di errore che non si può concedere. "
    "Il 5 maggio 2026 Anthropic ha annunciato dieci agenti precostituiti per il settore "
    "finanziario integrati in Claude, ognuno pensato per una fase specifica di questo "
    "tipo di lavoro."
)

para(doc,
    "Il lancio è avvenuto in contemporanea con altri due annunci: l'integrazione completa "
    "con Microsoft 365 e una partnership con Moody's per i dati di mercato. Tre mosse che "
    "disegnano una direzione precisa: Claude non vuole essere il chatbot che risponde alle "
    "domande, ma lo strumento che porta a termine compiti finanziari strutturati, dentro "
    "l'ambiente di lavoro già in uso."
)

heading(doc, "La differenza tra un chatbot e un agente")

para(doc,
    "Gli agenti non sono chatbot. Un agente AI è un sistema che, a partire da un obiettivo, "
    "pianifica una sequenza di azioni, usa strumenti come documenti, database e API, verifica "
    "i risultati e produce un output strutturato. Non risponde a domande: porta a termine un "
    "compito. Questa distinzione, che sembra tecnica, ha conseguenze molto pratiche per chi "
    "li usa in un contesto professionale."
)

para(doc,
    "L'agente \"month-end closer\" di Claude esegue la checklist di chiusura mensile, "
    "prepara le scritture di assestamento e produce il report di chiusura. Il \"general "
    "ledger reconciler\" riconcilia i conti del mastro, calcola il valore netto di portafoglio "
    "e lo confronta con i libri contabili. Il \"statement auditor\" verifica la coerenza e "
    "la completezza dei prospetti finanziari in vista della revisione. Tre agenti, tre fasi "
    "della vita di uno studio o di un ufficio contabile. Gli altri sette coprono il KYC e "
    "la compliance antiriciclaggio, la valutazione aziendale con comparables di mercato, "
    "la costruzione di modelli finanziari di proiezione, la preparazione di pitchbook, "
    "l'analisi di mercato e la sintesi di documenti per riunioni e negoziazioni."
)

heading(doc, "Cosa cambia per uno studio commercialista")

para(doc,
    "Per uno studio che affianca PMI italiane, l'impatto più immediato riguarda la chiusura "
    "mensile e la riconciliazione. Non perché questi agenti sostituiscano il lavoro di un "
    "collaboratore, ma perché possono eseguire la prima passata in modo più veloce e "
    "documentato di quanto faccia un operatore sotto pressione di fine mese. Il tempo "
    "risparmiato nella fase meccanica si può investire nella fase critica: capire perché "
    "il conto non quadra, non cercare dove non quadra."
)

para(doc,
    "L'integrazione con Microsoft 365 abbatte una barriera concreta che aveva frenato l'adozione "
    "di strumenti AI in molti studi. Claude può ora lavorare direttamente sui file Word, Excel "
    "e Outlook senza uscire dall'ambiente di lavoro. Per uno studio che usa già Microsoft 365 "
    "come infrastruttura, questo significa che non serve caricare file in piattaforme esterne, "
    "non serve copiare e incollare tra finestre diverse. Il flusso di lavoro rimane nell'ambiente "
    "conosciuto, con un agente che esegue la parte ripetitiva al suo interno."
)

para(doc,
    "La partnership con Moody's aggiunge un livello di profondità rilevante per chi si occupa "
    "di valutazioni aziendali o di analisi del merito creditizio. L'agente \"valuation reviewer\" "
    "può confrontare una valutazione interna con multipli di mercato aggiornati in tempo reale, "
    "segnalare scostamenti significativi e documentare il ragionamento alla base del giudizio. "
    "Per uno studio che prepara perizie di stima o supporta operazioni di M&A, questa "
    "integrazione è concreta."
)

heading(doc, "Il nodo della governance")

para(doc,
    "Il punto che nessun lancio commerciale risolve è quello della governance. Chi verifica "
    "che l'agente abbia eseguito correttamente la riconciliazione? Come si traccia il processo "
    "per dimostrare, in sede di revisione o di verifica fiscale, che la chiusura è stata "
    "fatta correttamente? I sistemi di audit trail degli agenti di Anthropic registrano le "
    "azioni eseguite, ma la responsabilità di conservare quella documentazione e di interpretarla "
    "resta nell'organizzazione che usa lo strumento. La Legge 132/2025 è chiara su questo: "
    "la firma sul documento finale resta in capo al professionista."
)

para(doc,
    "Per le aziende italiane che faticano a trovare collaboratori con competenze contabili "
    "sufficienti a gestire la chiusura mensile in modo autonomo, avere un agente che esegue "
    "la prima passata non è un'opzione teorica: è una risposta concreta a un problema di "
    "capacità. Ma funziona solo se c'è qualcuno in grado di interpretare i risultati dell'agente, "
    "di capire quando l'output è sospetto e di prendere la decisione finale. Dieci agenti "
    "per la finanza non cambiano la struttura del lavoro: cambiano la proporzione di tempo "
    "che si passa a fare cose che si potrebbero non fare a mano."
)

heading(doc, "Chi deve capire come funziona")

para(doc,
    "La domanda da farsi, prima di configurare uno qualsiasi di questi agenti, non è \"questo "
    "strumento fa quello che mi serve?\" ma \"chi nel mio team capisce abbastanza la logica "
    "dell'agente da sapere quando smettere di fidarsi?\" Un agente che esegue la chiusura "
    "mensile seguendo una checklist predefinita funziona bene finché la situazione è "
    "standard. Quando c'è un'operazione fuori schema, una scrittura di rettifica non "
    "prevista, un conto con movimentazione anomala, l'agente segue la regola che gli è "
    "stata data. Il professionista deve sapere che in quel caso la regola non basta."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "Anthropic — \"Agents for financial services\", anthropic.com/news/finance-agents (5 maggio 2026)",
    "The Register — \"Anthropic unleashes finance agents for Claude\" (5 maggio 2026)",
    "Fortune — \"Anthropic deepens push into Wall Street with new AI agents\" (5 maggio 2026)",
    "Winbuzzer — \"Anthropic Expands Claude With 10 Finance Workflow Agents\" (6 maggio 2026)",
    "Legge 132/2025 — Disciplina organica sull'intelligenza artificiale in Italia",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
