"""
Quattro articoli Ratio -- 24 luglio 2026

1. 2026-07-24_ai-act-trasparenza-chatbot-linee-guida-20-luglio.docx
2. 2026-07-24_ai-agentica-dalla-demo-alla-produzione.docx
3. 2026-07-24_pmi-italiane-ai-paradosso-maturita-aimix.docx
4. 2026-07-24_sonnet5-opus48-scegliere-bene-lavoro-professionale.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# AI Act art. 50: le linee guida del 20 luglio
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Normativa AI")
titolo(
    doc,
    "Il chatbot del vostro sito sa presentarsi?\nDal 2 agosto \xe8 obbligatorio.",
    "Il 20 luglio 2026 la Commissione europea ha pubblicato le linee guida finali "
    "sull’articolo 50 dell’AI Act: i sistemi conversazionali devono dichiarare la "
    "propria natura artificiale prima di ogni interazione. Mancano nove giorni. "
    "L’adempimento tecnico \xe8 tra i meno complessi dell’intero Regolamento, ma molte "
    "aziende non sanno ancora di rientrare nell’obbligo.",
    "A cura della Redazione Ratio  •  24 luglio 2026"
)

para(doc,
    "Il chatbot sul sito di uno studio commercialista risponde ai clienti in tempo "
    "reale, suggerisce le aree di specializzazione, raccoglie i dati per la prima "
    "valutazione del caso. Per chi ci interagisce, la distinzione tra un assistente "
    "umano e uno automatizzato pu\xf2 essere sfumata: il tono \xe8 calibrato, le risposte "
    "pertinenti, il ritmo simile a quello di un operatore in carne e ossa. Dal "
    "2 agosto 2026 questa ambiguit\xe0, che in molti casi \xe8 stata lasciata volutamente "
    "irrisolta, diventa una violazione delle norme europee."
)

para(doc,
    "Il 20 luglio 2026 la Commissione europea ha pubblicato le linee guida finali "
    "sull’articolo 50 del Regolamento UE 2024/1689, il capitolo dell’AI Act che "
    "disciplina gli obblighi di trasparenza per i sistemi di intelligenza artificiale "
    "conversazionali e per i contenuti sintetici. Il documento chiarisce cosa si "
    "intende per interazione avviata con un utente e come deve essere formulata la "
    "comunicazione che segnala la natura artificiale del sistema. Quattro giorni dopo "
    "la pubblicazione, le imprese italiane si trovano con nove giorni per essere conformi."
)

para(doc,
    "La responsabilit\xe0 dell’adempimento ricade sul deployer, cio\xe8 sull’azienda o il "
    "professionista che ha integrato il sistema AI nel proprio servizio, non sul "
    "fornitore della tecnologia. Il produttore del chatbot ha obblighi di progettazione "
    "(garantire che la disclosure sia tecnicamente possibile), ma chi lo ha attivato "
    "sul proprio sito web risponde dell’adempimento in sede ispettiva. Per uno studio "
    "legale che usa un assistente conversazionale sulla pagina contatti, o per un "
    "commercialista che ha implementato un bot per le prime informazioni fiscali, la "
    "questione \xe8 di immediata applicabilit\xe0."
)

heading(doc, "Tre sistemi, tre obblighi distinti")

para(doc,
    "Le linee guida classificano in modo preciso i sistemi soggetti all’obbligo di "
    "trasparenza. I chatbot e gli assistenti conversazionali devono comunicare "
    "all’utente, prima o all’avvio dell’interazione, di essere sistemi AI, a meno che "
    "questa natura non sia gi\xe0 evidente dal contesto. I sistemi che generano contenuti "
    "audiovisivi, audio o testuali capaci di somigliare a produzioni umane reali "
    "devono applicare una marcatura tecnica che segnali la provenienza artificiale, "
    "secondo le specifiche tecniche che gli organismi europei stanno finalizzando. I "
    "sistemi biometrici che identificano o categorizzano persone in spazi fisici "
    "accessibili al pubblico devono rendere visibile la propria presenza e funzione."
)

heading(doc, "La disclosure non sta nel footer")

para(doc,
    "Le linee guida chiariscono un punto che aveva generato interpretazioni divergenti: "
    "la disclosure deve avvenire prima che l’utente abbia fornito qualsiasi dato o "
    "avviato la conversazione, e deve essere visibile, non sepolta nei termini di "
    "servizio o in una nota a pi\xe8 di pagina. Un avviso nelle FAQ non \xe8 sufficiente. "
    "Un riferimento nella privacy policy non soddisfa il requisito. La specifica "
    "formulazione richiama l’idea che l’utente debba poter scegliere consapevolmente "
    "se interagire con un sistema AI prima di avviare l’interazione stessa. \xc8, in "
    "fondo, la stessa questione di identit\xe0 dichiarata che Pirandello aveva messo al "
    "centro di tutta la sua produzione: la realt\xe0 di una cosa non dipende da ci\xf2 che "
    "\xe8, ma da come si presenta. Cos\xec \xe8, se vi dico chi sono."
)

heading(doc, "Il nodo pratico: molte aziende non sanno di rientrare nell’obbligo")

para(doc,
    "Le sanzioni per le violazioni dell’articolo 50 arrivano fino a 15 milioni di "
    "euro o al 3% del fatturato mondiale annuo, se superiore. Per una PMI italiana "
    "con un fatturato di due milioni di euro, la soglia percentuale equivale a 60.000 "
    "euro per un chatbot che non si \xe8 presentato. Il paradosso \xe8 che l’adempimento "
    "tecnico \xe8 tra i meno complessi dell’intero AI Act: richiede una configurazione del "
    "sistema e un avviso visibile prima dell’apertura della chat, non una "
    "certificazione n\xe9 un’analisi del rischio elaborata. La difficolt\xe0 vera \xe8 che molte "
    "aziende hanno implementato assistenti AI tramite plugin o integrazioni di terze "
    "parti senza tracciare esattamente cosa gira sotto, e devono prima verificare "
    "quali dei loro strumenti digitali rientrano nella definizione di sistema "
    "conversazionale AI soggetto all’articolo 50."
)

para(doc,
    "Nove giorni. Per chi usa un chatbot sul proprio sito o un assistente AI per "
    "interagire con clienti e fornitori, la domanda da porsi non \xe8 se il sistema "
    "funzioni bene. \xc8 se sa presentarsi."
)

riferimenti(doc, [
    "Commissione europea, linee guida finali articolo 50 AI Act, 20 luglio 2026, EUR-Lex",
    "Regolamento UE 2024/1689 (AI Act), articolo 50 -- EUR-Lex",
    "AgendaDigitale.eu -- 'Obblighi di trasparenza AI Act: cosa devono fare le aziende dal 2 agosto 2026'",
    "LaborProject.it -- 'AI Act, dal 2 agosto 2026 l’intelligenza artificiale dovr\xe0 presentarsi'",
    "Cybersecurity360.it -- 'AI Act, dal 2 agosto cambiano gli obblighi di trasparenza dei sistemi di AI'",
    "Magellanopa.it -- 'AI Act: obblighi di trasparenza dal 2 agosto 2026'",
    "AI4Business.it -- 'AI Act trasparenza: le regole UE per chatbot e deep fake'",
])
doc.save(BASE + "2026-07-24_ai-act-trasparenza-chatbot-linee-guida-20-luglio.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# AI agentica: dal proof of concept alla produzione
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Strumenti e Processi AI")
titolo(
    doc,
    "L’AI agentica funziona nei demo.\nPortarla in produzione \xe8 un altro lavoro.",
    "Il 68% delle aziende italiane ha investito in strumenti AI agentica, ma la met\xe0 "
    "non \xe8 riuscita a trasformare il proof of concept in un processo operativo stabile. "
    "Il problema non \xe8 la tecnologia: \xe8 che un agente costruito su dati puliti incontra "
    "i dati reali dell’azienda e non \xe8 pronto per quello che trova.",
    "A cura della Redazione Ratio  •  24 luglio 2026"
)

para(doc,
    "DigitalAdrenalin ha pubblicato a luglio 2026 un’analisi sul divario tra "
    "sperimentazioni AI e deployment in ambienti produttivi nelle aziende italiane. "
    "Il dato pi\xf9 citato \xe8 che il 68% delle aziende ha gi\xe0 investito in strumenti AI "
    "per l’automazione dei processi, ma la met\xe0 di queste non \xe8 riuscita a trasformare "
    "il proof of concept in un flusso di lavoro operativo che giri da solo, senza "
    "supervisione manuale costante. Il gap non ha niente a che fare con la qualit\xe0 "
    "degli strumenti disponibili: ha a che fare con quello che gli ingegneri "
    "chiamano il ‘last mile problem’, il tratto finale che separa il laboratorio "
    "dal campo."
)

para(doc,
    "Icaro funzionava perfettamente finch\xe9 la cera reggeva. Il punto non era la "
    "qualit\xe0 delle piume o del progetto di Dedalo: era che il sistema era stato "
    "costruito e testato a terra, in condizioni controllate, senza mai stimare cosa "
    "sarebbe successo vicino al sole. Gli agenti AI aziendali hanno un problema "
    "analogo: vengono sviluppati e validati su set di dati puliti, strutturati, "
    "rappresentativi di uno scenario ideale, e poi vengono rilasciati in contesti "
    "dove i dati sono incompleti, le eccezioni sono frequenti e i processi che "
    "dovrebbero supportare sono stati costruiti in decenni senza pensare all’integrazione "
    "con un sistema automatizzato."
)

heading(doc, "Il problema non \xe8 l’agente, \xe8 il terreno in cui lo si mette")

para(doc,
    "Un agente AI che deve leggere le fatture in arrivo, classificarle per categoria "
    "di spesa e registrarle nel gestionale funziona in modo impeccabile quando le "
    "fatture hanno la struttura prevista nel test. Nel momento in cui si incontra una "
    "fattura con un formato insolito, un fornitore non censito, un codice prodotto "
    "ambiguo o un importo che non corrisponde all’ordine di acquisto, il comportamento "
    "dell’agente dipende interamente da come \xe8 stato progettato per gestire "
    "l’incertezza. Se la gestione dell’errore non \xe8 stata progettata con la stessa "
    "attenzione del caso standard, l’agente si blocca, produce un output errato "
    "senza segnalarlo, o esegue in modo parziale lasciando il file in uno stato "
    "inconsistente. Nessuno di questi tre risultati \xe8 accettabile in produzione."
)

para(doc,
    "Le aziende italiane che riescono a portare l’AI agentica in produzione in modo "
    "stabile condividono alcune caratteristiche nel modo in cui hanno affrontato il "
    "progetto. Prima di scegliere lo strumento, hanno mappato il processo da "
    "automatizzare in dettaglio, incluse le eccezioni: non solo il flusso standard, "
    "ma cosa succede quando arriva un documento atipico, un dato mancante, un errore "
    "di sistema a monte. Poi hanno definito esplicitamente cosa deve fare l’agente "
    "quando si trova in uno scenario non previsto: segnalare all’operatore, mettere in "
    "coda, rifiutare l’elaborazione o chiedere conferma. Questa fase di progettazione "
    "delle eccezioni richiede pi\xf9 tempo del training dell’agente vero e proprio, ed \xe8 "
    "quella che la maggior parte dei progetti falliti ha saltato."
)

heading(doc, "Partire dal processo circoscritto, non dalla trasformazione totale")

para(doc,
    "Per le PMI italiane che si avvicinano all’AI agentica, il rischio pi\xf9 comune \xe8 "
    "partire con obiettivi troppo ampi: automatizzare l’intero processo di acquisto, "
    "sostituire la gestione manuale di un reparto, costruire un assistente che risponda "
    "a qualsiasi domanda del cliente. Questi obiettivi sono teoricamente raggiungibili, "
    "ma richiedono infrastrutture di dati, qualit\xe0 di integrazione e capacit\xe0 di "
    "gestione degli errori che la maggior parte delle PMI non ha ancora costruito. "
    "Le implementazioni che mostrano i risultati pi\xf9 solidi riguardano processi "
    "circoscritti, con input ben definiti, output verificabili e un perimetro chiaro "
    "di cosa l’agente deve fare e cosa no. Un agente che classifica correttamente il "
    "94% dei documenti in ingresso su un processo specifico \xe8 un punto di partenza "
    "reale. Un agente che dovrebbe fare cinquantaquattro cose e ne fa la met\xe0 male \xe8 "
    "un progetto che verr\xe0 disabilitato nel giro di un mese."
)

heading(doc, "La governance dell’agente: chi controlla e cosa")

para(doc,
    "Un aspetto che emerge dall’analisi \xe8 la questione della governance degli agenti "
    "in produzione: chi \xe8 responsabile di monitorare il comportamento dell’agente nel "
    "tempo, di aggiornarlo quando il processo sottostante cambia, di intervenire "
    "quando produce output anomali? Nelle aziende che hanno avuto problemi, questa "
    "responsabilit\xe0 era rimasta implicita: il team IT aveva costruito l’agente, il "
    "team operativo lo usava, ma nessuno era formalmente incaricato di supervisionarne "
    "il funzionamento nel tempo. Gli agenti AI in produzione richiedono un owner, "
    "qualcuno che conosca sia il processo operativo sia il funzionamento dello "
    "strumento e sia in grado di valutare quando l’uno o l’altro cambia in modo da "
    "richiedere un aggiornamento del sistema."
)

para(doc,
    "Un agente AI non si schiaccia con il pollice come un’app. Richiede un progetto: "
    "non per il modello, ma per il processo. Le aziende che lo capiscono nella fase "
    "di design evitano di scoprirlo in produzione, quando il costo della scoperta \xe8 "
    "gi\xe0 diventato un problema operativo."
)

riferimenti(doc, [
    "DigitalAdrenalin.it -- 'AI agentica in azienda, il nodo \xe8 portarla davvero in produzione' (luglio 2026)",
    "AffarItaliani.it -- 'Agenti IA, violazione dei dati e incidenti informatici: ecco quanto rischiano le aziende italiane'",
    "EverestInnovation.it -- 'Agenti AI Autonomi: il 2026 \xe8 l’Anno della Rivoluzione Intelligente'",
    "BestTechPartner.ai -- 'AI agenti: Cosa significa davvero per le imprese (Guida Strategica 2026)'",
    "CorrierNazionale.it -- 'La Nuova Era dell’Efficienza: Come l’AI Agentica sta Trasformando il Lavoro in Italia nel 2026'",
])
doc.save(BASE + "2026-07-24_ai-agentica-dalla-demo-alla-produzione.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# PMI italiane e il paradosso AI: +30% possibile, 8,5% raggiunto
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "PMI e Adozione AI")
titolo(
    doc,
    "Potenziale al 30%, maturit\xe0 all’8,5%:\nle PMI italiane e il paradosso dell’AI.",
    "Il rapporto SME-AIMIX 2026 misura il divario tra ci\xf2 che l’AI potrebbe dare "
    "alle piccole e medie imprese italiane e quello che effettivamente restituisce. "
    "La distanza \xe8 reale, documentata, e non si chiude comprando lo strumento "
    "pi\xf9 avanzato disponibile.",
    "A cura della Redazione Ratio  •  24 luglio 2026"
)

para(doc,
    "Il rapporto SME-AIMIX 2026, presentato a inizio luglio, misura il divario tra "
    "ci\xf2 che l’intelligenza artificiale potrebbe dare alle piccole e medie imprese "
    "italiane e ci\xf2 che effettivamente restituisce. Il potenziale di incremento di "
    "produttivit\xe0 per le PMI che integrano l’AI in modo strutturato arriva fino al 30%. "
    "Il livello di maturit\xe0 digitale medio delle PMI italiane rispetto all’AI si "
    "attesta all’8,5%. Questi due numeri non vanno letti come una contraddizione: "
    "vanno letti come una mappa."
)

para(doc,
    "Nel 1987 l’economista Robert Solow scrisse una frase diventata celebre: 'Si vede "
    "l’era del computer dappertutto tranne che nelle statistiche sulla produttivit\xe0.' "
    "Il paradosso di Solow descriveva la distanza tra il clamore dei nuovi strumenti "
    "tecnologici e i dati reali sulle aziende che li avevano adottati: le macchine "
    "erano arrivate, la produttivit\xe0 misurata no. Il rapporto AIMIX fotografa qualcosa "
    "di simile applicato all’AI nelle PMI italiane trentanove anni dopo: gli strumenti "
    "ci sono, l’adozione dichiarata \xe8 in crescita, i risultati misurabili sulla "
    "produttivit\xe0 aggregata restano lontani dal potenziale teorico."
)

heading(doc, "La crescita c’\xe8, ma 'usare l’AI' non \xe8 'integrare l’AI'")

para(doc,
    "Tra il 2023 e il 2025 la quota di imprese italiane che dichiara di utilizzare "
    "almeno una tecnologia di intelligenza artificiale \xe8 pi\xf9 che triplicata, passando "
    "dal 6% al 16,4%. Il dato \xe8 reale, ma va letto con attenzione: molte di queste "
    "aziende usano ChatGPT per redigere comunicazioni o sintetizzare documenti, senza "
    "aver integrato l’AI nei processi operativi. La differenza tra 'usare uno strumento "
    "AI' e 'avere un processo AI-integrato' \xe8 la stessa che c’\xe8 tra comprare un "
    "martello e ristrutturare una casa: il martello \xe8 necessario, ma non \xe8 la "
    "ristrutturazione. Il gap tra il 16,4% di adozione dichiarata e l’8,5% di "
    "maturit\xe0 misurata racconta esattamente questa distanza."
)

heading(doc, "Il vero ostacolo non \xe8 il costo dello strumento")

para(doc,
    "Quando il rapporto AIMIX chiede alle PMI che hanno valutato l’AI senza adottarla "
    "quale sia l’ostacolo principale, la risposta ricorrente non \xe8 il prezzo degli "
    "abbonamenti, gi\xe0 accessibile per la maggior parte delle soluzioni di base, e "
    "nemmeno la complessit\xe0 tecnica percepita. \xc8 la mancanza di competenze interne "
    "per valutare quali processi beneficerebbero dell’AI, come configurare gli "
    "strumenti in modo utile e come misurare i risultati. Questa risposta rivela un "
    "problema strutturale: molte PMI cercano uno strumento quando avrebbero bisogno "
    "prima di una mappa. Senza sapere dove l’AI pu\xf2 fare la differenza nel proprio "
    "processo specifico, il rischio \xe8 comprare lo strumento pi\xf9 pubblicizzato e "
    "applicarlo dove capita, ottenendo risultati modesti che alimentano lo scetticismo "
    "sull’intera categoria."
)

heading(doc, "Dove l’AI sta davvero entrando nelle PMI")

para(doc,
    "L’AI viene adottata nelle aree dove generare efficienza \xe8 pi\xf9 immediato da "
    "misurare: il 33,1% dei casi riguarda marketing e vendite (redazione di contenuti, "
    "automazione di campagne, risposta ai lead), il 25,7% l’amministrazione (gestione "
    "documentale, sintesi, comunicazioni interne), il 20% la ricerca e sviluppo. "
    "Le aree dove l’AI potrebbe avere l’impatto pi\xf9 profondo, come la gestione della "
    "supply chain, la pianificazione della produzione o la manutenzione predittiva, "
    "restano le meno esplorate perch\xe9 richiedono integrazione con sistemi operativi "
    "legacy e qualit\xe0 dei dati che poche PMI hanno ancora costruito. Questo schema "
    "di adozione segue la logica del percorso di minima resistenza, non quella "
    "del massimo valore."
)

heading(doc, "La distanza si chiude con il metodo, non con il modello")

para(doc,
    "Le PMI che stanno effettivamente chiudendo il gap tra potenziale e risultati "
    "non si distinguono per aver scelto il modello AI pi\xf9 avanzato disponibile. Si "
    "distinguono per aver fatto tre cose nell’ordine giusto: hanno mappato i "
    "processi prima di scegliere lo strumento, hanno formato le persone prima di "
    "automatizzare i flussi, e hanno misurato i risultati su un perimetro circoscritto "
    "prima di scalare. Questo approccio richiede qualche settimana in pi\xf9 rispetto "
    "all’acquisto diretto di un abbonamento enterprise, ma produce implementazioni "
    "che reggono nel tempo invece di essere abbandonate dopo tre mesi perch\xe9 il "
    "risultato non corrispondeva alle aspettative."
)

para(doc,
    "Il 30% non \xe8 una promessa di marketing. \xc8 il delta che separa una PMI che ha "
    "fatto le cose nell’ordine giusto da una che ha fatto la fila davanti allo "
    "strumento sbagliato per il processo sbagliato."
)

riferimenti(doc, [
    "SME-AIMIX 2026 -- 'Fino al +30% di produttivit\xe0 con l’AI, ma le PMI italiane non sono ancora pronte' (luglio 2026)",
    "MediaKey.it -- 'SME-AIMIX 2026: le PMI italiane e l’AI' (luglio 2026)",
    "AdnKronos / laMilano -- 'AI: nelle PMI italiane il potenziale c’\xe8 ma la maturit\xe0 resta bassa' (giugno 2026)",
    "WTraining.it -- 'Rapporto Istat 2026: nelle PMI italiane le competenze sono il vero freno all’AI'",
    "AI4Business.it -- 'Intelligenza artificiale in Italia, crescita e ritardi nel 2026'",
    "SolutionBPartner.it -- 'Intelligenza artificiale per PMI nel 2026'",
    "BestTechPartner.ai -- 'Intelligenza artificiale in azienda: guida pratica 2026 per le PMI italiane'",
])
doc.save(BASE + "2026-07-24_pmi-italiane-ai-paradosso-maturita-aimix.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Sonnet 5 e Opus 4.8: due modelli, due lavori diversi
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Strumenti e Modelli AI")
titolo(
    doc,
    "La lineup Claude \xe8 cambiata:\nSonnet 5 per la routine, Opus 4.8 per quando conta.",
    "Da luglio 2026 Claude Sonnet 5 sostituisce Sonnet 4.6 come modello di riferimento "
    "predefinito. Per il lavoro professionale quotidiano copre quasi tutto. Opus 4.8 "
    "rimane la scelta giusta per un sottoinsieme preciso di task. Capire la differenza "
    "vale pi\xf9 che avere l’abbonamento pi\xf9 costoso.",
    "A cura della Redazione Ratio  •  24 luglio 2026"
)

para(doc,
    "La scena \xe8 comune: un professionista apre Claude, scrive un prompt per analizzare "
    "un contratto complesso, e usa il modello di default. Da luglio 2026 quel default "
    "\xe8 cambiato: Sonnet 5 ha sostituito Sonnet 4.6 come modello predefinito per i "
    "piani Claude Pro e Team. Il risultato dell’analisi \xe8 buono. Probabilmente avrebbe "
    "potuto essere migliore con Opus 4.8, su quel tipo di task specifico, ma per "
    "capirlo bisognerebbe provare entrambi sullo stesso documento. Quasi nessuno lo "
    "fa. Questa inerzia nel passare da un modello all’altro non \xe8 pigrizia: \xe8 che "
    "nessuno spiega chiaramente quali task cambiano davvero con il cambio di modello "
    "e quali no."
)

para(doc,
    "Benjamin Franklin invent\xf2 le lenti bifocali perch\xe9 si stanc\xf2 di portare due "
    "paia di occhiali, uno per leggere e uno per guardare lontano. Il punto non era "
    "avere la lente pi\xf9 sofisticata: era usare la lente giusta per la distanza giusta, "
    "nello stesso oggetto. La lineup Claude aggiornata a luglio 2026 funziona con una "
    "logica simile: quattro modelli, quattro fasce di uso, quattro profili di costo. "
    "Non tutti servono al professionista che usa Claude come strumento di lavoro "
    "diretto, e usare il modello sbagliato per il task sbagliato non produce risultati "
    "migliori, solo consumi pi\xf9 alti."
)

heading(doc, "La lineup di luglio 2026: cosa \xe8 cambiato")

para(doc,
    "Da luglio 2026 la gerarchia Claude \xe8 composta da quattro modelli con posizionamenti "
    "distinti. Fable 5 occupa la cima come modello flagship per i flussi agentici "
    "e il ragionamento complesso di lunga durata. Opus 4.8 \xe8 il modello per il "
    "ragionamento approfondito nel lavoro professionale diretto. Sonnet 5 \xe8 il nuovo "
    "modello di riferimento predefinito, con capacit\xe0 significativamente superiori "
    "alla versione precedente (Sonnet 4.6) a un costo API inferiore a Opus 4.8. "
    "Haiku 4.5 copre i task rapidi e ad alto volume. Per chi usa i piani Claude "
    "Pro o Team, il modello predefinito \xe8 cambiato automaticamente con "
    "l’aggiornamento di luglio: aprendo una nuova conversazione, si lavora con "
    "Sonnet 5 a meno che non si selezioni esplicitamente un’alternativa."
)

heading(doc, "Cosa fa bene Sonnet 5 nel lavoro professionale")

para(doc,
    "Sonnet 5 copre in modo solido l’analisi di documenti di media lunghezza, la "
    "redazione di comunicazioni strutturate, la revisione critica di testi, la "
    "sintesi di materiali e le risposte a quesiti che richiedono conoscenza normativa "
    "consolidata. Per la maggior parte dei compiti che un commercialista o un "
    "consulente svolge nella giornata tipo, Sonnet 5 produce risultati che non si "
    "distinguono in modo apprezzabile da quelli di Opus 4.8. I benchmark pubblicati "
    "da Anthropic mostrano che su knowledge work di livello professionale standard "
    "il gap tra i due modelli si \xe8 ridotto sensibilmente rispetto alle versioni "
    "precedenti, e il differenziale di latenza (Sonnet 5 risponde pi\xf9 velocemente) "
    "su task ripetuti nel corso della giornata diventa un vantaggio concreto."
)

heading(doc, "Quando Opus 4.8 resta la scelta giusta")

para(doc,
    "Il differenziale emerge in modo visibile su alcune categorie specifiche di task: "
    "il ragionamento su problemi che richiedono l’esplorazione di pi\xf9 ipotesi "
    "simultaneamente prima di convergere su una conclusione, l’analisi di documenti "
    "molto lunghi con istruzioni che richiedono inferenze su tutto il testo in modo "
    "coerente, e i task in cui il modello deve fare scelte tra alternative con "
    "implicazioni normative o giuridiche rilevanti che coinvolgono norme in interazione. "
    "Per un avvocato che chiede di analizzare un contratto lungo e identificare le "
    "clausole potenzialmente problematiche considerando pi\xf9 scenari processuali, "
    "Opus 4.8 mostra ancora un vantaggio nella capacit\xe0 di tenere insieme fili "
    "argomentativi complessi. La regola pratica: se il task richiede che il modello "
    "ragioni su pi\xf9 variabili in parallelo prima di convergere su una risposta, e "
    "il tempo aggiuntivo di elaborazione \xe8 un costo accettabile, Opus 4.8 \xe8 ancora "
    "la scelta migliore."
)

heading(doc, "Fable 5: per chi costruisce, non per chi usa direttamente")

para(doc,
    "Fable 5 \xe8 tornato disponibile in tutto il mondo dal 1\xb0 luglio 2026, dopo "
    "l’interruzione causata da un ordine di controllo sulle esportazioni statunitense. "
    "Le sue capacit\xe0 agentiche, la possibilit\xe0 di scomporre task complessi in "
    "sottotask, di delegare a strumenti e sotto-agenti e di gestire flussi di lavoro "
    "automatizzati di lunga durata, ne fanno il modello di riferimento per le aziende "
    "che costruiscono pipeline AI o automazioni articolate. Per il professionista che "
    "usa Claude direttamente nell’attivit\xe0 quotidiana, Fable 5 \xe8 meno rilevante: il "
    "suo valore emerge quando si lavora attraverso agenti, non quando si usa Claude "
    "come interlocutore diretto in una conversazione. Il prezzo sull’API, 10 dollari "
    "per milione di token di input e 50 in output, riflette un posizionamento pensato "
    "per chi costruisce infrastrutture AI, non per chi le usa."
)

para(doc,
    "Il modello migliore \xe8 quello che risolve il problema che si sta cercando di "
    "risolvere. Avere l’accesso a Fable 5 e usarlo per scrivere una email \xe8 come "
    "usare un tornio per appendere un quadro: funziona, forse, ma non era questo "
    "il punto."
)

riferimenti(doc, [
    "Anthropic -- 'Claude Sonnet 5 release notes' (luglio 2026) -- anthropic.com",
    "GiovanniliGuori.it -- 'Claude AI: Guida Completa 2026 per il Business'",
    "Cosimo.dev -- 'Claude Fable 5 vs Opus 4.8: quale scegliere'",
    "Vincos.it -- 'Fable 5: usare Claude al massimo' (6 luglio 2026)",
    "Mimir.bot -- 'Fable 5 di nuovo disponibile: capacit\xe0 e fine del blocco USA'",
    "Fenxi.fr -- 'GPT-5.5 vs Claude Opus 4.8 vs Gemini 3.1 Pro: 2026 Guide'",
    "FelloAI.com -- 'Best AI Models in July 2026: ChatGPT, Claude, Gemini & Grok'",
])
doc.save(BASE + "2026-07-24_sonnet5-opus48-scegliere-bene-lavoro-professionale.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
