#!/usr/bin/env python3
"""
Articolo 2: 2 agosto 2026: le scadenze dell'AI Act che nessuna impresa può ignorare
File: Ratio/articoli/2026-04-17_ai-act-scadenze-agosto-2026-imprese.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/home/user/amattavelli/Ratio/articoli/2026-04-17_ai-act-scadenze-agosto-2026-imprese.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(3)
sec.right_margin = Cm(3)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot)
    pPr.append(pBdr)


def para(doc, text, size=11, italic=False, color=None, sa=8, sb=0,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


# --- TESTATA ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.all_caps = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Normativa e Compliance")
r.font.name = 'Calibri'
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# --- TITOLO ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "2 agosto 2026: le scadenze dell\u2019AI Act "
    "che nessuna impresa italiana pu\u00f2 ignorare"
)
r.font.name = 'Calibri'
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Mancano meno di quattro mesi alla data pi\u00f9 critica del Regolamento europeo "
    "sull\u2019intelligenza artificiale. Ecco cosa devono fare le imprese e i professionisti "
    "per arrivare pronti, e cosa rischiano se non lo fanno."
)
r.font.name = 'Calibri'
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  \u2022  17 aprile 2026")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# --- CORPO ---
para(doc,
    "Il Regolamento UE 2024/1689, meglio noto come AI Act, \u00e8 entrato in vigore il "
    "1\u00b0 agosto 2024. Da allora, la sua applicazione \u00e8 stata graduale, con scadenze "
    "scaglionate nel tempo per permettere alle imprese di adeguarsi. La data del "
    "2 agosto 2026 rappresenta il momento pi\u00f9 significativo di questo calendario: "
    "da quel giorno entrano in vigore gli obblighi per i sistemi AI classificati "
    "ad alto rischio, le norme sulla trasparenza e tutta la governance documentale "
    "prevista dal Regolamento. Per chi non si \u00e8 ancora mosso, il tempo a disposizione "
    "si misura in settimane, non in mesi."
)

para(doc,
    "La logica dell\u2019AI Act si basa su una classificazione del rischio. In cima alla "
    "piramide ci sono i sistemi vietati in assoluto: riconoscimento biometrico di massa, "
    "sistemi di scoring sociale, manipolazione comportamentale subliminale. Subito sotto "
    "si trovano i sistemi ad alto rischio, che non sono vietati ma devono soddisfare "
    "requisiti precisi: trasparenza delle decisioni, documentazione tecnica, registro "
    "delle attivit\u00e0, supervisione umana, test di conformit\u00e0. \u00c8 questa categoria "
    "che \u00e8 al centro delle scadenze di agosto 2026."
)

h2(doc, "Quali sistemi rientrano nell\u2019alto rischio")

para(doc,
    "L\u2019allegato III del Regolamento elenca le categorie di sistemi AI classificati "
    "ad alto rischio. Tra questi: sistemi usati per la selezione del personale e la "
    "valutazione dei lavoratori, sistemi che determinano l\u2019accesso a servizi essenziali "
    "come il credito bancario o le assicurazioni, software utilizzati in ambito sanitario "
    "per diagnosi o trattamento, strumenti impiegati nell\u2019istruzione per valutare "
    "studenti, sistemi di profilazione dei clienti usati nel marketing personalizzato "
    "con impatto su decisioni economiche rilevanti."
)

para(doc,
    "Molte aziende italiane utilizzano gi\u00e0 sistemi che rientrano in queste categorie "
    "senza esserne consapevoli. Un software HR che filtra i curricula con algoritmi "
    "di ranking, un sistema di scoring del credito integrato in un gestionale, "
    "una piattaforma di analisi predittiva usata per decidere i prezzi assicurativi: "
    "tutti questi strumenti possono ricadere nella definizione di sistema AI ad alto "
    "rischio, con i conseguenti obblighi di conformit\u00e0."
)

h2(doc, "Gli obblighi concreti da rispettare")

para(doc,
    "Per i sistemi ad alto rischio, il Regolamento prevede sei famiglie di obblighi. "
    "La prima riguarda la documentazione tecnica: ogni sistema deve avere una scheda "
    "che descrive come funziona, su quali dati \u00e8 stato addestrato e quali limitazioni "
    "presenta. La seconda riguarda il registro delle attivit\u00e0: il sistema deve tenere "
    "traccia delle operazioni compiute, in modo che sia possibile ricostruire come "
    "\u00e8 arrivato a una decisione. La terza riguarda la trasparenza verso gli utenti: "
    "chi interagisce con un sistema AI deve sapere di farlo."
)

para(doc,
    "La quarta famiglia di obblighi riguarda la supervisione umana: deve esistere "
    "un meccanismo che permetta a una persona di intervenire, modificare o bloccare "
    "le decisioni del sistema. La quinta riguarda la precisione e la robustezza: "
    "il sistema deve essere testato per garantire che funzioni come dichiarato, "
    "anche in condizioni anomale. La sesta, infine, riguarda la sicurezza informatica: "
    "il sistema deve essere protetto da attacchi che potrebbero alterarne "
    "il comportamento."
)

h2(doc, "Il trattamento di favore per le PMI e le sanzioni")

para(doc,
    "L\u2019AI Act prevede un regime differenziato per le piccole e medie imprese. "
    "Le sanzioni si applicano nella misura pi\u00f9 bassa tra la percentuale del fatturato "
    "globale e una cifra assoluta, il che riduce l\u2019esposizione delle PMI rispetto "
    "alle grandi corporation. Le imprese pi\u00f9 piccole hanno anche accesso prioritario "
    "alle cosiddette regulatory sandbox, ambienti controllati in cui possono testare "
    "sistemi AI con supervisione delle autorit\u00e0 senza incorrere subito nelle sanzioni "
    "ordinarie. Ma il trattamento di favore non significa esenzione: gli obblighi "
    "sostanziali si applicano a tutte le imprese, indipendentemente dalle dimensioni."
)

para(doc,
    "Le sanzioni per chi viola i divieti assoluti arrivano fino a 35 milioni di euro "
    "o al 7% del fatturato globale annuo. Per le violazioni degli obblighi sui sistemi "
    "ad alto rischio, il massimo \u00e8 15 milioni di euro o il 3% del fatturato. "
    "Per le informazioni false fornite alle autorit\u00e0, fino a 7,5 milioni o 1,5% "
    "del fatturato. Numeri che rendono chiaro perch\u00e9 aspettare agosto per iniziare "
    "ad adeguarsi sia un rischio che poche imprese possono permettersi di correre."
)

h2(doc, "Il doppio livello: AI Act europeo e Legge 132 italiana")

para(doc,
    "Le imprese italiane devono navigare un doppio livello di compliance. Oltre all\u2019AI Act "
    "europeo, \u00e8 in vigore dal 10 ottobre 2025 la Legge nazionale 132/2025, che integra "
    "il Regolamento europeo con disposizioni specifiche per il contesto italiano. "
    "La vigilanza \u00e8 affidata all\u2019Agenzia per l\u2019Italia Digitale (AgID) per le notifiche "
    "e gli accreditamenti, e all\u2019Agenzia per la Cybersicurezza Nazionale (ACN) per "
    "i controlli. In settori regolamentati come finanza e assicurazioni, sono coinvolti "
    "anche CONSOB e IVASS. Il Garante Privacy interviene nei casi che riguardano dati sensibili."
)

para(doc,
    "La raccomandazione operativa per le imprese italiane \u00e8 quella di iniziare con una "
    "mappatura dei sistemi AI in uso, classificarli per livello di rischio e identificare "
    "quelli che richiedono interventi prima di agosto. Per i sistemi ad alto rischio gi\u00e0 "
    "in produzione, \u00e8 necessario avviare subito la documentazione tecnica e verificare "
    "che i requisiti di supervisione umana siano soddisfatti. Per chi non ha ancora "
    "avviato questo processo, i quattro mesi che separano dall\u2019agosto 2026 sono "
    "sufficienti, ma solo se si inizia adesso."
)

sep(doc)

# --- FONTI ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Fonti e riferimenti")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Regolamento UE 2024/1689 (AI Act) \u2014 Gazzetta Ufficiale dell\u2019Unione Europea, "
    "12 luglio 2024",
    "Paradigma \u2014 AI Act 2026: obblighi per le imprese e integrazione con GDPR "
    "(14 aprile 2026)",
    "Brain Computing \u2014 AI Act 2026: obblighi e conformit\u00e0 per le aziende italiane",
    "Sopra Steria \u2014 L\u2019AI Act \u00e8 realt\u00e0: come cambia il panorama per le aziende italiane",
    "EuAIAct.pro \u2014 Obblighi EU AI Act 2026 \u2014 Deployer, Provider e Livelli di Rischio",
    "Tom\u2019s Hardware Business \u2014 Agenti AI e nuove regolamentazioni: il quadro normativo "
    "italiano ed europeo 2026 per le imprese",
    "Legge 132/2025 \u2014 Disposizioni in materia di intelligenza artificiale, "
    "pubblicata in GU del 9 ottobre 2025",
]
for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"\u2022 {s}")
    r.font.name = 'Calibri'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte")
r.font.name = 'Calibri'
r.font.size = Pt(8)
r.italic = True
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

doc.save(OUTPUT)
print(f"Salvato: {OUTPUT}")
