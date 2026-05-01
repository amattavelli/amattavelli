"""
Articolo 4 — Il 65% dei commercialisti usa già l'AI: ma i numeri raccontano metà della storia
Ratio/articoli/2026-05-01_automazione-contabile-oltre-i-numeri.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-05-01_automazione-contabile-oltre-i-numeri.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Maggio 2026  |  Professione e Automazione")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Il 65% dei commercialisti usa già l'AI:\nma i numeri raccontano metà della storia")
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "L'adozione cresce, ma tra \"usare l'AI\" e \"automatizzare bene un processo\" "
    "la distanza è ancora larga. La differenza si vede nei risultati, non nelle statistiche."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  1 maggio 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Lo studio ha adottato un sistema AI per la contabilizzazione delle fatture diciotto "
    "mesi fa. I tempi di inserimento si sono ridotti del 40%. Poi qualcuno ha iniziato "
    "a contare quante ore si spendono a correggere gli errori di classificazione che il "
    "modello fa sistematicamente sulle fatture di fornitori esteri, sulle note spese con "
    "giustificativi non standard, sulle operazioni che escono dagli schemi abituali. "
    "Il saldo netto è ancora positivo, ma molto meno di quanto sembrava al momento "
    "dell'adozione."
)

para(doc,
    "Questa storia, con varianti, si sente spesso parlando con commercialisti che hanno "
    "adottato strumenti AI nell'ultimo anno. Il dato aggregato è impressionante: secondo "
    "le rilevazioni del 2025, il 65% dei commercialisti italiani usa già strumenti di "
    "intelligenza artificiale nella propria attività, con l'84% di chi li usa che percepisce "
    "un impatto positivo. Ma la media nasconde una dispersione molto ampia tra chi ha "
    "integrato l'AI in modo strutturato e chi la usa come uno strumento puntuale senza "
    "averci costruito intorno un processo ridisegnato."
)

heading(doc, "Cosa si automatizza davvero (e cosa no)")

para(doc,
    "Le applicazioni di AI che funzionano meglio in ambito contabile sono quelle con "
    "caratteristiche specifiche: input strutturato, regole ben definite, volume alto, "
    "varianza bassa. La contabilizzazione di fatture passive da fornitori ricorrenti rientra "
    "perfettamente in questo schema. Stessa azienda, stesso formato di documento, stessi "
    "centri di costo, stessa aliquota IVA. Con una buona fase di configurazione iniziale, "
    "i modelli di classificazione automatica raggiungono tassi di accuratezza tra il 90 e "
    "il 97% su questo tipo di documenti."
)

para(doc,
    "Le difficoltà emergono con le eccezioni: fatture con imputazioni a più centri di "
    "costo, note spese con giustificativi non standard, documenti esteri con layout "
    "inusuali, operazioni non ricorrenti che richiedono una valutazione contestuale. "
    "Queste eccezioni rappresentano spesso meno del 10% del volume totale ma assorbono "
    "il 40-50% del tempo di gestione. Un sistema AI che gestisce bene il 90% del volume "
    "liberando tempo per il 10% residuo è comunque un guadagno netto, ma richiede una "
    "governance chiara di questo confine. Il professionista deve sapere con precisione "
    "dove il modello è affidabile e dove serve la revisione umana, e organizzare il "
    "processo di conseguenza."
)

heading(doc, "L'errore più comune nell'adozione")

para(doc,
    "Il problema più frequente che emerge dall'esperienza degli studi che hanno adottato "
    "AI in contabilità non è di natura tecnica, ma organizzativa. Molti studi hanno "
    "sovrapposto lo strumento AI al processo esistente senza ridisegnarlo. Il risultato "
    "è che l'AI fa la parte meccanica di quello che faceva prima un operatore umano, ma "
    "il flusso di verifica, approvazione e correzione rimane identico a prima. Il risparmio "
    "di tempo c'è, ma è molto inferiore a quello che si potrebbe ottenere."
)

para(doc,
    "Un processo contabile ridisegnato intorno alle capacità reali dell'AI funziona "
    "diversamente. Le attività ad alta standardizzazione (inserimento, classificazione, "
    "riconciliazione su documenti standard) passano completamente sotto il controllo del "
    "modello, con un check umano finale su campione o su anomalie segnalate. Le attività "
    "ad alta varianza o con implicazioni interpretative (operazioni straordinarie, rettifiche "
    "di fine anno, gestione di eccezioni normative) rimangono integralmente sotto il "
    "controllo del professionista. Questa separazione richiede una mappa del processo, "
    "non solo un abbonamento a un software."
)

heading(doc, "Cosa chiedere ai fornitori di software gestionale")

para(doc,
    "La maggior parte degli studi professionali italiani accede all'AI attraverso i propri "
    "software gestionali: TeamSystem, Wolters Kluwer, Zucchetti, Sistemi hanno tutti "
    "integrato funzionalità AI nei loro prodotti negli ultimi due anni. Prima di attivare "
    "queste funzionalità, vale la pena fare tre domande concrete ai fornitori."
)

para(doc,
    "La prima riguarda l'elaborazione dei dati: i dati vengono trattati su server in "
    "Europa o vengono inviati a infrastrutture extra-UE? Per gli studi che trattano dati "
    "personali dei clienti (e tutti lo fanno), questa informazione è necessaria per "
    "valutare la conformità al GDPR. La seconda domanda riguarda la gestione delle "
    "eccezioni: come vengono segnalate le situazioni che il modello non riesce a "
    "classificare con sufficiente confidenza? La terza, la più trascurata, riguarda il "
    "monitoring: quali strumenti sono disponibili per misurare l'accuratezza del modello "
    "nel tempo?"
)

para(doc,
    "Un modello di classificazione che parte con il 94% di accuratezza può degradare nel "
    "tempo se il profilo delle fatture cambia, se vengono aggiunti nuovi fornitori con "
    "formati diversi, o se cambiano le regole contabili applicabili. Senza un sistema "
    "di monitoring, il professionista non sa quando l'accuratezza è scesa sotto la soglia "
    "accettabile, e scopre il problema solo quando un cliente segnala un errore nel bilancio."
)

heading(doc, "Usare bene è diverso da usare")

para(doc,
    "Il 65% di adozione è un numero che racconta l'ingresso di massa di questi strumenti "
    "negli studi italiani. La domanda più interessante, adesso, è quanti di questi studi "
    "stanno usando l'AI per ridisegnare i processi, e quanti la stanno usando per fare "
    "più velocemente quello che facevano già. La differenza, nei prossimi due anni, sarà "
    "visibile nei margini operativi e nella capacità di dedicarsi alla consulenza che i "
    "clienti pagano per la competenza professionale, non per il tempo di inserimento dati. "
    "Uno studio che ha liberato il 30% del tempo dall'inserimento meccanico ha il problema "
    "concreto di decidere come usarlo: questa è la scelta strategica che i dati sull'adozione "
    "non raccontano."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "TeamSystem — AI e Automazione Contabile Studio Commercialisti (2026)",
    "Wolters Kluwer — L'Intelligenza Artificiale Trasforma la Contabilita': Opportunita' e Sfide (2026)",
    "Sibill — Contabilita' e AI: guida all'intelligenza artificiale in studio (2026)",
    "Studio AZ — Intelligenza Artificiale e Studi Professionali: le prospettive e le strategie (aprile 2026)",
    "Professionista Digitale — Contabilita' 2026: Trend e Digitalizzazione Studi Professionali (2026)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
