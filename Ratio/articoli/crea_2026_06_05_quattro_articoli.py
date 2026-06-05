"""
Quattro articoli Ratio — 5 giugno 2026

1. 2026-06-05_cassazione-allucinazioni-ai-atto-giudiziario.docx
2. 2026-06-05_agente-ai-risponde-due-ore.docx
3. 2026-06-05_ai-act-agosto-chatbot-deve-dichiararsi.docx
4. 2026-06-05_iperammortamento-180-ai-investimenti-studio.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# La Cassazione e le allucinazioni AI nell'atto giudiziario
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Responsabilità professionale")
titolo(
    doc,
    "Il ricorso lo ha scritto la macchina.\nIl professionista ha pagato.",
    "Con l'ordinanza n. 11431/2026 la Cassazione ha dichiarato inammissibile "
    "un ricorso che citava sentenze reali ma con principi sbagliati, frutto di allucinazione AI. "
    "La diligenza non si delega a nessun algoritmo.",
    "A cura della Redazione Ratio  •  5 giugno 2026"
)

para(doc,
    "Un ricorso penale viene presentato in Cassazione. Le sentenze citate esistono, i numeri "
    "di registro sono corretti, le sezioni sono quelle giuste. Solo che i principi di diritto "
    "attribuiti a quelle sentenze non corrispondono a quanto quelle sentenze hanno effettivamente "
    "deciso. Lo strumento di intelligenza artificiale generativa usato per redigere l'atto ha "
    "costruito un ragionamento plausibile, ha trovato precedenti reali, li ha collegati alla tesi "
    "da sostenere, e ha sbagliato su quello che conta di più: il significato. La Corte di "
    "Cassazione, Sezione Settima Penale, con l'ordinanza n. 11431 del 26 marzo 2026 ha dichiarato "
    "il ricorso inammissibile, condannando il ricorrente al pagamento di tremila euro a favore "
    "della cassa delle ammende e attribuendo la causa dell'errore a probabile allucinazione "
    "informatica conseguente all'uso di applicativi di intelligenza artificiale generativa."
)

para(doc,
    "La sentenza ha fatto discutere negli ambienti forensi perché è la prima volta che la "
    "Suprema Corte italiana indica esplicitamente l'allucinazione AI come elemento causale "
    "di un'inammissibilità. Ma il problema che solleva non riguarda solo gli avvocati. "
    "Riguarda chiunque, commercialista, consulente del lavoro, revisore, usi strumenti "
    "di intelligenza artificiale per produrre documenti che poi firma e consegna a un cliente "
    "o deposita davanti a un'autorità."
)

heading(doc, "Cos'è un'allucinazione e perché è diversa da un errore di battitura")

para(doc,
    "Il termine allucinazione, nel contesto dei modelli linguistici, indica qualcosa di più "
    "sottile di un errore di fatto. Un errore di fatto è verificabile: il modello scrive "
    "che una norma è entrata in vigore il 3 luglio quando la data corretta è il 3 giugno. "
    "Basta controllare. Un'allucinazione giurisprudenziale funziona diversamente: il modello "
    "cita una sentenza che esiste, la nomina con il numero corretto, la attribuisce alla sezione "
    "giusta, ma ricostruisce il ragionamento giuridico in modo diverso da quello effettivo. "
    "Chi non conosce quella sentenza, o chi si fida del fatto che i riferimenti siano verificabili, "
    "non ha un segnale immediato che qualcosa non va. La plausibilità formale dell'output nasconde "
    "l'errore sostanziale."
)

para(doc,
    "Questo meccanismo non è specifico del diritto. Si manifesta ogni volta che un modello "
    "linguistico viene usato per ragionare su un corpus di conoscenze specialistiche. Un "
    "commercialista che usa uno strumento AI per costruire un'argomentazione su una norma "
    "fiscale ottiene una risposta che cita correttamente la norma, la colloca nell'articolo "
    "giusto, ma ne ricostruisce il campo di applicazione in modo parzialmente sbagliato. "
    "Se l'errore è coerente con le aspettative del lettore, passare inosservato è la cosa "
    "più facile del mondo."
)

heading(doc, "La responsabilità rimane dove è sempre stata")

para(doc,
    "La Corte ha chiarito un principio che molti avevano già intuito ma che adesso ha un "
    "riferimento giurisprudenziale preciso: l'uso di strumenti AI per la ricerca, la sintesi "
    "e la strutturazione di documenti professionali non esonera il professionista dai propri "
    "doveri di diligenza e competenza. Selezionare, verificare e validare le fonti è un'attività "
    "intellettuale che non si può affidare a sistemi automatizzati. Chi firma è responsabile "
    "di quello che firma, indipendentemente da come lo ha prodotto."
)

para(doc,
    "Detto in altri termini: usare uno strumento AI non è una causa di esclusione della colpa "
    "processuale, è semmai una spiegazione del meccanismo con cui l'errore si è prodotto. "
    "La Corte ha usato questa spiegazione per descrivere l'accaduto, non per attenuare la "
    "responsabilità del difensore. Anzi, l'ha quantificata in tremila euro."
)

heading(doc, "Cosa cambia nella pratica professionale")

para(doc,
    "Per un professionista che già usa strumenti AI nel proprio lavoro quotidiano, questa "
    "ordinanza non chiede di smettere di usarli. Chiede di usarli con un livello di verifica "
    "proporzionato al rischio dell'output. Un documento prodotto con l'AI che viene "
    "consegnato a un cliente come analisi interna ha un livello di rischio diverso da uno "
    "che viene depositato davanti a un organo giurisdizionale o inviato a un ente pubblico. "
    "Il grado di verifica richiesto deve essere proporzionato."
)

para(doc,
    "Sul piano pratico, questo significa due cose. La prima è che qualsiasi riferimento "
    "normativo, giurisprudenziale o dottrinale prodotto da uno strumento AI va verificato "
    "alla fonte primaria prima di essere usato in un documento ufficiale. Non è sufficiente "
    "che il numero di sentenza esista: va controllato che il principio attribuito corrisponda "
    "a quello effettivo. La seconda è che il flusso di lavoro con l'AI deve includere un "
    "momento esplicito di revisione critica, non come formalità, ma come passaggio del processo "
    "in cui il professionista esercita il proprio giudizio. Quel momento è quello che distingue "
    "un assistente utile da una fonte di rischio."
)

para(doc,
    "L'ordinanza della Cassazione arriva in un momento in cui molti studi stanno adottando "
    "strumenti AI senza avere ancora definito procedure interne su come usarli. Definire quelle "
    "procedure adesso, prima che arrivi un episodio simile, è la differenza tra usare la "
    "tecnologia come vantaggio e subirla come problema."
)

riferimenti(doc, [
    "Corte di Cassazione, Sez. VII Penale, ordinanza n. 11431, 26 marzo 2026",
    "Sistema Penale — commento all'ordinanza: allucinazioni giurisprudenziali e inammissibilita' del ricorso",
    "Studio Cervellino Avvocati e Commercialisti — analisi dell'ordinanza (21 maggio 2026)",
    "Blasto Online — \"Quando il ricorso lo scrive la macchina\" (Cass. pen., Sez. VII, ord. n. 11431/2026)",
    "Milano Post — \"Cassazione e IA: la diligenza digitale contrapposta alla falsa autorevolezza\" (29 maggio 2026)",
    "Fisco 7 — \"Intelligenza artificiale negli studi professionali: responsabilita' operative\" (maggio 2026)",
])
doc.save(BASE + "2026-06-05_cassazione-allucinazioni-ai-atto-giudiziario.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# Agenti AI per PMI: il divario che costa clienti
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Adozione AI nelle PMI")
titolo(
    doc,
    "L'agente AI risponde in due ore.\nIl tuo commerciale rientra domani.",
    "Il divario nell'adozione degli agenti AI tra grandi imprese e PMI non e' piu' "
    "una questione di produttivita' interna. Sta gia' costando ordini.",
    "A cura della Redazione Ratio  •  5 giugno 2026"
)

para(doc,
    "Supponi di gestire un'impresa che produce componenti su misura. Un cliente ti chiede "
    "disponibilita' per un ordine urgente, con variante tecnica rispetto allo standard. "
    "La risposta del tuo concorrente arriva entro due ore: preventivo, stima di consegna, "
    "proposta alternativa per la variante tecnica. La tua risposta arriva il giorno dopo, "
    "perche' il responsabile commerciale era fuori sede e hai dovuto aspettare che rientrasse "
    "per elaborare i dati dal gestionale. Non hai perso quell'ordine per il prezzo o la "
    "qualita'. Lo hai perso perche' il concorrente ha un agente AI collegato al suo sistema "
    "di gestione degli ordini. L'agente legge la richiesta, consulta le disponibilita' di "
    "magazzino, calcola i tempi e compone la risposta. Il commerciale la controlla e la manda."
)

para(doc,
    "Questo scenario non e' piu' un caso di scuola. Lo raccontano imprenditori di tutta "
    "Italia in questo primo semestre del 2026, mentre il divario nell'adozione "
    "dell'intelligenza artificiale tra grandi imprese e PMI smette di essere una questione "
    "di efficienza interna e diventa una questione di competitivita' commerciale. Secondo i "
    "dati disponibili, il 71% delle grandi imprese italiane ha avviato almeno un progetto "
    "AI. La quota scende sotto il 10% per le PMI. ISTAT rilevava a fine 2025 che solo il "
    "16% delle imprese con almeno dieci dipendenti usava soluzioni AI, e appena il 7% tra "
    "quelle piu' piccole aveva avviato qualcosa di concreto. Nel frattempo il mercato si "
    "stava spostando."
)

heading(doc, "Cosa fa un agente AI che un chatbot non fa")

para(doc,
    "Un chatbot risponde alle domande. Un agente AI completa compiti. La differenza e' "
    "operativamente rilevante. Un chatbot sul sito di uno studio risponde alle FAQ, fornisce "
    "orari, raccoglie dati di contatto. Un agente AI puo' leggere un documento inviato dal "
    "cliente, estrarne le informazioni rilevanti, verificare se corrispondono a quelle gia' "
    "presenti in archivio, identificare le discrepanze e preparare una nota sintetica per "
    "il professionista, tutto prima che il professionista apra la mail. La differenza non e' "
    "nella tecnologia sottostante, che in entrambi i casi puo' usare lo stesso modello "
    "linguistico. La differenza e' nel fatto che l'agente ha accesso agli strumenti dell'impresa: "
    "il gestionale, l'archivio documenti, il sistema di pianificazione."
)

para(doc,
    "Il Piano Transizione 5.0, con il suo iperammortamento al 180% per investimenti in "
    "software e sistemi digitali fino a 2,5 milioni di euro, rende questo tipo di "
    "investimento piu' accessibile di quanto sembri. Ma il problema che blocca molte PMI "
    "non e' il costo dello strumento. E' la mancanza di qualcuno che si assuma la "
    "responsabilita' di capire da dove partire."
)

heading(doc, "Le tre ragioni per cui si aspetta")

para(doc,
    "La prima ragione e' la percezione del costo. Molti imprenditori associano l'intelligenza "
    "artificiale a progetti infrastrutturali impegnativi, a consulenze specialistiche, a "
    "implementazioni che richiedono mesi. Questa percezione e' rimasta ancorata a una "
    "realta' che e' cambiata. Oggi esistono agenti configurabili con strumenti no-code "
    "o low-code, integrabili con i software gestionali gia' in uso, a costi mensili "
    "accessibili a qualsiasi PMI con un minimo di fatturato."
)

para(doc,
    "La seconda ragione e' la mancanza di un interlocutore interno. Nelle grandi aziende "
    "esiste una figura che ha nel mandato la valutazione degli strumenti digitali. In una "
    "PMI con venti persone, questa responsabilita' e' diffusa e quindi non e' di nessuno. "
    "Il titolare delega all'IT, l'IT non ha visione sui processi commerciali, chi gestisce "
    "i processi commerciali non conosce le opzioni disponibili. Il risultato e' l'inerzia."
)

para(doc,
    "La terza ragione e' l'incertezza normativa. Molti imprenditori hanno sentito parlare "
    "dell'AI Act, delle sanzioni, della compliance, e hanno concluso che aspettare fosse "
    "la scelta piu' prudente. Il paradosso e' che aspettare ha lasciato campo libero ai "
    "concorrenti che si sono mossi comunque, con meno attenzione alla conformita' ma con "
    "un vantaggio competitivo reale."
)

heading(doc, "Da dove partire senza trasformazioni radicali")

para(doc,
    "Le PMI italiane che hanno avviato progetti AI con risultati concreti non sono partite "
    "da grandi trasformazioni. Sono partite da un processo ripetitivo che consumava tempo "
    "e aveva basso valore decisionale. Il Politecnico di Milano stima che circa il 40% "
    "del tempo del reparto amministrativo di una PMI media sia occupato da attivita' "
    "automatizzabili: categorizzazione delle fatture, riconciliazioni, preparazione di "
    "report periodici, risposta a email standardizzate. Un processo, uno strumento, "
    "tre mesi di misurazione. Se il risparmio di tempo e' reale, si allarga ad altri "
    "processi. Se non funziona, si cambia approccio senza aver perso molto."
)

para(doc,
    "Il concorrente che risponde in due ore non ha investito in una piattaforma enterprise. "
    "Ha preso un processo che faceva perdere tempo, ha scelto uno strumento che lo "
    "automatizza parzialmente, ha lasciato al commerciale il controllo finale. "
    "Quello e' il modello da copiare, non il caso eccezionale da ammirare da lontano."
)

riferimenti(doc, [
    "ISTAT — Indagine sull'uso dell'intelligenza artificiale nelle imprese italiane, 2025",
    "Osservatorio Digital Innovation, Politecnico di Milano — Automazione dei processi nelle PMI",
    "Best Tech Partner — \"AI agenti: cosa significa davvero per le imprese\" (29 maggio 2026)",
    "Ivemind — \"Agenti AI per PMI: Guida Pratica 2026\"",
    "Kinetikon — \"Agenti IA nelle PMI italiane: adozione e casi d'uso\"",
    "MIMIT — Piano Transizione 5.0, nuove aliquote iperammortamento 2026",
])
doc.save(BASE + "2026-06-05_agente-ai-risponde-due-ore.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# AI Act agosto 2026: il chatbot sul sito deve dichiararsi
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Normativa AI")
titolo(
    doc,
    "Due agosto 2026: il chatbot\nsul tuo sito deve dichiararsi.",
    "L'articolo 50 dell'AI Act non e' stato rinviato dal Digital Omnibus. "
    "Dal 2 agosto ogni sistema che interagisce con le persone deve comunicare "
    "in modo inequivocabile la propria natura artificiale.",
    "A cura della Redazione Ratio  •  5 giugno 2026"
)

para(doc,
    "Un cliente entra sul sito di uno studio professionale e trova una finestra di chat. "
    "Scrive una domanda sulla scadenza per la presentazione della dichiarazione dei redditi "
    "e riceve una risposta precisa, cordiale, formulata in modo professionale. Non sa se "
    "sta parlando con un operatore dello studio o con un sistema automatico. Non glielo "
    "viene detto. Dal 2 agosto 2026, quella situazione configura una violazione dell'articolo "
    "50 del Regolamento UE sull'intelligenza artificiale. La norma e' in vigore senza rinvii, "
    "senza deroghe per le PMI, senza soglie dimensionali. Si applica a qualsiasi soggetto "
    "che utilizzi un sistema AI per interagire con persone fisiche."
)

para(doc,
    "L'accordo Digital Omnibus del 7 maggio 2026 ha spostato in avanti alcune delle "
    "scadenze piu' impegnative dell'AI Act, quelle sui sistemi ad alto rischio, sulla "
    "documentazione tecnica e sulla supervisione obbligatoria. Molte imprese hanno letto "
    "questa notizia come un segnale di allentamento generale. La lettura e' incompleta. "
    "L'articolo 50 era in una parte diversa del regolamento e non e' stato toccato. "
    "La scadenza rimane il 2 agosto 2026."
)

heading(doc, "Cosa dice l'articolo 50 nella pratica")

para(doc,
    "L'obbligo di trasparenza si articola in due situazioni distinte. La prima riguarda "
    "i sistemi che interagiscono direttamente con le persone: chatbot, assistenti virtuali, "
    "risponditori automatici via email o messaggio. Tutti questi sistemi devono comunicare, "
    "in modo chiaro e comprensibile, che si tratta di un sistema artificiale. La comunicazione "
    "deve avvenire prima che l'interazione abbia inizio, non in un footnote a fondo pagina. "
    "La seconda situazione riguarda i contenuti generati con strumenti AI e destinati a "
    "circolare come documenti, comunicazioni o pubblicazioni: testi, immagini, audio, video. "
    "Questi contenuti devono essere marcati come generati o modificati da AI in modo leggibile "
    "a macchina, con tecnologie come il watermarking o i metadati di provenienza."
)

para(doc,
    "L'obbligo vale per chi mette in uso il sistema, non per chi lo ha sviluppato. "
    "Uno studio professionale che usa un chatbot acquistato da un fornitore e' responsabile "
    "di assicurarsi che quel sistema rispetti l'articolo 50. Se il fornitore non lo ha "
    "configurato correttamente, il problema e' dello studio, non del fornitore. "
    "Questo richiede di verificare, prima del 2 agosto, come funziona ogni strumento AI "
    "che interagisce con clienti o terzi."
)

heading(doc, "Chi controlla e cosa rischia chi non si adegua")

para(doc,
    "In Italia l'autorita' di vigilanza designata dalla Legge 132/2025 e' l'ACN, "
    "l'Agenzia per la Cybersicurezza Nazionale. L'ACN puo' avviare verifiche su "
    "segnalazione o d'ufficio. Le sanzioni previste per la violazione degli obblighi "
    "di trasparenza possono arrivare fino a 15 milioni di euro o al 3% del fatturato "
    "annuo globale, con la soglia piu' alta che si applica. Per una PMI con fatturato "
    "nella media, il 3% del fatturato e' una cifra concreta. Ma la sanzione pecuniaria "
    "non e' l'unico rischio."
)

para(doc,
    "Un cliente che scopre di aver interagito con un sistema AI senza saperlo, e che "
    "ritiene di aver ricevuto informazioni non accurate, ha un argomento in piu' per "
    "contestare la qualita' del servizio ricevuto. Per uno studio professionale, la "
    "fiducia del cliente e' il bene piu' difficile da ricostruire. Dichiarare la natura "
    "artificiale di uno strumento non indebolisce quella fiducia: al contrario, la "
    "gestione trasparente degli strumenti che si usano e' un elemento di professionalita'."
)

heading(doc, "Tre cose da fare entro la fine di giugno")

para(doc,
    "Il primo passo e' fare un inventario degli strumenti AI che interagiscono con l'esterno. "
    "Chatbot sul sito, risponditori automatici, sistemi di gestione delle email dei clienti: "
    "vanno mappati e va verificato se e come comunicano la propria natura. Il secondo passo "
    "e' contattare il fornitore di ogni strumento e chiedere come viene gestita la "
    "trasparenza ai sensi dell'articolo 50. Se il fornitore non sa rispondere, e' necessario "
    "valutare se il contratto in essere include clausole di conformita' normativa e come "
    "gestire il rischio residuo. Il terzo passo e' verificare se i documenti prodotti con "
    "strumenti AI, pareri, preventivi, comunicazioni commerciali, vengono inviati a clienti "
    "o terzi senza alcuna indicazione della loro origine."
)

para(doc,
    "Agosto e' tra otto settimane. Per chi non ha ancora avviato questo tipo di verifica, "
    "il tempo e' sufficiente, ma non e' illimitato. Gli adempimenti dell'articolo 50 non "
    "richiedono trasformazioni tecnologiche profonde: richiedono consapevolezza di quello "
    "che si usa e un livello minimo di configurazione degli strumenti gia' in uso."
)

riferimenti(doc, [
    "Regolamento UE 2024/1689 (AI Act) art. 50 — obblighi di trasparenza",
    "Agenda Digitale — \"AI Act, dal 2 agosto scatta la trasparenza obbligatoria: cosa cambia e per chi\"",
    "Legge 10 ottobre 2025, n. 132 — designazione ACN come autorita' nazionale di vigilanza AI",
    "Business Intelligence Group — \"AI Act, cosa cambia il 2 agosto 2026\"",
    "Cyberness — \"AI Act 2026: divieti, sanzioni e scadenze per le aziende italiane\"",
    "Truescreen — \"EU AI Act: obblighi di trasparenza per le aziende nel 2026\"",
])
doc.save(BASE + "2026-06-05_ai-act-agosto-chatbot-deve-dichiararsi.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Iperammortamento 5.0 e AI: l'incentivo che pochi conoscono
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Incentivi e Fiscalita'")
titolo(
    doc,
    "Iperammortamento al 180%: il software AI\nentra in bilancio con gli incentivi 5.0.",
    "Dal 1 gennaio 2026 il credito d'imposta Transizione 5.0 e' stato sostituito "
    "da un iperammortamento che copre software e sistemi AI fino a 2,5 milioni. "
    "La maggior parte delle PMI non lo sa ancora.",
    "A cura della Redazione Ratio  •  5 giugno 2026"
)

para(doc,
    "Se un cliente ti porta il contratto di abbonamento annuale per un sistema AI "
    "che ha appena introdotto in azienda e ti chiede come trattarlo fiscalmente, "
    "la risposta non e' piu' quella di qualche mese fa. Dal 1 gennaio 2026 la "
    "Legge di Bilancio ha sostituito il credito d'imposta Transizione 5.0 con un "
    "meccanismo di super-ammortamento per gli investimenti in beni digitali. "
    "L'aliquota per la fascia fino a 2,5 milioni di euro e' del 180%. Un software "
    "AI acquistato o sottoscritto in abbonamento per ottimizzare un processo aziendale, "
    "se rientra tra i beni ammissibili, si porta in deduzione per un importo pari a "
    "una volta e mezza il suo costo. Su dieci euro spesi, se ne deducono diciotto."
)

para(doc,
    "La novita' non e' solo quantitativa. Il meccanismo precedente, quello del credito "
    "d'imposta Transizione 5.0, aveva requisiti di interconnessione con il sistema "
    "produttivo e di rendicontazione energetica che ne complicavano l'accesso per le "
    "aziende di servizi e gli studi professionali. Il nuovo iperammortamento ha una "
    "struttura piu' vicina al classico super-ammortamento del piano Industria 4.0, "
    "con aggiornamenti che includono esplicitamente le tecnologie legate all'intelligenza "
    "artificiale, alla cybersecurity e al software gestionale avanzato. Il decreto "
    "attuativo era ancora in fase di iter normativo a fine maggio 2026, ma il quadro "
    "di riferimento e' gia' operativo."
)

heading(doc, "Cosa rientra negli investimenti ammissibili")

para(doc,
    "La lista dei beni ammissibili e' in aggiornamento, ma il perimetro attuale include "
    "i software pagati per canone (subscription), il che e' rilevante perche' la maggior "
    "parte degli strumenti AI di uso professionale e' distribuita in questo modo. "
    "Un abbonamento mensile a un sistema di automazione documentale, un agente AI "
    "integrato nel gestionale, un software di analisi predittiva per la gestione della "
    "tesoreria: questi sono candidati all'agevolazione, a condizione che siano "
    "interconnessi con il sistema aziendale e che l'investimento sia tracciato con "
    "la documentazione prevista."
)

para(doc,
    "Per i professionisti autonomi, lavoratori autonomi e studi individuali, c'e' una "
    "precisazione importante: il nuovo iperammortamento, come il precedente credito "
    "d'imposta 5.0, e' riservato alle imprese. I professionisti con partita IVA in "
    "regime di lavoro autonomo non rientrano tra i beneficiari. Questo non esclude "
    "che le associazioni professionali costituite in forma societaria possano accedere "
    "all'agevolazione, ma richiede una verifica della forma giuridica del soggetto "
    "richiedente prima di qualsiasi pianificazione fiscale."
)

heading(doc, "Come si documenta l'investimento")

para(doc,
    "La documentazione richiesta per fruire dell'iperammortamento segue lo schema "
    "consolidato del piano 4.0: perizia tecnica asseverata o dichiarazione del fornitore "
    "che attesta le caratteristiche tecnologiche del bene, comunicazione al MIMIT "
    "secondo le modalita' previste dal decreto attuativo, e tracciabilita' contabile "
    "dell'investimento. Per i beni immateriali, come i software in abbonamento, "
    "la documentazione si concentra sulla prova dell'interconnessione con il sistema "
    "aziendale e sull'utilizzo effettivo nell'attivita' produttiva o gestionale."
)

para(doc,
    "Per uno studio professionale che assiste una PMI nell'acquisto di un sistema AI, "
    "il momento piu' utile per impostare correttamente la documentazione e' quello "
    "precedente l'acquisto, non quello successivo. Una volta che l'investimento e' "
    "stato contabilizzato senza la prospettiva dell'agevolazione, recuperare i requisiti "
    "formali a posteriori e' piu' complesso. La pianificazione fiscale va fatta prima "
    "della firma del contratto, non al momento della chiusura del bilancio."
)

heading(doc, "L'opportunita' che si chiude nel 2028")

para(doc,
    "La finestra temporale per l'iperammortamento 5.0 copre gli investimenti effettuati "
    "dal 1 gennaio 2026 al 30 settembre 2028. Per una PMI che sta valutando se e quando "
    "introdurre strumenti AI nei propri processi, questo orizzonte temporale coincide "
    "quasi esattamente con la fase piu' critica di adozione: quella in cui gli strumenti "
    "sono abbastanza maturi da essere affidabili, abbastanza diffusi da essere comprensibili, "
    "e abbastanza regolamentati da poter essere introdotti con un quadro normativo chiaro."
)

para(doc,
    "L'incentivo fiscale non e' la ragione per cui vale la pena introdurre un sistema AI "
    "in azienda. Ma se la valutazione operativa e' gia' positiva, avere un'aliquota del "
    "180% sulla deduzione degli investimenti e' un argomento che entra nel piano finanziario "
    "e riduce il periodo di payback. Per chi non l'ha ancora fatto, vale la pena farlo "
    "entrare nel prossimo budget planning."
)

riferimenti(doc, [
    "Legge di Bilancio 2026 — iperammortamento beni digitali, art. 1 commi 445-464",
    "MIMIT — Piano Transizione 5.0: nuove aliquote e beni ammissibili 2026",
    "PMI.it — \"Transizione 5.0: scadenze e novita' 2026 crediti d'imposta\"",
    "PMI.it — \"Transizione 5.0: lavoratori autonomi esclusi dagli incentivi\"",
    "Beneggi & Associati — \"Piano Transizione 5.0: incentivi, requisiti e documenti richiesti\"",
    "IPSOA — Bonus Transizione 5.0: condizioni e adempimenti per la fruizione",
])
doc.save(BASE + "2026-06-05_iperammortamento-180-ai-investimenti-studio.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
