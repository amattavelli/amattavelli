"""
Articolo 1: Tre modelli e nessuna scelta ovvia
Ratio/articoli/2026-04-03_tre-modelli-nessuna-scelta-ovvia.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-03_tre-modelli-nessuna-scelta-ovvia.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_paragraph(doc, text, size=11, italic=False, color=None,
                  space_after=8, space_before=0,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

# ---- TESTATA ----
p_rivista = doc.add_paragraph()
p_rivista.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = p_rivista.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
run_r.font.name = 'Calibri'
run_r.font.size = Pt(9)
run_r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run_r.font.bold = True
run_r.font.all_caps = True

add_separator(doc)

p_data = doc.add_paragraph()
p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_d = p_data.add_run("Aprile 2026  |  Strumenti e Mercato")
run_d.font.name = 'Calibri'
run_d.font.size = Pt(8.5)
run_d.italic = True
run_d.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_after = Pt(6)
run_t = p_title.add_run("Tre modelli e nessuna scelta ovvia: come orientarsi tra GPT-5.4, Claude 4.6 e Gemini 3.1")
run_t.font.name = 'Calibri'
run_t.font.size = Pt(24)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_after = Pt(10)
run_s = p_sub.add_run(
    "In 45 giorni tre grandi fornitori hanno aggiornato i loro modelli quasi in simultanea. "
    "Per i professionisti, la vera notizia non \u00e8 chi ha vinto il benchmark."
)
run_s.font.name = 'Calibri'
run_s.font.size = Pt(13)
run_s.italic = True
run_s.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_separator(doc)

p_autore = doc.add_paragraph()
p_autore.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_autore.paragraph_format.space_after = Pt(16)
run_a = p_autore.add_run("A cura della Redazione Ratio  \u2022  3 aprile 2026")
run_a.font.name = 'Calibri'
run_a.font.size = Pt(9)
run_a.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- CORPO ----

add_paragraph(doc,
    "Tra febbraio e marzo 2026, in meno di quarantacinque giorni, i tre principali fornitori di modelli "
    "linguistici hanno rilasciato aggiornamenti significativi quasi in simultanea. Anthropic ha pubblicato "
    "Claude Opus 4.6 il 5 febbraio, con nuove capacit\u00e0 agentiche e un record di performance sull'indice "
    "SWE-bench (80,8%). Google DeepMind ha risposto a fine febbraio con Gemini 3.1 Pro, che ha superato "
    "tredici dei sedici principali benchmark di settore e si colloca oggi in cima alla classifica complessiva. "
    "OpenAI ha chiuso il ciclo l'11 marzo con GPT-5.4, disponibile in due varianti \u2014 Thinking e Pro \u2014 "
    "con una finestra di contesto di un milione di token via API."
)

add_paragraph(doc,
    "Chi lavora in un contesto professionale \u2014 uno studio, un ufficio amministrativo, un team di controllo "
    "di gestione \u2014 pu\u00f2 essere tentato di leggere questa sequenza come un'ulteriore conferma che la "
    "tecnologia corre troppo veloce per essere seguita. La lettura pi\u00f9 utile \u00e8 diversa: la distanza "
    "di performance tra i modelli si \u00e8 assottigliata al punto che la scelta del modello, in molti casi, "
    "ha smesso di essere la domanda pi\u00f9 importante."
)

add_heading2(doc, "Cosa distingue davvero i tre modelli")

add_paragraph(doc,
    "Le differenze tra GPT-5.4, Claude 4.6 e Gemini 3.1 esistono, ma sono diventate pi\u00f9 sottili e pi\u00f9 "
    "dipendenti dal contesto d'uso. Claude Opus 4.6 eccelle nei compiti che richiedono precisione sequenziale "
    "e correzione autonoma degli errori: \u00e8 il modello pi\u00f9 affidabile per catene di ragionamento lunghe, "
    "dove ogni passo dipende dal precedente. Gemini 3.1 Pro ha una finestra di contesto molto ampia e si "
    "integra nativamente con l'ecosistema Google Workspace, rendendolo particolarmente efficace per chi "
    "lavora con Drive, Docs e Sheets. GPT-5.4 offre la finestra di contesto pi\u00f9 ampia del mercato "
    "(un milione di token) e si integra con l'ecosistema Microsoft tramite Copilot, che \u00e8 gi\u00e0 "
    "presente in Teams, Word ed Excel per milioni di aziende italiane."
)

add_paragraph(doc,
    "Il fattore costo \u00e8 cambiato in modo rilevante. La concorrenza tra i tre player ha compresso i "
    "prezzi: GPT-5.4 costa oggi significativamente meno di quanto costasse GPT-4 Turbo dodici mesi fa. "
    "Per uno studio professionale o una PMI che usa questi strumenti in modo non sistematico, il costo "
    "per token \u00e8 diventato una variabile marginale. Diventa rilevante solo quando si costruiscono "
    "processi automatizzati che girano decine di migliaia di volte al mese."
)

add_heading2(doc, "Il cambiamento che conta: dagli strumenti ai processi")

add_paragraph(doc,
    "La novit\u00e0 pi\u00f9 rilevante di questo ciclo di aggiornamenti non \u00e8 nei benchmark. \u00c8 nel "
    "consolidamento degli agentic workflow: sistemi in cui l'AI non risponde a una domanda ma esegue un "
    "processo complesso in autonomia, usando strumenti esterni come database, caselle di posta, gestionali "
    "e portali web. Claude 4.6 ha migliorato in modo specifico questa capacit\u00e0, correggendo da solo "
    "gli errori che emergono durante l'esecuzione di sequenze multi-step. GPT-5.4 pu\u00f2 analizzare "
    "un intero fascicolo documentale \u2014 centinaia di pagine \u2014 in un singolo passaggio, grazie "
    "alla finestra di contesto estesa."
)

add_paragraph(doc,
    "Per un commercialista o un controller, questo si traduce in qualcosa di concreto. Un agente "
    "configurato sul gestionale dello studio pu\u00f2 raccogliere i dati del cliente, confrontarli con "
    "le scadenze fiscali del mese, identificare le anomalie nei movimenti contabili e preparare una "
    "bozza di comunicazione al cliente \u2014 senza che il professionista abbia dovuto fare nulla di "
    "pi\u00f9 che definire l'obiettivo iniziale. Non si tratta di fantascienza: queste configurazioni "
    "sono operative oggi, su piattaforme come Microsoft Copilot, n8n e Make, e nei software verticali "
    "che stanno integrando l'AI nei flussi gi\u00e0 esistenti."
)

add_heading2(doc, "Come scegliere, in pratica")

add_paragraph(doc,
    "Il criterio pi\u00f9 efficace per orientarsi non \u00e8 cercare il modello con il punteggio pi\u00f9 alto "
    "nei benchmark, ma partire dall'ecosistema gi\u00e0 in uso. Chi lavora prevalentemente su Microsoft "
    "365 trover\u00e0 in Copilot \u2014 alimentato da GPT-5.x \u2014 il punto di ingresso pi\u00f9 naturale, "
    "senza dover imparare nuove interfacce. Chi \u00e8 gi\u00e0 dentro Google Workspace avr\u00e0 Gemini "
    "integrato nativamente. Chi cerca la massima precisione in compiti di ragionamento complesso, come "
    "la revisione critica di contratti o la costruzione di modelli finanziari articolati, trover\u00e0 "
    "in Claude un interlocutore pi\u00f9 affidabile."
)

add_paragraph(doc,
    "Vale la pena anche valutare la questione della localizzazione dei dati. Per gli studi che trattano "
    "informazioni riservate di clienti, la scelta di un fornitore con data center europei e contratti "
    "conformi al GDPR non \u00e8 un dettaglio tecnico ma una precondizione. Su questo fronte, Microsoft "
    "Azure e Google Cloud offrono opzioni di residenza dei dati in Europa; \u00e8 necessario verificare "
    "caso per caso che la specifica configurazione usata garantisca quella residenza."
)

add_heading2(doc, "La domanda che vale la pena farsi")

add_paragraph(doc,
    "Tre modelli eccellenti, costi in calo, capacit\u00e0 agentiche sempre pi\u00f9 mature. Il rischio, "
    "paradossalmente, \u00e8 che l'abbondanza di opzioni diventi un alibi per non decidere. Molti "
    "professionisti e manager italiani usano ancora questi strumenti come motori di ricerca "
    "leggermente pi\u00f9 intelligenti: si chiede una risposta, si legge, si chiude la scheda. "
    "Il salto di valore arriva quando si smette di usare l'AI per rispondere a domande e si "
    "comincia a usarla per costruire processi. La scelta del modello diventa secondaria rispetto "
    "a questa: stiamo costruendo qualcosa che funziona, o stiamo solo aggiornando il modo in "
    "cui cerchiamo informazioni?"
)

add_separator(doc)

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(10)
run_n = p_note.add_run("Fonti e riferimenti")
run_n.font.name = 'Calibri'
run_n.font.size = Pt(9)
run_n.font.bold = True
run_n.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Anthropic \u2014 Claude Opus 4.6 e Claude Sonnet 4.6: note di rilascio (febbraio 2026)",
    "Google DeepMind \u2014 Gemini 3.1 Pro: benchmark e documentazione tecnica (febbraio 2026)",
    "OpenAI \u2014 GPT-5.4 Thinking e Pro: note di rilascio e pricing (marzo 2026)",
    "LM Council Benchmarks \u2014 Classifica comparativa modelli AI, aprile 2026",
    "Microsoft \u2014 Documentazione Copilot for Microsoft 365 (2026)",
]
for s in sources:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_after = Pt(2)
    run_s = p_s.add_run(f"\u2022 {s}")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(8.5)
    run_s.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(16)
run_f = p_footer.add_run(
    "\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte"
)
run_f.font.name = 'Calibri'
run_f.font.size = Pt(8)
run_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_f.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
