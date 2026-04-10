"""
Articolo 18: AI verticale per lo studio professionale — Normo, MIA e Dyogene a confronto
Ratio/articoli/2026-04-10_ai-verticale-studio-professionale-normo-mia-dyogene.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-10_ai-verticale-studio-professionale-normo-mia-dyogene.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_paragraph(doc, text, size=11, italic=False, color=None,
                  space_after=8, space_before=0,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


# ---- TESTATA ----
p_rivista = doc.add_paragraph()
p_rivista.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = p_rivista.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
run_r.font.name = 'Calibri'
run_r.font.size = Pt(9)
run_r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run_r.font.bold = True
run_r.font.all_caps = True

add_separator(doc)

p_data = doc.add_paragraph()
p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_d = p_data.add_run("Aprile 2026  |  Strumenti per Professionisti")
run_d.font.name = 'Calibri'
run_d.font.size = Pt(8.5)
run_d.italic = True
run_d.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_after = Pt(6)
run_t = p_title.add_run(
    "AI verticale per lo studio: Normo, MIA e Dyogene a confronto"
)
run_t.font.name = 'Calibri'
run_t.font.size = Pt(24)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_after = Pt(10)
run_s = p_sub.add_run(
    "Tre architetture diverse, tre filosofie diverse. Scegliere l\u2019AI verticale giusta per lo studio "
    "non significa trovare quella con il punteggio pi\u00f9 alto nei test, ma quella che si adatta "
    "meglio alla struttura dei propri processi e all\u2019ecosistema gestionale gi\u00e0 in uso."
)
run_s.font.name = 'Calibri'
run_s.font.size = Pt(13)
run_s.italic = True
run_s.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_separator(doc)

p_autore = doc.add_paragraph()
p_autore.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_autore.paragraph_format.space_after = Pt(16)
run_a = p_autore.add_run("A cura della Redazione Ratio  \u2022  10 aprile 2026")
run_a.font.name = 'Calibri'
run_a.font.size = Pt(9)
run_a.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- CORPO ----

add_paragraph(doc,
    "Il panorama degli strumenti AI dedicati ai professionisti economico-giuridici italiani si \u00e8 "
    "rapidamente popolato. Accanto ai modelli generalisti \u2014 ChatGPT, Claude, Gemini \u2014 sono "
    "emersi strumenti verticali pensati specificamente per commercialisti, avvocati e consulenti del "
    "lavoro: sistemi che attingono a basi normative certificate, che citano la fonte esatta di ogni "
    "risposta, che si aggiornano al ritmo dei provvedimenti dell\u2019Agenzia delle Entrate. Nel 2026, "
    "tre piattaforme in particolare si contendono lo spazio negli studi professionali italiani: Normo, "
    "MIA (MySolution Intelligence Advisor) e Dyogene."
)

add_paragraph(doc,
    "La scelta tra queste tre soluzioni non \u00e8 semplicemente una questione di funzionalit\u00e0: "
    "\u00e8 una questione di architettura. Ognuna \u00e8 costruita su una logica diversa, risponde a "
    "un\u2019esigenza diversa e si integra in modo diverso con i flussi di lavoro dello studio. "
    "Capire questa differenza \u00e8 il presupposto per una scelta utile."
)

add_heading2(doc, "Normo: il copilota integrato nell\u2019ecosistema gestionale")

add_paragraph(doc,
    "Normo nasce come assistente conversazionale verticale per l\u2019ambito fiscale, contabile e del "
    "lavoro. Risponde a quesiti in linguaggio naturale attingendo a una base documentale aggiornata "
    "quotidianamente: leggi, circolari, prassi dell\u2019Agenzia delle Entrate, giurisprudenza "
    "tributaria. Ogni risposta cita la fonte normativa esatta, riducendo il rischio di allucinazioni "
    "che caratterizza i modelli generalisti quando operano su testi di legge."
)

add_paragraph(doc,
    "La svolta strategica di Normo \u00e8 arrivata con l\u2019acquisizione da parte di TeamSystem, "
    "che ne ha spostato il baricentro verso l\u2019integrazione con i flussi gestionali. Per chi "
    "lavora su Profis o Euroconference, Normo diventa un copilota che opera dentro l\u2019ambiente "
    "gi\u00e0 in uso, senza cambiare strumento. \u00c8 la scelta pi\u00f9 naturale per chi \u00e8 "
    "gi\u00e0 radicato nell\u2019ecosistema TeamSystem e cerca un assistente conversazionale che "
    "risponda a quesiti fiscali con citazione delle fonti e si integri nei processi esistenti."
)

add_heading2(doc, "MIA: la banca dati potenziata")

add_paragraph(doc,
    "MIA \u2014 MySolution Intelligence Advisor \u2014 \u00e8 concepita come un sistema di Information "
    "Retrieval avanzato: non un chatbot generalista, ma un motore di ricerca intelligente su una banca "
    "dati fiscale, societaria e del lavoro certificata. Il punto di forza \u00e8 la profondit\u00e0 "
    "documentale e la capacit\u00e0 di eseguire calcoli specifici \u2014 imposte, acconti, scadenze "
    "\u2014 integrati nelle risposte. Per chi cerca uno strumento di supporto quotidiano per quesiti "
    "di routine \u2014 verifica di una scadenza, calcolo di un\u2019imposta, ricerca di una circolare "
    "specifica \u2014 MIA offre velocit\u00e0 e precisione."
)

add_paragraph(doc,
    "Il limite di MIA \u00e8 la stessa cosa che ne \u00e8 il pregio: \u00e8 ottimizzata per la ricerca "
    "e il calcolo, non per il ragionamento su casi complessi o la costruzione di analisi multi-fonte. "
    "Per un quesito strutturato con pi\u00f9 variabili \u2014 un\u2019operazione straordinaria con "
    "implicazioni fiscali e giuslavoristiche intrecciate \u2014 la profondit\u00e0 analitica di MIA "
    "pu\u00f2 non essere sufficiente."
)

add_heading2(doc, "Dyogene: dall\u2019analisi all\u2019esecuzione")

add_paragraph(doc,
    "Dyogene ha un\u2019architettura radicalmente diversa. Non \u00e8 un assistente conversazionale: "
    "\u00e8 una piattaforma di Business Intelligence e automazione di processo orientata all\u2019esecuzione. "
    "Carica direttamente file PDF o XBRL, processa bilanci e produce in automatico la riclassificazione "
    "con indicatori finanziari \u2014 ROE, ROI, EBITDA margin \u2014 pi\u00f9 metriche avanzate come "
    "il Rating MCC e l\u2019Altman Z-Score, con benchmark settoriali integrati. Il professionista "
    "carica il documento; Dyogene restituisce un\u2019analisi strutturata."
)

add_paragraph(doc,
    "La differenza rispetto a Normo e MIA \u00e8 sintetizzabile cos\u00ec: mentre quelle due eccellono "
    "nel \u201csapere\u201d \u2014 rispondere a domande, trovare norme, spiegare concetti \u2014 "
    "Dyogene eccelle nel \u201cfare\u201d. Processa dati strutturati, produce output pronti all\u2019uso, "
    "automatizza fasi del lavoro che altrimenti richiederebbero ore. Per uno studio che fa molta "
    "analisi di bilancio, consulenza finanziaria o supporto a operazioni straordinarie, Dyogene "
    "\u00e8 uno strumento che cambia concretamente la produttivit\u00e0."
)

add_heading2(doc, "Come scegliere: tre domande pratiche")

add_paragraph(doc,
    "Prima di valutare qualsiasi strumento, vale la pena rispondere a tre domande. La prima: su quale "
    "gestionale lavoriamo? Se l\u2019ecosistema \u00e8 TeamSystem, l\u2019integrazione con Normo \u00e8 "
    "gi\u00e0 disponibile e riduce i costi di adozione. La seconda: qual \u00e8 l\u2019attivit\u00e0 "
    "che assorbe pi\u00f9 tempo nello studio? Se \u00e8 rispondere a quesiti fiscali di routine, "
    "MIA o Normo; se \u00e8 produrre analisi finanziarie e di bilancio, Dyogene. La terza: vogliamo "
    "uno strumento che risponde o uno che esegue? La distinzione \u00e8 fondamentale e determina "
    "l\u2019architettura giusta."
)

add_paragraph(doc,
    "C\u2019\u00e8 poi un quarto elemento, spesso trascurato: la conformit\u00e0 alla Legge 132/2025, "
    "che dal 10 ottobre 2025 impone ai professionisti di comunicare ai clienti l\u2019utilizzo di "
    "sistemi AI con linguaggio chiaro e comprensibile. Tutti e tre gli strumenti citati operano su "
    "dati forniti dallo studio e non vengono condivisi con terzi per l\u2019addestramento dei modelli, "
    "un requisito che \u00e8 necessario verificare caso per caso con il fornitore del servizio. "
    "La compliance AI Act non riguarda solo le grandi imprese: riguarda ogni studio che usa l\u2019AI "
    "in modo professionale."
)

add_separator(doc)

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(10)
run_n = p_note.add_run("Fonti e riferimenti")
run_n.font.name = 'Calibri'
run_n.font.size = Pt(9)
run_n.font.bold = True
run_n.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Tuttotek.it \u2014 AI per commercialisti nel 2026: Normo, Dyogene, MIA, quale \u00e8 la migliore?",
    "ItaliaOggi \u2014 Intelligenza artificiale per commercialisti: la specializzazione verticale (2026)",
    "Techbusiness.it \u2014 Dyogene: la nuova AI per contabili e consulenti",
    "Normo.ai \u2014 Documentazione prodotto e aggiornamenti 2026",
    "Dyogene.ai \u2014 Documentazione prodotto e aggiornamenti 2026",
    "Legge 23 settembre 2025, n. 132 \u2014 art. 13, obblighi di trasparenza per i professionisti",
]
for s in sources:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_after = Pt(2)
    run_s = p_s.add_run(f"\u2022 {s}")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(8.5)
    run_s.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(16)
run_f = p_footer.add_run(
    "\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte"
)
run_f.font.name = 'Calibri'
run_f.font.size = Pt(8)
run_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_f.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
