"""
Quattro articoli Ratio — 12 giugno 2026

1. 2026-06-12_decreti-attuativi-ai-italia-prima-europa.docx
2. 2026-06-12_paradosso-produttivita-ai-aziende-italiane.docx
3. 2026-06-12_sme-ai-accelerator-openai-confartigianato-pmi.docx
4. 2026-06-12_microsoft-copilot-agentico-aziende-italiane.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# Decreti attuativi L. 132/2025: l'Italia prima in Europa
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Normativa AI")
titolo(
    doc,
    "L'Italia vara i decreti sull'AI.\nSiamo primi in Europa, ma c'e' da fare.",
    "Il 10 giugno il Consiglio dei Ministri ha approvato i due decreti attuativi "
    "della Legge 132/2025. Governance, responsabilita' professionale, formazione: "
    "il quadro normativo si completa. Ecco cosa cambia concretamente per professionisti e imprese.",
    "A cura della Redazione Ratio  •  12 giugno 2026"
)

para(doc,
    "Il 10 giugno 2026 il Consiglio dei Ministri ha approvato in esame preliminare "
    "i due decreti legislativi attuativi della Legge 23 settembre 2025 n. 132, "
    "la legge italiana sull'intelligenza artificiale. L'Italia diventa cosi' il primo "
    "paese europeo ad adottare un pacchetto attuativo organico che integra la normativa "
    "nazionale con il Regolamento UE 2024/1689 (AI Act). Non e' un traguardo formale: "
    "significa che le imprese italiane si trovano a operare con un quadro regolatorio "
    "piu' definito rispetto ai concorrenti europei, ma anche con obblighi piu' precisi "
    "e scadenze ravvicinate. I decreti ora passano all'esame parlamentare e alla "
    "Conferenza delle Regioni, ma i contenuti sono sostanzialmente consolidati."
)

para(doc,
    "Il primo decreto riguarda la governance nazionale dell'intelligenza artificiale. "
    "Conferma AgID, l'Agenzia per l'Italia Digitale, come autorita' di notifica, "
    "e ACN, l'Agenzia per la Cybersicurezza Nazionale, come autorita' di vigilanza del "
    "mercato. Stabilisce il funzionamento della sandbox regolatoria, cioe' lo spazio "
    "controllato in cui le imprese potranno sperimentare sistemi AI con requisiti "
    "semplificati prima di immetterli nel mercato ordinario. E definisce il regime "
    "sanzionatorio per le violazioni del regolamento europeo, con importi che possono "
    "arrivare fino a 35 milioni di euro o al 7% del fatturato annuo globale per le "
    "infrazioni piu' gravi. Il secondo decreto affronta i settori: formazione, lavoro, "
    "sanita', pubblica amministrazione, responsabilita' civile e penale."
)

heading(doc, "Responsabilita' professionale: il principio che cambia tutto")

para(doc,
    "Per i professionisti, il passaggio normativo piu' rilevante e' quello sulla "
    "responsabilita' nell'uso dell'AI. La legge 132 aveva introdotto il principio "
    "antropocentrico: l'AI supporta il professionista, non lo sostituisce nel giudizio. "
    "I decreti attuativi precisano le conseguenze operative di quel principio. "
    "Chi usa uno strumento AI per produrre documenti, pareri, analisi o comunicazioni "
    "destinate a clienti o a terzi ha obblighi informativi precisi: deve comunicare "
    "se e come lo strumento e' stato usato, in quali fasi, con quale supervisione. "
    "Non si tratta di una dichiarazione generica nell'informativa privacy. Si tratta "
    "di trasparenza sull'output specifico, adeguata alla natura del documento e al "
    "contesto professionale."
)

para(doc,
    "Questo non significa che ogni parere fiscale debba portare un bollino 'prodotto "
    "con AI'. Significa che se il professionista usa uno strumento AI per costruire "
    "l'argomentazione e il cliente lo chiede, deve essere in grado di rispondergli "
    "in modo chiaro. Il confine non e' ancora definito nei dettagli applicativi, "
    "che dipenderanno anche dagli orientamenti degli ordini professionali. Ma il "
    "principio e' netto: trasparenza sull'uso, responsabilita' sul risultato."
)

heading(doc, "Formazione obbligatoria: la scadenza che si avvicina")

para(doc,
    "Dal 1 agosto 2026 entrano in vigore gli obblighi di AI literacy previsti "
    "dall'articolo 14 della Legge 132. I datori di lavoro devono garantire che "
    "i dipendenti che usano sistemi AI nell'ambito della propria attivita' lavorativa "
    "abbiano ricevuto una formazione adeguata, proporzionata al livello di interazione "
    "con lo strumento. Non si tratta di un corso generico sull'intelligenza artificiale: "
    "si tratta di formazione specifica sugli strumenti effettivamente usati in azienda, "
    "sulle loro limitazioni, sui rischi e sulle responsabilita' connesse. "
    "Gli ordini professionali sono chiamati a integrare questi contenuti nei programmi "
    "di formazione continua obbligatoria."
)

para(doc,
    "Per uno studio professionale o un'impresa di medie dimensioni, questo significa "
    "fare un censimento degli strumenti AI in uso, identificare i ruoli che li usano, "
    "e verificare se i programmi formativi gia' esistenti siano sufficienti o debbano "
    "essere aggiornati. Chi non ha ancora avviato questa ricognizione ha meno di due "
    "mesi per farlo. Chi l'ha gia' avviata deve verificare che la documentazione "
    "sia sufficiente in caso di controllo."
)

heading(doc, "La sandbox e i nuovi spazi per sperimentare")

para(doc,
    "Uno degli elementi piu' interessanti del pacchetto attuativo e' la sandbox "
    "regolatoria, la possibilita' per le imprese di sviluppare e testare sistemi AI "
    "in un ambiente controllato, con requisiti semplificati rispetto a quelli previsti "
    "per l'immissione sul mercato ordinario. Il meccanismo e' pensato per ridurre "
    "l'asimmetria tra grandi aziende, che possono permettersi i costi della compliance "
    "completa fin dall'inizio, e PMI, che rischiano di essere escluse dall'innovazione "
    "proprio per l'impatto dei requisiti normativi. L'accesso alla sandbox richiede "
    "una richiesta formale ad AgID e il rispetto di condizioni minime di sicurezza "
    "e trasparenza, ma senza i requisiti documentali completi del regime ordinario."
)

para(doc,
    "Il quadro che emerge da questi decreti non e' quello di una normativa punitiva. "
    "E' quello di un sistema che cerca di bilanciare tre esigenze che in molti contesti "
    "europei sono ancora in tensione: abilitare l'innovazione, proteggere i cittadini, "
    "dare certezza alle imprese. Il fatto che l'Italia arrivi prima degli altri paesi "
    "europei su questo fronte e' un vantaggio competitivo potenziale, a patto che "
    "le imprese usino questa chiarezza normativa come punto di partenza per muoversi, "
    "non come ragione per aspettare ulteriori precisazioni."
)

riferimenti(doc, [
    "Consiglio dei Ministri n. 177 del 10 giugno 2026 — comunicato stampa ufficiale",
    "Legge 23 settembre 2025, n. 132 sull'intelligenza artificiale — Normattiva",
    "Regolamento UE 2024/1689 (AI Act) — Gazzetta Ufficiale UE",
    "Prima Pagina News — 'Il governo vara i decreti sull'intelligenza artificiale: l'Italia prima in Europa'",
    "Punto Informatico — 'Legge italiana sull'AI: governo approva decreti attuativi'",
    "Altalex — 'Decreti attuativi legge AI 132/2025: governance e responsabilita' professionale'",
    "Gility Mag — 'Legge 132/2025: guida completa alla conformita' (con riferimenti all'AI Act)'",
])
doc.save(BASE + "2026-06-12_decreti-attuativi-ai-italia-prima-europa.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# Il paradosso della produttività AI in Italia
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Adozione AI")
titolo(
    doc,
    "L'AI la usiamo tutti.\nNessuno e' ancora piu' produttivo.",
    "Il 32% delle imprese italiane usa gia' l'intelligenza artificiale, "
    "ma 7 aziende su 10 dichiarano che non ha ancora inciso sulla produttivita'. "
    "Il problema non e' la tecnologia. E' il modo in cui viene adottata.",
    "A cura della Redazione Ratio  •  12 giugno 2026"
)

para(doc,
    "Il mercato italiano dell'intelligenza artificiale vale 1,8 miliardi di euro "
    "e cresce del 50% all'anno. La quota di imprese che dichiara di usare AI nei "
    "propri processi e' salita dal 27% nel 2025 al 32% all'inizio del 2026. "
    "I titolari di studi professionali parlano di ChatGPT nelle riunioni. "
    "I responsabili amministrativi usano Copilot per le email. I team commerciali "
    "generano presentazioni in dieci minuti invece di due ore. Eppure, quando si "
    "chiede alle stesse imprese se l'intelligenza artificiale ha inciso sulla "
    "produttivita', sette su dieci rispondono no. Benvenuti nel paradosso della "
    "produttivita' AI."
)

para(doc,
    "Il fenomeno ha un nome nella letteratura economica: il productivity paradox. "
    "E' gia' successo con i personal computer negli anni Ottanta, con internet negli "
    "anni Novanta, con i sistemi ERP nei Duemila. In tutti questi casi, le nuove "
    "tecnologie sono entrate nelle imprese, sono state adottate da milioni di "
    "lavoratori, e per anni non si sono viste nei dati di produttivita'. Poi, a "
    "un certo punto, si sono viste tutte insieme. L'economista Robert Solow lo "
    "aveva detto nel 1987: 'Vedo i computer ovunque tranne che nelle statistiche "
    "della produttivita''. Con l'AI stiamo rivivendo lo stesso schema, con una "
    "differenza: questa volta lo stiamo riconoscendo mentre accade."
)

heading(doc, "Perche' l'AI non compare ancora nei risultati")

para(doc,
    "La prima ragione e' strutturale. L'impatto macroeconomico dell'AI richiede "
    "tempo per materializzarsi nei dati ufficiali. Le statistiche della produttivita' "
    "misurano output per ora lavorata a livello aggregato: ci vogliono anni perche' "
    "i guadagni individuali si traducano in variazioni statisticamente rilevabili. "
    "Ma la spiegazione strutturale non esaurisce il problema. Ci sono ragioni "
    "specifiche per cui molte imprese italiane non stanno ancora vedendo risultati, "
    "e queste ragioni sono gestibili."
)

para(doc,
    "La seconda ragione e' che solo il 5% delle imprese italiane ha integrato l'AI "
    "in profondita' nei propri processi. Il restante 95% la usa in modo superficiale: "
    "uno strumento in piu', non un cambiamento nel modo di lavorare. La differenza "
    "e' concreta. Usare ChatGPT per riscrivere un'email e' utile ma non cambia "
    "il modo in cui un'impresa gestisce i clienti. Usare un agente AI che legge "
    "le richieste in arrivo, le categorizza, recupera le informazioni dal gestionale "
    "e prepara una bozza di risposta prima che l'operatore apra la casella di posta: "
    "quello cambia il processo. Il risparmio non e' di cinque minuti a email. "
    "E' di un'ora al giorno per persona."
)

heading(doc, "Il tempo risparmiato viene immediatamente riempito")

para(doc,
    "C'e' un terzo meccanismo, piu' sottile, che emerge dalla ricerca sul campo. "
    "Gli strumenti AI fanno risparmiare tempo, ma quel tempo viene immediatamente "
    "occupato con altro lavoro. Non si traduce in riposo, non si traduce in "
    "riduzione dell'orario, non si traduce in capacita' di gestire piu' clienti. "
    "Si traduce in giornate piu' dense. Il risultato e' che il lavoratore si "
    "sente piu' occupato di prima, non piu' produttivo. Il report ADP People "
    "at Work 2026 rileva un dato paradossale: chi usa l'AI quotidianamente tende "
    "a percepirsi meno produttivo di chi la usa raramente, probabilmente perche' "
    "il carico di aspettative e di lavoro cresce in proporzione alla velocita' "
    "degli strumenti."
)

para(doc,
    "Questo e' un problema manageriale prima che tecnologico. Se si introduce "
    "uno strumento che rende un processo il 40% piu' veloce senza ridefinire "
    "i carichi di lavoro, si ottiene il 40% di lavoro in piu', non il 40% "
    "di risparmio. Il beneficio va catturato con una scelta esplicita: "
    "meno tempo sullo stesso volume, o stesso tempo con un volume maggiore. "
    "Senza questa scelta, il guadagno evaporazione."
)

heading(doc, "Chi sta gia' vedendo i risultati e perche'")

para(doc,
    "Le imprese che stanno gia' misurando impatti concreti hanno in comune "
    "alcune caratteristiche. La prima e' che hanno identificato un processo "
    "specifico da ottimizzare, non hanno introdotto l'AI in modo generico. "
    "La seconda e' che hanno ridefinito le aspettative su quel processo dopo "
    "l'introduzione dello strumento: meno ore, stessa qualita'. La terza e' "
    "che hanno formato le persone non solo su come usare lo strumento, ma su "
    "come verificarne l'output e su quali decisioni non possono essere delegate "
    "allo strumento."
)

para(doc,
    "I dati dell'iniziativa SME AI Accelerator di OpenAI e Confartigianato, "
    "su un campione di PMI italiane che hanno partecipato al programma, "
    "mostrano un risparmio medio di oltre cinque ore a settimana per imprenditore. "
    "Non e' poco: e' piu' di una giornata lavorativa al mese. Ma il risparmio "
    "si ottiene solo quando l'uso e' sistematico, non episodico. "
    "La differenza tra cinque ore risparmiate e zero non e' nella qualita' "
    "dello strumento. E' nella disciplina con cui viene usato."
)

riferimenti(doc, [
    "Banca d'Italia — Indagine Invind: adozione AI nelle imprese italiane, 2026",
    "Osservatorio Digital Innovation, Politecnico di Milano — Mercato AI in Italia 2026",
    "ADP People at Work 2026 — Report sull'uso dell'AI nel lavoro in Italia",
    "Fortune Italia — 'Il paradosso della produttivita': perche' l'AI non compare ancora nei dati' (28 maggio 2026)",
    "QuiFinanza — 'Paradosso intelligenza artificiale in Italia: non sta creando guadagni'",
    "Agenda Digitale — 'Tanta tecnologia ma la produttivita' e' bloccata: il paradosso Italia'",
    "Confartigianato — 'PMI italiane accelerano sull'AI: oltre cinque ore risparmiate a settimana' (maggio 2026)",
])
doc.save(BASE + "2026-06-12_paradosso-produttivita-ai-aziende-italiane.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# SME AI Accelerator: OpenAI + Confartigianato per le PMI
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Strumenti e risorse")
titolo(
    doc,
    "OpenAI porta l'AI nelle PMI italiane.\nE' gratis e puoi iscriverti ora.",
    "Il programma SME AI Accelerator, lanciato da OpenAI e Confartigianato ad aprile, "
    "offre formazione pratica gratuita alle piccole imprese italiane. "
    "Chi l'ha gia' usato risparmia in media cinque ore di lavoro a settimana.",
    "A cura della Redazione Ratio  •  12 giugno 2026"
)

para(doc,
    "Ad aprile 2026 OpenAI e Confartigianato Imprese hanno firmato un memorandum "
    "d'intesa per lanciare in Italia lo SME AI Accelerator, il programma europeo "
    "che OpenAI ha avviato in sei paesi — Italia, Francia, Germania, Polonia, "
    "Irlanda e Regno Unito — con l'obiettivo di portare l'intelligenza artificiale "
    "concretamente nelle piccole e medie imprese. Il programma e' gratuito, aperto "
    "a tutte le categorie produttive, e non richiede nessuna competenza tecnica "
    "preliminare. Obiettivo dichiarato: raggiungere 10.000 imprese europee, "
    "con l'Italia come uno dei mercati prioritari."
)

para(doc,
    "Il 15 maggio a Milano si e' tenuto l'evento principale del programma in Italia, "
    "con workshop pratici, dimostrazioni di casi d'uso concreti e sessioni "
    "hands-on con imprenditori e titolari di piccole imprese. I risultati del "
    "sondaggio condotto su chi ha gia' partecipato sono chiari: le PMI italiane "
    "che usano sistemi AI in modo regolare risparmiano in media piu' di cinque "
    "ore a settimana. Sono piu' di una giornata lavorativa al mese. Per un piccolo "
    "studio professionale con tre persone, sono quindici ore al mese che tornano "
    "disponibili per il lavoro a piu' alto valore aggiunto."
)

heading(doc, "Cosa offre concretamente il programma")

para(doc,
    "Lo SME AI Accelerator non e' un corso teorico sull'intelligenza artificiale. "
    "E' un percorso operativo che parte dai processi reali dell'impresa e insegna "
    "a usare gli strumenti disponibili, in primo luogo quelli di OpenAI, "
    "per automatizzare, velocizzare o migliorare quei processi. Il programma "
    "si articola in tre componenti. La prima e' la formazione online attraverso "
    "l'OpenAI Academy: moduli brevi e pratici, fruibili in autonomia, organizzati "
    "per settore e tipo di attivita'. La seconda sono i workshop dal vivo, "
    "organizzati nelle sedi territoriali di Confartigianato, con dimostrazioni "
    "di casi concreti e spazio per sperimentare direttamente. La terza sono le "
    "risorse scritte: guide, template, esempi di prompt per i casi d'uso piu' "
    "frequenti nelle PMI italiane."
)

para(doc,
    "L'aspetto piu' interessante del programma, dal punto di vista di un professionista "
    "che assiste PMI, e' che affronta un problema che ostacola l'adozione AI "
    "nelle piccole imprese: la mancanza di un interlocutore interno. Nelle grandi "
    "aziende esiste qualcuno con il mandato di valutare e implementare strumenti "
    "digitali. Nelle PMI quella responsabilita' non e' di nessuno in modo esplicito. "
    "Il programma offre un punto di riferimento esterno, strutturato, che aiuta "
    "l'imprenditore a capire da dove partire senza dover fare tutto da solo."
)

heading(doc, "I casi d'uso piu' rilevanti per le PMI italiane")

para(doc,
    "Tra i casi d'uso piu' frequenti nelle PMI che partecipano al programma "
    "emergono alcune categorie ricorrenti. La gestione delle comunicazioni con "
    "i clienti: risposta alle richieste via email, gestione delle FAQ, "
    "follow-up commerciali. La produzione di contenuti: offerte commerciali, "
    "presentazioni, newsletter. L'analisi dei dati: riconciliazioni, report "
    "periodici, sintesi di documenti lunghi. E la ricerca: normative, "
    "aggiornamenti di settore, benchmark di mercato. "
    "Nessuno di questi casi richiede competenze tecniche specifiche. "
    "Richiedono la volonta' di cambiare il modo in cui si svolge un'attivita' "
    "che si e' sempre svolta allo stesso modo."
)

para(doc,
    "Per un commercialista o un consulente che segue PMI, il programma e' anche "
    "uno strumento di conversazione con i clienti. Sapere cosa offre, come "
    "funziona e quali risultati stanno ottenendo le imprese che partecipano "
    "consente di avere una risposta concreta quando un imprenditore chiede "
    "'ma noi da dove partiamo con l'AI?'. La risposta non e' piu' solo teorica: "
    "e' un programma gratuito, con casi pratici, accessibile subito."
)

heading(doc, "Come accedere al programma")

para(doc,
    "L'iscrizione al programma avviene attraverso la pagina dedicata sul sito "
    "di Confartigianato o direttamente attraverso la piattaforma OpenAI Academy. "
    "Non e' necessario essere associati a Confartigianato per partecipare: "
    "il programma e' aperto a tutte le imprese italiane. "
    "I workshop in presenza vengono organizzati nelle sedi territoriali di "
    "Confartigianato su tutto il territorio nazionale: il calendario aggiornato "
    "e' disponibile sul sito dell'associazione. Chi non puo' partecipare in "
    "presenza puo' accedere ai materiali online in autonomia, senza scadenze."
)

riferimenti(doc, [
    "Confartigianato Imprese — 'OpenAI e Confartigianato lanciano lo SME AI Accelerator per le PMI italiane' (aprile 2026)",
    "Confartigianato Imprese — 'PMI italiane accelerano sull'intelligenza artificiale: oltre cinque ore risparmiate a settimana grazie all'AI' (maggio 2026)",
    "ANSA — 'OpenAI con Confartigianato per accelerare l'IA nelle PMI' (20 aprile 2026)",
    "Arena Digitale — 'OpenAI, Confartigianato e Booking.com lanciano l'SME AI Accelerator per le PMI italiane'",
    "Borsa Italiana / Radiocor — 'AI: OpenAI, Confartigianato e Booking lanciano SME AI Accelerator' (20 aprile 2026)",
    "La Mia Finanza — 'Confartigianato: le PMI italiane accelerano sull'AI' (maggio 2026)",
])
doc.save(BASE + "2026-06-12_sme-ai-accelerator-openai-confartigianato-pmi.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Microsoft Copilot diventa agentico: cosa cambia per le imprese
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Strumenti AI")
titolo(
    doc,
    "Copilot non suggerisce piu'.\nAdesso agisce.",
    "Microsoft 365 Copilot ha compiuto un passo ulteriore: da assistente a sistema "
    "agentico che esegue compiti in autonomia. Per le imprese italiane che usano "
    "Office la differenza non e' tecnica. E' organizzativa.",
    "A cura della Redazione Ratio  •  12 giugno 2026"
)

para(doc,
    "All'AI Tour del marzo 2026 Microsoft ha presentato quella che chiama la "
    "Frontier Transformation: un cambio di paradigma nel modo in cui Copilot "
    "si integra nei processi aziendali. Il salto non e' nella quantita' di cose "
    "che il sistema sa fare, ma nella modalita' con cui le fa. Il Copilot che "
    "molte aziende italiane conoscono era un assistente: suggeriva, completava, "
    "rispondeva. Il nuovo Copilot e' un sistema agentico: pianifica sequenze "
    "di azioni, usa gli strumenti disponibili in Office 365, delega sottofasi "
    "ad agenti specializzati e restituisce un risultato completo invece di "
    "un suggerimento da completare manualmente."
)

para(doc,
    "La differenza pratica e' significativa. Un Copilot assistivo risponde "
    "alla domanda 'come scrivo questo contratto?' con un testo da correggere "
    "e integrare. Un Copilot agentico risponde alla richiesta 'prepara la "
    "bozza di contratto con il fornitore Rossi usando il template standard, "
    "inserendo le condizioni discusse nell'ultima email e segnalando eventuali "
    "discrepanze rispetto al contratto precedente'. Il sistema legge le email, "
    "recupera il template, produce la bozza, confronta con il documento "
    "precedente e segnala le differenze. Il professionista revisiona e approva. "
    "Non e' fantascienza: e' gia' disponibile per le aziende con licenza "
    "Microsoft 365 Copilot."
)

heading(doc, "Perche' questo cambia le cose per le imprese italiane")

para(doc,
    "L'adozione di Microsoft 365 in Italia e' capillare. Quasi il 90% delle "
    "grandi aziende italiane usa gia' strumenti Microsoft, e la diffusione "
    "nelle PMI e' altrettanto ampia grazie ai piani in abbonamento a partire "
    "da pochi euro al mese per utente. Questo significa che la transizione "
    "verso le funzionalita' agentiche non richiede un cambio di piattaforma: "
    "richiede l'aggiornamento della licenza e, soprattutto, un cambiamento "
    "nel modo in cui le persone lavorano con lo strumento. "
    "Il secondo aspetto e' quello che spesso viene sottovalutato."
)

para(doc,
    "Le licenze di Copilot a livello globale sono cresciute del 160% su base "
    "annua, ma i dati di utilizzo attivo raccontano una storia piu' articolata: "
    "molte aziende hanno acquistato le licenze senza trasformare effettivamente "
    "i flussi di lavoro. Il Copilot viene usato come motore di ricerca avanzato "
    "o come generatore di testi, non come agente autonomo. Il potenziale rimane "
    "non sfruttato. La ragione principale, secondo i dati dell'osservatorio "
    "Microsoft, e' la mancanza di cultura digitale e di change management, "
    "citata dal 72% delle imprese come ostacolo principale."
)

heading(doc, "La nuova suite E7 e cosa significa per chi decide gli acquisti")

para(doc,
    "Con la Frontier Transformation, Microsoft ha introdotto anche la suite "
    "Microsoft 365 E7, la versione piu' avanzata dell'offerta enterprise, "
    "che integra funzionalita' agentiche avanzate, strumenti di sicurezza "
    "potenziati e capacita' di automazione dei flussi di lavoro. E7 non "
    "e' una versione incrementale di E5: e' una ridefinizione del concetto "
    "di produttivita' office. Per chi si occupa di pianificazione degli "
    "acquisti IT nelle aziende o assiste i clienti in decisioni di questo "
    "tipo, e' importante capire che la scelta tra le diverse versioni di "
    "licenza non e' piu' solo una questione di prezzo: e' una scelta su "
    "quali processi si vogliono automatizzare e con quale livello di "
    "autonomia si e' pronti a operare."
)

para(doc,
    "Per uno studio professionale o una PMI che usa Office 365 Business, "
    "il percorso verso le funzionalita' agentiche e' piu' graduale. "
    "Le funzionalita' di base di Copilot, come l'assistenza nella redazione "
    "di documenti, la sintesi di email e riunioni, la ricerca nei file "
    "aziendali, sono accessibili con le licenze standard. Le funzionalita' "
    "agentiche avanzate richiedono licenze piu' elevate. Vale la pena "
    "fare una valutazione concreta: identificare i tre processi che "
    "consumano piu' tempo nell'attivita' quotidiana, verificare se "
    "le funzionalita' agentiche di Copilot li coprono, e calcolare "
    "se il risparmio di tempo giustifica il costo incrementale della licenza."
)

heading(doc, "La domanda che vale la pena porsi prima di aggiornare la licenza")

para(doc,
    "Prima di aggiornare qualsiasi licenza, c'e' una domanda da porsi: "
    "i processi per cui si vuole usare il sistema agentico sono sufficientemente "
    "strutturati da poter essere delegati a un agente? Un agente AI funziona "
    "bene su processi con regole chiare, input definiti e output verificabili. "
    "Funziona male su processi che dipendono molto dal contesto relazionale, "
    "dalla conoscenza implicita o dal giudizio professionale in situazioni "
    "nuove. La transizione verso l'AI agentica non e' una questione di "
    "aggiornare la tecnologia: e' una questione di decidere consapevolmente "
    "quali compiti si e' pronti a delegare e quali si vuole mantenere in "
    "mano alle persone. Quella decisione vale piu' di qualsiasi licenza."
)

riferimenti(doc, [
    "Microsoft — AI Tour 2026: Frontier Transformation e nuove funzionalita' agentiche di Copilot (marzo 2026)",
    "Industria Italiana — 'Microsoft Frontier Transformation: Copilot diventa agentico, nasce la suite E7'",
    "Microsoft Source EMEA — 'Microsoft accelera la Frontier Transformation all'AI Tour 2026' (marzo 2026)",
    "HDBlog Business — 'Microsoft 365 Copilot: un anno di successi e crescita per l'AI in Italia'",
    "DeepElse — 'Microsoft Copilot per aziende italiane: vale la pena?'",
    "Il Sole 24 Ore — 'Microsoft accelerates AI agentica: all the new features of 365 Copilot'",
    "Osservatorio Digital Innovation, Politecnico di Milano — AI nelle grandi imprese italiane 2026",
])
doc.save(BASE + "2026-06-12_microsoft-copilot-agentico-aziende-italiane.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
