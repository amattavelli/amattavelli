"""
Quattro articoli Ratio -- 10 luglio 2026

1. 2026-07-10_gpt56-sol-terra-luna-quale-scegliere-professionisti.docx
2. 2026-07-10_ai-literacy-23-giorni-cosa-documentare-entro-agosto.docx
3. 2026-07-10_siracusa-allucinazioni-ai-responsabilita-professionale.docx
4. 2026-07-10_dati-cliente-chatgpt-gdpr-studio-professionale.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# GPT-5.6: Sol, Terra o Luna, quale scegliere
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Strumenti e Modelli AI")
titolo(
    doc,
    "GPT-5.6 e' uscito ieri:\nSol, Terra o Luna, e come scegliere quello giusto.",
    "Il 9 luglio 2026 OpenAI ha rilasciato tre modelli contemporaneamente, "
    "con nomi, prezzi e capacita' distinti. Per chi usa gia' ChatGPT nel "
    "lavoro quotidiano, la questione non e' se aggiornare: e' capire quale "
    "versione serve davvero, per quali compiti, e se vale il costo aggiuntivo.",
    "A cura della Redazione Ratio  •  10 luglio 2026"
)

para(doc,
    "Un consulente fiscale ha aperto ChatGPT ieri mattina e ha trovato tre "
    "modelli dove prima ce n'era uno scelto automaticamente dal sistema. Sol, "
    "Terra e Luna: nomi che sembrano le tappe di una missione spaziale, ma che "
    "OpenAI ha scelto per rappresentare tre livelli di capacita' e costo molto "
    "diversi tra loro. La reazione piu' comune in questi casi e' duplice: c'e' "
    "chi vuole subito il modello piu' potente senza valutarne il costo, e chi "
    "ignora il cambiamento sperando che le cose rimangano come prima. GPT-5.6, "
    "reso disponibile al pubblico il 9 luglio 2026 dopo un'anteprima di fine "
    "giugno, introduce pero' una distinzione che vale la pena capire, perche' "
    "cambia concretamente le scelte di configurazione per chi usa questi "
    "strumenti nel lavoro quotidiano."
)

para(doc,
    "La famiglia GPT-5.6 e' composta da tre modelli costruiti con obiettivi "
    "distinti. Luna e' il modello piu' rapido e meno costoso: risponde in "
    "frazioni di secondo, gestisce compiti di scrittura e sintesi di routine "
    "con ottimi risultati e costa un dollaro per milione di token in input e "
    "sei in output. Terra e' il modello di equilibrio, con capacita' di "
    "ragionamento piu' approfondite rispetto a Luna e un costo intermedio, "
    "due dollari e mezzo in input e quindici in output: copre la grande "
    "maggioranza dei compiti professionali che richiedono analisi di documenti "
    "di media lunghezza, strutturazione di pareri, ricerca su temi consolidati. "
    "Sol e' il modello flagship, pensato per il ragionamento complesso, il "
    "lavoro agentivo di lunga durata e i flussi che richiedono il coordinamento "
    "di piu' passaggi: cinque dollari in input, trenta in output."
)

para(doc,
    "Tutti e tre condividono la stessa finestra di contesto, un milione di "
    "token, e lo stesso limite di output massimo, 128.000 token. La data di "
    "conoscenza e' identica per tutti: 16 febbraio 2026. La differenza non "
    "e' nella capacita' di gestire documenti lunghi: e' nella profondita' del "
    "ragionamento applicato a quei documenti. Per chi ha un piano ChatGPT Pro "
    "o Team, l'accesso ai tre modelli e' gia' incluso nel canone mensile, con "
    "limiti d'uso differenziati: Sol e' disponibile ma con quote piu' "
    "restrittive rispetto a Terra e Luna. Per chi usa l'API per sviluppare "
    "applicazioni o automatizzare processi, il differenziale di costo tra Sol "
    "e Luna su volumi alti diventa una variabile di progetto concreta."
)

heading(doc, "Sol, Terra o Luna: la mappa per scegliere senza sprechi")

para(doc,
    "La regola piu' utile per i professionisti che devono decidere quale modello "
    "usare per quali compiti e' partire dalla natura del task, non dalla fascia "
    "di prezzo. Luna copre la redazione di comunicazioni standard, la "
    "formattazione di documenti, la sintesi di testi brevi, le risposte a "
    "quesiti ricorrenti e la produzione di bozze destinate a revisione. Terra "
    "copre l'analisi di contratti di media lunghezza, la strutturazione di "
    "pareri su tematiche consolidate, la ricerca su normativa non recentissima "
    "e la revisione critica di documenti redatti da altri. Sol serve per i "
    "compiti in cui si vuole che il modello ragioni in modo autonomo su problemi "
    "aperti, esplori piu' ipotesi prima di proporre una risposta, o coordini "
    "sottoagenti paralleli per completare flussi di lavoro articolati. Per uno "
    "studio legale o fiscale di medie dimensioni, Sol e' probabilmente utile "
    "per meno del venti per cento dei compiti correnti."
)

para(doc,
    "Sol introduce due modalita' di inferenza aggiuntive rispetto ai modelli "
    "precedenti. La prima e' il max reasoning, che consente al modello di "
    "dedicare piu' tempo alla deliberazione su un singolo problema, esplorando "
    "piu' percorsi logici prima di convergere su una risposta. La seconda e' "
    "l'ultra mode, che coordina piu' subagenti in parallelo per completare "
    "flussi di lavoro articolati: il modello scompone il task in sottocompiti, "
    "li assegna a istanze parallele e ricombina i risultati. Questa seconda "
    "modalita' e' la novita' piu' significativa per le aziende che lavorano "
    "su processi automatizzati: permette di costruire pipeline di analisi che "
    "in precedenza richiedevano infrastrutture piu' complesse. Per un uso "
    "professionale di tipo consulenziale, max reasoning e' la modalita' piu' "
    "rilevante, non ultra mode."
)

heading(doc, "Il rischio di affidarsi al modello sbagliato per il motivo sbagliato")

para(doc,
    "Un errore che emerge spesso nei primi giorni dopo il lancio di una nuova "
    "generazione di modelli e' aspettarsi che un sistema piu' potente produca "
    "risultati migliori con gli stessi prompt usati in precedenza. Con GPT-5.6 "
    "questo vale in parte, ma non in modo automatico. Sol in ultra mode "
    "applicato a un task che non lo richiede produce un output ampio, costoso "
    "in termini di crediti, e non necessariamente piu' utile di una risposta "
    "di Terra. Al contrario, Luna applicato a un'analisi normativa che richiede "
    "ragionamento comparativo tra norme di anni diversi produce spesso un "
    "risultato approssimativo che il professionista deve ricostruire da capo. "
    "La scelta del modello non sostituisce la costruzione di un prompt preciso: "
    "li combina, e il secondo conta quanto il primo."
)

para(doc,
    "Per chi vuole un punto di partenza concreto, il suggerimento e' testare "
    "Terra sui compiti professionali abituali nelle prossime due settimane e "
    "usare Sol solo nei casi in cui Terra mostra limiti evidenti. Il confronto "
    "diretto, non la scheda tecnica dei benchmark, e' il modo piu' affidabile "
    "per capire dove il differenziale di qualita' giustifica il differenziale "
    "di costo. OpenAI ha costruito una famiglia che offre granularita' reale: "
    "usarla con criterio e' il lavoro che spetta al professionista, non all'AI."
)

riferimenti(doc, [
    "OpenAI -- 'GPT-5.6: Frontier intelligence that scales with your ambition' (9 luglio 2026) -- openai.com/index/gpt-5-6/",
    "OpenAI -- 'Previewing GPT-5.6 Sol: a next-generation model' -- openai.com/index/previewing-gpt-5-6-sol/",
    "MarkTechPost -- 'OpenAI Releases GPT-5.6 (Sol, Terra, Luna): A Three-Tier Model Family' (9 luglio 2026)",
    "Simon Willison -- 'The new GPT-5.6 family: Luna, Terra, Sol' -- simonwillison.net (9 luglio 2026)",
    "QCode.cc -- 'GPT-5.6 Sol, Terra & Luna: Benchmarks, Pricing & Access (GA July 2026)'",
    "TechTimes -- 'GPT-5.6 Goes Public Today: Sol, Terra, Luna and the Return of Base Model Wars' (9 luglio 2026)",
    "CNBC -- 'OpenAI expanding GPT-5.6 AI model release' (8 luglio 2026)",
    "GitHub Changelog -- 'OpenAI GPT-5.6 Sol, Terra, and Luna now available in GitHub Copilot' (9 luglio 2026)",
])
doc.save(BASE + "2026-07-10_gpt56-sol-terra-luna-quale-scegliere-professionisti.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# AI Literacy: 23 giorni alla scadenza, cosa documentare
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Normativa AI")
titolo(
    doc,
    "L'obbligo di AI literacy scade tra 23 giorni:\ncosa deve avere in carta la tua organizzazione.",
    "Dal 2 agosto 2026 le autorita' di vigilanza avranno poteri ispettivi "
    "pieni sull'articolo 4 dell'AI Act, che impone alle aziende di documentare "
    "la formazione sull'AI dei propri collaboratori. Non basta aver fatto una "
    "riunione: serve evidenza concreta. Ecco cosa cercano i controllori e "
    "come prepararsi in meno di un mese.",
    "A cura della Redazione Ratio  •  10 luglio 2026"
)

para(doc,
    "Una societa' di consulenza con dodici collaboratori ha ricevuto a giugno "
    "una comunicazione dell'associazione di categoria che ricordava la scadenza "
    "del 2 agosto per l'obbligo di AI literacy previsto dall'AI Act. Il "
    "responsabile dello studio ha risposto che avevano gia' fatto una riunione "
    "interna sull'intelligenza artificiale a marzo. Quando gli e' stato chiesto "
    "se avesse un attestato di partecipazione per ogni collaboratore, un registro "
    "della formazione e una policy aziendale scritta sull'uso degli strumenti AI, "
    "la risposta e' stata che probabilmente avevano abbastanza, ma non era "
    "sicuro. Quella risposta riflette la situazione di migliaia di PMI italiane "
    "e studi professionali: hanno fatto qualcosa, ma non in modo documentato, "
    "e tra 23 giorni la documentazione e' cio' che conta."
)

para(doc,
    "L'articolo 4 del Regolamento UE 2024/1689 stabilisce che i fornitori e "
    "i deployer di sistemi AI, cioe' le aziende e gli studi che li utilizzano "
    "nell'attivita' operativa, devono adottare misure per garantire un livello "
    "sufficiente di AI literacy alle persone che gestiscono o operano questi "
    "sistemi per loro conto. La norma non fissa un monte ore minimo, non "
    "prescrive un programma specifico e non impone una certificazione esterna. "
    "Chiede che l'organizzazione abbia affrontato il tema in modo proporzionato "
    "alla propria dimensione e al tipo di AI che utilizza, e che sia in grado "
    "di dimostrarlo. Dal 2 agosto le autorita' nazionali di vigilanza avranno "
    "i poteri ispettivi e sanzionatori pieni, e la domanda che puo' arrivare "
    "in occasione di un controllo e' semplice: cosa ha fatto la vostra "
    "organizzazione per formare chi usa l'AI?"
)

para(doc,
    "La parola proporzionale e' la chiave del meccanismo. Un'azienda di dieci "
    "persone che usa ChatGPT per scrivere email commerciali e sintetizzare "
    "documenti non deve costruire lo stesso sistema formativo di un operatore "
    "sanitario che usa AI per il supporto alla diagnosi. Il rischio dell'attivita', "
    "l'impatto delle decisioni supportate o automatizzate su persone terze, "
    "e il numero di collaboratori coinvolti determinano il livello di presidio "
    "formativo richiesto. Per la grande maggioranza delle PMI italiane e degli "
    "studi professionali, un percorso formativo di alcune ore per tutti i "
    "collaboratori che usano strumenti AI, documentato e accompagnato da una "
    "policy interna sull'uso, copre l'obbligo."
)

heading(doc, "I tre elementi che le autorita' cercheranno in un controllo")

para(doc,
    "Chi lavora con esperti di compliance e' concorde su tre elementi minimi "
    "che un'organizzazione deve avere per dimostrare la conformita' all'articolo 4. "
    "Il primo e' la traccia della formazione: un attestato nominativo per ogni "
    "collaboratore che ha partecipato a un corso o a una sessione interna, con "
    "data, durata e contenuto. Non basta un link a un video inviato via email "
    "o una comunicazione interna che invita alla lettura di un documento. "
    "Il secondo e' un registro interno che riporti chi ha ricevuto formazione, "
    "quando e su che cosa: anche un foglio di calcolo con nomi, date e titolo "
    "del corso soddisfa questo requisito se compilato con precisione. "
    "Il terzo e' una policy aziendale sull'uso degli strumenti AI, un documento "
    "anche breve che definisca quali strumenti sono autorizzati, per quali "
    "compiti e con quali limiti."
)

para(doc,
    "La policy non deve essere un manuale tecnico. Deve rispondere a tre "
    "domande: quali strumenti AI possono essere usati nell'attivita' lavorativa, "
    "quali categorie di dati non vanno mai inseriti in questi strumenti (dati "
    "personali di clienti, informazioni riservate, credenziali di accesso), "
    "e chi e' responsabile di verificare che le regole vengano rispettate. "
    "Una policy cosi' strutturata richiede meno di mezza giornata di lavoro "
    "per essere redatta, ma deve essere formalmente adottata dall'organo di "
    "gestione e comunicata a tutti i collaboratori interessati. Senza questa "
    "formalizzazione, rimane un documento interno senza valore probatorio "
    "in caso di contestazione."
)

heading(doc, "Cosa fare concretamente nei prossimi 23 giorni")

para(doc,
    "Il percorso minimo per le organizzazioni che non hanno ancora affrontato "
    "il tema si articola in quattro passi da completare prima del 2 agosto. "
    "Il primo e' un censimento degli strumenti AI effettivamente in uso: "
    "non solo quelli acquistati dall'azienda, ma anche quelli usati "
    "autonomamente dai singoli collaboratori per compiti lavorativi, il "
    "cosiddetto shadow AI. Molte PMI scopriranno in questa fase che l'uso "
    "e' piu' diffuso di quanto il management pensasse. Il secondo passo e' "
    "identificare chi usa quali strumenti e per quali processi. "
    "Il terzo e' organizzare una sessione formativa, anche di due ore, che "
    "copra i concetti base del funzionamento dei sistemi AI, i rischi principali "
    "come le allucinazioni e la gestione dei dati, e le regole di utilizzo "
    "interne. Il quarto passo e' raccogliere gli attestati e redigere la policy."
)

para(doc,
    "Esistono corsi online certificati da poche decine di euro che rilasciano "
    "attestati nominativi e coprono i contenuti richiesti. Per i professionisti "
    "iscritti a ordini che riconoscono crediti formativi, vale la pena verificare "
    "se esistono moduli sull'AI gia' accreditati: si recupera la formazione "
    "obbligatoria dell'ordine e si soddisfa l'obbligo AI Act nello stesso momento. "
    "La scadenza del 2 agosto e' anche un'occasione per fare un censimento "
    "reale di cosa si usa in azienda e con quali criteri. Chi la tratta come "
    "un semplice adempimento ne trae un beneficio minimo. Chi la usa per "
    "costruire un presidio ragionato sull'uso dell'AI ottiene anche un vantaggio "
    "operativo: sa cosa fanno i suoi collaboratori con questi strumenti, "
    "e puo' governarlo."
)

riferimenti(doc, [
    "Regolamento UE 2024/1689 (AI Act), articolo 4 -- EUR-Lex",
    "AgendaDigitale.eu -- 'Formazione AI obbligatoria per le imprese: guida completa'",
    "AntonioSinibaldi.com -- 'AI literacy e obbligo formativo: cosa fare entro agosto 2026'",
    "AntonioSinibaldi.com -- 'AI literacy: cos'e', perche' e' obbligatoria e come formare i dipendenti prima del 2 agosto 2026'",
    "Randstad.it -- 'AI literacy e obbligo di formazione nelle aziende: come muoversi'",
    "ICDQ.it -- 'AI Act: la formazione sull'intelligenza artificiale non e' piu' facoltativa'",
    "InforeLeA Academy -- 'AI Act dal 2 agosto: iniziano i controlli, cosa devono fare le aziende'",
    "Tinexta Innovation Hub -- 'AI Act: i passi necessari per la conformita' entro il 2 agosto 2026'",
])
doc.save(BASE + "2026-07-10_ai-literacy-23-giorni-cosa-documentare-entro-agosto.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# Siracusa: allucinazioni AI e responsabilita' professionale
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Professione e Responsabilita'")
titolo(
    doc,
    "Le citazioni inventate dall'AI in un atto giudiziario:\ncosa ha stabilito il Tribunale di Siracusa.",
    "Con la sentenza del 20 febbraio 2026, il Tribunale di Siracusa ha "
    "qualificato come colpa grave l'inserimento in un atto difensivo di "
    "citazioni giurisprudenziali prodotte da un sistema AI e non verificate "
    "sulla fonte originale. Il principio vale per avvocati, ma si estende "
    "a tutti i professionisti che sottoscrivono documenti elaborati con "
    "il supporto dell'intelligenza artificiale.",
    "A cura della Redazione Ratio  •  10 luglio 2026"
)

para(doc,
    "Nel gennaio 2026, un atto difensivo depositato davanti al Tribunale di "
    "Siracusa conteneva quattro citazioni di sentenze della Corte di Cassazione, "
    "complete di massime e brani riportati tra virgolette come citazioni "
    "letterali. Il giudice, insospettito da alcune formulazioni inusuali nella "
    "terminologia giuridica, ha verificato le sentenze indicate. Le citazioni "
    "non corrispondevano: alcuni passaggi non esistevano affatto, altri erano "
    "attribuiti a decisioni diverse da quelle indicate. L'avvocato aveva usato "
    "un sistema di intelligenza artificiale generativa per la ricerca "
    "giurisprudenziale, aveva copiato i risultati nell'atto senza verificarli "
    "sulle fonti primarie e aveva sottoscritto il documento. Con la sentenza "
    "del 20 febbraio 2026 il Tribunale di Siracusa ha qualificato quella "
    "condotta come colpa grave, aprendo un fronte che interessa l'intera "
    "classe professionale, non solo gli avvocati."
)

para(doc,
    "La decisione siracusana si inserisce in un filone giurisprudenziale che "
    "si sta consolidando a livello europeo e che la magistratura italiana ha "
    "cominciato ad affrontare con una linea interpretativa coerente. Il "
    "principio nella sua formulazione e' chiaro: la sottoscrizione di un atto "
    "professionale comporta la piena responsabilita' del professionista su "
    "tutto il contenuto, indipendentemente da quale strumento sia stato usato "
    "per produrlo. Quando il professionista usa un sistema AI per la ricerca "
    "o la redazione e inserisce i risultati in un documento firmato senza "
    "verifica autonoma, si configura un uso acritico dello strumento che "
    "la giurisprudenza qualifica come negligenza grave."
)

para(doc,
    "Il caso delle citazioni inventate dai sistemi AI, definito tecnicamente "
    "come allucinazione giurisprudenziale, non e' una rarità. I modelli di "
    "linguaggio producono testi plausibili, non testi verificati: quando vengono "
    "interrogati su casi giurisprudenziali specifici, specialmente su decisioni "
    "meno note o su materie con giurisprudenza frammentata, tendono a costruire "
    "citazioni verosimili ma non reali, combinando elementi di sentenze diverse "
    "o inventando passaggi che non esistono. Un professionista che riceve "
    "questa risposta e non la verifica sulla fonte primaria sta esponendo il "
    "proprio cliente e se stesso a un rischio che cresce con la complessita' "
    "del caso trattato."
)

heading(doc, "Come cambia la valutazione della diligenza professionale")

para(doc,
    "Per avvocati e commercialisti, la sentenza di Siracusa ha un'implicazione "
    "che va oltre il caso specifico delle citazioni. Il criterio della diligenza "
    "professionale qualificata, che l'ordinamento italiano richiede a chiunque "
    "eserciti una professione intellettuale regolamentata, include ora la "
    "verifica dell'adeguatezza degli strumenti usati e la revisione critica "
    "dei loro output. Non e' sufficiente dimostrare che uno strumento e' "
    "ampiamente diffuso o che altri professionisti lo utilizzano: serve "
    "dimostrare che il professionista ha applicato il proprio giudizio al "
    "contenuto prodotto dallo strumento, verificandone l'accuratezza nei punti "
    "critici. Per un commercialista che usa l'AI per strutturare un parere "
    "fiscale, questo significa controllare le aliquote sulla fonte ufficiale, "
    "verificare i riferimenti normativi su Normattiva o EUR-Lex e validare "
    "i calcoli sui sistemi certificati. Per un avvocato, significa riscontrare "
    "ogni citazione giurisprudenziale sul database originale prima di inserirla "
    "in un atto."
)

para(doc,
    "La Legge 132/2025, entrata in vigore nel settembre 2025, ha aggiunto un "
    "ulteriore livello normativo a questa questione. L'articolo dedicato alla "
    "responsabilita' dei professionisti che usano sistemi AI stabilisce che "
    "l'utilizzo di strumenti di intelligenza artificiale nella prestazione "
    "professionale non riduce la responsabilita' del professionista nei "
    "confronti del cliente: se l'output dell'AI e' errato e il professionista "
    "non lo ha verificato con la dovuta diligenza, risponde come se l'errore "
    "fosse interamente suo, perche' in senso giuridico lo e'. La responsabilita' "
    "non si trasferisce al fornitore dello strumento, salvo casi di difetto "
    "del prodotto che richiedono una prova separata e di difficile costruzione."
)

heading(doc, "Costruire un metodo di revisione che regge in sede disciplinare")

para(doc,
    "Gli studi che hanno affrontato piu' seriamente il tema della responsabilita' "
    "AI hanno adottato flussi di revisione che distinguono nettamente tra l'uso "
    "dell'AI come strumento di ricerca e l'uso dell'AI come fornitore di "
    "contenuto finale. Nel primo caso, il sistema suggerisce piste, identifica "
    "norme potenzialmente pertinenti e propone strutture argomentative: tutto "
    "questo viene poi verificato dal professionista sulle fonti primarie prima "
    "di essere incorporato nell'elaborato finale. Nel secondo caso, il "
    "professionista accetta il testo dell'AI senza revisione autonoma e lo "
    "incorpora direttamente nell'atto o nel parere: e' qui che si materializza "
    "il rischio identificato dalla sentenza di Siracusa. La distinzione non "
    "dipende dalla qualita' del modello usato: dipende dal processo di lavoro "
    "del professionista."
)

para(doc,
    "Alcune misure pratiche si sono affermate come standard tra i professionisti "
    "piu' attenti. Le citazioni giurisprudenziali vengono sistematicamente "
    "verificate su database certificati come Dejure, Italgiure o le banche "
    "dati degli ordini, prima di essere inserite in qualsiasi documento. "
    "I riferimenti normativi vengono confrontati con il testo ufficiale, "
    "non lasciati alla formulazione proposta dall'AI. I calcoli fiscali prodotti "
    "da strumenti generativi vengono validati sui parametri chiave, in "
    "particolare quando riguardano aliquote, scadenze o regole di competenza "
    "territoriale oggetto di modifiche recenti. Queste verifiche rallentano "
    "il processo rispetto a una delega totale all'AI, ma molto meno di quanto "
    "rallenti la gestione di un contenzioso con il cliente per un errore "
    "non rilevato in tempo."
)

para(doc,
    "La sentenza di Siracusa non chiede ai professionisti di smettere di "
    "usare l'AI. Chiede di usarla in modo compatibile con gli standard di "
    "diligenza che l'ordinamento ha sempre richiesto. Chi costruisce adesso "
    "un metodo di revisione documentato, con passaggi di verifica tracciabili, "
    "si mette in una posizione difensiva solida se la questione dovesse "
    "emergere in sede disciplinare o in un contenzioso con il cliente. "
    "Chi aspetta che il problema si manifesti nel proprio dossier avra' "
    "meno tempo e meno argomenti a disposizione."
)

riferimenti(doc, [
    "Legge 9 settembre 2025, n. 132 (Legge sull'Intelligenza Artificiale) -- Normattiva",
    "Tribunale di Siracusa, sentenza 20 febbraio 2026 (responsabilita' professionale e AI)",
    "PaganiniBellini.it -- 'Intelligenza artificiale e avvocati: quando l'errore diventa responsabilita' professionale'",
    "Altalex.com -- 'Responsabilita' del professionista per uso improvvido dell'IA' (aprile 2026)",
    "DirittoDellaInformazione.it -- 'AI e professione forense: quando il software sbaglia risponde l'avvocato'",
    "DirittoEGiustizia.it -- 'Chi risponde se l'AI sbaglia?'",
    "ConvieneOnline.it -- 'AI e avvocati: responsabilita' e colpa grave (sentenza 2026)'",
    "Regolamento UE 2024/1689 (AI Act), considerando 47 e articoli sulla responsabilita' -- EUR-Lex",
])
doc.save(BASE + "2026-07-10_siracusa-allucinazioni-ai-responsabilita-professionale.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Dati del cliente in ChatGPT: GDPR e piani business
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Privacy e Dati")
titolo(
    doc,
    "I dati del cliente in ChatGPT:\ncosa permette il piano business e cosa rischia lo studio.",
    "Migliaia di professionisti italiani usano ogni giorno strumenti AI con "
    "documenti che contengono dati personali dei clienti. Pochi sanno che il "
    "piano individuale di ChatGPT non include un accordo sul trattamento dei "
    "dati conforme al GDPR. La differenza tra un abbonamento e l'altro non e' "
    "solo di prezzo: e' di responsabilita' legale.",
    "A cura della Redazione Ratio  •  10 luglio 2026"
)

para(doc,
    "Uno studio di consulenza aziendale ha ricevuto a maggio una richiesta "
    "di informazioni dal Garante della Privacy, avviata su segnalazione di "
    "un ex dipendente. L'oggetto dell'accertamento riguardava l'uso di "
    "ChatGPT per elaborare documenti che contenevano dati personali di "
    "clienti: contratti, analisi finanziarie, corrispondenza con istituti "
    "di credito. Lo studio usava il piano individuale di ChatGPT, non il "
    "piano Team o Enterprise, e caricava i documenti come allegati alle "
    "conversazioni per chiedere sintesi e analisi. Il Garante ha rilevato "
    "che i dati personali trattati in questo modo erano trasferiti a OpenAI "
    "senza una base giuridica adeguata e senza che i clienti fossero stati "
    "informati del trattamento. La questione non e' tecnica: e' contrattuale."
)

para(doc,
    "Il discrimine tra un uso lecito e un uso potenzialmente problematico "
    "degli strumenti AI nel contesto professionale italiano dipende da un "
    "elemento spesso trascurato: il tipo di piano sottoscritto con il "
    "fornitore. OpenAI, Anthropic e Google hanno strutture contrattuali "
    "diverse per gli utenti individuali e per quelli business, e queste "
    "differenze hanno conseguenze dirette sul trattamento dei dati. "
    "Il piano individuale di ChatGPT, quello che si attiva con l'abbonamento "
    "standard da circa venti euro al mese, non include di default un "
    "Data Processing Agreement (DPA) con l'utente: il trattamento dei dati "
    "inseriti nelle conversazioni e' regolato dai termini di servizio generali, "
    "non da un accordo specifico che soddisfi i requisiti del GDPR per "
    "il trattamento dati per conto terzi."
)

para(doc,
    "Il piano ChatGPT Team e il piano ChatGPT Enterprise includono invece "
    "un DPA standard che qualifica OpenAI come responsabile del trattamento "
    "ai sensi dell'articolo 28 del GDPR. Questo significa che i dati inseriti "
    "nelle conversazioni non vengono usati di default per addestrare i modelli, "
    "che OpenAI si impegna a trattarli secondo le istruzioni del titolare "
    "(l'azienda o lo studio), e che esiste una base contrattuale formale "
    "per il trasferimento verso infrastrutture extra-UE. La stessa struttura "
    "vale per Claude for Teams e Claude Enterprise di Anthropic, e per "
    "Gemini Business e Google Workspace di Google. Chi usa le versioni "
    "individuali di questi strumenti con dati di terzi opera in una zona "
    "grigia che, come mostra il caso del Garante, puo' diventare oggetto "
    "di contestazione."
)

heading(doc, "Cosa non va mai inserito in un piano AI senza DPA")

para(doc,
    "La regola pratica piu' sicura per chi usa strumenti AI nell'attivita' "
    "professionale e' non inserire mai dati personali di clienti in sistemi "
    "privi di DPA, indipendentemente da quanto i dati sembrino anonimizzati. "
    "Anche un documento in cui vengono sostituiti i nomi con iniziali o codici "
    "puo' contenere elementi che rendono identificabile la persona fisica: "
    "la combinazione di localita', importi specifici, date e tipologia di "
    "operazione e' spesso sufficiente per la re-identificazione. "
    "Il rischio e' duplice: da un lato il trasferimento di dati personali "
    "senza base giuridica adeguata, dall'altro l'impossibilita' di rispettare "
    "i diritti dell'interessato in materia di accesso, cancellazione e "
    "portabilita' su dati che l'azienda non controlla, perche' sono finiti "
    "nei sistemi di un fornitore estero senza accordo contrattuale formale."
)

para(doc,
    "La soluzione non e' evitare gli strumenti AI nel lavoro con i clienti. "
    "Per gli studi che usano ChatGPT regolarmente con documenti dei clienti, "
    "il passaggio al piano Team costa intorno a trenta euro per utente al mese "
    "e include il DPA: il differenziale rispetto al piano individuale e' modesto "
    "rispetto al rischio che elimina. Per chi preferisce strumenti con "
    "data residency europea, Microsoft Copilot integrato in Microsoft 365 "
    "gestisce i dati su infrastruttura cloud europea per i tenant configurati "
    "in UE, e Claude for Teams di Anthropic prevede clausole analoghe. "
    "La scelta dipende anche dal tipo di uso: chi sintetizza solo documenti "
    "pubblici o testi privi di dati personali ha un profilo di rischio "
    "molto diverso da chi elabora fascicoli clienti con informazioni "
    "finanziarie o giuridiche."
)

heading(doc, "L'informativa ai clienti e la Legge 132/2025")

para(doc,
    "Al quadro GDPR si sovrappone l'obbligo introdotto dalla Legge 132/2025, "
    "che impone ai professionisti che esercitano attivita' regolamentate di "
    "informare il cliente quando strumenti di intelligenza artificiale vengono "
    "usati nell'elaborazione di pareri, perizie o consulenze. Questo obbligo "
    "e' distinto da quello privacy: riguarda la trasparenza sull'uso dell'AI "
    "come strumento professionale, non il solo trattamento dei dati personali. "
    "In pratica, la maggior parte degli studi sta inserendo una clausola "
    "standard nell'incarico professionale o nel preventivo che specifica "
    "l'utilizzo di strumenti di supporto digitale basati su intelligenza "
    "artificiale per alcune fasi del lavoro, che questi sistemi operano su "
    "infrastrutture conformi al GDPR e che i dati del cliente sono trattati "
    "secondo le finalita' indicate nell'informativa privacy allegata all'incarico."
)

para(doc,
    "La verifica da fare adesso, prima della scadenza di agosto, e' semplice: "
    "quale piano hanno i collaboratori dello studio per gli strumenti AI che "
    "usano piu' frequentemente? Se la risposta e' il piano individuale, "
    "o peggio non lo si sa, serve un'azione concreta. Aggiornare il piano "
    "e aggiungere una clausola nell'incarico professionale richiede meno di "
    "un pomeriggio di lavoro. Affrontarlo dopo un accertamento del Garante "
    "richiede molto di piu': tempo, risorse legali e una risposta a domande "
    "che nessuno vuole ricevere da un cliente che non sapeva di essere "
    "nel mezzo di un accertamento sulla privacy."
)

riferimenti(doc, [
    "Regolamento UE 2016/679 (GDPR), articoli 28 e 46 -- EUR-Lex",
    "Legge 9 settembre 2025, n. 132 (Legge sull'Intelligenza Artificiale) -- Normattiva",
    "Garante per la Protezione dei Dati Personali -- Provvedimento su ChatGPT (2023 e aggiornamenti 2025-2026)",
    "OpenAI -- 'Data processing agreement for ChatGPT Team and Enterprise' -- openai.com",
    "Anthropic -- 'Claude for Teams: privacy and data processing' -- anthropic.com",
    "Compliance-GDPR-LLM.vercel.app -- 'GDPR e IA 2025-2026: uso lecito in studi professionali e PMI'",
    "AgendaDigitale.eu -- 'Strumenti AI in azienda e GDPR: guida pratica 2026'",
    "IlGarante.it -- 'Intelligenza artificiale e protezione dei dati: domande frequenti'",
])
doc.save(BASE + "2026-07-10_dati-cliente-chatgpt-gdpr-studio-professionale.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
