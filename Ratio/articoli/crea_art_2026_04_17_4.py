#!/usr/bin/env python3
"""
Articolo 4: Automazione contabile AI
File: Ratio/articoli/2026-04-17_automazione-contabile-ai-dove-siamo.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/home/user/amattavelli/Ratio/articoli/2026-04-17_automazione-contabile-ai-dove-siamo.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(3)
sec.right_margin = Cm(3)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot)
    pPr.append(pBdr)


def para(doc, text, size=11, italic=False, color=None, sa=8, sb=0,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


# --- TESTATA ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.all_caps = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Contabilit\u00e0 e Digitale")
r.font.name = 'Calibri'
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# --- TITOLO ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "La fattura che si registra quasi da sola: "
    "a che punto \u00e8 davvero l\u2019automazione contabile"
)
r.font.name = 'Calibri'
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "I principali software per studi professionali integrano da mesi funzioni AI "
    "per la registrazione automatica delle fatture. Cosa funziona, cosa richiede "
    "ancora supervisione, e cosa cambia nel lavoro quotidiano."
)
r.font.name = 'Calibri'
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  \u2022  17 aprile 2026")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# --- CORPO ---
para(doc,
    "Se lavori in uno studio che ha adottato TeamSystem Studio, Genya o un sistema simile "
    "negli ultimi dodici mesi, probabilmente hai gi\u00e0 visto la funzione di registrazione "
    "automatica delle fatture. La demo del venditore mostrava un documento che si "
    "registrava nel giro di pochi secondi, con il conto corretto, il centro di costo "
    "giusto, il regime IVA appropriato. In molti casi, funziona cos\u00ec anche nella realt\u00e0. "
    "In altri casi, la situazione \u00e8 pi\u00f9 sfumata. Capire la differenza \u00e8 utile per "
    "calibrare le aspettative e usare questi strumenti in modo produttivo."
)

para(doc,
    "I sistemi AI per la contabilit\u00e0 automatica lavorano bene su fatture ricorrenti: "
    "stesso fornitore, stessa categoria merceologica, importi coerenti con la storia "
    "del cliente. In questi casi, l\u2019accuratezza della registrazione automatica si "
    "avvicina al 90-95% nei test interni dei principali fornitori. Il sistema ha gi\u00e0 "
    "visto quel fornitore centinaia di volte, conosce i conti associati, e replica il "
    "pattern con affidabilit\u00e0. La velocit\u00e0 \u00e8 reale, il risparmio di tempo \u00e8 misurabile."
)

para(doc,
    "Dove i sistemi iniziano a vacillare \u00e8 sulle situazioni non standard: primo accesso "
    "a un nuovo fornitore, fatture che includono componenti miste (parte manutenzione, "
    "parte acquisto bene), operazioni intracomunitarie con codici IVA specifici, note di "
    "credito con storni parziali. In questi casi, l\u2019AI tende a proporre la soluzione "
    "statisticamente pi\u00f9 probabile, che non \u00e8 necessariamente quella corretta per il "
    "caso specifico. Il professionista che sa riconoscere questi contesti interviene "
    "tempestivamente; quello che ha smesso di guardare le registrazioni rischia di "
    "accumulare errori silenziosi."
)

h2(doc, "Il vero cambiamento: dalla registrazione alla gestione delle eccezioni")

para(doc,
    "Il modello operativo che emerge dall\u2019automazione contabile non \u00e8 la sparizione "
    "del lavoro umano, ma il suo spostamento. Chi si occupava di registrare fatture "
    "ora si occupa prevalentemente di verificare le registrazioni automatiche, correggere "
    "le anomalie e gestire le situazioni che il sistema non sa classificare. \u00c8 un "
    "cambiamento significativo: richiede competenze analitiche, non operative. "
    "Il collaboratore che eccelleva nell\u2019inserimento rapido e preciso dei dati non \u00e8 "
    "necessariamente quello che eccelle nel riconoscere perch\u00e9 una registrazione "
    "automatica \u00e8 plausibile ma scorretta. Formare il team su questo tipo di lettura "
    "critica \u00e8 uno degli investimenti pi\u00f9 concreti che uno studio pu\u00f2 fare oggi."
)

para(doc,
    "I dati del mercato confermano questa direzione. Secondo le ultime rilevazioni, "
    "il 34% dei commercialisti italiani usa gi\u00e0 strumenti AI nella propria attivit\u00e0, "
    "con una proiezione di crescita al 72% entro i prossimi tre anni. I sondaggi mostrano "
    "che il tempo risparmiato nelle attivit\u00e0 di inserimento viene reinvestito in analisi, "
    "consulenza e relazione con il cliente. Il che corrisponde esattamente al tipo di "
    "valore aggiunto che giustifica la parcella di un professionista, e che i software "
    "non potranno mai replicare autonomamente."
)

para(doc,
    "Vale la pena anche considerare l\u2019impatto sul personale dello studio. Chi si \u00e8 "
    "formato su lavori di inserimento dati trover\u00e0 la transizione verso la gestione "
    "delle eccezioni naturalmente difficile, non perch\u00e9 manchi di capacit\u00e0 ma perch\u00e9 "
    "richiede un tipo diverso di attenzione. Il lavoro di inserimento \u00e8 ripetitivo ma "
    "prevedibile; la gestione delle anomalie richiede giudizio contestuale. Gli studi "
    "che hanno affrontato questa transizione con consapevolezza, dedicando tempo alla "
    "formazione e al ridisegno dei flussi interni, riferiscono risultati pi\u00f9 solidi "
    "rispetto a quelli che hanno semplicemente attivato la funzione AI e atteso che "
    "il team si adattasse da solo."
)

h2(doc, "La variabile spesso trascurata: la qualit\u00e0 dei documenti in ingresso")

para(doc,
    "La qualit\u00e0 dei dati in ingresso resta la variabile critica dell\u2019intero sistema. "
    "Uno studio che riceve documentazione cartacea scansionata male, fatture con layout "
    "non standard o file PDF immagine non leggibili dall\u2019OCR ottiene risultati scadenti "
    "anche dai migliori sistemi AI. Prima di investire in automazione contabile, vale la "
    "pena analizzare la qualit\u00e0 del flusso documentale con i clienti: spesso il collo di "
    "bottiglia non \u00e8 il software ma il modo in cui arrivano i documenti. Una conversazione "
    "con i clienti sulla qualit\u00e0 degli allegati che inviano pu\u00f2 migliorare l\u2019efficacia "
    "dell\u2019automazione pi\u00f9 di qualsiasi aggiornamento del gestionale."
)

para(doc,
    "L\u2019automazione contabile basata sull\u2019AI non elimina il bisogno di competenza: "
    "lo trasforma. Il valore professionale si sposta dalla capacit\u00e0 di registrare "
    "correttamente alla capacit\u00e0 di riconoscere dove l\u2019automazione \u00e8 affidabile e "
    "dove richiede supervisione. Chi impara a fare questa distinzione oggi costruisce "
    "un modo di lavorare pi\u00f9 efficiente e pi\u00f9 sostenibile. Il software non sar\u00e0 mai "
    "perfetto, ma sar\u00e0 progressivamente migliore: il professionista che lo conosce bene "
    "sa dove fidarsene e dove no."
)

sep(doc)

# --- FONTI ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Fonti e riferimenti")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "TeamSystem \u2014 AI e automazione contabile per studi commercialisti (2026)",
    "Datalog \u2014 Digitalizzazione dei commercialisti: evoluzioni e trend 2026",
    "Wolters Kluwer \u2014 L\u2019intelligenza artificiale trasforma la contabilit\u00e0: "
    "opportunit\u00e0 e sfide",
    "Professionista Digitale \u2014 Contabilit\u00e0 2026: trend e digitalizzazione "
    "studi professionali",
    "Scuola di Amministrazione Intrapresa \u2014 Le 5 novit\u00e0 pi\u00f9 importanti nel "
    "settore contabile aziendale nel 2026",
]
for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"\u2022 {s}")
    r.font.name = 'Calibri'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte")
r.font.name = 'Calibri'
r.font.size = Pt(8)
r.italic = True
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

doc.save(OUTPUT)
print(f"Salvato: {OUTPUT}")
