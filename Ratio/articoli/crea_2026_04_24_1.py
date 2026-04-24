"""
Articolo 1 — Il 2 agosto si avvicina
Ratio/articoli/2026-04-24_ai-act-agosto-2026-imprese-cento-giorni.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-24_ai-act-agosto-2026-imprese-cento-giorni.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Normativa e Compliance AI")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Il 2 agosto si avvicina: cosa devono fare le imprese\nche usano sistemi AI")
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "La piena applicazione dell’AI Act europeo è a meno di cento giorni. "
    "La maggioranza delle PMI italiane non ha ancora avviato l’inventario dei propri sistemi."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  24 aprile 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Se la tua azienda usa un chatbot per rispondere ai clienti, un sistema di analisi "
    "delle fatture, un software di selezione del personale che include un punteggio "
    "automatico, o anche solo un assistente virtuale integrato in Microsoft 365, il 2 agosto "
    "2026 è una data che ti riguarda direttamente. Da quel giorno entra in piena applicazione "
    "il Regolamento (UE) 2024/1689, l’AI Act europeo, con tutti gli obblighi relativi ai "
    "sistemi di intelligenza artificiale classificati ad alto rischio. Il livello di "
    "preparazione delle PMI italiane, stando ai dati disponibili, è preoccupante."
)

heading(doc, "Chi è coinvolto: più aziende di quanto si pensi")

para(doc,
    "Il punto critico che molte aziende sottovalutano riguarda il campo di applicazione. "
    "L’AI Act non riguarda solo chi sviluppa software: si applica anche a chi acquista e "
    "utilizza sistemi AI già pronti. Se il tuo gestionale HR valuta automaticamente le "
    "candidature, se il tuo software di credito assegna uno scoring ai clienti, se il tuo "
    "strumento di analisi dei dati prende decisioni con impatto sulle persone, la normativa "
    "ti riguarda come “deployer”, ovvero come utilizzatore finale. Molte PMI si sentono al "
    "riparo perché non sviluppano software. Questo è un errore di lettura della norma che può "
    "avere conseguenze concrete."
)

para(doc,
    "L’Italia ha recepito l’AI Act con la Legge 132/2025, entrata in vigore il 10 ottobre "
    "2025, che introduce una doppia compliance particolarmente rilevante per i professionisti. "
    "Oltre agli obblighi europei sui sistemi ad alto rischio, la legge italiana prevede già da "
    "subito un obbligo di informativa ai clienti sull’uso di sistemi AI nello svolgimento "
    "dell’incarico professionale. Si tratta di un adempimento a sé stante, indipendente dalla "
    "classificazione del sistema: qualunque professionista che usi strumenti AI nello studio "
    "(anche ChatGPT per redigere una lettera) deve informare il cliente in modo chiaro e "
    "comprensibile."
)

heading(doc, "Le scadenze concrete del secondo semestre 2026")

para(doc,
    "Le scadenze di questo secondo semestre si articolano su tre livelli operativi. Fino al "
    "2 agosto, le aziende devono completare l’inventario dei sistemi AI in uso, classificarli "
    "secondo i criteri del Regolamento (rischio inaccettabile, alto rischio, rischio limitato, "
    "rischio minimo) e avviare le procedure di conformità per i sistemi ad alto rischio. Dal "
    "2 agosto in poi, chi utilizza sistemi ad alto rischio deve garantire supervisione umana "
    "adeguata, conservare i log delle decisioni automatizzate, assicurare la qualità dei dati "
    "di input e disporre di documentazione tecnica aggiornata. Le sanzioni per la non "
    "conformità possono arrivare fino al 7% del fatturato globale annuo o 35 milioni di euro, "
    "con la soglia più alta che si applica tra le due."
)

para(doc,
    "Un aspetto che genera confusione pratica riguarda la classificazione. Non tutti i sistemi "
    "AI rientrano nell’alto rischio: la stragrande maggioranza degli strumenti generativi usati "
    "quotidianamente negli studi professionali (assistenti per la redazione di testi, riassunti "
    "di documenti, supporto alla ricerca normativa) ricade nella categoria a rischio limitato, "
    "con obblighi molto più leggeri, essenzialmente di trasparenza verso l’utente finale. I "
    "sistemi ad alto rischio sono quelli che incidono su decisioni con impatto significativo "
    "sulle persone: accesso al credito, selezione del personale, istruzione, sicurezza di "
    "prodotti, assistenza sanitaria. Per la maggior parte degli studi professionali il rischio "
    "principale è nella categoria limitato, non nell’alto rischio, ma occorre saperlo con "
    "certezza e non per ipotesi."
)

heading(doc, "Da dove iniziare: l’inventario come primo atto concreto")

para(doc,
    "Per la maggior parte degli studi professionali e delle PMI italiane, il percorso verso "
    "la conformità parte da un’azione concreta e non costosa: l’inventario. Si tratta di "
    "elencare tutti i sistemi AI in uso (compresi quelli integrati in software già acquistati), "
    "verificare se il fornitore li ha classificati secondo l’AI Act, richiedere la "
    "documentazione tecnica e valutare se esiste supervisione umana sui processi decisionali "
    "automatizzati. Questa attività, che richiede qualche giornata di lavoro con il supporto "
    "di un consulente esperto, è la base da cui tutto il resto dipende. Non si può gestire la "
    "conformità di sistemi che non si sa di avere."
)

para(doc,
    "Vale la pena ricordare che molti fornitori di software gestionale e di piattaforme "
    "business hanno già avviato i propri percorsi di conformità all’AI Act e mettono a "
    "disposizione dei clienti la documentazione necessaria. Il lavoro richiesto all’azienda "
    "utilizzatrice non è quindi duplicare quello del fornitore, ma verificare che la propria "
    "configurazione e il proprio uso del sistema siano conformi, e che la supervisione umana "
    "richiesta sia effettivamente garantita nei processi operativi."
)

heading(doc, "Cosa succede se non si è pronti il 2 agosto")

para(doc,
    "Chi arriva al 2 agosto senza aver svolto questa analisi si trova in una posizione di "
    "rischio operativo difficile da gestire rapidamente. L’AI Act, a differenza di molte "
    "normative precedenti, non prevede un lungo periodo di tolleranza dopo la scadenza. Le "
    "autorità nazionali competenti (in Italia, AGID e ACN per i profili tecnici, il Garante "
    "Privacy per quelli relativi ai dati personali) hanno già avviato le attività preparatorie "
    "per la vigilanza. Le prime verifiche potranno riguardare le categorie più esposte: "
    "aziende del settore finanziario, HR, sanitario e della pubblica amministrazione, ma il "
    "perimetro si allargherà progressivamente."
)

para(doc,
    "Cominciare adesso significa avere ancora il tempo di fare l’inventario con metodo, "
    "coinvolgere i fornitori nella raccolta della documentazione e, se necessario, adeguare "
    "i processi prima della scadenza, senza delegare a consulenti di urgenza decisioni che "
    "dovrebbero essere scelte strategiche. Chi aspetta luglio si troverà a fare in fretta "
    "ciò che avrebbe potuto fare bene."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "Regolamento (UE) 2024/1689 — AI Act (GU UE, 12 luglio 2024)",
    "Legge 132/2025 (Italia) — Disposizioni nazionali in materia di intelligenza artificiale",
    "AGID / ACN — Comunicazioni sulle autorità nazionali competenti AI Act",
    "Garante per la Protezione dei Dati Personali — Orientamenti sull’uso dell’AI",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
