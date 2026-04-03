"""
Articolo 4: Obbligo informativa AI per professionisti - Art. 13 L.132/2025
Ratio/articoli/2026-04-03_obbligo-informativa-ai-professionisti.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-03_obbligo-informativa-ai-professionisti.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_paragraph(doc, text, size=11, italic=False, color=None,
                  space_after=8, space_before=0,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

# ---- TESTATA ----
p_rivista = doc.add_paragraph()
p_rivista.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = p_rivista.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
run_r.font.name = 'Calibri'
run_r.font.size = Pt(9)
run_r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run_r.font.bold = True
run_r.font.all_caps = True

add_separator(doc)

p_data = doc.add_paragraph()
p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_d = p_data.add_run("Aprile 2026  |  Normativa e Studio Professionale")
run_d.font.name = 'Calibri'
run_d.font.size = Pt(8.5)
run_d.italic = True
run_d.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_after = Pt(6)
run_t = p_title.add_run(
    "Dal 10 ottobre 2025, chi usa l\u2019AI nello studio deve dirlo al cliente per iscritto"
)
run_t.font.name = 'Calibri'
run_t.font.size = Pt(24)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_after = Pt(10)
run_s = p_sub.add_run(
    "L\u2019articolo 13 della Legge 132/2025 \u00e8 gi\u00e0 in vigore. Molti studi non lo sanno ancora. "
    "Cosa prevede, quali rischi comporta ignorarlo, e come adeguarsi senza stravolgere i contratti d\u2019incarico."
)
run_s.font.name = 'Calibri'
run_s.font.size = Pt(13)
run_s.italic = True
run_s.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_separator(doc)

p_autore = doc.add_paragraph()
p_autore.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_autore.paragraph_format.space_after = Pt(16)
run_a = p_autore.add_run("A cura della Redazione Ratio  \u2022  3 aprile 2026")
run_a.font.name = 'Calibri'
run_a.font.size = Pt(9)
run_a.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- CORPO ----

add_paragraph(doc,
    "Dal 10 ottobre 2025 \u00e8 in vigore in Italia un obbligo che riguarda ogni professionista "
    "intellettuale \u2014 commercialista, avvocato, consulente del lavoro, ingegnere, architetto, "
    "notaio \u2014 che utilizzi strumenti di intelligenza artificiale nell\u2019esecuzione di un "
    "incarico professionale. L\u2019obbligo \u00e8 semplice: prima di iniziare il lavoro, il "
    "professionista deve dichiararlo al cliente per iscritto. Lo stabilisce l\u2019articolo 13 "
    "della Legge 23 settembre 2025, n. 132. A sei mesi dall\u2019entrata in vigore, la norma "
    "\u00e8 ancora sconosciuta alla maggior parte degli studi che usano quotidianamente ChatGPT, "
    "Claude o Gemini per redigere bozze, sintetizzare documenti, rispondere a quesiti normativi."
)

add_paragraph(doc,
    "Non si tratta di un adempimento burocratico minore. Se un professionista usa un sistema AI "
    "senza averlo dichiarato e il cliente subisce un danno riconducibile a un errore dell\u2019AI, "
    "l\u2019omissione dell\u2019informativa diventa un elemento aggravante della responsabilit\u00e0 "
    "contrattuale. Le implicazioni con le polizze di responsabilit\u00e0 civile professionale "
    "sono altrettanto concrete."
)

add_heading2(doc, "Cosa dice esattamente l\u2019articolo 13")

add_paragraph(doc,
    "La norma prevede quattro obblighi distinti. Al momento del conferimento dell\u2019incarico, "
    "il professionista deve comunicare al cliente, in forma scritta e con linguaggio chiaro, semplice "
    "ed esaustivo, se intende utilizzare sistemi AI nell\u2019esecuzione della prestazione. Deve "
    "specificare quali sistemi AI user\u00e0 e per quali finalit\u00e0. L\u2019informativa deve "
    "rendere esplicito che l\u2019AI opera come strumento di supporto e che il giudizio professionale "
    "e l\u2019attivit\u00e0 intellettuale umana restano prevalenti. Infine, la responsabilit\u00e0 "
    "civile e disciplinare rimane interamente in capo al professionista, indipendentemente dal fatto "
    "che l\u2019output sia stato generato con il supporto di un sistema AI."
)

add_paragraph(doc,
    "La norma non vieta l\u2019uso dell\u2019AI. Non impone una certificazione. Non richiede che "
    "il cliente dia il consenso (salvo casi specifici legati alla normativa privacy). Richiede "
    "trasparenza preventiva. Nella pratica, questo si traduce in un\u2019aggiunta al mandato "
    "professionale, redatta in modo comprensibile anche per chi non conosce il funzionamento "
    "dei modelli linguistici."
)

add_heading2(doc, "I rischi concreti di non adeguarsi")

add_paragraph(doc,
    "L\u2019articolo 13 non prevede sanzioni amministrative dirette per la violazione dell\u2019obbligo "
    "informativo, ma questo non significa che l\u2019omissione sia priva di conseguenze. Il primo "
    "rischio \u00e8 disciplinare: il Consiglio Nazionale Forense, il CNDCEC e gli altri ordini "
    "professionali possono rilevare il mancato rispetto di un obbligo di legge come violazione delle "
    "norme deontologiche. Diversi ordini, tra cui l\u2019Ordine degli Avvocati di Milano e quello "
    "di Verona, hanno gi\u00e0 emesso circolari operative sul tema a marzo 2026."
)

add_paragraph(doc,
    "Il secondo rischio riguarda la responsabilit\u00e0 contrattuale. Se il cliente lamenta un errore "
    "professionale e scopre che il lavoro era stato prodotto con il supporto di un\u2019AI, senza che "
    "gliene fosse stato fatto cenno, l\u2019omissione dell\u2019informativa pu\u00f2 essere contestata "
    "come inadempimento del contratto d\u2019incarico. Il terzo rischio \u00e8 assicurativo: alcune "
    "compagnie che coprono la RC professionale stanno gi\u00e0 inserendo clausole che condizionano "
    "la copertura alla corretta dichiarazione al cliente dell\u2019uso di strumenti AI. Vale la pena "
    "verificare con la propria compagnia la situazione della propria polizza."
)

add_heading2(doc, "Come adeguarsi senza complicare i contratti d\u2019incarico")

add_paragraph(doc,
    "L\u2019adeguamento non richiede uno stravolgimento dei contratti in uso. La strada pi\u00f9 "
    "semplice \u00e8 predisporre un\u2019appendice standard \u2014 una o due pagine \u2014 da allegare "
    "al mandato professionale. L\u2019appendice descrive i sistemi AI usati dallo studio (con il "
    "grado di dettaglio ritenuto opportuno), le finalit\u00e0 del loro utilizzo (supporto alla "
    "ricerca normativa, redazione di bozze, sintesi di documenti), le garanzie di supervisione "
    "umana adottate e il principio di responsabilit\u00e0 integrale del professionista."
)

add_paragraph(doc,
    "Diversi ordini professionali stanno mettendo a disposizione fac-simile di informativa. "
    "Prima di predisporre un testo autonomo, conviene verificare se il proprio ordine ha gi\u00e0"
    " pubblicato un modello sul proprio sito istituzionale. Chi si avvale di collaboratori o "
    "praticanti che usano AI deve assicurarsi che anche il loro utilizzo sia coperto dall\u2019"
    "informativa, poich\u00e9 la responsabilit\u00e0 ricade sul titolare dello studio o "
    "sull\u2019iscritto all\u2019ordine che firma il lavoro."
)

add_heading2(doc, "Un obbligo che pu\u00f2 diventare una scelta di posizionamento")

add_paragraph(doc,
    "La norma nasce con un obiettivo di tutela del cliente. Ma per il professionista che la "
    "affronta con metodo, pu\u00f2 diventare qualcosa di pi\u00f9. Spiegare al cliente come si usa "
    "l\u2019AI nello studio \u2014 quali strumenti, con quale logica, con quali verifiche \u2014 "
    "\u00e8 un\u2019occasione per costruire fiducia su un tema che molti clienti guardano con "
    "incertezza. Il professionista che lo fa in modo proattivo e trasparente, prima ancora "
    "che il cliente lo chieda, si distingue da chi tratta l\u2019AI come uno strumento "
    "invisibile da non menzionare."
)

add_paragraph(doc,
    "La norma non chiede ai professionisti di smettere di usare l\u2019AI. Chiede di usarla "
    "in modo trasparente. Chi trasforma questo obbligo in un elemento esplicito del proprio "
    "modo di lavorare scoprir\u00e0 che non \u00e8 solo un adempimento: \u00e8 anche una "
    "conversazione che il cliente apprezza."
)

add_separator(doc)

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(10)
run_n = p_note.add_run("Fonti e riferimenti")
run_n.font.name = 'Calibri'
run_n.font.size = Pt(9)
run_n.font.bold = True
run_n.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Legge 23 settembre 2025, n. 132, art. 13 \u2014 Obblighi informativi dei professionisti",
    "Informazione Fiscale \u2014 Fac-simile informativa cliente professionista AI (2026)",
    "Agenda Digitale \u2014 Intelligenza artificiale e avvocati: come cambia l\u2019obbligo di trasparenza",
    "Il Fatto Quotidiano \u2014 I professionisti che usano l\u2019AI hanno l\u2019obbligo di informare i clienti (28 marzo 2026)",
    "Dagostinolex \u2014 Legge 132/2025 sulla IA: diritto d\u2019autore e sanzioni penali",
    "Ordine degli Avvocati di Milano \u2014 Circolare operativa sull\u2019uso dell\u2019AI (marzo 2026)",
]
for s in sources:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_after = Pt(2)
    run_s = p_s.add_run(f"\u2022 {s}")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(8.5)
    run_s.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(16)
run_f = p_footer.add_run(
    "\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte"
)
run_f.font.name = 'Calibri'
run_f.font.size = Pt(8)
run_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_f.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
