"""
Quattro articoli Ratio -- 21 agosto 2026

1. 2026-08-21_competenze-ai-621mila-introvabili-italia.docx
2. 2026-08-21_ai-act-tre-settimane-dopo-cosa-fare-ora.docx
3. 2026-08-21_multimodalita-ai-vede-sente-agisce-professioni.docx
4. 2026-08-21_roi-ai-dalla-promessa-al-bilancio.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# 621.000 tecnici AI cercati, 316.000 introvabili
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Mercato del Lavoro e Competenze AI")
titolo(
    doc,
    "621.000 tecnici AI cercati. 316.000 introvabili.",
    "Confartigianato fotografa il paradosso del mercato del lavoro italiano nell'era dell'AI: "
    "le imprese vogliono assumere esperti di intelligenza artificiale, "
    "ma più della metà delle posizioni resta scoperta. "
    "Il problema non è l'AI. È chi la sa usare.",
    "A cura della Redazione Ratio  •  21 agosto 2026"
)

para(doc,
    "Nel 1844, mentre le ferrovie stavano cambiando la geografia economica dell'Europa, "
    "John Stuart Mill scrisse che il progresso tecnico non crea disoccupazione, "
    "ma crea una domanda di competenze diverse che il mercato del lavoro impiega tempo "
    "ad assorbire. Il tempo di adattamento, scriveva, è sempre il costo nascosto "
    "dell'innovazione. Quasi duecento anni dopo, i numeri pubblicati da Confartigianato "
    "il 18 agosto 2026 descrivono lo stesso fenomeno in forma quantitativa: "
    "le imprese italiane cercano 621.000 lavoratori con competenze in intelligenza artificiale, "
    "cloud computing, data analytics e tecnologie digitali avanzate. "
    "Di questi, 316.000 — il 50,8% — sono risultati di difficile o impossibile reperimento."
)

para(doc,
    "Il dato va letto con attenzione perché non descrive un problema di offerta formativa "
    "in astratto: descrive una frattura concreta tra la velocità con cui le imprese "
    "adottano strumenti AI e la velocità con cui il sistema forma persone in grado "
    "di usarli. Il 71,8% delle aziende italiane ha già investito in almeno un ambito "
    "della transizione digitale. Ma investire in tecnologia non equivale ad avere "
    "le persone che la fanno funzionare. La Lombardia guida la classifica delle difficoltà "
    "di reperimento con 60.960 posizioni coperte difficilmente su 121.440 richieste. "
    "Milano, Roma e Napoli sono i capoluoghi con il maggior numero di posizioni scoperte. "
    "Non sono i territori più arretrati: sono i più avanzati, e la scarsità si sente di più."
)

heading(doc, "Il paradosso della soglia bassa")

para(doc,
    "Uno degli aspetti meno discussi del dato Confartigianato è che una quota rilevante "
    "delle posizioni cercate non richiede competenze da ricercatore o ingegnere di sistema. "
    "Richiede quello che potremmo definire 'AI literacy operativa': la capacità di usare "
    "strumenti AI in un contesto professionale specifico, valutarne l'output, "
    "integrarlo nei processi esistenti e identificare i limiti del sistema. "
    "Un operatore contabile che sa usare l'AI per riconciliare movimenti bancari e "
    "sa quando l'AI sbaglia è già una figura rara e ricercata. "
    "Un addetto al commerciale che sa costruire un prompt efficace per analizzare "
    "un portafoglio clienti è già più produttivo della media. "
    "Le posizioni più difficili da coprire non sono necessariamente le più tecniche: "
    "sono quelle ibride, a metà tra il dominio professionale e la competenza digitale."
)

heading(doc, "Il ruolo del professionista come intermediario")

para(doc,
    "Per gli studi professionali che seguono PMI e imprese di medie dimensioni, "
    "la scarsità di competenze AI nel mercato del lavoro apre uno spazio "
    "di consulenza che non esisteva tre anni fa. Le aziende che non trovano "
    "il responsabile AI interno stanno cercando figure esterne in grado di "
    "guidare l'adozione: scegliere gli strumenti, formare il personale, "
    "impostare i processi, verificare la conformità normativa. "
    "Il commercialista o il consulente del lavoro che ha già sviluppato "
    "competenze AI è in una posizione privilegiata: conosce il cliente, "
    "conosce i processi, e può offrire una guida contestualizzata "
    "che un consulente tecnologico generalista non può replicare. "
    "Il tempo di adattamento di Mill è esattamente questo: "
    "una finestra che si apre per chi è già dentro e si chiude "
    "quando il mercato si equilibra."
)

heading(doc, "Cosa possono fare le imprese nel breve")

para(doc,
    "In assenza di figure dedicate sul mercato, le imprese italiane hanno "
    "due leve immediate. La prima è la formazione interna: identificare "
    "2-3 figure già presenti in azienda con propensione digitale e investire "
    "su di loro in modo mirato, non con corsi generici sull'AI, "
    "ma con percorsi su casi d'uso specifici del proprio settore. "
    "La seconda è affidarsi a professionisti esterni che già integrano l'AI "
    "nel proprio servizio: non come acquisto di software, ma come consulenza "
    "di processo. L'obbligo di AI literacy introdotto dall'AI Act non è "
    "solo un adempimento: è un'occasione per documentare ciò che si fa già "
    "e strutturare quello che manca. Le 316.000 posizioni introvabili "
    "non si risolvono tutte con l'assunzione. Alcune si risolvono con "
    "la formazione di chi c'è già."
)

para(doc,
    "Mill aveva ragione: il costo nascosto dell'innovazione è il tempo. "
    "Ma aveva anche ragione su un'altra cosa: chi impara prima a stare "
    "nel mezzo del cambiamento finisce per guidarlo. "
    "Le 316.000 posizioni introvabili sono anche 316.000 ragioni "
    "per formarsi adesso."
)

riferimenti(doc, [
    "ANSA – 'Confartigianato, 621mila tecnici IA da assumere, più della metà è introvabile' (18 agosto 2026)",
    "QuiFinanza – 'IA, le aziende cercano esperti ma non si trovano: mancano 316mila tecnici' (agosto 2026)",
    "Il Metropolitano – 'Intelligenza artificiale, l’allarme di Confartigianato: alle imprese italiane manca il 50% delle professionalità' (19 agosto 2026)",
    "Corriere delle Comunicazioni – 'Competenze digitali, allarme Confartigianato: mancano all’appello 362mila esperti di AI'",
    "Il Sole 24 Ore – 'Companies are looking for AI experts, more than half of whom are impossible to find' (agosto 2026)",
    "Regolamento UE 2024/1689 (AI Act), art. 4 – Obbligo di alfabetizzazione AI per il personale",
])
doc.save(BASE + "2026-08-21_competenze-ai-621mila-introvabili-italia.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# AI Act tre settimane dopo il 2 agosto: cosa fare ora
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Normativa AI e Compliance")
titolo(
    doc,
    "Tre settimane dopo il 2 agosto: ora cosa si fa.",
    "Il 2 agosto 2026 è passato. Gli obblighi di trasparenza dell'AI Act sono in vigore. "
    "Le autorità di vigilanza hanno i poteri per ispezionare. "
    "Per la maggior parte delle imprese italiane, il lavoro vero inizia adesso.",
    "A cura della Redazione Ratio  •  21 agosto 2026"
)

para(doc,
    "Nell'agosto del 79 d.C., i Pompeiani sapevano che il Vesuvio fumava. "
    "Plinio il Vecchio aveva scritto di eruzioni. La letteratura romana "
    "aveva già registrato i segni premonitori. Eppure la risposta prevalente "
    "fu l'attesa: il vulcano aveva sempre fumato, il peggio era sempre stato "
    "rimandato, le case erano costruite, i negozi aperti. "
    "La data del 2 agosto 2026 non era un segreto: l'AI Act è in vigore "
    "dal 2024, il calendario degli obblighi era pubblico da due anni. "
    "Tre settimane dopo quella data, l'85% delle PMI italiane non ha ancora "
    "avviato una verifica sistematica della propria conformità. "
    "Il vulcano, per ora, fuma ancora."
)

para(doc,
    "Dal 2 agosto 2026 sono operativi tre gruppi di obblighi dell'AI Act "
    "direttamente rilevanti per imprese e professionisti. Il primo è "
    "la trasparenza: i sistemi AI che interagiscono con le persone devono "
    "dichiararsi tali, e i contenuti generati da AI devono essere marcati. "
    "Il secondo è la governance dei modelli di uso generale: i fornitori "
    "di modelli come GPT, Claude e Gemini sono soggetti a nuovi obblighi "
    "di documentazione e alle verifiche dell'AGID e dell'ACN come "
    "autorità nazionali di vigilanza. Il terzo è l'alfabetizzazione: "
    "le imprese devono garantire che il personale che usa sistemi AI "
    "abbia un livello adeguato di competenza per farlo. "
    "Le sanzioni per le violazioni degli obblighi di trasparenza "
    "arrivano al 3% del fatturato mondiale."
)

heading(doc, "Il nodo dell’alfabetizzazione")

para(doc,
    "L'obbligo di AI literacy è quello che più frequentemente viene "
    "sottovalutato nei piani di compliance aziendali. L'AI Act non richiede "
    "che tutti i dipendenti diventino esperti: richiede che chi usa sistemi AI "
    "in modo sistematico abbia una comprensione adeguata di cosa fa il sistema, "
    "dei suoi limiti e dei rischi associati. 'Adeguata' non è definita "
    "quantitativamente dalla norma, ma le linee guida dell'AGID pubblicate "
    "in luglio orientano verso la documentazione: chi usa l'AI, "
    "su quali compiti, con quale formazione ricevuta. "
    "Per uno studio professionale che usa strumenti AI per attività "
    "che impattano i clienti — redazione di documenti, analisi, consulenza — "
    "la domanda operativa è: se oggi arrivasse un'ispezione, "
    "saremmo in grado di mostrare che chi usa questi strumenti "
    "sa valutarne l'output?"
)

heading(doc, "La checklist pratica post-2 agosto")

para(doc,
    "Per le imprese e gli studi professionali che non hanno ancora avviato "
    "una verifica, la sequenza pratica è la seguente. Primo passo: "
    "mappare i sistemi AI in uso, inclusi quelli consumer come ChatGPT, "
    "Copilot e Claude, distinguendo tra uso personale occasionale "
    "e uso sistematico nei processi. Secondo passo: verificare che i "
    "sistemi rivolti ai clienti — chatbot sul sito, risponditori automatici, "
    "email generate da AI — siano configurati per dichiarare la propria "
    "natura artificiale. Terzo passo: documentare la formazione ricevuta "
    "dal personale che usa AI in modo sistematico, anche sotto forma "
    "di una nota interna che registri chi ha fatto cosa. "
    "Quarto passo: nominare un referente interno che abbia in carico "
    "il monitoraggio della compliance AI, non necessariamente a tempo pieno, "
    "ma con un ruolo definito. Questi quattro passaggi non esauriscono "
    "la compliance per i sistemi ad alto rischio, che rientrano "
    "in un regime più stringente, ma coprono la base per la grande "
    "maggioranza delle PMI e degli studi professionali."
)

heading(doc, "Cosa aspettarsi nei prossimi mesi")

para(doc,
    "L'attività di vigilanza dell'AGID e dell'ACN partirà verosimilmente "
    "dai soggetti più grandi e dai casi più visibili: i fornitori di sistemi "
    "AI sul mercato, le piattaforme con milioni di utenti, le grandi aziende "
    "che usano AI in processi ad alto impatto. Per le PMI e gli studi "
    "professionali, il rischio di un'ispezione diretta nel 2026 è limitato. "
    "Ma il rischio indiretto è già reale: un cliente che contesta "
    "un documento generato da AI senza disclosure, un dipendente che segnala "
    "l'uso di AI senza formazione adeguata, un fornitore che chiede "
    "attestazioni di compliance per continuare un rapporto commerciale. "
    "La compliance non è solo un obbligo verso l'autorità: "
    "è anche una protezione verso i clienti e i collaboratori."
)

para(doc,
    "Pompei non fu distrutta dall'eruzione in un giorno solo. "
    "Fu sepolta lentamente, in ore. Chi usò bene quel tempo si salvò. "
    "Tre settimane dopo il 2 agosto, il tempo è ancora disponibile."
)

riferimenti(doc, [
    "Altalex – 'AI Act: che cosa cambia davvero dal 2 agosto 2026 per imprese e professionisti' (31 luglio 2026)",
    "UniverseIT – 'AI Act 2 agosto 2026: cosa cambia per le imprese'",
    "Confartigianato Lecce – 'AI ACT | dal 2 agosto 2026 nuovi obblighi per imprese e operatori economici'",
    "Meta Communications – 'Regolamentazione dell’AI: cosa cambia da agosto 2026'",
    "LaP.A.M. – 'AI Act 2026: nuovi obblighi dal 2 agosto'",
    "AGID – Linee guida nazionali sull’obbligo di alfabetizzazione AI (luglio 2026)",
    "Regolamento UE 2024/1689 (AI Act), artt. 4, 50, 99",
])
doc.save(BASE + "2026-08-21_ai-act-tre-settimane-dopo-cosa-fare-ora.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# Multimodalità: l'AI che vede, sente e agisce
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Strumenti AI per Professionisti")
titolo(
    doc,
    "L’AI vede le fatture, ascolta le riunioni, legge i contratti.",
    "La multimodalità è la capacità dei modelli AI di lavorare su testo, "
    "immagini, audio e documenti nella stessa sessione. "
    "Per uno studio professionale, significa che l’AI si adatta al flusso di lavoro reale, "
    "non il contrario.",
    "A cura della Redazione Ratio  •  21 agosto 2026"
)

para(doc,
    "Thomas Edison, quando presentò il fonografo nel 1877, elencò dieci possibili "
    "applicazioni dell’invenzione. All’ottavo posto metteva 'preservare le ultime "
    "parole dei moribondi'. Al decimo 'giocattoli parlanti'. La registrazione della "
    "musica, che sarebbe diventata l’applicazione economicamente dominante, "
    "non era nella lista. Chi usa strumenti AI multimodali nel lavoro professionale "
    "sta facendo la stessa scoperta: le applicazioni che cambiano di più "
    "il lavoro quotidiano spesso non sono quelle previste quando si è "
    "abbonati al servizio."
)

para(doc,
    "La multimodalità è la capacità di un sistema AI di elaborare "
    "contemporaneamente diverse tipologie di input: testo, immagini, PDF, "
    "fogli di calcolo, file audio, screenshot, fotografie. I principali modelli "
    "disponibili nel 2026 — GPT-5.5, Claude Sonnet 5, Gemini 3.1 Pro — "
    "gestiscono tutti questi formati in modo nativo. "
    "Per uno studio professionale o un'azienda, questo significa che "
    "non è più necessario trascrivere il contenuto di un documento prima "
    "di analizzarlo: si carica il PDF e si chiede. Non è più necessario "
    "riscrivere i dati di una fattura in un foglio di calcolo: "
    "si mostra la fattura e si chiede di estrarre i campi. "
    "Non è più necessario leggere l’intera trascrizione di una riunione: "
    "si carica l’audio e si chiede il riassunto per punti d’azione."
)

heading(doc, "Le applicazioni più usate negli studi professionali")

para(doc,
    "I casi d’uso multimodali che stanno entrando negli studi professionali "
    "si concentrano su tre aree. La prima è l’analisi documentale: "
    "contratti, bilanci, visure, dichiarazioni, estratti conto. "
    "Il professionista carica il documento, pone domande specifiche "
    "e riceve risposte che rimandano al testo originale. "
    "Il tempo di lettura di un contratto di 40 pagine per identificare "
    "le clausole rilevanti si riduce da ore a minuti, "
    "con la verifica umana che si concentra sui punti segnalati dall’AI. "
    "La seconda area è l’elaborazione di dati visivi: fotografie di ricevute "
    "e scontrini per la nota spese, scansioni di documenti cartacei, "
    "screenshot di gestionali. La terza è la sintesi di riunioni "
    "e call registrate: i principali strumenti transcrivono e riassumono "
    "in italiano con precisione sufficiente per l’uso operativo."
)

heading(doc, "I limiti che contano")

para(doc,
    "La multimodalità non è infallibile e i suoi limiti sono diversi "
    "da quelli del testo puro. I modelli che leggono immagini possono "
    "sbagliare cifre, soprattutto in documenti con layout complessi, "
    "tabelle sovrapposte o qualità di scansione bassa. "
    "I modelli che trascrivono audio perdono precisione su dialetti, "
    "voci sovrapposte e terminologia tecnica di nicchia. "
    "I modelli che analizzano PDF strutturati come bilanci XBRL "
    "o dichiarazioni XML possono avere difficoltà con i metadati "
    "che un revisore esperto leggerebbe automaticamente. "
    "La regola operativa è la stessa che vale per il testo: "
    "l’AI produce una prima elaborazione di qualità, "
    "il professionista verifica i punti critici. "
    "Su documenti dove un errore ha conseguenze rilevanti, "
    "la verifica non è opzionale."
)

heading(doc, "Privacy e dati del cliente")

para(doc,
    "Il caricamento di documenti contenenti dati personali dei clienti "
    "su piattaforme AI esterne è soggetto alle disposizioni del GDPR "
    "e, dal 2 agosto 2026, anche all’AI Act. "
    "Per gli studi professionali, la domanda pratica è se il fornitore "
    "del servizio AI ha sottoscritto un Data Processing Agreement adeguato "
    "e se i dati caricati vengono usati per addestrare i modelli. "
    "I principali fornitori enterprise offrono opzioni di non-training "
    "e storage dei dati in Europa. I piani consumer standard "
    "hanno spesso condizioni meno favorevoli. "
    "Prima di caricare un contratto con i dati di un cliente "
    "su uno strumento AI, vale la pena leggere le condizioni "
    "del proprio piano."
)

para(doc,
    "Edison non aveva previsto la musica registrata "
    "perché stava pensando a cosa il fonografo potesse conservare, "
    "non a cosa le persone avessero piacere di ascoltare. "
    "L’AI multimodale funziona meglio quando si parte dal problema reale "
    "del lavoro quotidiano, non dalla lista delle funzionalità."
)

riferimenti(doc, [
    "Bleap Finance – 'Claude vs GPT vs Gemini: Confronto tra i Migliori Modelli di IA del 2026'",
    "Alessio Pomaro – 'Generative AI: novità e riflessioni #4/2026'",
    "Culture Digitali – 'Tutti i Modelli Gemini AI Disponibili nel 2026: Guida Completa'",
    "OpenAI Help Center – Note di rilascio dei modelli (agosto 2026)",
    "Garante per la protezione dei dati personali – Provvedimento sull’uso di strumenti AI nei contesti professionali",
    "Regolamento UE 2016/679 (GDPR), art. 28 – Responsabili del trattamento",
    "Regolamento UE 2024/1689 (AI Act), artt. 4, 13, 50",
])
doc.save(BASE + "2026-08-21_multimodalita-ai-vede-sente-agisce-professioni.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# ROI dell'AI: dalla promessa al bilancio
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Gestione e Strategia AI")
titolo(
    doc,
    "Il ROI dell’AI: come si misura quello che conta.",
    "Il mercato AI è pieno di promesse di efficienza. "
    "Molte aziende investono, poche misurano i risultati in modo rigoroso. "
    "Per un imprenditore o un professionista, la domanda non è se l’AI funziona: "
    "è se funziona per me, su questo processo, con questi numeri.",
    "A cura della Redazione Ratio  •  21 agosto 2026"
)

para(doc,
    "Lord Kelvin, il fisico vittoriano che misurò la temperatura assoluta, "
    "aveva una massima che i suoi studenti impararono a memoria: "
    "'Se non puoi misurarlo, non puoi migliorarlo.' "
    "La frase viene spesso attribuita erroneamente a Peter Drucker, "
    "il che dice qualcosa sull’affidabilità delle citazioni non verificate, "
    "tema che ci tornerà utile tra poco. Ma il principio regge: "
    "gli investimenti in AI che non vengono misurati sistematicamente "
    "producono percezioni di valore difficili da giustificare "
    "di fronte a un consiglio di amministrazione, a un socio "
    "o a un cliente che paga la consulenza. "
    "Nel 2026, mentre le aziende italiane dichiarano investimenti AI "
    "in crescita, la quota di quelle che misurano il ritorno "
    "in modo strutturato resta sotto il 30%."
)

para(doc,
    "Il calcolo del ROI per gli strumenti AI è più complicato "
    "di quanto sembri perché i benefici non sono tutti monetizzabili "
    "facilmente e i costi non si esauriscono nell’abbonamento mensile. "
    "Sul lato dei costi entrano: il costo della piattaforma, "
    "il tempo di formazione del personale, il tempo di supervisione "
    "dell’output (che non è zero, anche quando l’AI lavora bene), "
    "i costi di governance e compliance introdotti dall’AI Act. "
    "Sul lato dei benefici entrano: il tempo risparmiato su compiti "
    "ricorrenti, la riduzione degli errori su processi ad alto volume, "
    "la capacità di gestire un maggior numero di pratiche "
    "con lo stesso organico, la qualità migliorata degli output "
    "su compiti di analisi e redazione."
)

heading(doc, "Come costruire una misura credibile")

para(doc,
    "Il metodo più semplice per costruire una misura del ROI su un caso "
    "d’uso AI specifico parte da tre dati. Primo: il tempo medio "
    "impiegato per completare il processo prima dell’AI, "
    "espresso in ore per unità (per pratica, per fattura, per documento). "
    "Secondo: il tempo medio impiegato dopo l’introduzione dell’AI, "
    "inclusa la verifica dell’output. Terzo: il costo orario "
    "della figura professionale coinvolta. "
    "La differenza in ore per il volume annuo di pratiche, "
    "moltiplicata per il costo orario, è il risparmio diretto. "
    "A questo si sottrae il costo annuo dello strumento "
    "e una stima del tempo di governance. "
    "Questo calcolo è rozzo, ma è verificabile e confrontabile nel tempo. "
    "Un calcolo rozzo e verificato vale infinitamente più "
    "di una stima sofisticata non controllata."
)

heading(doc, "I casi d’uso con ROI più misurabile")

para(doc,
    "Nei contesti professionali italiani, i processi dove il ROI dell’AI "
    "è più facilmente misurabile sono quelli ad alto volume e bassa variabilità: "
    "la riconciliazione di estratti conto bancari, "
    "la classificazione di movimenti contabili ricorrenti, "
    "la redazione di comunicazioni standard ai clienti su temi definiti, "
    "la sintesi di documenti lunghi per estrarne i dati rilevanti, "
    "la risposta a quesiti normativi su argomenti ben perimetrati. "
    "Su questi processi, studi professionali che hanno misurato "
    "riferiscono risparmi di tempo tra il 30% e il 60% per pratica, "
    "con variazioni significative legate alla qualità del prompt "
    "e alla curva di apprendimento iniziale. "
    "I processi ad alta variabilità e giudizio — "
    "la consulenza strategica, la rappresentanza in contenzioso, "
    "la negoziazione — mostrano benefici meno lineari "
    "e più difficili da isolare."
)

heading(doc, "Quando il ROI non chiude")

para(doc,
    "Alcuni investimenti AI non producono il ROI atteso. "
    "Le ragioni più frequenti sono tre. La prima è la scelta sbagliata "
    "del caso d’uso: si introduce l’AI su un processo dove il collo di bottiglia "
    "non è il tempo di elaborazione ma la raccolta dei dati, "
    "e il risparmio è marginale. La seconda è la sopravvalutazione "
    "dell’autonomia del sistema: si abbandona la supervisione dell’output "
    "troppo presto, si accumulano errori, e il tempo di correzione "
    "annulla il risparmio. La terza è la mancanza di adattamento "
    "del processo: si usa l’AI come sostituto dell’operatore "
    "senza ridisegnare il flusso, e il guadagno rimane sulla carta. "
    "In tutti e tre i casi, la misura aiuta a diagnosticare: "
    "un ROI che non chiude su un caso d’uso specifico "
    "è un dato utile, non un fallimento."
)

para(doc,
    "Lord Kelvin aveva anche un altro principio, meno citato: "
    "'La misura è utile non quando conferma ciò che già sapevi, "
    "ma quando ti dice ciò che non ti aspettavi.' "
    "Il ROI dell’AI misurato seriamente dirà dove investire di più "
    "e dove smettere. Entrambe le informazioni valgono."
)

riferimenti(doc, [
    "SME-AI Maturity Index 2026 – Webidoo Insight Lab (giugno 2026)",
    "AdnKronos – 'AI: nelle PMI italiane il potenziale c’è ma la maturità resta bassa' (30 giugno 2026)",
    "AgendaDigitale.eu – 'AI nelle PMI italiane: competenze e dati frenano la svolta digitale'",
    "MediaKey.it – 'SME-AIMIX 2026: fino al +30% di produttività con l’AI, ma le PMI italiane non sono ancora pronte'",
    "Yellow Tech – 'AI Italia: trend e dati 2026'",
    "Istat – Rapporto ICT e AI nelle imprese italiane 2025-2026",
])
doc.save(BASE + "2026-08-21_roi-ai-dalla-promessa-al-bilancio.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
