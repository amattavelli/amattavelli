#!/usr/bin/env python3
"""
Articolo 2: Responsabilità professionale e AI
File: Ratio/articoli/2026-04-17_responsabilita-professionale-errori-ai.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/home/user/amattavelli/Ratio/articoli/2026-04-17_responsabilita-professionale-errori-ai.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(3)
sec.right_margin = Cm(3)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot)
    pPr.append(pBdr)


def para(doc, text, size=11, italic=False, color=None, sa=8, sb=0,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


# --- TESTATA ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.all_caps = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Professione e Responsabilit\u00e0")
r.font.name = 'Calibri'
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# --- TITOLO ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "L\u2019AI ha sbagliato la detrazione, il cliente ha pagato la sanzione: "
    "la responsabilit\u00e0 non si divide"
)
r.font.name = 'Calibri'
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Quando gli strumenti di intelligenza artificiale producono un errore che arriva "
    "fino al cliente, la responsabilit\u00e0 professionale resta intera. "
    "Nessun software \u00e8 un\u2019attenuante."
)
r.font.name = 'Calibri'
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  \u2022  17 aprile 2026")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# --- CORPO ---
para(doc,
    "Un software AI integrato in un gestionale tributario suggerisce il codice errato per "
    "una deduzione. Il professionista lo accetta senza verificare, il modello viene firmato "
    "e trasmesso. Qualche mese dopo, l\u2019Agenzia delle Entrate contesta la deduzione e "
    "irroga una sanzione al cliente. In sede di rivalsa, il professionista prova a indicare "
    "l\u2019errore del software come causa concorrente. Non funziona. La giurisprudenza "
    "italiana, consolidatasi nel 2025 e 2026, \u00e8 chiara: l\u2019errore prodotto da un sistema "
    "AI non attenua la responsabilit\u00e0 professionale, non la divide, non la trasferisce "
    "al fornitore del software."
)

para(doc,
    "Il principio ha radici nella struttura stessa della responsabilit\u00e0 professionale. "
    "Il commercialista, l\u2019avvocato, il consulente del lavoro sono chiamati a rispondere "
    "del risultato della propria prestazione intellettuale. Quando decidono di avvalersi "
    "di strumenti automatizzati, inclusi i sistemi AI, non trasferiscono questa "
    "responsabilit\u00e0 al produttore dello strumento: la mantengono, aggiungendo per\u00f2 "
    "l\u2019obbligo di verificare che l\u2019output dello strumento sia corretto. L\u2019ordinanza "
    "della Cassazione del 12 marzo 2026, n. 5635, ha ribadito questo principio in un caso "
    "di tenuta della contabilit\u00e0, chiarendo che l\u2019automazione del processo non esime "
    "il professionista dalla responsabilit\u00e0 sull\u2019esattezza del risultato."
)

h2(doc, "Il problema concreto della verifica")

para(doc,
    "La questione pratica \u00e8 questa: come si verifica l\u2019output di un sistema AI in modo "
    "che la verifica sia effettiva e non solo formale? Non \u00e8 sufficiente leggere il "
    "risultato e considerarlo plausibile. Una supervisione adeguata implica comprendere "
    "il ragionamento seguito dallo strumento quando \u00e8 reso visibile, confrontare l\u2019output "
    "con fonti normative primarie nei casi dubbi, e mantenere la propria capacit\u00e0 di "
    "giudizio autonomo anche sulle operazioni routinarie. Il paradosso \u00e8 che l\u2019AI, "
    "usata in modo superficiale, pu\u00f2 degradare proprio quella capacit\u00e0 di giudizio: "
    "se il professionista smette di fare certe verifiche perch\u00e9 ci pensa il software, "
    "perde progressivamente la competenza necessaria per accorgersi quando il software "
    "sbaglia."
)

para(doc,
    "Alcune compagnie assicurative stanno aggiornando le polizze di responsabilit\u00e0 "
    "professionale per includere esplicitamente l\u2019uso di strumenti AI tra i fattori di "
    "rischio valutati. Dichiarare di utilizzare strumenti AI senza un protocollo di "
    "verifica strutturato potrebbe influire sulle condizioni della copertura. Non si "
    "tratta ancora di una prassi uniforme di mercato, ma la direzione \u00e8 chiara e "
    "i rinnovi 2026 stanno gi\u00e0 iniziando a riflettere questa sensibilit\u00e0."
)

para(doc,
    "Un protocollo di verifica non deve essere elaborato: pu\u00f2 essere semplice, purch\u00e9 "
    "sistematico. Definire quali categorie di output AI richiedono verifica su fonte "
    "primaria (interpretazioni normative, aliquote, codici specifici), quali possono "
    "essere accettati dopo una lettura attenta, e quali vanno sempre rielaborati dal "
    "professionista indipendentemente da quanto sembri attendibile il suggerimento "
    "automatico. Documentare questa procedura internamente allo studio \u00e8 il primo passo "
    "per trasformarla da intenzione in pratica."
)

para(doc,
    "Ci sono per\u00f2 differenze rilevanti tra le varie categorie di strumenti. Un software "
    "tributario verticale, addestrato su normativa italiana e aggiornato periodicamente "
    "dal produttore, offre garanzie diverse rispetto a un modello linguistico generalista "
    "usato per interpretare una circolare. Nel primo caso, il professionista pu\u00f2 contare "
    "su un sistema che il fornitore certifica e aggiorna. Nel secondo caso, la verifica "
    "\u00e8 completamente a suo carico. Sapere quale tipo di strumento si sta usando, e cosa "
    "garantisce ciascuno, \u00e8 il prerequisito per costruire un protocollo di verifica "
    "sensato."
)

h2(doc, "Il dato che mette a fuoco il rischio")

para(doc,
    "Ricerche recenti indicano che i modelli linguistici generalisti producono risposte "
    "errate nel 50% o pi\u00f9 dei casi quando interrogati su questioni giuridiche specifiche. "
    "I software verticali, addestrati su dati fiscali italiani, hanno tassi di errore "
    "pi\u00f9 bassi, ma non trascurabili, specialmente in presenza di interpretazioni "
    "normative recenti o di situazioni non standard. Il professionista che tratta "
    "questi strumenti come assistenti infallibili sta assumendo un rischio che la "
    "normativa e la giurisprudenza scaricano interamente su di lui."
)

para(doc,
    "L\u2019AI usata bene nello studio non riduce la responsabilit\u00e0 professionale: "
    "la redistribuisce verso le attivit\u00e0 in cui il giudizio umano fa davvero la "
    "differenza. Il valore non sta nell\u2019automatizzare tutto, ma nell\u2019automatizzare "
    "quello che pu\u00f2 essere automatizzato con sicurezza, mantenendo un presidio attivo "
    "su tutto il resto. Capire dove sta questo confine, nel proprio contesto specifico, "
    "\u00e8 la competenza professionale pi\u00f9 rilevante per chiunque lavori oggi con "
    "strumenti AI."
)

sep(doc)

# --- FONTI ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Fonti e riferimenti")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Cassazione civile, ordinanza 12 marzo 2026 n. 5635 \u2014 Responsabilit\u00e0 del "
    "commercialista nella tenuta della contabilit\u00e0",
    "Il Fatto Quotidiano \u2014 I professionisti che usano l\u2019AI hanno l\u2019obbligo di informare "
    "i clienti e sono responsabili degli errori (28 marzo 2026)",
    "EC News \u2014 Intelligenza Artificiale e professioni intellettuali: l\u2019obbligo di "
    "trasparenza verso il cliente",
    "KTS Finance \u2014 Assicurazione professionale commercialisti e IA: i rischi del 730/2026",
    "Stanford Institute for Human-Centered AI \u2014 Accuracy of LLMs on legal questions (2025)",
]
for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"\u2022 {s}")
    r.font.name = 'Calibri'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte")
r.font.name = 'Calibri'
r.font.size = Pt(8)
r.italic = True
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

doc.save(OUTPUT)
print(f"Salvato: {OUTPUT}")
