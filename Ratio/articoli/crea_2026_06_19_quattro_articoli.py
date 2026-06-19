"""
Quattro articoli Ratio — 19 giugno 2026

1. 2026-06-19_fable5-blocco-usa-rischio-dipendenza-ai.docx
2. 2026-06-19_ai-act-omnibus-rinvio-alto-rischio-imprese.docx
3. 2026-06-19_lisa-bologna-supercomputer-ai-generativa.docx
4. 2026-06-19_dipendenza-piattaforme-ai-rischio-operativo.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# La regola e' europea, l'interruttore e' a Washington
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Geopolitica AI")
titolo(
    doc,
    "La regola e' europea.\nL'interruttore e' a Washington.",
    "Il 12 giugno il governo USA ha bloccato Claude Fable 5 per tutti gli utenti "
    "non americani. Al G7 di Evian si e' discusso di 'trusted partners' per l'accesso "
    "ai modelli di frontiera. Per i professionisti italiani, la questione non e' "
    "geopolitica. E' operativa.",
    "A cura della Redazione Ratio  •  19 giugno 2026"
)

para(doc,
    "Il 12 giugno uno studio milanese che aveva integrato Claude come strumento "
    "quotidiano per la redazione di pareri e la sintesi di documenti ha ricevuto, "
    "al primo accesso della mattina, un messaggio di errore. Il servizio non era "
    "disponibile. Non si trattava di un problema tecnico temporaneo, ne' di un "
    "disservizio da segnalare all'assistenza. Era una direttiva del governo degli "
    "Stati Uniti. Tre giorni dopo il lancio di Claude Fable 5, il modello piu' "
    "avanzato mai rilasciato da Anthropic, Washington aveva ordinato alla societa' "
    "di sospendere l'accesso per tutti i cittadini non americani, ovunque si "
    "trovassero nel mondo, invocando l'Export Control Reform Act e il rischio di "
    "'uso militare-intelligence da parte di soggetti stranieri'. Anthropic ha "
    "rispettato la direttiva entro ventiquattro ore, perche' non aveva alternative."
)

para(doc,
    "L'episodio ha avuto una certa copertura nella stampa tecnica italiana, ma e' "
    "stato inquadrato prevalentemente come una notizia sul nuovo modello di Anthropic, "
    "sulla sua potenza e sulle sue capacita'. L'aspetto che merita attenzione per "
    "chi usa questi strumenti professionalmente e' un altro: molte imprese e studi "
    "italiani si trovano in una condizione di dipendenza operativa da piattaforme la "
    "cui disponibilita' non dipende ne' dall'utente ne' dalle norme europee. Dipende "
    "da decisioni prese a Washington. Questa non e' una critica alle piattaforme "
    "americane, che restano le piu' avanzate disponibili. E' una descrizione di come "
    "funziona la struttura del mercato AI nel 2026."
)

heading(doc, "Al G7 di Evian: il club dei partner fidati")

para(doc,
    "Al vertice del G7 di Evian-les-Bains, il 16 giugno, il tema ha trovato spazio "
    "ufficiale nelle discussioni tra i leader. Reuters ha riferito che si e' valutato "
    "un piano per consentire a un gruppo limitato di paesi alleati di accedere ai "
    "modelli di frontiera americani anche in regime di export control. La logica e' "
    "quella dei 'trusted partners': paesi con cui gli USA condividono sufficienti "
    "garanzie di sicurezza da potersi fidare dell'uso che ne fanno. L'Italia, come "
    "membro del G7 e paese NATO, rientrerebbe in questo perimetro. Ma la struttura "
    "del meccanismo e' significativa: l'accesso ai modelli di punta diventa una "
    "concessione politica, negoziabile, subordinata a criteri che si definiscono a "
    "Washington, non a Bruxelles o a Roma."
)

para(doc,
    "Sam Altman di OpenAI, Dario Amodei di Anthropic e Demis Hassabis di Google "
    "hanno partecipato al summit come relatori e hanno sostenuto una governance "
    "globale dell'AI guidata dagli USA. La posizione e' che i modelli americani "
    "siano piu' sicuri e debbano rimanere il punto di riferimento mondiale. La "
    "sintesi piu' efficace di questa situazione l'ha offerta MilanoFinanza: "
    "'nell'intelligenza artificiale la regola e' europea ma l'interruttore e' "
    "americano'. L'Europa ha elaborato il quadro normativo piu' avanzato al mondo, "
    "dall'AI Act alla Legge 132/2025. Ma i modelli che le aziende europee usano "
    "ogni giorno girano su infrastrutture americane, con termini di servizio "
    "americani, sotto la giurisdizione americana."
)

heading(doc, "La domanda operativa che vale la pena farsi adesso")

para(doc,
    "Per uno studio professionale o una PMI italiana, la risposta corretta a questo "
    "scenario non e' smettere di usare strumenti AI americani. La qualita' di questi "
    "strumenti e' reale e verificata sul campo, e non trova oggi equivalente nelle "
    "alternative europee per la maggior parte dei casi d'uso professionali. La "
    "risposta e' costruire i propri flussi di lavoro con consapevolezza delle "
    "dipendenze. Se un processo lavorativo dipende in modo critico da uno strumento "
    "specifico, val la pena chiedersi con franchezza cosa succederebbe se quello "
    "strumento diventasse indisponibile per quarantotto ore. Esiste una modalita' "
    "alternativa, anche se meno efficiente? Si ha familiarita' con almeno uno "
    "strumento sostitutivo per le funzioni piu' critiche? Il processo e' documentato "
    "in modo che possa essere spiegato a un nuovo collaboratore senza presupporre "
    "l'accesso allo strumento specifico?"
)

para(doc,
    "Queste domande si pongono per gli stessi motivi per cui uno studio documenta "
    "le procedure operative: non perche' si preveda un disastro, ma perche' la "
    "continuita' del servizio ai clienti e' una responsabilita' professionale, e "
    "le dipendenze da strumenti digitali non sono diverse da qualsiasi altra "
    "dipendenza da un fornitore critico. Chi ha gia' in uso un gestionale cloud o "
    "un'applicazione SaaS per la fatturazione sa bene quanto possono essere "
    "scomodi anche solo i downtime programmati. Con gli strumenti AI, la posta in "
    "gioco e' simile, ma i fattori di rischio sono diversi: non si tratta di "
    "problemi tecnici prevedibili, ma di decisioni esterne su cui non si ha "
    "influenza."
)

para(doc,
    "L'inaugurazione di LISA al Tecnopolo di Bologna, l'11 giugno, come primo "
    "supercomputer europeo dedicato all'AI generativa, segnala che l'Europa sta "
    "lavorando a ridurre questa asimmetria strutturale nel medio periodo. Ma ci "
    "vorranno anni prima che le alternative europee raggiungano la qualita' dei "
    "modelli americani di punta. Nel frattempo, la lezione pratica dell'episodio "
    "Fable 5 e' semplice: usare con convinzione i migliori strumenti disponibili, "
    "sapendo dove sono gli interruttori, e avendo almeno pensato a cosa fare se "
    "qualcuno li spegne."
)

riferimenti(doc, [
    "CorriereNerd.it — 'Claude Fable 5: il mistero dell'IA Anthropic spenta dopo tre giorni' (giugno 2026)",
    "Cosimo.dev — 'Fable 5 e Mythos 5 spenti dopo 3 giorni: cosa e' successo' (giugno 2026)",
    "Simon Willison — 'Statement on the US government directive to suspend access to Fable 5 and Mythos 5' (13 giugno 2026)",
    "Internazionale (Reuters) — 'G7 leaders discuss trusted partners access for cutting-edge US AI models' (16 giugno 2026)",
    "Key4biz — 'G7 e AI, nasce il club dei Paesi fidati' (giugno 2026)",
    "MilanoFinanza — 'Nell'intelligenza artificiale la regola e' europea ma l'interruttore e' americano' (15 giugno 2026)",
    "AI4Business — 'G7 e modelli AI Usa, scontro su accesso e partner fidati' (giugno 2026)",
])
doc.save(BASE + "2026-06-19_fable5-blocco-usa-rischio-dipendenza-ai.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# Il rinvio dell'AI Act non vale per tutti
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Normativa AI")
titolo(
    doc,
    "Il rinvio non cancella gli obblighi.\nEcco cosa resta ad agosto.",
    "Il 7 maggio l'accordo Omnibus ha spostato al dicembre 2027 le scadenze per "
    "i sistemi AI ad alto rischio. Per molte imprese italiane, la notizia ha "
    "generato un rilassamento prematuro. Gli obblighi di agosto rimangono, e "
    "quelli della Legge 132 non li ha toccati nessuno.",
    "A cura della Redazione Ratio  •  19 giugno 2026"
)

para(doc,
    "A meta' maggio molte imprese italiane hanno letto una notizia che sembrava "
    "positiva: la Commissione Europea e il Parlamento avevano raggiunto un accordo "
    "sull'AI Omnibus che spostava le scadenze per i sistemi AI classificati ad "
    "alto rischio dal 2 agosto 2026 al dicembre 2027. Sedici mesi in piu'. "
    "Per molte imprese, e per molti consulenti che le assistono, questo ha "
    "significato allentare la pressione sui progetti di compliance. Se il termine "
    "si allunga, il problema puo' aspettare. Questa lettura e' comprensibile, "
    "ma e' imprecisa, e l'imprecisione ha conseguenze pratiche che conviene "
    "conoscere prima di agosto."
)

para(doc,
    "L'accordo Omnibus del 7 maggio 2026 e' il risultato di un percorso politico "
    "che ha visto la Commissione confrontarsi con le pressioni delle industrie, "
    "degli stati membri piu' attenti alla competitivita' e del nuovo contesto "
    "geopolitico che vede gli USA ridurre le proprie restrizioni sull'AI verso i "
    "paesi alleati. Le modifiche approvate sono reali e rilevanti. I sistemi AI "
    "classificati ad alto rischio dall'Allegato III del regolamento, che includono "
    "quelli usati nella selezione del personale, nel credito, nella gestione delle "
    "infrastrutture critiche, non dovranno essere conformi ai requisiti documentali "
    "e di supervisione umana previsti dal regolamento fino al 2 dicembre 2027. "
    "Le esenzioni per le PMI sono state estese alle cosiddette small mid-caps, "
    "cioe' le aziende con meno di 500 dipendenti e fatturato fino a 100 milioni "
    "di euro."
)

heading(doc, "Gli obblighi di agosto che sono rimasti intatti")

para(doc,
    "La proroga riguarda pero' solo i sistemi ad alto rischio. Gli obblighi di "
    "trasparenza previsti dall'articolo 50 del Regolamento UE, che entrano in "
    "vigore il 2 agosto 2026, non sono stati toccati dall'accordo Omnibus. "
    "Dal 2 agosto, qualsiasi sistema di IA che interagisce con persone fisiche, "
    "come un chatbot sul sito aziendale o un assistente virtuale al telefono, "
    "deve comunicare all'utente che sta interagendo con un sistema automatizzato, "
    "a meno che non sia gia' ovvio dal contesto. I contenuti generati da AI "
    "destinati al pubblico devono essere identificabili come tali in formato "
    "leggibile da macchina. I deepfake con persone reali devono portare "
    "un'etichetta visibile. Queste norme si applicano a tutte le imprese, "
    "indipendentemente dalle dimensioni."
)

para(doc,
    "Secondo le stime dell'Osservatorio Digital Innovation del Politecnico di "
    "Milano, piu' del 40% delle imprese italiane con sito web usa oggi qualche "
    "forma di chatbot o assistente automatizzato per il primo contatto con i "
    "clienti. Per molte di queste imprese, agosto non e' una scadenza astratta: "
    "e' la data entro cui il chatbot deve dichiarare di essere un chatbot. "
    "Chi non si e' ancora mosso ha meno di cinquanta giorni. La questione non "
    "richiede una consulenza specialistica per essere risolta: richiede che "
    "qualcuno in azienda sappia quali strumenti AI sono in uso e verifichi "
    "se rispettano i requisiti minimi di trasparenza verso gli utenti."
)

heading(doc, "Il livello aggiuntivo della Legge 132")

para(doc,
    "Per le imprese italiane, a rendere il quadro ulteriormente articolato c'e' "
    "la Legge n. 132/2025, la legge nazionale sull'intelligenza artificiale "
    "entrata in vigore il 10 ottobre 2025. A differenza del regolamento europeo, "
    "la legge italiana non e' stata modificata dall'accordo Omnibus: gli obblighi "
    "che introduce rimangono con le loro scadenze originali. Il piu' rilevante per "
    "studi professionali e aziende e' quello dell'AI literacy: a partire dal "
    "1 agosto 2026, i datori di lavoro devono garantire che i dipendenti che usano "
    "sistemi AI nell'attivita' lavorativa abbiano ricevuto una formazione adeguata, "
    "proporzionata al livello di interazione con lo strumento. La formazione non "
    "puo' essere generica: deve riguardare gli strumenti specificamente in uso e "
    "le loro limitazioni operative."
)

para(doc,
    "La legge introduce poi il principio antropocentrico, gia' chiarito dai decreti "
    "attuativi approvati il 10 giugno: l'AI supporta il professionista nel giudizio, "
    "non lo sostituisce. Chi usa uno strumento AI per produrre documenti, pareri o "
    "comunicazioni destinate a clienti deve essere in grado di spiegare come e' "
    "stato usato e con quale supervisione. L'Italia si trova in una posizione "
    "singolare: ha anticipato l'Europa con una legge nazionale, e questa legge non "
    "e' stata alleggerita dal processo di revisione europeo. Le imprese italiane "
    "operano con un doppio regime normativo che il rinvio dell'Omnibus non ha "
    "modificato."
)

para(doc,
    "La mappa degli obblighi che si applica a un'impresa italiana nel 2026 e' piu' "
    "articolata di quanto una singola notizia sul rinvio possa far sembrare. "
    "La domanda utile non e' 'quando e' la scadenza?' ma 'quali obblighi si "
    "applicano alla mia situazione specifica?'. Quella mappatura richiede di "
    "distinguere il tipo di sistema AI in uso, il ruolo dell'impresa come "
    "provider o deployer, e le norme applicabili sia a livello europeo sia a "
    "livello nazionale. Chi ha gia' avviato questa ricognizione si trova in "
    "una posizione migliore, non perche' debba necessariamente consegnare un "
    "documento entro agosto, ma perche' capisce il proprio profilo di rischio "
    "e sa su quali aspetti concentrare l'attenzione nei mesi successivi."
)

riferimenti(doc, [
    "DeepElse Blog — 'AI Act Omnibus, accordo del 7 maggio 2026: cosa cambia per le aziende italiane'",
    "Tom's Hardware — 'AI Act rinviato al 2027, ma le norme sulla trasparenza partono subito'",
    "Agenda Digitale — 'Obblighi di trasparenza AI Act: cosa devono fare le aziende dal 2 agosto 2026'",
    "Money.it — 'Dal 2 agosto chi usa l'AI deve dirlo. Cosa cambia per le imprese italiane con l'AI Act'",
    "AiPolicy.it — 'Rinvio AI Act: nuove scadenze 2027-2028 per PMI'",
    "Key4biz — 'AI Act, l'Ue prende tempo e rinvia gli obblighi per i sistemi ad alto rischio'",
    "Legge 23 settembre 2025, n. 132 sull'intelligenza artificiale — Normattiva",
])
doc.save(BASE + "2026-06-19_ai-act-omnibus-rinvio-alto-rischio-imprese.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# LISA a Bologna: il primo supercomputer europeo per l'AI
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Infrastruttura AI")
titolo(
    doc,
    "LISA a Bologna: il primo supercomputer\neuropeo per l'AI generativa.",
    "L'11 giugno al Tecnopolo di Bologna e' stata inaugurata LISA, prima "
    "infrastruttura europea di calcolo ad alte prestazioni dedicata all'AI "
    "generativa nell'ambito dell'EuroHPC. Per le imprese italiane che lavorano "
    "su dati sensibili o vogliono sviluppare capacita' AI proprietarie, "
    "cambia qualcosa.",
    "A cura della Redazione Ratio  •  19 giugno 2026"
)

para(doc,
    "L'11 giugno al Tecnopolo di Bologna, lo stesso polo che ospita il "
    "supercomputer Leonardo, e' stata inaugurata LISA: Learning Infrastructure "
    "for Scalable AI. E' la prima infrastruttura di calcolo ad alte prestazioni "
    "in Europa specificamente dedicata allo sviluppo di sistemi di intelligenza "
    "artificiale generativa. Centosei server GPU di fascia alta, completamente "
    "interconnessi in rete ad alta velocita', progettati per supportare "
    "l'addestramento di modelli linguistici di grandi dimensioni, sistemi di "
    "computer vision e applicazioni multimodali. L'infrastruttura e' finanziata "
    "con 28,2 milioni di euro e si inserisce nell'ambito dell'EuroHPC Joint "
    "Undertaking, il programma europeo per il supercalcolo. E' il primo sistema "
    "per applicazioni AI installato in Europa all'interno dell'infrastruttura "
    "EuroHPC."
)

para(doc,
    "La notizia e' passata quasi esclusivamente nelle pagine di tecnologia e "
    "ricerca. Ha pero' implicazioni che riguardano anche le imprese italiane "
    "che usano o stanno valutando di usare l'intelligenza artificiale in modo "
    "serio, non come strumento accessorio ma come componente di un processo "
    "produttivo o professionale. Per capirle, val la pena partire da una domanda "
    "semplice: quando un'azienda italiana vuole sviluppare o addestrare un modello "
    "AI propriamente inteso, dove lo fa oggi?"
)

heading(doc, "Il problema dell'infrastruttura per le imprese con dati sensibili")

para(doc,
    "La risposta, nella quasi totalita' dei casi, e': su cloud americano. "
    "Amazon Web Services, Microsoft Azure o Google Cloud forniscono "
    "l'infrastruttura computazionale su cui girano i modelli, su cui si "
    "addestrano i sistemi, su cui vengono elaborati i dati. Per la maggior "
    "parte delle applicazioni aziendali, questo non e' un problema: i servizi "
    "cloud americani sono affidabili, ben documentati e accessibili anche per "
    "le PMI. Ma per un sottoinsieme rilevante di casi d'uso, la localizzazione "
    "geografica dei dati e l'infrastruttura su cui vengono elaborati hanno "
    "implicazioni legali e competitive che non si possono ignorare."
)

para(doc,
    "Le aziende che operano in settori regolamentati, come sanita', finanza, "
    "servizi legali e revisione contabile, lavorano spesso con dati personali "
    "o riservati che non possono essere trasmessi su server extra-europei senza "
    "vincoli precisi. Il GDPR e l'AI Act richiedono governance dei dati "
    "verificabile. Un'impresa che voglia addestrare un modello sul proprio "
    "archivio di documenti contrattuali, sui fascicoli dei propri clienti o "
    "sui dati finanziari dei propri assistiti si trova oggi in una posizione "
    "di compromesso: o accetta i termini dei cloud americani e gestisce la "
    "compliance con strumenti contrattuali, o rinuncia a sviluppare capacita' "
    "AI proprietarie. LISA introduce una terza opzione: infrastruttura ad alte "
    "prestazioni in Europa, sotto giurisdizione europea, con accesso regolato "
    "da istituzioni pubbliche italiane ed europee."
)

heading(doc, "L'AI Factory e il progetto da 430 milioni")

para(doc,
    "LISA non e' un progetto isolato. Si inserisce nell'AI Factory di CINECA, "
    "il consorzio interuniversitario che gestisce Leonardo, un programma da "
    "430 milioni di euro finanziato dal PNRR e dalle istituzioni europee, "
    "che mira a creare a Bologna un ecosistema completo per lo sviluppo di AI "
    "generativa a servizio delle imprese italiane, con particolare attenzione "
    "al settore manifatturiero. CINECA e' gia' il referente tecnico di circa "
    "3.000 istituti tra universita', centri di ricerca e enti pubblici italiani. "
    "Il piano e' estendere progressivamente l'accesso anche alle imprese private, "
    "in particolare alle PMI innovative che non hanno accesso alle risorse di "
    "calcolo delle grandi aziende."
)

para(doc,
    "Per ora, l'accesso diretto a LISA e' riservato a universita', centri di "
    "ricerca e progetti finanziati da istituzioni pubbliche. Ma la traiettoria "
    "e' verso un'apertura alle imprese, attraverso bandi specifici e accordi "
    "di collaborazione con il sistema universitario. Le imprese che stanno "
    "sviluppando capacita' AI e che operano in settori con vincoli di "
    "localizzazione dei dati farebbero bene a monitorare i programmi di accesso "
    "che verranno pubblicati nei prossimi mesi. Non si tratta di un'alternativa "
    "immediata ai servizi cloud americani: si tratta di un'opzione strutturale "
    "che potrebbe diventare rilevante per casi d'uso specifici, specialmente "
    "dove la riservatezza dei dati e' un requisito non negoziabile."
)

para(doc,
    "L'aspetto piu' interessante di LISA non e' tecnico ma strategico. Mostra "
    "che l'Europa sta investendo in infrastruttura, non solo in normativa. "
    "L'AI Act e la Legge 132 hanno stabilito le regole. Leonardo e LISA stanno "
    "costruendo la capacita' di giocare la partita con strumenti propri. "
    "Per un commercialista o un consulente che assiste imprese innovative, "
    "la domanda da porre ai propri clienti che stanno valutando investimenti "
    "in AI non e' solo 'quale strumento usate?' ma 'su quale infrastruttura "
    "girano i vostri dati, e avete verificato che sia compatibile con i vincoli "
    "normativi del vostro settore?'. La risposta a quella domanda diventa piu' "
    "articolata da quando LISA e' operativa."
)

riferimenti(doc, [
    "Il Resto del Carlino — 'Supercomputer Leonardo, si chiama Lisa la prima infrastruttura in Europa dedicata all'AI' (giugno 2026)",
    "Meteoweb — 'Il supercomputer Leonardo cresce: arriva Lisa, primo sistema europeo dedicato all'IA' (maggio-giugno 2026)",
    "Il Giornale — 'Il supercomputer Leonardo si rafforza sul fronte AI con l'arrivo di Lisa' (giugno 2026)",
    "Il Nordest Quotidiano — 'Il supercomputer Leonardo cresce e al Tecnopolo arriva Lisa' (28 maggio 2026)",
    "Digital World Italia — 'LISA e' il supercomputer universitario per l'IA che affianchera' Leonardo a Bologna'",
    "Industria Italiana — 'AI Factory: 430 milioni per un'IA tutta italiana dedicata al manifatturiero'",
    "Il Sole 24 Ore — 'Computing power and artificial intelligence: how Leonardo enables innovation'",
])
doc.save(BASE + "2026-06-19_lisa-bologna-supercomputer-ai-generativa.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Quando la piattaforma cambia le regole a meta' partita
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Strategia AI")
titolo(
    doc,
    "Quando la piattaforma cambia\nle regole a meta' partita.",
    "Il 15 giugno Anthropic ha annunciato che l'uso programmatico di Claude "
    "sarebbe passato a fatturazione a consumo. Tre giorni dopo ha fatto marcia "
    "indietro. Ma la domanda che l'episodio pone non ha ancora risposta: cosa "
    "succede quando lo strumento su cui hai costruito un processo decide di "
    "cambiare le condizioni?",
    "A cura della Redazione Ratio  •  19 giugno 2026"
)

para(doc,
    "Il 15 giugno Anthropic ha inviato una comunicazione ai propri utenti Pro, "
    "Max, Team ed Enterprise con un annuncio che, per chi usa Claude in modo "
    "programmatico, aveva l'aspetto di una doccia fredda. A partire da quella "
    "data, l'uso di Claude attraverso l'Agent SDK, attraverso script automatizzati "
    "o attraverso applicazioni di terze parti collegate all'account Claude avrebbe "
    "consumato un credito mensile separato, fatturato a tariffe API complete. "
    "Per chi aveva costruito flussi di lavoro automatizzati sul piano Pro da "
    "venti euro al mese, l'aumento potenziale era di dieci volte o piu'. "
    "Per chi aveva investito settimane a costruire un processo aziendale su Claude "
    "come infrastruttura, era uno scenario che richiedeva di riconsiderare "
    "l'intera architettura."
)

para(doc,
    "Tre giorni dopo, Anthropic ha fatto marcia indietro. Il 18 giugno ha "
    "comunicato che 'per ora non cambia nulla', che l'uso degli agenti continua "
    "ad essere coperto dai piani di abbonamento esistenti, e che eventuali "
    "modifiche future verranno comunicate con anticipo sufficiente. La crisi, "
    "se cosi' si vuole chiamare, e' rientrata rapidamente. Ma l'episodio lascia "
    "in sospeso una domanda che riguarda non solo Claude, ma qualsiasi impresa "
    "o studio professionale che abbia costruito processi su piattaforme AI "
    "esterne: cosa succede quando la piattaforma decide di cambiare le condizioni "
    "a meta' partita?"
)

heading(doc, "I tre rischi che l'episodio ha reso visibili")

para(doc,
    "Il primo rischio e' quello di prezzo. I piani di abbonamento alle "
    "piattaforme AI sono stati, negli ultimi due anni, significativamente "
    "sovvenzionati dai provider per accelerare l'adozione di massa. Il costo "
    "reale del calcolo computazionale per un uso intensivo supera di molto "
    "quello dei piani flat mensili. Prima o poi, la struttura economica si "
    "aggiustera'. L'annuncio di Anthropic era un tentativo in quella direzione, "
    "poi ritirato per motivi competitivi. Ma la direzione e' chiara. Chi ha "
    "costruito un processo aziendale sulla base di un costo di venti euro al "
    "mese deve chiedersi quanto sarebbe disposto a pagare se quel costo "
    "diventasse duecento, e se la risposta influenza la sostenibilita' del "
    "processo stesso."
)

para(doc,
    "Il secondo rischio e' quello di accesso, reso ancora piu' evidente "
    "dall'episodio di Fable 5 nello stesso periodo: una piattaforma AI puo' "
    "diventare indisponibile per motivi tecnici, normativi o commerciali con "
    "un preavviso brevissimo. Per un uso occasionale, questo e' un inconveniente. "
    "Per un processo che elabora documenti, risponde ai clienti o produce "
    "reportistica su base quotidiana, e' un blocco operativo. Quante ore "
    "lavorative sarebbero perse se lo strumento AI su cui si basa un processo "
    "smettesse di rispondere per ventiquattro ore? Quella cifra vale la pena "
    "calcolarla adesso, non quando si verifica il problema."
)

para(doc,
    "Il terzo rischio e' piu' sottile: il comportamento dei modelli. Le "
    "piattaforme AI aggiornano i propri modelli in modo silenzioso e frequente. "
    "Un system prompt che produce output coerenti e utili oggi puo' produrre "
    "risultati diversi dopo un aggiornamento del modello sottostante, senza "
    "che l'utente riceva alcuna notifica. Per un uso esplorativo, questo non "
    "e' un problema. Per un processo automatizzato che produce documenti da "
    "consegnare ai clienti o dati da inserire in un gestionale, la variabilita' "
    "del comportamento del modello e' un rischio operativo reale. Le API di "
    "Anthropic e OpenAI consentono di bloccare la versione del modello, ma "
    "richiede di gestirla attivamente, non di assumere che il comportamento "
    "rimanga stabile nel tempo."
)

heading(doc, "Come costruire senza diventare ostaggi")

para(doc,
    "Nessuno di questi rischi significa che non si debbano costruire processi "
    "su piattaforme AI. Significa che quei processi vanno costruiti con la "
    "stessa consapevolezza con cui si valuta qualsiasi fornitore critico. "
    "Se un commercialista affida la tenuta della contabilita' a un software "
    "gestionale, verifica che il fornitore sia solido, che i dati siano "
    "esportabili, che esista un percorso di migrazione se il rapporto si "
    "interrompe. La stessa logica si applica a un processo costruito su Claude "
    "o ChatGPT: non come ostacolo all'adozione, ma come parte di una "
    "progettazione responsabile."
)

para(doc,
    "In pratica, questo significa alcune cose concrete. Documentare il processo "
    "e non solo lo strumento: sapere cosa fa il sistema AI in ogni fase, quali "
    "input riceve e quali output produce, in modo da poterlo descrivere e "
    "riprodurre anche con strumenti diversi. Testare regolarmente l'output su "
    "casi di riferimento fissi, per verificare che il comportamento del modello "
    "non sia cambiato silenziosamente. Calcolare il costo unitario per singola "
    "transazione o documento prodotto, per capire la sostenibilita' a diverse "
    "tariffe. Assicurarsi che esista, nello studio o nell'impresa, qualcuno "
    "che sappia svolgere il processo manualmente, anche se con meno velocita'."
)

para(doc,
    "La marcia indietro di Anthropic e' una buona notizia per chi usa Claude "
    "in abbonamento. Ma non risponde alla domanda di fondo: quando quella "
    "modifica arrivera', e prima o poi arrivera', si sara' pronti? La risposta "
    "non richiede di cambiare piattaforma o di rinunciare ai benefici degli "
    "strumenti AI. Richiede di costruire con cognizione di causa, sapendo che "
    "la tecnologia disponibile oggi potrebbe non esserlo, o non alle stesse "
    "condizioni, domani. Questa non e' una posizione conservatrice. E' la "
    "stessa prudenza che si applica a qualsiasi rapporto con un fornitore "
    "di cui si dipende strutturalmente."
)

riferimenti(doc, [
    "Tom's Hardware — 'Claude diventa a consumo (per alcuni) dal 15 giugno. Addio abbonamento flat' (giugno 2026)",
    "Tom's Hardware — 'Claude a consumo via SDK, Anthropic fa un passo indietro' (giugno 2026)",
    "Tecnoandroid — 'Anthropic fa marcia indietro: stop alla fatturazione a token per Claude Agent SDK' (giugno 2026)",
    "Infotelematico — 'Abbonamento Claude: cosa cambia dal 15 giugno 2026'",
    "TechSy.io — 'Credito Claude Agent SDK spiegato (15 giugno 2026)'",
    "CorriereNerd.it — 'Claude Fable 5: il mistero dell'IA Anthropic spenta dopo tre giorni' (giugno 2026)",
    "Cosimo.dev — 'Fable 5 e Mythos 5 spenti dopo 3 giorni: cosa e' successo' (giugno 2026)",
])
doc.save(BASE + "2026-06-19_dipendenza-piattaforme-ai-rischio-operativo.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
