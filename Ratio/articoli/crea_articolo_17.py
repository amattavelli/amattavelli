"""
Articolo 17: Il ROI dell'AI generativa per le PMI italiane — dal benchmark alla cassa
Ratio/articoli/2026-04-10_roi-ai-generativa-pmi-italiane.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-10_roi-ai-generativa-pmi-italiane.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_paragraph(doc, text, size=11, italic=False, color=None,
                  space_after=8, space_before=0,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


# ---- TESTATA ----
p_rivista = doc.add_paragraph()
p_rivista.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = p_rivista.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
run_r.font.name = 'Calibri'
run_r.font.size = Pt(9)
run_r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run_r.font.bold = True
run_r.font.all_caps = True

add_separator(doc)

p_data = doc.add_paragraph()
p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_d = p_data.add_run("Aprile 2026  |  Gestione e Strategia")
run_d.font.name = 'Calibri'
run_d.font.size = Pt(8.5)
run_d.italic = True
run_d.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_after = Pt(6)
run_t = p_title.add_run(
    "Il ROI dell\u2019AI per le PMI italiane: dai benchmark alla cassa"
)
run_t.font.name = 'Calibri'
run_t.font.size = Pt(24)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_after = Pt(10)
run_s = p_sub.add_run(
    "Il 92% delle aziende early adopter riporta un ritorno positivo sull\u2019investimento in AI. "
    "Ma cosa significano questi numeri per una PMI italiana? Come si misura davvero il valore "
    "generato, e da dove conviene cominciare?"
)
run_s.font.name = 'Calibri'
run_s.font.size = Pt(13)
run_s.italic = True
run_s.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_separator(doc)

p_autore = doc.add_paragraph()
p_autore.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_autore.paragraph_format.space_after = Pt(16)
run_a = p_autore.add_run("A cura della Redazione Ratio  \u2022  10 aprile 2026")
run_a.font.name = 'Calibri'
run_a.font.size = Pt(9)
run_a.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- CORPO ----

add_paragraph(doc,
    "Nel dibattito sull\u2019intelligenza artificiale in azienda, i numeri globali sono ormai ovunque. "
    "Il 92% delle aziende che ha adottato AI generativa in modo strutturato riporta un ROI positivo, "
    "con un ritorno medio del 49% sull\u2019investimento. Il break-even si raggiunge in media tra gli "
    "otto e i quattordici mesi. Su un orizzonte triennale, la media aggregata dei progetti AI nelle PMI "
    "supera il 340% di ritorno. Sono dati diffusi dalle principali societ\u00e0 di analisi di mercato "
    "e confermati da survey su campioni di migliaia di aziende in Europa e Nord America."
)

add_paragraph(doc,
    "Il problema di questi numeri \u00e8 che non dicono nulla di utile finch\u00e9 non vengono "
    "contestualizzati. Un ROI del 49% su cosa? In quale settore? Con quale uso specifico dell\u2019AI? "
    "Per un imprenditore o un consulente che deve decidere se e dove investire, la domanda non \u00e8 "
    "\u201cl\u2019AI fa guadagnare le aziende?\u201d \u2014 la risposta \u00e8 s\u00ec, ma non \u00e8 "
    "questa la domanda giusta. La domanda \u00e8: \u201ccosa funziona concretamente in un\u2019azienda "
    "come la nostra, con le nostre dimensioni, i nostri processi, le nostre competenze interne?\u201d"
)

add_heading2(doc, "Dove il ROI \u00e8 pi\u00f9 visibile: i casi italiani")

add_paragraph(doc,
    "I casi documentati di PMI italiane che hanno misurato il ritorno dell\u2019AI mostrano un pattern "
    "ricorrente: il valore pi\u00f9 alto si concentra nei processi ad alto volume e bassa variabilit\u00e0. "
    "Uno studio di consulenza fiscale e contabile che ha automatizzato l\u2019estrazione dei dati dalle "
    "fatture elettroniche ha misurato un risparmio di quindici ore settimanali di lavoro operativo, "
    "l\u2019eliminazione sostanziale degli errori di inserimento e la capacit\u00e0 di gestire il 30% "
    "di clienti in pi\u00f9 senza aumentare il personale. Con un investimento iniziale di poche migliaia "
    "di euro in una soluzione AI dedicata all\u2019estrazione documentale, il break-even \u00e8 stato "
    "raggiunto in meno di sei mesi."
)

add_paragraph(doc,
    "Un\u2019azienda manifatturiera nella lavorazione meccanica ha introdotto un sistema di controllo "
    "qualit\u00e0 visivo basato su AI: riduzione dei difetti del 78%, calo degli scarti del 23%, "
    "ritorno sull\u2019investimento raggiunto in otto mesi. Una societ\u00e0 di distribuzione ha usato "
    "AI generativa per automatizzare la stesura delle risposte alle email commerciali standardizzate: "
    "il team vendite ha recuperato due ore al giorno a persona, reinvestite in attivit\u00e0 di "
    "sviluppo commerciale. Casi diversi, dimensioni diverse, settori diversi \u2014 ma un elemento "
    "comune: si \u00e8 partiti da un processo specifico con un problema misurabile, non da una "
    "strategia AI aziendale."
)

add_heading2(doc, "Il divario che ancora esiste")

add_paragraph(doc,
    "A fronte di questi risultati, il divario tra grandi imprese e PMI nell\u2019adozione dell\u2019AI "
    "rimane significativo. Il 71% delle grandi imprese italiane ha avviato almeno un progetto AI "
    "strutturato; tra le PMI con meno di cinquanta addetti, la percentuale scende all\u20198%. "
    "Le ragioni sono note: mancanza di competenze interne, incertezza sui costi, difficolt\u00e0 "
    "nell\u2019identificare i fornitori giusti, e \u2014 forse la pi\u00f9 sottovalutata \u2014 "
    "assenza di qualcuno che faccia da traduttore tra la tecnologia e i processi aziendali specifici."
)

add_paragraph(doc,
    "Questo \u00e8 esattamente lo spazio in cui il commercialista o il consulente aziendale pu\u00f2 "
    "aggiungere valore. Non vendendo soluzioni AI, ma aiutando l\u2019imprenditore a rispondere a tre "
    "domande precise: qual \u00e8 il processo che consuma pi\u00f9 ore a basso valore aggiunto? Esiste "
    "gi\u00e0 uno strumento AI che lo automatizza, o \u00e8 necessario costruire qualcosa? Come si "
    "misurer\u00e0 il risultato? Se queste tre domande trovano risposta concreta prima di qualsiasi "
    "investimento, la probabilit\u00e0 di un ROI positivo aumenta considerevolmente."
)

add_heading2(doc, "Come misurare il ROI: un approccio pratico")

add_paragraph(doc,
    "Misurare il ritorno dell\u2019AI non richiede modelli econometrici complessi. Per la maggior parte "
    "dei casi d\u2019uso tipici di una PMI, bastano tre metriche: ore di lavoro risparmiate (convertite "
    "in costo orario del personale coinvolto), tasso di errore prima e dopo l\u2019introduzione dello "
    "strumento, e capacit\u00e0 aggiuntiva generata \u2014 ovvero quante unit\u00e0 di output in pi\u00f9 "
    "si riesce a produrre con le stesse risorse. A queste si aggiunge, dove applicabile, la velocit\u00e0 "
    "di risposta al cliente, un fattore sempre pi\u00f9 rilevante in settori dove i tempi di risposta "
    "sono diventati un fattore competitivo."
)

add_paragraph(doc,
    "Il punto di partenza \u00e8 sempre la baseline: misurare il processo cos\u00ec com\u2019\u00e8 oggi, "
    "prima di introdurre qualsiasi strumento. Senza una baseline, non si pu\u00f2 misurare nulla. "
    "\u00c8 un passaggio che molte aziende saltano, e che poi rende impossibile dimostrare il valore "
    "generato. Il consulente che aiuta il cliente a costruire questa baseline prima dell\u2019investimento "
    "fa un lavoro che vale molto pi\u00f9 di qualsiasi raccomandazione tecnologica."
)

add_heading2(doc, "Il 2026 come anno di svolta per le PMI")

add_paragraph(doc,
    "Il 2026 \u00e8 l\u2019anno in cui l\u2019AI esce dalla fase sperimentale per diventare uno "
    "strumento operativo quotidiano nelle aziende che scelgono di usarla. I costi dei modelli sono "
    "scesi del 90% negli ultimi due anni. Le interfacce sono diventate accessibili. I fornitori di "
    "software gestionale stanno integrando funzionalit\u00e0 AI nei prodotti gi\u00e0 in uso dalle "
    "aziende. Chi aspetta che \u201cla tecnologia maturi\u201d rischia di scoprire che \u00e8 gi\u00e0 "
    "matura \u2014 e che i concorrenti lo sanno gi\u00e0 da un anno."
)

add_separator(doc)

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(10)
run_n = p_note.add_run("Fonti e riferimenti")
run_n.font.name = 'Calibri'
run_n.font.size = Pt(9)
run_n.font.bold = True
run_n.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Lineaedp \u2014 AI generativa: il 92% delle aziende registra ROI positivo, ritorni medi del 49% (2026)",
    "BitMat \u2014 AI generativa: il 92% delle aziende registra un ROI positivo",
    "DeepElse Blog \u2014 AI per PMI italiane: guida completa 2026",
    "Paradigma.it \u2014 AI generativa nelle PMI: linee guida 2026 per un uso responsabile",
    "ManagementCue \u2014 Gli agenti AI in azienda: cosa cambier\u00e0 davvero nel 2026 per le PMI italiane",
]
for s in sources:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_after = Pt(2)
    run_s = p_s.add_run(f"\u2022 {s}")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(8.5)
    run_s.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(16)
run_f = p_footer.add_run(
    "\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte"
)
run_f.font.name = 'Calibri'
run_f.font.size = Pt(8)
run_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_f.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
