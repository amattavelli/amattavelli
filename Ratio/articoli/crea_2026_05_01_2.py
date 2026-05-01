"""
Articolo 2 — L'obbligo di AI literacy che agosto renderà visibile
Ratio/articoli/2026-05-01_ai-literacy-obbligo-formazione-agosto-2026.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-05-01_ai-literacy-obbligo-formazione-agosto-2026.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Maggio 2026  |  Normativa e Compliance AI")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("L'obbligo di formazione AI\nche agosto renderà visibile")
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "L'articolo 4 dell'AI Act è in vigore dal febbraio 2025. "
    "La maggioranza degli studi e delle PMI italiane non ha ancora documentato nulla. "
    "Dal 3 agosto 2026 partono i controlli."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  1 maggio 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Il titolare dello studio apre ChatGPT ogni mattina. Lo usa per stendere le bozze "
    "delle comunicazioni ai clienti, riassumere le circolari dell'Agenzia delle Entrate, "
    "preparare le slide per le assemblee dei soci. Ha convinto anche due colleghi ad "
    "adottarlo, e tutti e tre lo trovano utile. Nessuno, in nessuna delle conversazioni "
    "che hanno avuto su questo strumento, ha mai nominato la parola formazione."
)

para(doc,
    "Eppure, dal 2 febbraio 2025, ogni soggetto che utilizza sistemi di intelligenza "
    "artificiale in un contesto professionale strutturato è tenuto ad assicurare un livello "
    "adeguato di competenza sull'AI al proprio personale. Lo stabilisce l'articolo 4 del "
    "Regolamento (UE) 2024/1689, l'AI Act europeo. La norma è già in vigore da oltre un "
    "anno. Quello che cambia il 3 agosto 2026 è che le autorità nazionali competenti "
    "inizieranno le attività di vigilanza e potranno irrogare sanzioni per le violazioni "
    "accertate."
)

heading(doc, "Cosa prevede l'obbligo in concreto")

para(doc,
    "L'articolo 4 non prescrive ore minime di formazione né il superamento di un esame. "
    "Non richiede certificazioni esterne né la partecipazione a corsi accreditati da enti "
    "pubblici. Chiede che il personale che utilizza o sviluppa sistemi AI abbia un livello "
    "sufficiente di conoscenza delle caratteristiche, delle capacità e dei limiti degli "
    "strumenti in uso, e che l'organizzazione sappia dimostrarlo in modo documentato."
)

para(doc,
    "La Commissione Europea ha pubblicato nel gennaio 2026 una serie di FAQ chiarificatori "
    "sull'obbligo. Il punto più rilevante per le PMI è questo: l'adeguatezza della "
    "formazione si misura in modo proporzionato alle dimensioni dell'organizzazione e alla "
    "complessità dei sistemi utilizzati. Un'impresa con cinque dipendenti che usa solo un "
    "assistente virtuale per le email non è trattata come una società finanziaria con "
    "modelli predittivi di scoring del credito. L'obbligo esiste per entrambe, ma il "
    "livello di documentazione richiesto è calibrato sulla realtà operativa."
)

para(doc,
    "La documentazione raccomandata (non imposta come requisito formale, ma fortemente "
    "indicata nelle FAQ della Commissione) include un registro interno delle attività "
    "formative, la descrizione dei sistemi AI effettivamente in uso, l'indicazione del "
    "personale coinvolto e la periodicità dell'aggiornamento previsto."
)

heading(doc, "Il rischio di trattarlo come un adempimento da spuntare")

para(doc,
    "Molte aziende stanno reagendo all'obbligo di AI literacy come hanno reagito in passato "
    "ai corsi obbligatori sulla sicurezza nei luoghi di lavoro: un'ora di video, un click "
    "per attestare la partecipazione, un file archiviato nel cloud. Questo approccio "
    "soddisfa formalmente l'obbligo, ma non produce nessuna delle competenze che la norma "
    "intende generare."
)

para(doc,
    "Il problema non è solo legale. Un team che usa strumenti AI senza capire cosa può "
    "andare storto è un team che delegherà compiti sbagliati, non verificherà gli output "
    "nelle situazioni in cui servirebbe farlo, e si accorgerà degli errori quando le "
    "conseguenze sono già visibili ai clienti. La formazione sull'AI serve prima di tutto "
    "a costruire un'abitudine alla verifica critica degli output, non alla fiducia cieca "
    "nello strumento. Il modello genera testo plausibile: decidere se quel testo è anche "
    "corretto e adeguato al contesto è ancora compito del professionista."
)

heading(doc, "Come costruire un percorso che abbia senso")

para(doc,
    "Per uno studio professionale o una PMI che vuole affrontare questo tema con metodo, "
    "il punto di partenza è l'inventario degli strumenti: quali sistemi AI sono in uso, "
    "chi li usa, per quali attività specifiche. Questa ricognizione, che richiede qualche "
    "ora di lavoro interno, è anche il prerequisito per la classificazione dei sistemi ai "
    "fini dell'AI Act più in generale."
)

para(doc,
    "A partire dall'inventario si costruisce un percorso formativo calibrato. I contenuti "
    "minimi utili per ogni persona che usa strumenti AI in ambito professionale includono "
    "quattro aree: come funziona un modello linguistico e perché può produrre errori "
    "plausibili; come verificare un output generato dall'AI senza accettarlo per buono; "
    "quali tipi di decisioni non possono essere delegate a un modello senza revisione umana; "
    "quali sono i rischi legali dell'uso di dati riservati su piattaforme non conformi al "
    "GDPR europeo."
)

para(doc,
    "Questi contenuti non richiedono un corso lungo. Due o tre ore di formazione ben "
    "strutturata, documentata con un verbale e archiviata insieme all'inventario dei "
    "sistemi, è sufficiente per la maggior parte degli studi professionali e delle PMI "
    "italiane. La periodicità minima consigliata dalla Commissione è annuale, con "
    "aggiornamenti ogni volta che vengono adottati nuovi strumenti significativi."
)

heading(doc, "Il valore che rimane dopo la scadenza")

para(doc,
    "Chi ha ancora qualche mese prima del 3 agosto ha il tempo di fare questa cosa bene, "
    "senza fretta. Un percorso formativo costruito adesso, con attenzione ai sistemi "
    "realmente in uso e ai rischi concreti che ne derivano, vale molto di più di una "
    "documentazione compilata nell'ultima settimana di luglio per rispettare una scadenza "
    "regolamentare. Quello che rimane, dopo la formazione fatta con cura, è un team più "
    "capace di usare questi strumenti con piena consapevolezza delle conseguenze, non solo "
    "più al sicuro da sanzioni. La differenza, nel lavoro quotidiano, si vede."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "Regolamento (UE) 2024/1689 — AI Act, art. 4 (GU UE, 12 luglio 2024)",
    "Commissione Europea — FAQ sull’obbligo di AI literacy (gennaio 2026)",
    "ADVANT Nctm — Le nuove FAQ della Commissione sull’obbligo di AI literacy (2026)",
    "Randstad — AI literacy e obbligo di formazione nelle aziende: come muoversi? (2026)",
    "Agenda Digitale — AI Literacy, l’obbligo in Europa: ecco in cosa consiste (2026)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
