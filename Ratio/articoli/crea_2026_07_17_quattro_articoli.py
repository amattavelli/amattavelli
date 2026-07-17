"""
Quattro articoli Ratio -- 17 luglio 2026

1. 2026-07-17_ai-act-articolo-50-chatbot-obbligo-dichiarazione.docx
2. 2026-07-17_fable5-sol-guerra-modelli-come-scegliere-studio.docx
3. 2026-07-17_shadow-operation-agenti-ai-senza-governance-pmi.docx
4. 2026-07-17_fiducia-ai-decisioni-finanziarie-paradosso-pmi.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# AI Act articolo 50: il chatbot deve dichiararsi
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Normativa AI")
titolo(
    doc,
    "Dal 2 agosto ogni chatbot aziendale\ndeve dichiararsi: 16 giorni per adeguarsi.",
    "L'articolo 50 dell'AI Act entra in piena applicazione il 2 agosto 2026. "
    "Chi ha un assistente virtuale sul sito, un bot su WhatsApp o un sistema "
    "automatico di risposta ai clienti ha sedici giorni per adeguarsi. Non "
    "si tratta di un obbligo tecnico complesso: si tratta di comunicare "
    "chiaramente che dall'altra parte c'e' una macchina, non una persona.",
    "A cura della Redazione Ratio  •  17 luglio 2026"
)

para(doc,
    "Uno studio di consulenza tributaria ha installato a gennaio un chatbot "
    "sul proprio sito per gestire le domande ricorrenti dei clienti: orari, "
    "documenti da portare, scadenze fiscali del mese corrente. Il titolare "
    "ha investito circa ottocento euro nell'implementazione e lo considera "
    "uno strumento utile, che scarica il centralino nelle ore di punta. "
    "Quando gli e' stato chiesto dell'articolo 50 dell'AI Act e della "
    "scadenza del 2 agosto, la risposta e' stata che quell'obbligo riguardava "
    "le grandi piattaforme tecnologiche, non un piccolo studio professionale. "
    "L'obbligo si applica invece esattamente a questo tipo di utilizzo: "
    "qualsiasi sistema che interagisce con persone fisiche per conto di "
    "un'organizzazione deve comunicare, in modo visibile all'inizio "
    "dell'interazione, che non e' una persona."
)

para(doc,
    "L'articolo 50 del Regolamento UE 2024/1689 stabilisce un obbligo di "
    "trasparenza per i sistemi AI che generano contenuti sintetici o "
    "interagiscono con persone fisiche. Per le aziende e gli studi professionali "
    "che hanno attivato chatbot sui propri siti, su WhatsApp Business o su "
    "altri canali di messaggistica, l'obbligo diventa pienamente applicabile "
    "il 2 agosto 2026: il soggetto che interagisce con il sistema deve essere "
    "informato, all'inizio della conversazione, che sta parlando con un "
    "sistema automatizzato basato su intelligenza artificiale. L'esenzione "
    "esiste solo quando dal contesto risulta evidente che si tratta di un "
    "sistema automatico, ma questa condizione non copre la maggior parte "
    "dei chatbot su siti professionali, che per loro natura tendono a "
    "simulare risposte di tipo umano."
)

para(doc,
    "Il soggetto obbligato dall'articolo 50 e' il deployer, ovvero "
    "l'organizzazione che usa un sistema AI in un'attivita' rivolta a persone "
    "fisiche. In termini pratici: uno studio o un'azienda che ha acquistato "
    "o configurato un chatbot e lo ha messo online e' deployer di quel "
    "sistema, anche se la tecnologia sottostante e' stata costruita da "
    "un fornitore esterno. L'obbligo non ricade sul fornitore del chatbot: "
    "ricade su chi lo usa per interagire con i propri clienti o utenti. "
    "Questo e' un elemento che molte PMI fraintendono: pensando che sia "
    "il fornitore del software a dover adempiere, si trovano in una situazione "
    "in cui nessuno ha curato la conformita'."
)

heading(doc, "Cosa deve comparire e dove: la regola minima per essere conformi")

para(doc,
    "L'obbligo di disclosure si soddisfa con un messaggio chiaro all'inizio "
    "di ogni conversazione che informa l'utente di stare interagendo con un "
    "sistema AI automatizzato. Il messaggio non deve essere lungo: una frase "
    "come 'Stai conversando con un assistente automatico basato su intelligenza "
    "artificiale. Per esigenze urgenti puoi contattare lo studio direttamente "
    "al numero...' soddisfa il requisito. L'obbligo si applica dal primo "
    "messaggio di interazione, non dopo alcuni scambi. Molte piattaforme "
    "(Tidio, Intercom, Freshchat, Crisp) hanno gia' un'opzione configurabile "
    "per mostrare una disclosure automatica all'avvio della conversazione; "
    "per chi usa WhatsApp Business con integrazioni di risposta automatica, "
    "la configurazione e' leggermente piu' articolata ma gestibile dal pannello "
    "della piattaforma di messaggistica. Il punto centrale e' che la disclosure "
    "deve essere proattiva, non attivata da una richiesta esplicita dell'utente."
)

para(doc,
    "Il regime sanzionatorio per la violazione dell'articolo 50 prevede "
    "sanzioni amministrative fino a quindici milioni di euro o al tre per cento "
    "del fatturato mondiale annuo, a seconda di quale valore sia piu' elevato, "
    "applicate dalle autorita' di vigilanza nazionali. Per uno studio di piccole "
    "dimensioni o una PMI con un chatbot locale, il rischio concreto di una "
    "sanzione di quella portata e' marginale nel breve periodo: le autorita' "
    "si concentreranno in primo luogo sui soggetti di grandi dimensioni e "
    "sui casi sistematici. Cio' che e' gia' operativo e' il diritto del "
    "cittadino a segnalare una violazione: un cliente o utente che scopre "
    "di aver interagito con un sistema AI senza essere stato informato puo' "
    "presentare un reclamo all'autorita' nazionale. In Italia, l'autorita' "
    "di vigilanza designata nell'ambito dell'AI Act e' l'AgID, Agenzia "
    "per l'Italia Digitale."
)

heading(doc, "I sedici giorni che restano: cosa fare e in che ordine")

para(doc,
    "La lista di azioni per studi e aziende con chatbot o sistemi di "
    "interazione automatizzata e' gestibile in poche ore. Il primo passo "
    "e' mappare i sistemi AI che interagiscono con terzi per conto "
    "dell'organizzazione: chatbot su siti web, risponditori automatici su "
    "WhatsApp o altri canali, sistemi di risposta email automatizzata, "
    "assistenti virtuali integrati in applicazioni rivolte ai clienti. "
    "Il secondo passo e' verificare le impostazioni di ciascuna piattaforma "
    "utilizzata e controllare se sia configurabile un messaggio di disclosure "
    "all'avvio. Il terzo passo e' scrivere il messaggio, in italiano, "
    "chiaro e posizionato all'inizio di ogni conversazione. Il quarto "
    "passo e' documentare quanto fatto: data di configurazione, testo del "
    "messaggio e uno screenshot della configurazione attiva."
)

para(doc,
    "L'obbligo dell'articolo 50 e' una delle poche disposizioni dell'AI Act "
    "che richiede un'azione operativa concreta, non solo documentazione. "
    "Uno studio che non ha ancora provveduto puo' mettersi in regola in meno "
    "di un pomeriggio. Chi affronta questa scadenza con metodo si trova anche "
    "con un modello applicabile alle future implementazioni AI, perche' "
    "requisiti di trasparenza analoghi si estenderanno via via che piu' "
    "sistemi AI vengono integrati nei processi di relazione con il cliente. "
    "Costruire adesso la procedura significa non doverla reinventare "
    "ogni volta che si aggiunge uno strumento."
)

riferimenti(doc, [
    "Regolamento UE 2024/1689 (AI Act), articolo 50 (Obblighi di trasparenza) -- EUR-Lex",
    "AgID -- Autorita' nazionale di vigilanza AI Act -- agid.gov.it",
    "LaborProject.it -- 'AI Act dal 2 agosto 2026: cosa devono fare aziende, consulenti e DPO' (14 luglio 2026)",
    "UniverseIT.blog -- 'AI Act 2 agosto 2026: cosa cambia per le imprese'",
    "Certifico.com -- 'AI Act: obblighi per le imprese dal 2 agosto 2026'",
    "AscenSys.it -- 'AI Act 2026 PMI: obblighi, scadenze e sanzioni'",
    "YellowTech.it -- 'AI Act 2026: compliance per aziende italiane'",
    "LeanBet.eu -- 'AI Act ultima fase: obblighi, scadenze e checklist per le aziende'",
])
doc.save(BASE + "2026-07-17_ai-act-articolo-50-chatbot-obbligo-dichiarazione.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# Fable 5 e GPT-5.6 Sol: come scegliere il modello giusto
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Strumenti e Modelli AI")
titolo(
    doc,
    "Fable 5 gratis fino al 19 luglio, Sol senza limiti:\ncome scegliere il modello giusto per lo studio.",
    "Anthropic ha prolungato l'accesso senza costi aggiuntivi a Fable 5 "
    "fino al 19 luglio, mentre OpenAI ha rimosso i limiti a GPT-5.6 Sol. "
    "Due mosse nella stessa settimana che aprono una finestra concreta per "
    "testare i modelli piu' capaci sul mercato senza costi aggiuntivi. "
    "Ma la scelta del modello giusto non si decide sui benchmark.",
    "A cura della Redazione Ratio  •  17 luglio 2026"
)

para(doc,
    "Uno studio di consulenza aziendale ha dedicato due settimane a luglio "
    "a valutare quale strumento AI adottare come assistente principale. "
    "Ha testato GPT-5.6 Terra, Gemini Ultra 2.5 e Claude Opus 4.8 su "
    "compiti reali: revisione di contratti commerciali, strutturazione di "
    "pareri su temi fiscali, sintesi di relazioni finanziarie. "
    "Il 13 luglio e' arrivata la notizia: Anthropic ha esteso fino al 19 "
    "luglio l'accesso senza costi aggiuntivi a Claude Fable 5 per gli "
    "abbonati Pro, Max, Team ed Enterprise, e contemporaneamente OpenAI "
    "ha rimosso i limiti di utilizzo a GPT-5.6 Sol. Due mosse che, "
    "nella stessa settimana, hanno cambiato il panorama competitivo "
    "tra i modelli AI di punta e aperto una finestra che vale la pena "
    "usare con metodo."
)

para(doc,
    "Fable 5 e' il modello piu' capace rilasciato da Anthropic a maggio 2026, "
    "posizionato al di sopra di Claude Opus 4.8 per ragionamento e "
    "capacita' analitica. Fino al 13 luglio l'accesso era limitato a "
    "specifici piani a pagamento con quote di consumo. Il prolungamento "
    "al 19 luglio e' una risposta competitiva alla mossa di OpenAI con Sol: "
    "Anthropic sta consentendo ai propri abbonati di testare il modello "
    "piu' potente senza costo aggiuntivo, con l'obiettivo implicito di "
    "influenzare le scelte di adozione di chi sta valutando le piattaforme. "
    "Il meccanismo e' consolidato nei mercati software, ma per gli studi "
    "professionali che usano questi strumenti quotidianamente l'effetto "
    "pratico e' concreto: ci sono sei giorni, da oggi al 19 luglio, "
    "per testare Fable 5 senza costi aggiuntivi sul proprio piano esistente."
)

para(doc,
    "Claude Science e' un'iniziativa separata che Anthropic ha lanciato "
    "nelle stesse settimane: una versione di Claude specializzata per "
    "la ricerca scientifica, con accesso a database scientifici, strumenti "
    "di analisi dati e ottimizzazioni specifiche per la revisione della "
    "letteratura. Per la maggior parte degli studi professionali e delle "
    "PMI, questo non e' il prodotto rilevante oggi. Segnala pero' una "
    "direzione di sviluppo che vale la pena seguire: modelli non "
    "generalisti ma specializzati per settore. La stessa traiettoria, "
    "applicata all'analisi legale, fiscale o finanziaria, produrrebbe "
    "modelli con prestazioni molto piu' affidabili su compiti di dominio "
    "rispetto ai modelli generalisti attuali. Non e' una scelta che "
    "riguarda il presente immediato, ma quella dei prossimi dodici mesi."
)

heading(doc, "Come valutare Fable 5 e Sol senza perdersi nei benchmark")

para(doc,
    "L'approccio che funziona per gli studi professionali non e' leggere "
    "i benchmark pubblicati dai produttori, ma testare sui compiti "
    "effettivamente svolti ogni giorno. Un benchmark sul ragionamento "
    "matematico dice poco su quanto bene un modello struttura una "
    "consulenza fiscale su una norma italiana specifica. Un benchmark "
    "sulla generazione di codice non dice nulla su come il modello gestisce "
    "la revisione di un contratto commerciale in italiano. La sola "
    "valutazione significativa e' su casi d'uso reali: prendere tre o "
    "quattro compiti tipici dello studio, eseguirli sui modelli in "
    "competizione nella stessa settimana, confrontare la qualita' "
    "dell'output e il tempo necessario per la revisione. La finestra "
    "aperta fino al 19 luglio e' esattamente l'occasione per fare "
    "questo confronto senza costi aggiuntivi."
)

para(doc,
    "La scelta tra Fable 5 e GPT-5.6 Sol non si decide sulla sola "
    "capacita' per la maggior parte dei professionisti. Entra in gioco "
    "anche la gestione dei dati, la struttura del DPA, l'integrazione "
    "con gli strumenti gia' in uso e l'ergonomia dell'interfaccia. "
    "Uno studio che utilizza principalmente Microsoft 365 troverà "
    "un'integrazione naturale con Copilot. Uno studio che lavora "
    "principalmente via browser ha piu' liberta' nella scelta del modello. "
    "Il momento di luglio 2026 e' uno dei pochi in cui il differenziale "
    "di costo tra i modelli piu' capaci e' temporaneamente ridotto: "
    "usare questo periodo per costruire una visione concreta su quale "
    "strumento funziona meglio per il proprio lavoro specifico e' "
    "piu' utile che aspettare il modello 'definitivo' che uscira' "
    "tra tre mesi e riavvierà la stessa incertezza."
)

heading(doc, "Cosa rimane stabile mentre i modelli si avvicendano")

para(doc,
    "L'esperienza degli ultimi diciotto mesi insegna che il ciclo di "
    "rilascio dei modelli si e' accorciato: ogni tre o quattro mesi "
    "c'e' un nuovo flagship con benchmark migliori del precedente. "
    "Costruire un metodo di lavoro su un modello specifico e doverlo "
    "ricostruire ogni volta che ne esce uno nuovo e' costoso in termini "
    "di tempo e di adattamento. Quello che ha senso e' investire nel "
    "processo, non nel modello: definire come strutturare i prompt, "
    "quali passi di verifica applicare agli output AI, quali compiti "
    "e' sicuro delegare e quali richiedono revisione. Questo approccio "
    "e' portabile tra modelli diversi e rende ogni transizione piu' "
    "rapida e meno disruptiva."
)

para(doc,
    "Per uno studio che sta ancora valutando quale piattaforma scegliere, "
    "la finestra di questa settimana offre un vantaggio raro: due tra "
    "i modelli piu' capaci sul mercato sono accessibili ai piani "
    "esistenti senza costi aggiuntivi per pochi giorni. Chi usa questa "
    "opportunita' per fare una valutazione sistematica su compiti reali "
    "avra' una base di confronto solida per la scelta. Chi aspetta "
    "che la situazione si stabilizzi troverà, tra qualche mese, "
    "un nuovo insieme di modelli e la stessa decisione da prendere, "
    "senza aver costruito nel frattempo criteri propri di valutazione."
)

riferimenti(doc, [
    "Anthropic -- 'Extending Fable 5 access through July 19' (13 luglio 2026) -- anthropic.com",
    "Forbes -- 'AI Model Wars: Anthropic Extends Fable Access Again After OpenAI Sol Release' (13 luglio 2026)",
    "MatriceDigitale.it -- 'OpenAI limiti GPT-5.6 Sol, Anthropic prorroga accesso a Claude Fable 5' (13 luglio 2026)",
    "RivistaAI.it -- 'Claude cambia personalita: i nuovi modelli e la crescita dell'AI' (14 luglio 2026)",
    "LucaCazzaniga.substack.com -- 'OpenAI regala il suo modello migliore. Anthropic non puo.' (luglio 2026)",
    "Anthropic -- 'Claude Science: AI for research workflows' -- anthropic.com",
    "OpenAI -- 'GPT-5.6 Sol: expanded access and lifted rate limits' -- openai.com (luglio 2026)",
])
doc.save(BASE + "2026-07-17_fable5-sol-guerra-modelli-come-scegliere-studio.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# Shadow operation: agenti AI senza governance nelle PMI
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Governance e Sicurezza AI")
titolo(
    doc,
    "Il problema non e' l'AI che sbaglia:\ne' quella che agisce senza che nessuno sappia.",
    "Il 91% delle organizzazioni usa gia' agenti AI autonomi, ma solo "
    "il 10% ha una strategia per governarli. Nelle imprese italiane il "
    "31% ha avuto almeno un incidente legato all'AI. Si chiama 'shadow "
    "operation' e non e' un problema tecnico: e' un problema "
    "organizzativo che ogni azienda deve affrontare prima che si "
    "manifesti in modo costoso.",
    "A cura della Redazione Ratio  •  17 luglio 2026"
)

para(doc,
    "Un'azienda italiana di medie dimensioni operante nella logistica "
    "ha scoperto ad aprile che undici diversi strumenti AI erano in uso "
    "attivo all'interno dell'organizzazione. Il reparto IT ne conosceva "
    "tre, quelli ufficialmente approvati e acquistati dall'azienda. "
    "Gli altri otto erano stati attivati autonomamente da dipendenti "
    "e responsabili di reparto, a volte con carte di credito aziendali, "
    "a volte con account personali riutilizzati per lavoro. Nessuno "
    "di questi strumenti aveva un accordo sul trattamento dei dati "
    "con l'azienda. Due di essi elaboravano documenti contenenti "
    "dati personali di clienti. La scoperta non e' arrivata da "
    "un audit interno: e' arrivata dal reclamo di un cliente che "
    "aveva ricevuto una risposta generata da AI per conto dell'azienda "
    "e l'aveva riconosciuta dallo stile della scrittura."
)

para(doc,
    "Questa situazione ha un nome: shadow operation. Per analogia con "
    "lo shadow IT, che descrive l'uso di strumenti non autorizzati "
    "all'interno delle aziende, lo shadow operation descrive il "
    "deployment di sistemi AI, compresi quelli agentici, nei processi "
    "aziendali senza il controllo del reparto IT o del management. "
    "Una ricerca di luglio 2026 di Agenda Digitale, confermata da "
    "rilevazioni tra le imprese italiane, trova che il 31% delle "
    "aziende italiane ha avuto un incidente sospetto o confermato "
    "legato all'uso dell'AI, e nella maggior parte dei casi l'incidente "
    "e' originato da strumenti messi in esercizio senza governance formale."
)

para(doc,
    "La dimensione agentica aggiunge un rischio specifico. Un agente AI "
    "non e' uno strumento che attende un prompt e risponde: e' un sistema "
    "che puo' ricevere istruzioni, eseguire sequenze di operazioni nel "
    "tempo, interagire con servizi esterni, inviare email, compilare "
    "moduli, interrogare database. Quando un dipendente mette in "
    "esercizio un sistema agentivo senza il controllo dell'azienda, "
    "crea un attore operativo autonomo che agisce in nome dell'azienda "
    "senza alcuna supervisione. La ricerca di European Affairs del luglio "
    "2026 calcola che il costo medio di una violazione riconducibile "
    "ad agenti AI non governati per le imprese italiane si attesta "
    "intorno ai 4,24 milioni di euro, una cifra che include costi "
    "legali, sanzioni regolamentari e danno reputazionale."
)

heading(doc, "Come accade davvero: i percorsi verso lo shadow operation")

para(doc,
    "Lo shadow operation nelle aziende non avviene attraverso manovre "
    "sofisticate: avviene attraverso decisioni quotidiane che sembrano "
    "innocue a chi le prende. Un responsabile commerciale si abbona "
    "a uno strumento AI che redige automaticamente proposte commerciali "
    "a partire dai dati del CRM e le invia ai prospect. Un'addetta "
    "alla contabilita' attiva un agente che estrae dati riepilogativi "
    "dalle fatture e popola fogli di calcolo, compresi quelli con "
    "dati personali dei fornitori. Un responsabile della comunicazione "
    "collega i canali social aziendali a un sistema AI autonomo che "
    "pubblica contenuti generati a partire dai briefing settimanali. "
    "Ciascuno di questi usi ha una razionale operativa immediata. "
    "Nessuno di essi ha attraversato un processo di approvazione, "
    "valutazione legale o configurazione di adeguati controlli di sicurezza."
)

para(doc,
    "Il rischio specifico di questi scenari non e' solo la fuga di dati, "
    "che pure e' seria. E' che l'azienda compie operazioni in modo "
    "automatizzato, a volte operazioni irreversibili, senza che nessun "
    "essere umano le abbia consapevolmente autorizzate. Un agente che "
    "invia email ai clienti per conto dell'azienda, un agente che "
    "negozia con i fornitori tramite risposte automatizzate, un agente "
    "che accede ai database aziendali per generare report: ciascuno "
    "di questi sistemi puo' causare danni che il dipendente che lo ha "
    "attivato non aveva previsto e che nessuno in azienda sapeva "
    "stessero accadendo. Il paradosso e' che, in molti casi, "
    "i sistemi funzionano correttamente per mesi prima che un "
    "problema renda visibile l'assenza di governance."
)

heading(doc, "Da dove iniziare per governare quello che gia' succede")

para(doc,
    "La governance degli agenti AI non inizia da una policy aziendale "
    "comprensiva: inizia da un censimento. Il primo passo concreto e' "
    "mappare quali strumenti sono effettivamente in uso, inclusi quelli "
    "attivati senza autorizzazione. Questo censimento non puo' essere "
    "fatto solo attraverso interviste ai responsabili di reparto: "
    "deve includere un'analisi dei software installati sui dispositivi "
    "aziendali, degli abbonamenti attivi visibili nelle note spese "
    "o sulle carte di credito aziendali, e delle integrazioni attive "
    "su piattaforme di terze parti come Google Workspace, "
    "Microsoft 365 o Slack. Solo dopo aver completato il censimento "
    "un'azienda ha una visione realistica del rischio che gia' porta."
)

para(doc,
    "Le aziende che gestiscono meglio questa transizione non sono quelle "
    "che proibiscono l'uso degli agenti AI: sono quelle che creano "
    "un percorso di autorizzazione rapido. Un processo snello che "
    "consenta a un dipendente di proporre un uso AI, ottenere una "
    "valutazione tecnica e legale entro una settimana e ricevere "
    "un'autorizzazione formale riduce lo shadow operation togliendo "
    "la ragione principale della sua esistenza: l'impossibilita' "
    "pratica di fare le cose nei canali ufficiali. Chi rende troppo "
    "complicato fare le cose correttamente crea le condizioni per un "
    "uso informale diffuso. L'obiettivo non e' bloccare l'utilizzo: "
    "e' renderlo visibile, e governabile da chi ha la responsabilita' "
    "di farlo."
)

riferimenti(doc, [
    "AgendaDigitale.eu -- 'Shadow operation, il nuovo rischio degli agenti AI in azienda' (luglio 2026)",
    "EuropeanAffairs.it -- 'Agenti AI, il nuovo rischio cyber da milioni di euro per le imprese italiane' (9 luglio 2026)",
    "AffarItaliani.it -- 'Agenti IA, violazione dei dati e incidenti informatici: quanto rischiano le aziende italiane'",
    "DataManager.it -- 'Il 31% delle aziende italiane ha subito incidenti legati all'AI' (aprile 2026)",
    "TomHardware.it -- 'Gli agenti AI sono la nuova minaccia interna per le aziende nel 2026'",
    "ZeroUnoWeb.it -- 'Agenti AI nelle imprese italiane: dove creano valore, dove possono fallire'",
    "FortuneIta.com -- 'Agenti AI, rischio di un'autonomia mal gestita?' (aprile 2026)",
    "Regolamento UE 2024/1689 (AI Act), considerando 12 e articoli 28-29 (obblighi deployer) -- EUR-Lex",
])
doc.save(BASE + "2026-07-17_shadow-operation-agenti-ai-senza-governance-pmi.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Il paradosso della fiducia AI nelle decisioni finanziarie
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Strategia e Adozione AI")
titolo(
    doc,
    "Otto imprenditori su dieci non delegano all'AI\nle decisioni finanziarie. Ma lo fanno gia'.",
    "Una ricerca di luglio 2026 mostra che la grande maggioranza degli "
    "imprenditori italiani non si fida dell'AI per le decisioni economiche. "
    "Il paradosso e' che molti di loro usano gia' strumenti AI nei "
    "processi che alimentano quelle stesse decisioni, senza averlo "
    "formalizzato e senza aver definito chi controlla gli output.",
    "A cura della Redazione Ratio  •  17 luglio 2026"
)

para(doc,
    "Qonto, il neobank per PMI e professionisti, ha pubblicato a luglio 2026 "
    "i risultati di una rilevazione su come gli imprenditori italiani "
    "usano l'AI nelle decisioni finanziarie. I numeri principali: il 78% "
    "degli intervistati ha dichiarato che non delegherebbe decisioni "
    "finanziarie all'AI, il 36% per mancanza di fiducia nel sistema "
    "e il 42% perche' vuole mantenere il controllo diretto. "
    "La stessa ricerca, nella sua analisi qualitativa, riporta che "
    "molti di questi imprenditori usano gia' strumenti AI per l'analisi "
    "dei flussi di cassa, per il confronto tra offerte di fornitori, "
    "per la redazione di offerte commerciali con impatto sui prezzi. "
    "Il gap tra cio' che si dichiara di non voler delegare e cio' "
    "che si delega gia' nella pratica non e' piccolo: e' strutturale."
)

para(doc,
    "La distinzione che conta, e che la ricerca Qonto evidenzia senza "
    "trarne la conclusione esplicitamente, e' tra usare l'AI come "
    "strumento di supporto alla decisione e delegare la decisione "
    "all'AI. Nella pratica, il confine e' molto piu' sfumato di come "
    "appare. Un imprenditore che chiede a un sistema AI di confrontare "
    "tre opzioni di finanziamento con il costo totale stimato e "
    "l'impatto mensile sul cash flow, e sceglie poi l'opzione che "
    "l'AI presenta come piu' favorevole, non ha delegato la decisione "
    "in senso formale: ha firmato lui l'approvazione finale. "
    "Ma l'analisi che ha determinato la scelta e' stata prodotta dal "
    "sistema. Se il sistema ha commesso un errore, chi risponde "
    "della scelta?"
)

para(doc,
    "Il dato generazionale aggiunge una dimensione che ha effetti "
    "pratici sul funzionamento degli studi e delle aziende. Tra gli "
    "imprenditori tra i 18 e i 34 anni, il 69,8% usa strumenti AI; "
    "tra gli over 55, solo il 37,5%. Negli studi o nelle aziende "
    "in cui il socio anziano o il titolare ha piu' di cinquantacinque "
    "anni, la scelta di quali strumenti usare e con quali criteri "
    "e' spesso presa da persone che usano l'AI solo marginalmente. "
    "I collaboratori piu' giovani o i colleghi che la usano "
    "quotidianamente non hanno un mandato formale per governare "
    "queste scelte. Il risultato e' un divario cognitivo interno "
    "alle organizzazioni: le persone che conoscono meglio gli "
    "strumenti hanno il minor potere decisionale su di essi."
)

heading(doc, "Il rischio specifico di chi delega solo le operazioni")

para(doc,
    "Il pattern che emerge dall'osservazione sul campo e' una "
    "divisione ricorrente: le decisioni strategiche ad alta posta "
    "(acquisizioni, finanziamenti, ristrutturazioni) rimangono "
    "pienamente sotto controllo umano, mentre i singoli passaggi "
    "che portano a quelle decisioni sono sempre piu' automatizzati. "
    "Analisi dei costi, ricerche di mercato, redazione di contratti, "
    "confronto tra fornitori: tutti compiti che alimentano la "
    "decisione strategica. Se questi passaggi vengono gestiti da "
    "sistemi AI senza un controllo adeguato, la qualita' della "
    "decisione strategica dipende dalla qualita' degli output AI "
    "che nessuno ha verificato correttamente. Un imprenditore che "
    "si considera in pieno controllo perche' prende lui la decisione "
    "finale, mentre delega tutta l'analisi a sistemi che nessuno "
    "governa, e' in una posizione piu' precaria di quanto percepisca."
)

para(doc,
    "Il rischio specifico che l'AI introduce nelle decisioni finanziarie "
    "delle PMI non e' che il sistema faccia scelte palesemente errate: "
    "e' che produca scelte plausibili a partire da dati parziali o "
    "non corretti, e che nessuno verifichi i dati perche' l'output "
    "sembra credibile. Una proiezione di cash flow basata su dati "
    "contabili importati in modo non corretto non e' ovviamente sbagliata: "
    "ha lo stesso formato di una corretta, usa le stesse formule, "
    "produce numeri coerenti tra loro ma comunque errati. Il problema "
    "non emerge finche' una linea di credito non viene esercitata "
    "su quelle proiezioni, o un fornitore non viene pagato prima che "
    "un cliente abbia effettivamente saldato il dovuto."
)

heading(doc, "Costruire fiducia nel modo corretto: non evitare l'AI, governarla")

para(doc,
    "Il 78% che dichiara di non fidarsi dell'AI per le decisioni "
    "finanziarie non ha torto ad essere cauto. Ha pero' torto se "
    "quella cautela lo porta a rifiutare gli strumenti senza costruire "
    "un framework di utilizzo controllato. Il percorso che funziona "
    "non e' 'usa l'AI liberamente' ne' 'non usare l'AI per le cose "
    "importanti': e' 'usa l'AI con passaggi di verifica definiti'. "
    "Per l'analisi finanziaria questo significa stabilire quali dati "
    "il sistema puo' accedere, quali output richiedono una revisione "
    "umana prima di informare una decisione, e chi nell'organizzazione "
    "e' responsabile di quella revisione."
)

para(doc,
    "Le aziende che costruiscono questo approccio nel 2026 non avranno "
    "necessariamente usato l'AI piu' delle altre: l'avranno usata con "
    "piu' consapevolezza. Il vantaggio competitivo non sta nello "
    "strumento, che e' disponibile per tutti: sta nel processo che "
    "rende lo strumento affidabile. Questo processo non si compra "
    "pronto all'uso, ma si costruisce in pochi mesi da parte di "
    "un'organizzazione che lo affronta con serietà, partendo "
    "dalle decisioni che gia' dipendono dall'analisi AI senza che "
    "nessuno abbia formalizzato quella dipendenza. Riconoscere cosa "
    "e' gia' delegato, senza saperlo, e' il primo passo per "
    "governarlo davvero."
)

riferimenti(doc, [
    "Qonto -- 'PMI, AI e finanza: non e' la tecnologia il vero ostacolo, e' la fiducia' -- ArenaDigitale.it (14 luglio 2026)",
    "QuotidianoAlessandria.it -- 'Qonto: 8 imprenditori su 10 non delegano decisioni finanziarie ad AI' (14 luglio 2026)",
    "EM-EasyMobile.it -- 'AI nelle PMI italiane: competenze e dati frenano la svolta digitale' (1 luglio 2026)",
    "LaMillano.it -- 'AI nelle PMI italiane: il potenziale c'e' ma la maturita' resta bassa'",
    "LucaSammarco.com -- 'Intelligenza Artificiale per PMI: Guida Operativa 2026'",
    "PIComputers.it -- 'Intelligenza artificiale 2026: dove siamo e cosa cambia'",
    "IncentivImpresa.it -- 'AI Act 2026: cosa cambia per bandi e PMI italiane'",
    "UNI 11621-8:2026 -- 'Ruoli dell'AI nelle organizzazioni: Chief AI Officer e profili professionali'",
])
doc.save(BASE + "2026-07-17_fiducia-ai-decisioni-finanziarie-paradosso-pmi.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
