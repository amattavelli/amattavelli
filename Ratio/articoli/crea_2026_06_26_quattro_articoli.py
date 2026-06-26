"""
Quattro articoli Ratio — 26 giugno 2026

1. 2026-06-26_iperammortamento-180-software-ai-pmi.docx
2. 2026-06-26_agenti-ai-autonomi-pmi-rischi-operativi.docx
3. 2026-06-26_gpt55-claude-gemini-quale-scegliere-studio.docx
4. 2026-06-26_multimodalita-fatture-riunioni-cambia.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# Il software AI torna agevolabile: iperammortamento 180%
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Fiscalita' e Agevolazioni")
titolo(
    doc,
    "Il software AI torna agevolabile.\nCosa prevede l'iperammortamento 2026.",
    "La legge di bilancio 2026 ha reintrodotto la maxi deduzione fiscale per "
    "gli investimenti in beni digitali. Dopo anni di esclusione, il software "
    "torna tra i beni agevolabili, con una maggiorazione che arriva al 180% "
    "sui primi 2,5 milioni. Per le PMI che stanno investendo in strumenti AI, "
    "il calendario conta.",
    "A cura della Redazione Ratio  •  26 giugno 2026"
)

para(doc,
    "Un imprenditore del settore manifatturiero che a marzo ha firmato un "
    "contratto per un software di pianificazione della produzione basato su AI "
    "si e' sentito dire dal proprio commercialista qualcosa che non si "
    "aspettava: l'acquisto potrebbe essere deducibile con una maggiorazione del "
    "180%. Non e' un errore, non e' un'offerta promozionale del fornitore. "
    "E' la norma introdotta dalla legge di bilancio 2026, che ha reintrodotto "
    "l'iperammortamento dopo che il meccanismo era stato sostituito dal credito "
    "d'imposta Industria 4.0 nel 2020. Per capire perche' questa novita' e' "
    "rilevante per chi investe in strumenti AI nel 2026, conviene partire dalla "
    "sostanza dell'agevolazione."
)

heading(doc, "Come funziona la maggiorazione")

para(doc,
    "L'iperammortamento consente alle imprese di dedurre dal reddito imponibile "
    "una quota del costo degli investimenti che supera il costo effettivo "
    "sostenuto. Con una maggiorazione del 180%, un'impresa che acquista un "
    "software AI al costo di 100.000 euro puo' dedurre 280.000 euro, cioe' "
    "100.000 piu' il 180% aggiuntivo. Il vantaggio fiscale effettivo dipende "
    "dall'aliquota IRES o IRPEF dell'impresa, ma per le societa' soggette ad "
    "IRES corrisponde a un risparmio fiscale di circa 66.000 euro su un "
    "investimento di 100.000. Le aliquote variano in base all'entita' "
    "dell'investimento: il 180% si applica fino a 2,5 milioni di euro; "
    "tra 2,5 e 10 milioni scende al 100%; tra 10 e 20 milioni si ferma al 50%."
)

para(doc,
    "Il meccanismo e' simile a quello dell'iperammortamento Industria 4.0 in "
    "vigore fino al 2019, con alcune differenze importanti. La prima, per chi "
    "si occupa di strumenti digitali, e' che il software torna tra i beni "
    "agevolabili. Nel regime del credito d'imposta 4.0 e 5.0 che ha "
    "caratterizzato gli anni 2020-2025, il software poteva essere agevolato "
    "solo se abbinato a beni materiali 4.0. Dal 2026, i beni immateriali "
    "dell'Allegato V del decreto sono nuovamente autonomamente agevolabili, "
    "il che include le applicazioni software di gestione aziendale con funzioni "
    "di intelligenza artificiale certificate come beni 4.0."
)

heading(doc, "Quali investimenti AI rientrano nell'agevolazione")

para(doc,
    "La distinzione rilevante e' tra software genericamente 'AI' e software "
    "classificabile come bene immateriale 4.0. Non tutti i software che usano "
    "intelligenza artificiale rientrano nell'agevolazione. I criteri stabiliti "
    "dal decreto fanno riferimento a caratteristiche tecniche specifiche: "
    "integrazione con sistemi di raccolta dati in tempo reale, capacita' di "
    "analisi predittiva o ottimizzazione automatica di processi produttivi, "
    "interoperabilita' con sistemi di supervisione industriale. Un software di "
    "marketing automation basato su AI non rientra. Un sistema MES con moduli "
    "AI per la pianificazione predittiva della produzione, che raccoglie dati "
    "da macchine collegate, verosimilmente si'."
)

para(doc,
    "Per le PMI che si trovano in questa zona grigia, la documentazione tecnica "
    "e' il punto critico. L'accesso all'agevolazione richiede una comunicazione "
    "al Gestore dei Servizi Energetici (GSE) attraverso la piattaforma "
    "telematica predisposta, e il GSE puo' richiedere documentazione che "
    "dimostri il possesso dei requisiti tecnici. Farsi assistere da un "
    "professionista con esperienza in perizie tecniche 4.0, o chiedere "
    "direttamente al fornitore del software di fornire la documentazione "
    "certificata, e' la strada piu' sicura per evitare contestazioni in sede "
    "di verifica."
)

para(doc,
    "L'agevolazione copre gli investimenti effettuati tra il 1 gennaio 2026 e "
    "il 30 settembre 2028. Per chi ha gia' acquistato software nel corso del "
    "2026, il punto da verificare con il proprio commercialista e' se "
    "l'investimento rispetta tutti i requisiti e come impostare correttamente "
    "la comunicazione GSE. L'iperammortamento non e' una scorciatoia per "
    "trasformare qualsiasi spesa in software in un vantaggio fiscale: funziona "
    "su categorie precise e richiede procedure specifiche. Ma per le PMI che "
    "stanno investendo in strumenti digitali AI nell'ambito di processi "
    "produttivi o gestionali strutturati, e' uno strumento reale che vale la "
    "pena verificare concretamente prima di chiudere il bilancio."
)

riferimenti(doc, [
    "Legge 30 dicembre 2025, n. 207 (Legge di Bilancio 2026) — Normattiva",
    "Finera.it — 'Iperammortamento 2026: guida completa sul ritorno della maxi deduzione per gli investimenti delle imprese'",
    "IncentivImpresa — 'Iperammortamento 2026: Decreto Attuativo, GSE e Aliquote 220%'",
    "SoluzioneTasse — 'Iperammortamento 2026: come funziona la maxi deduzione fiscale'",
    "FiscoeTasse — 'Legge finanziaria 2026: l'iperammortamento'",
    "MerlinConnect — 'Iperammortamento Industria 4.0 2026 — Guida aggiornata'",
    "PremierConsulting — 'Iperammortamento 2026: esempio di calcolo'",
])
doc.save(BASE + "2026-06-26_iperammortamento-180-software-ai-pmi.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# Agenti AI autonomi: i rischi operativi non calcolati
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Strategia AI")
titolo(
    doc,
    "Quando l'agente AI decide da solo:\ni rischi che non si stanno calcolando.",
    "Gli agenti AI possono gestire email, aggiornare gestionali, inviare "
    "preventivi e aprire ticket senza supervisione umana. La produttivita' "
    "promessa e' reale. Ma la struttura del rischio e' diversa da qualsiasi "
    "altro strumento digitale che le imprese hanno adottato finora.",
    "A cura della Redazione Ratio  •  26 giugno 2026"
)

para(doc,
    "Una PMI del settore edile ha configurato a gennaio un agente AI collegato "
    "al proprio gestionale per gestire le richieste di preventivo ricevute via "
    "email. L'agente legge il messaggio, identifica il tipo di lavoro richiesto, "
    "recupera i listini dal gestionale e risponde con un preventivo automatico "
    "entro pochi minuti. Per tre mesi ha funzionato bene. Poi, a fine marzo, "
    "un cliente ha inviato un'email con una richiesta non standard, formulata "
    "in modo ambiguo. L'agente ha interpretato la richiesta, ha costruito un "
    "preventivo, lo ha inviato. Il preventivo era sbagliato: includeva una voce "
    "di costo errata e impegnava l'impresa su tempi che non corrispondevano "
    "alle disponibilita' reali. Il cliente aveva gia' risposto con un'accettazione "
    "prima che qualcuno in azienda se ne accorgesse."
)

para(doc,
    "Non e' un caso ipotetico. Varianti di questo scenario si verificano in "
    "molte imprese che hanno introdotto agenti AI in produzione senza aver "
    "definito con chiarezza dove termina l'autonomia dell'agente e dove inizia "
    "la supervisione umana. La quota di imprese italiane con almeno dieci "
    "addetti che utilizza almeno una tecnologia di intelligenza artificiale e' "
    "oggi al 16,4%, raddoppiata rispetto all'8,2% del 2024. Una parte di questa "
    "crescita viene dall'adozione di agenti AI su processi operativi. "
    "Il problema e' che la velocita' di deploy sta superando quella con cui le "
    "imprese sviluppano i controlli."
)

heading(doc, "La differenza strutturale rispetto agli strumenti AI passivi")

para(doc,
    "Un modello AI che risponde alle domande degli utenti, genera testi o "
    "analizza documenti e' uno strumento passivo: non fa nulla che l'utente "
    "non abbia richiesto esplicitamente. Un agente AI e' diverso: opera in modo "
    "autonomo su un insieme di compiti predefiniti, prende decisioni senza "
    "attendere conferma per ogni passo, e puo' interagire con sistemi esterni, "
    "email, gestionali, CRM, piattaforme di acquisto. Il rischio non e' che "
    "l'agente 'si rompa': e' che funzioni esattamente come e' stato configurato, "
    "ma in un contesto che non era stato previsto al momento della configurazione."
)

para(doc,
    "Questo trasforma la struttura del rischio in modo significativo. Con uno "
    "strumento passivo, l'errore e' quasi sempre visibile prima che produca "
    "conseguenze: l'utente legge il testo generato, valuta la risposta, decide "
    "se usarla. Con un agente, il ciclo di retroazione e' piu' lungo. L'agente "
    "ha gia' agito prima che qualcuno riveda l'output. Se gestisce email, le ha "
    "gia' inviate. Se aggiorna un database, i dati sono gia' stati modificati. "
    "Se ha aperto un ordine di acquisto, il fornitore ha gia' ricevuto la "
    "conferma."
)

heading(doc, "Le tre categorie di rischio che le imprese stanno sottovalutando")

para(doc,
    "Il primo e' il rischio di interpretazione. Gli agenti AI sono progettati "
    "per completare compiti, non per fermarsi davanti all'ambiguita'. Quando un "
    "input non corrisponde ai pattern su cui sono stati configurati, tendono a "
    "fare la cosa piu' plausibile invece di chiedere chiarimento. Per un uso "
    "esplorativo, un output sbagliato viene semplicemente scartato. In un agente "
    "che agisce su sistemi reali, 'la cosa piu' plausibile' puo' generare "
    "conseguenze difficili da reversire: un preventivo inviato, un ordine "
    "confermato, un'email spedita a un indirizzo sbagliato."
)

para(doc,
    "Il secondo e' il rischio di identita' e credenziali. Ogni agente AI che "
    "opera su sistemi aziendali ha bisogno di accesso a quei sistemi, il che "
    "significa credenziali, token di autenticazione, permessi. Piu' funzioni si "
    "affidano all'agente, piu' permessi accumula. Un agente configurato per "
    "gestire email, aggiornare il CRM e accedere al gestionale ha, di fatto, le "
    "stesse credenziali di un dipendente con mansioni simili. Se l'account "
    "dell'agente viene compromesso, le conseguenze sono quelle di una "
    "compromissione di un account aziendale con ampi privilegi. I team di "
    "sicurezza che subiscono pressione per il deploy immediato degli agenti "
    "senza i controlli necessari stanno gia' segnalando questo come il rischio "
    "principale per il 2026."
)

para(doc,
    "Il terzo e' la variabilita' silenziosa del modello sottostante. I provider "
    "aggiornano i modelli frequentemente, spesso senza notifica esplicita. Un "
    "agente che funziona in modo prevedibile su una versione del modello puo' "
    "comportarsi diversamente dopo un aggiornamento silenzioso. Le API di "
    "Anthropic e OpenAI consentono di bloccare la versione specifica del modello, "
    "ma richiede di gestirla attivamente. Per un processo automatizzato che "
    "gestisce comunicazioni o transazioni per conto dell'impresa, la variabilita' "
    "del comportamento e' un rischio operativo che va monitorato, non dato per "
    "scontato."
)

para(doc,
    "Introdurre un agente AI in un processo aziendale senza aver mappato questi "
    "rischi non significa che il problema si verifichera': significa non sapere "
    "se si sta verificando. La domanda da farsi prima di mettere un agente in "
    "produzione non e' solo 'funziona?' ma 'cosa fa quando riceve un input che "
    "non si aspetta?', 'chi vede cosa fa in tempo reale?', 'chi puo' fermarlo?'. "
    "Le aziende che rispondono a queste domande prima di avviare trovano i "
    "problemi prima che diventino incidenti. Quelle che le affrontano dopo, "
    "di solito, lo fanno su un caso concreto."
)

riferimenti(doc, [
    "Tom's Hardware — 'Gli agenti AI sono la nuova minaccia interna per le aziende nel 2026'",
    "DataMagazine — 'Agenti AI autonomi: opportunita' e rischi della nuova automazione' (marzo 2026)",
    "DataMagazine — 'Agenti AI e rischi per l'identita': come cambiera' la sicurezza' (gennaio 2026)",
    "AI4Business — 'Economia agent-to-agent 2026: A2A, AP2 e B2B'",
    "MilanoFinanza — 'Il grande salto dell'AI nel 2026: dati, agenti autonomi e nuove superfici d'attacco'",
    "BestTechPartner — 'AI agenti: cosa significa davvero per le imprese' (maggio 2026)",
    "CorrerieComunicazioni — 'Agenti AI e regolamentazione 2026: cosa cambia per le imprese'",
])
doc.save(BASE + "2026-06-26_agenti-ai-autonomi-pmi-rischi-operativi.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# GPT-5.5, Claude, Gemini: come scegliere
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Strumenti AI")
titolo(
    doc,
    "GPT-5.5, Claude e Gemini:\ntre modelli di frontiera, zero certezze su quale sia meglio.",
    "Nei benchmark del 2026 le differenze tra i modelli di punta si misurano "
    "in frazioni di punto. Eppure la scelta del modello per un uso professionale "
    "non e' mai stata cosi' rilevante. Il problema non e' quale vince: "
    "e' capire per quale compito.",
    "A cura della Redazione Ratio  •  26 giugno 2026"
)

para(doc,
    "Un responsabile amministrativo di una societa' di consulenza racconta di "
    "aver cambiato strumento AI quattro volte in sei mesi. Prima ChatGPT, poi "
    "Claude, poi di nuovo ChatGPT su consiglio di un collega, poi Gemini perche' "
    "si integrava meglio con i Google Workspace gia' in uso. In ogni caso, il "
    "nuovo strumento funzionava bene per alcune cose e deludeva su altre. "
    "Il ciclo non finisce perche' nessuno dei tre modelli e' migliore degli "
    "altri su tutto. Sono diversi. La differenza conta, specialmente per chi li "
    "usa in modo sistematico su compiti professionali ripetitivi."
)

para(doc,
    "Il 2026 ha portato a maturita' tre modelli di frontiera: GPT-5.5 di OpenAI, "
    "rilasciato il 24 aprile, Claude Opus 4.7 di Anthropic e Gemini 3.1 di "
    "Google. Nei benchmark tecnici, i gap si sono ridotti al punto da rendere i "
    "confronti numerici poco informativi per un utente professionale. La "
    "differenza reale non emerge dai test standardizzati: emerge dall'uso "
    "quotidiano su compiti specifici, in un contesto di lavoro reale, con "
    "documenti e dati reali."
)

heading(doc, "Cosa distingue davvero i tre modelli nel lavoro professionale")

para(doc,
    "GPT-5.5 e' orientato al lavoro che richiede azioni su sistemi esterni. "
    "Pianifica sequenze di operazioni, usa strumenti come browser, esecuzione "
    "di codice e lettura di file in modo piu' fluido degli altri, ed e' lo "
    "strumento piu' adatto quando si vuole costruire flussi automatizzati o "
    "agenti che operano su piu' applicazioni in sequenza. La sua integrazione "
    "con l'ecosistema Microsoft, da Outlook a Excel, e con le API di OpenAI e' "
    "matura e ben documentata. Per un controller o un CFO che lavora su fogli "
    "di calcolo complessi, analisi di dati e reportistica automatizzata, GPT-5.5 "
    "offre il miglior supporto operativo disponibile oggi."
)

para(doc,
    "Claude Opus 4.7 esprime il meglio su testi lunghi e complessi che "
    "richiedono coerenza, sfumature e ragionamento esperto. Se il compito e' "
    "redigere un parere, sintetizzare decine di documenti mantenendo la "
    "struttura logica, o revisionare un contratto con attenzione alle "
    "implicazioni non esplicite, Claude tende a produrre output di qualita' "
    "superiore su quella categoria specifica. Ha anche la finestra di contesto "
    "piu' ampia tra i tre, il che lo rende utile quando si lavora su documenti "
    "di grandi dimensioni senza dover operare tagli o suddivisioni."
)

para(doc,
    "Gemini 3.1 e' il piu' integrato con i prodotti Google, il che lo rende "
    "conveniente per chi usa Gmail, Drive, Docs e Calendar come infrastruttura "
    "lavorativa quotidiana. La multimodalita' nativa, la capacita' di elaborare "
    "contemporaneamente testo, immagini, audio e video nello stesso modello, "
    "e' tra le piu' avanzate, con applicazioni pratiche nella gestione di "
    "documentazione mista come report con grafici, verbali, presentazioni. "
    "Per gli studi e le aziende gia' nell'ecosistema Google Workspace, "
    "l'integrazione riduce significativamente la frizione operativa rispetto a "
    "strumenti di terze parti."
)

heading(doc, "La questione dei dati, del costo e della dipendenza")

para(doc,
    "La scelta del modello ha anche una dimensione che riguarda i dati. Quando "
    "si usa un modello AI per elaborare documenti aziendali, contratti o dati "
    "di clienti, questi transitano attraverso le infrastrutture del provider. "
    "Le policy di OpenAI, Anthropic e Google differiscono nei dettagli su come "
    "i dati vengono trattati, per quanto tempo vengono conservati, e se possono "
    "essere usati per addestrare i modelli futuri. I piani business e "
    "enterprise di tutti e tre i provider offrono garanzie adeguate per la "
    "maggior parte degli usi aziendali. Per dati particolarmente sensibili, "
    "soggetti a vincoli di confidenzialita' contrattuale o a normative di "
    "settore, leggere le policy specifiche del piano in uso e' necessario."
)

para(doc,
    "Sul costo, i tre modelli di frontiera hanno prezzi API simili nella fascia "
    "alta, ma differenze significative nelle opzioni di abbonamento mensile. "
    "Chi usa l'AI in modo intensivo ma non programmatico puo' trovare differenze "
    "rilevanti tra i piani flat e le tariffe a consumo. L'episodio di giugno "
    "con Claude, che aveva annunciato e poi ritirato una modifica alla "
    "fatturazione degli agenti, e' un segnale che i modelli di prezzo sono "
    "ancora in evoluzione. Per gli usi professionali strutturati, la "
    "diversificazione tra due modelli diversi per funzioni diverse offre anche "
    "un margine di negoziazione sui rinnovi contrattuali e riduce l'esposizione "
    "a interruzioni di servizio."
)

para(doc,
    "La domanda 'qual e' il modello migliore?' e' la domanda sbagliata. "
    "Quella giusta e' 'per quale compito specifico, con quali dati, a quale "
    "costo?'. Un uso professionale strutturato risponde a questo identificando "
    "una o due funzioni ad alto volume su cui il modello verra' usato "
    "sistematicamente, e scegliendo su quella base. Cambiare strumento ogni "
    "mese in cerca del migliore in assoluto e' il modo piu' efficace per "
    "non diventare mai davvero bravi con nessuno."
)

riferimenti(doc, [
    "AiPia.it — 'GPT-5.5 vs Gemini 3.1 vs Claude Opus 4.7: confronto 2026'",
    "OpenAI — 'Introducing GPT-5.5' (aprile 2026)",
    "IntuitionLabs — 'Claude vs ChatGPT vs Copilot vs Gemini: 2026 Enterprise Guide'",
    "Mimir.bot — 'GPT-5.5: il modello piu' potente di OpenAI per coding, ricerca e analisi dati'",
    "Stob.AI — 'Best AI Model 2026: GPT-5.5 vs Claude 4.8 vs Gemini 3.5 vs Llama 4'",
    "AiPia.it — 'Benchmark AI 2026: classifica aggiornata dei migliori modelli'",
    "Emanuele-Ricci.it — 'GPT-5.5: il nuovo modello frontier di OpenAI che ridefinisce l'intelligenza artificiale agentiva'",
])
doc.save(BASE + "2026-06-26_gpt55-claude-gemini-quale-scegliere-studio.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Multimodalita': l'AI che legge fatture e trascrive riunioni
# ============================================================

doc = new_doc()
testata(doc, "Giugno 2026", "Operativita' AI")
titolo(
    doc,
    "Fotografa la fattura, registra la riunione:\ncosa cambia con i modelli multimodali.",
    "I modelli AI del 2026 leggono documenti fotografati, trascrivono riunioni "
    "e analizzano grafici. Per un commercialista o un consulente, questo non "
    "e' solo una comodita' tecnica: cambia quali operazioni ha senso automatizzare "
    "e dove si devono tenere i controlli.",
    "A cura della Redazione Ratio  •  26 giugno 2026"
)

para(doc,
    "Un tributarista racconta di aver passato l'ultimo anno a fotografare con "
    "il telefono le fatture cartacee dei propri clienti e a inviarle a un "
    "modello AI per estrarne i dati. Il processo funziona: il modello "
    "identifica fornitore, importo, data, numero fattura, IVA, con un tasso di "
    "errore inferiore a quello di un inserimento manuale su grandi volumi. "
    "Quello che ancora lo sorprende non e' che l'AI riesca a leggere le fatture, "
    "ma che ci riesca partendo da una fotografia mediocre, scattata con il "
    "telefono in condizioni di luce non ottimali, di una fattura parzialmente "
    "sgualcita. Questo e' quello che significa multimodalita' in pratica: non "
    "la capacita' di fare piu' cose diverse, ma la capacita' di ragionare su "
    "input di tipo diverso, testo, immagini, audio, video, nello stesso spazio, "
    "senza che l'utente debba gestire la conversione tra formati."
)

heading(doc, "Da chatbot a strumento che 'vede' e 'ascolta'")

para(doc,
    "Fino al 2024, l'uso professionale dei modelli AI si basava quasi "
    "esclusivamente su testo. Si copiava un testo in un'area di input, si "
    "riceveva una risposta testuale. Questo limitava le applicazioni ai "
    "documenti gia' in formato digitale, escludendo gran parte del materiale "
    "cartaceo, fotografico o audio con cui lavorano commercialisti, consulenti "
    "e aziende ogni giorno. I modelli multimodali del 2026, GPT-5.5, Claude "
    "Opus 4.7 e Gemini 3.1 in primo piano, gestiscono immagini, PDF, audio e "
    "in alcuni casi video come input nativi, aprendo applicazioni pratiche che "
    "fino a poco tempo fa richiedevano sistemi specializzati separati o "
    "integrazioni complesse."
)

para(doc,
    "Per la gestione documentale, significa poter inviare direttamente la "
    "fotografia di una lettera, una fattura, un estratto conto, un contratto "
    "firmato a mano, e ricevere un'analisi strutturata del contenuto. "
    "Non e' OCR tradizionale: il modello non si limita a trascrivere il testo "
    "ma interpreta il documento nel suo contesto, identifica le informazioni "
    "chiave, segnala anomalie o elementi incoerenti. Per chi gestisce archivi "
    "cartacei o riceve documentazione fisica dai clienti, questo abbassa "
    "drasticamente il costo di ingresso per processi di digitalizzazione che "
    "prima richiedevano scanner, software di riconoscimento testo e verifica "
    "manuale."
)

para(doc,
    "Per la gestione delle riunioni, la multimodalita' significa registrare una "
    "call, inviare l'audio al modello e ricevere non solo una trascrizione ma "
    "un riassunto strutturato con i punti di decisione, le azioni assegnate e "
    "le questioni aperte. Funziona su registrazioni in italiano, con qualita' "
    "accettabile anche in presenza di piu' voci sovrapposte o rumori di fondo. "
    "Per uno studio con molte riunioni di aggiornamento con i clienti, la "
    "possibilita' di avere un verbale strutturato automatico riduce un'attivita' "
    "spesso affidata al collaboratore piu' giovane a pochi minuti di verifica "
    "su un output gia' elaborato."
)

heading(doc, "I limiti operativi da conoscere prima di affidarsi")

para(doc,
    "La multimodalita' funziona meglio sui formati per cui i modelli sono stati "
    "addestrati: documenti stampati chiari, audio di qualita' accettabile, "
    "immagini a risoluzione sufficiente. Con materiale di qualita' bassa, "
    "fotografie sfocate, registrazioni con eco forte, documenti scritti a mano "
    "in grafia non standard, il tasso di errore aumenta in modo non lineare. "
    "Il limite operativo piu' rilevante e' che il modello non segnala "
    "necessariamente l'incertezza: produce un output che sembra plausibile "
    "anche quando ha estratto informazioni errate. Un importo di fattura "
    "sbagliato che entra in un sistema di contabilita' senza essere verificato "
    "e' un errore. Un estratto conto mal letto che alimenta una riconciliazione "
    "automatizzata e' un errore potenzialmente difficile da trovare a posteriori."
)

para(doc,
    "La soluzione non e' rinunciare allo strumento: e' progettare il flusso di "
    "lavoro in modo che gli output ad alto impatto vengano verificati prima di "
    "essere usati. Spot check regolari su un campione dei documenti elaborati "
    "aiutano a tenere il tasso di errore sotto controllo senza dover ricontrollare "
    "tutto manualmente, che vanificherebbe il vantaggio dello strumento. "
    "Una regola pratica: i dati che entrano direttamente in un sistema "
    "gestionale o che vengono comunicati al cliente meritano un controllo "
    "umano. Quelli che servono a costruire una bozza, una sintesi o un "
    "promemoria interno possono tollerare un margine di imprecisione piu' ampio."
)

para(doc,
    "C'e' poi una questione di riservatezza che riguarda specificatamente i dati "
    "inviati in formato immagine o audio. I documenti che contengono dati "
    "personali di terzi, clienti, dipendenti, controparti, rientrano nell'ambito "
    "GDPR quando vengono inviati a sistemi AI cloud. L'analisi di documenti "
    "sensibili va fatta con i piani aziendali che offrono le garanzie adeguate, "
    "non con account personali o gratuiti. Il passaggio pratico piu' immediato "
    "per uno studio professionale non e' riprogettare i processi: e' identificare "
    "uno o due flussi dove oggi si perdono ore in attivita' manuali su documenti "
    "fisici o audio, e verificare se uno strumento multimodale gestisce quel "
    "caso specifico con una qualita' accettabile. La verifica richiede "
    "trenta minuti, non un progetto."
)

riferimenti(doc, [
    "AiPia.it — 'GPT-5.5 vs Gemini 3.1 vs Claude Opus 4.7: multimodalita' nativa a confronto, 2026'",
    "Osservatori.net — 'Gemini: cos'e' e come si usa. Limiti e opportunita' in azienda'",
    "Mimir.bot — 'ChatGPT: guida completa alle funzionalita' multimodali 2026'",
    "WebalchLab — 'Modelli multimodali: come sfruttare testo, immagini e voce in un'unica strategia'",
    "AgendaDigitale — 'Gemini AI di Google entra in azienda: ecco i vantaggi'",
    "CultureDigitali — 'Tutti i modelli Gemini AI disponibili nel 2026: guida completa'",
    "Osservatorio Digital Innovation, Politecnico di Milano — 'AI nelle imprese italiane 2026'",
])
doc.save(BASE + "2026-06-26_multimodalita-fatture-riunioni-cambia.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
