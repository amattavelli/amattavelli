#!/usr/bin/env python3
"""
Articolo 3: Agenti AI autonomi per studi e PMI
File: Ratio/articoli/2026-04-17_agenti-ai-delegare-sistema-autonomo.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/home/user/amattavelli/Ratio/articoli/2026-04-17_agenti-ai-delegare-sistema-autonomo.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(3)
sec.right_margin = Cm(3)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot)
    pPr.append(pBdr)


def para(doc, text, size=11, italic=False, color=None, sa=8, sb=0,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


# --- TESTATA ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.all_caps = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Strumenti e Processi")
r.font.name = 'Calibri'
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# --- TITOLO ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "Quando l\u2019agente AI agisce senza aspettare: "
    "quello che cambia se deleghi a un sistema autonomo"
)
r.font.name = 'Calibri'
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Gli agenti AI non rispondono a domande: eseguono processi in autonomia, "
    "interagendo con sistemi esterni. Prima di adottarli in studio o in azienda, "
    "conviene capire cosa significa davvero delegare."
)
r.font.name = 'Calibri'
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  \u2022  17 aprile 2026")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# --- CORPO ---
para(doc,
    "Luned\u00ec pomeriggio configuri un agente AI collegato al gestionale dello studio. "
    "Marted\u00ec mattina, quando accendi il computer, trovi che ha verificato le scadenze "
    "di quaranta clienti, segnalato tre anomalie nei flussi di cassa, preparato due bozze "
    "di email e avviato l\u2019upload di documentazione su un portale esterno. Ottimo. "
    "Poi ti accorgi che una delle email \u00e8 gi\u00e0 stata inviata. Il portale usato era quello "
    "sbagliato per quel tipo di documento. L\u2019agente ha fatto esattamente quello che gli "
    "era stato chiesto, solo che le istruzioni non erano abbastanza precise."
)

para(doc,
    "La differenza tra un chatbot e un agente AI \u00e8 precisamente questa: il chatbot "
    "risponde a una domanda e aspetta la prossima, l\u2019agente persegue un obiettivo "
    "complesso eseguendo una sequenza di azioni, spesso su sistemi esterni, senza "
    "richiedere approvazione a ogni passo. Pu\u00f2 accedere a database, inviare email, "
    "aprire file, compilare moduli, effettuare ricerche su portali web. \u00c8 una "
    "caratteristica potente, non un difetto. Richiede per\u00f2 una logica di governance "
    "diversa da quella che si applica a un assistente che attende sempre la tua risposta."
)

para(doc,
    "Nel mercato italiano, gli agenti AI stanno arrivando da due direzioni. Le piattaforme "
    "generaliste come n8n, Make e Zapier consentono di costruire workflow automatizzati "
    "collegando strumenti AI a sistemi gi\u00e0 in uso nello studio. Sul versante delle "
    "soluzioni verticali, ad aprile 2026 \u00e8 stata lanciata AgenVIO, pensata "
    "specificamente per le PMI italiane con funzioni AI agentiche integrate nei processi "
    "amministrativi. Il confine tra automazione e agente autonomo \u00e8 sottile ma rilevante: "
    "un\u2019automazione esegue una regola fissa, un agente adatta il proprio comportamento "
    "al contesto."
)

h2(doc, "Il nodo della governance: chi decide cosa")

para(doc,
    "Prima di adottare un agente AI in un contesto professionale o aziendale, la domanda "
    "pi\u00f9 utile non \u00e8 \u201ccosa riesce a fare?\u201d ma \u201csu quali azioni voglio che chieda "
    "conferma, e su quali pu\u00f2 procedere autonomamente?\u201d. Questa distinzione deve essere "
    "decisa prima della configurazione, non dopo il primo incidente. Le azioni reversibili "
    "(preparare una bozza, raccogliere dati, generare un report) si prestano bene "
    "all\u2019autonomia. Le azioni irreversibili o ad alto impatto (inviare comunicazioni "
    "a clienti, trasmettere documenti, effettuare operazioni su sistemi fiscali) richiedono "
    "un punto di conferma umana obbligatorio."
)

para(doc,
    "L\u2019AI Act europeo classifica come ad alto rischio molti sistemi agentici che operano "
    "in ambito professionale, gestionale e decisionale. Questo significa che, entro agosto "
    "2026, le aziende e gli studi che utilizzano agenti AI in questi contesti dovranno "
    "disporre di documentazione tecnica, sistema di gestione del rischio e supervisione "
    "umana dimostrabile. Non \u00e8 un requisito astratto: \u00e8 la formalizzazione di qualcosa "
    "che andrebbe fatto comunque per ragioni operative, indipendentemente dalla normativa."
)

para(doc,
    "Un agente configurato male non \u00e8 pericoloso nel senso drammatico del termine. "
    "Pu\u00f2 per\u00f2 consumare risorse, produrre output inconsistenti, creare confusione con "
    "i clienti o con i sistemi interni, e farlo in modo silenzioso perch\u00e9, per "
    "definizione, agisce quando tu non stai guardando. La potenza degli agenti AI si "
    "realizza pienamente quando le istruzioni sono precise, il perimetro di azione \u00e8 "
    "definito, e i punti di controllo umano sono previsti in anticipo."
)

h2(doc, "Il modello operativo che funziona")

para(doc,
    "Gli studi e le aziende che ottengono risultati concreti dall\u2019AI agentica hanno in "
    "comune un approccio: prima ridisegnano il flusso di lavoro, poi configurano l\u2019agente. "
    "Chi fa il contrario, aggiungendo l\u2019agente su processi gi\u00e0 disordinati, ottiene "
    "automazione del caos. La fase di analisi preliminare non richiede competenze "
    "tecniche avanzate: richiede la capacit\u00e0 di rispondere a domande semplici. Cosa deve "
    "fare l\u2019agente, esattamente? Con quali sistemi deve interagire? Quali output produce "
    "e chi li legge? Cosa succede se si blocca o produce un risultato inatteso?"
)

para(doc,
    "La buona notizia \u00e8 che la differenza tra un\u2019adozione ben riuscita e una "
    "fallimentare non dipende dalla tecnologia: dipende dalla chiarezza con cui si "
    "decide cosa si vuole automatizzare e a quali condizioni. Uno studio di cinque "
    "persone con processi ben definiti pu\u00f2 ottenere pi\u00f9 valore da un agente AI di "
    "uno studio di trenta persone che lo adotta senza una visione precisa. "
    "La domanda da farsi non \u00e8 \u201csiamo pronti per l\u2019AI agentica?\u201d ma \u201cconosciamo "
    "abbastanza bene i nostri processi per sapere cosa delegare?\u201d."
)

sep(doc)

# --- FONTI ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Fonti e riferimenti")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Innovami News \u2014 AI nelle PMI in Italia: AgenVIO automatizza i processi aziendali "
    "(15 aprile 2026)",
    "Management CUE \u2014 Gli agenti AI in azienda: cosa cambier\u00e0 davvero nel 2026 "
    "per le PMI italiane",
    "AI4Business \u2014 Agenti AI e nuove regolamentazioni: il quadro normativo italiano "
    "ed europeo 2026 per le imprese",
    "BitMat \u2014 Agenti AI: qual \u00e8 l\u2019impatto sulle PMI italiane",
    "Regolamento (UE) 2024/1689 (AI Act) \u2014 Requisiti per sistemi AI ad alto rischio, "
    "articoli 8-15",
]
for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"\u2022 {s}")
    r.font.name = 'Calibri'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte")
r.font.name = 'Calibri'
r.font.size = Pt(8)
r.italic = True
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

doc.save(OUTPUT)
print(f"Salvato: {OUTPUT}")
