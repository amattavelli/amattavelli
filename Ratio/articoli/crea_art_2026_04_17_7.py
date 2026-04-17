#!/usr/bin/env python3
"""
Articolo 3: L'AI nei gestionali dei commercialisti: da Genya a TeamSystem, cosa cambia davvero
File: Ratio/articoli/2026-04-17_ai-gestionali-commercialisti-2026.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/home/user/amattavelli/Ratio/articoli/2026-04-17_ai-gestionali-commercialisti-2026.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(3)
sec.right_margin = Cm(3)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot)
    pPr.append(pBdr)


def para(doc, text, size=11, italic=False, color=None, sa=8, sb=0,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


# --- TESTATA ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.all_caps = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Strumenti e Tecnologia")
r.font.name = 'Calibri'
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# --- TITOLO ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "L\u2019AI nei gestionali dei commercialisti: "
    "da Genya a TeamSystem, cosa cambia davvero"
)
r.font.name = 'Calibri'
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "I principali software per studi professionali italiani hanno integrato funzioni "
    "di intelligenza artificiale. Non \u00e8 marketing: alcune funzionalit\u00e0 stanno gi\u00e0 "
    "cambiando il modo in cui si lavora. Una guida pratica per orientarsi."
)
r.font.name = 'Calibri'
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  \u2022  17 aprile 2026")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# --- CORPO ---
para(doc,
    "Secondo l\u2019indagine della Fondazione Nazionale di Ricerca dei Commercialisti "
    "condotta tra luglio e settembre 2025, circa il 34% dei commercialisti italiani "
    "utilizza gi\u00e0 strumenti di intelligenza artificiale nella propria attivit\u00e0 quotidiana. "
    "Una quota che, stando alle previsioni degli stessi professionisti, potrebbe "
    "salire fino al 72% nei prossimi tre anni. Il salto in avanti non avverr\u00e0 per "
    "conversione spontanea: avverr\u00e0 perch\u00e9 i software che questi professionisti gi\u00e0"
    "usano ogni giorno stanno incorporando l\u2019AI direttamente nei flussi di lavoro, "
    "senza richiedere l\u2019adozione di strumenti nuovi o separati."
)

para(doc,
    "Il mercato del software contabile e gestionale per studi professionali ha subito "
    "una trasformazione rapida. I principali player italiani \u2014 Wolters Kluwer con Genya, "
    "TeamSystem, Sistemi con PROFIS, Datalog \u2014 hanno tutti rilasciato o annunciato "
    "funzioni AI integrate nei loro prodotti principali. Non si tratta di add-on "
    "opzionali: in molti casi l\u2019AI \u00e8 diventata parte del workflow ordinario, "
    "dalla categorizzazione automatica delle fatture alla compilazione assistita "
    "delle dichiarazioni fiscali."
)

h2(doc, "Genya Expert AI: l\u2019AI nelle dichiarazioni fiscali")

para(doc,
    "Wolters Kluwer ha lanciato a marzo 2026 Genya Dichiarativi Expert AI, "
    "un\u2019integrazione che porta l\u2019intelligenza artificiale direttamente nel workflow "
    "delle dichiarazioni fiscali. Il sistema analizza i dati contabili del cliente, "
    "identifica anomalie rispetto agli anni precedenti, segnala potenziali errori "
    "prima dell\u2019invio e suggerisce compilazioni coerenti con il profilo fiscale "
    "storico dell\u2019impresa. Non sostituisce il commercialista nel giudizio finale, "
    "ma riduce drasticamente il tempo necessario per la verifica e il controllo "
    "incrociato dei dati."
)

para(doc,
    "Un aspetto particolarmente rilevante di Genya Expert AI \u00e8 la funzione di "
    "controllo intelligente sulle deduzioni e detrazioni. Il sistema incrocia "
    "automaticamente le voci inserite con le normative vigenti e con la casistica "
    "di accertamento fiscale, segnalando le aree di maggiore esposizione al rischio. "
    "Per gli studi che gestiscono un alto volume di dichiarazioni, questo tipo di "
    "controllo automatizzato pu\u00f2 fare la differenza tra una stagione fiscale "
    "gestibile e una sotto pressione continua."
)

h2(doc, "PROFIS e TeamSystem: automazione contabile quotidiana")

para(doc,
    "Sistemi con il gestionale PROFIS ha integrato funzioni AI per accelerare "
    "la registrazione delle fatture elettroniche passive e dei movimenti bancari. "
    "Il sistema apprende nel tempo le abitudini di categorizzazione dello studio "
    "e propone automaticamente la classificazione contabile corretta, riducendo "
    "l\u2019intervento manuale alle eccezioni e ai casi ambigui. La riclassificazione "
    "automatica dei bilanci esterni, una delle operazioni pi\u00f9 time-intensive "
    "nella consulenza societaria, viene assistita da un modello che riconosce "
    "le strutture di bilancio e propone le corrispondenze con il piano dei conti "
    "dello studio."
)

para(doc,
    "TeamSystem, orientato principalmente alle aziende di dimensioni medio-grandi, "
    "ha sviluppato funzioni di classificazione automatica delle spese e gestione "
    "predittiva dei flussi di cassa. Il modulo AI analizza i movimenti storici, "
    "identifica stagionalit\u00e0 e anomalie, e genera previsioni di liquidit\u00e0 a "
    "trenta e novanta giorni. Per gli studi che offrono consulenza finanziaria "
    "alle imprese clienti, questa funzionalit\u00e0 diventa un argomento concreto "
    "da portare in fase di sviluppo del mandato."
)

h2(doc, "La questione della privacy e della responsabilit\u00e0")

para(doc,
    "L\u2019adozione dell\u2019AI nei gestionali non \u00e8 priva di implicazioni normative. "
    "Quando un software elabora dati fiscali e contabili dei clienti attraverso "
    "modelli di intelligenza artificiale, si pone il tema della catena del "
    "trattamento dei dati. Lo studio deve verificare che il fornitore del software "
    "abbia sottoscritto un accordo di responsabile del trattamento (DPA) conforme "
    "al GDPR, che i dati non vengano trasmessi fuori dallo Spazio Economico Europeo "
    "senza garanzie adeguate, e che il modello AI non venga addestrato sui dati "
    "dei clienti senza consenso esplicito."
)

para(doc,
    "I principali fornitori italiani hanno dichiarato che i loro modelli AI "
    "non si addestrano sui dati dei clienti e che il trattamento avviene "
    "interamente su infrastrutture certificate in Europa. Ma la verifica di queste "
    "dichiarazioni compete allo studio professionale, che rimane responsabile "
    "nei confronti dei propri clienti. Il consiglio operativo \u00e8 di richiedere "
    "al fornitore documentazione esplicita sul trattamento dei dati, integrarla "
    "nel registro delle attivit\u00e0 di trattamento e aggiornarla ogni volta "
    "che il software viene aggiornato con nuove funzioni AI."
)

h2(doc, "Il cambiamento di ruolo che nessuno nomina")

para(doc,
    "Il dibattito sull\u2019AI nei gestionali si concentra spesso sulla produttivit\u00e0: "
    "quanto tempo si risparmia, quanti errori si evitano, quanto aumenta "
    "la capacit\u00e0 di gestire clienti. Ma il cambiamento pi\u00f9 profondo \u00e8 di natura "
    "professionale. Quando la categorizzazione contabile, il controllo delle "
    "dichiarazioni e la riconciliazione bancaria vengono gestiti in automatico, "
    "il commercialista smette di essere prevalentemente un operatore contabile "
    "e diventa prevalentemente un consulente. Le ore liberate dall\u2019automazione "
    "devono essere reindirizzate verso attivit\u00e0 ad alto valore aggiunto: "
    "pianificazione fiscale, consulenza societaria, supporto alle decisioni "
    "finanziarie dei clienti."
)

para(doc,
    "Questa transizione non \u00e8 automatica. Richiede che lo studio ridefinisca "
    "la propria proposta di valore, comunichi ai clienti che cosa offre oggi "
    "che prima non poteva offrire, e prezzi di conseguenza il proprio lavoro. "
    "L\u2019AI nei gestionali \u00e8 gi\u00e0 qui, ma il suo impatto reale dipende da come "
    "gli studi professionali scelgono di usare il tempo che essa libera."
)

sep(doc)

# --- FONTI ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Fonti e riferimenti")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Fondazione Nazionale di Ricerca dei Commercialisti \u2014 Indagine sull\u2019uso dell\u2019AI "
    "negli studi professionali, luglio-settembre 2025",
    "Wolters Kluwer Italia \u2014 Comunicato stampa: Genya Dichiarativi Expert AI "
    "(marzo 2026)",
    "Sistemi \u2014 Sistemi AI per PROFIS: registrazione fatture e riclassificazione bilanci",
    "TeamSystem \u2014 AI e Automazione Contabile Studio Commercialisti (2026)",
    "Datalog \u2014 Commercialisti 2026: evoluzione e trend della digitalizzazione "
    "negli studi professionali",
    "Wolters Kluwer \u2014 L\u2019intelligenza artificiale trasforma la contabilit\u00e0: "
    "opportunit\u00e0 e sfide (2026)",
    "Professionista Digitale \u2014 Contabilit\u00e0 2026: trend e digitalizzazione "
    "degli studi professionali",
]
for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"\u2022 {s}")
    r.font.name = 'Calibri'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte")
r.font.name = 'Calibri'
r.font.size = Pt(8)
r.italic = True
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

doc.save(OUTPUT)
print(f"Salvato: {OUTPUT}")
