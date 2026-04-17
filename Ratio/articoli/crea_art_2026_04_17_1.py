#!/usr/bin/env python3
"""
Articolo 1: Il cliente che chiede se hai usato l'AI
File: Ratio/articoli/2026-04-17_informativa-ai-clienti-legge-132.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/home/user/amattavelli/Ratio/articoli/2026-04-17_informativa-ai-clienti-legge-132.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(3)
sec.right_margin = Cm(3)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot)
    pPr.append(pBdr)


def para(doc, text, size=11, italic=False, color=None, sa=8, sb=0,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


# --- TESTATA ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.all_caps = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Normativa e Professione")
r.font.name = 'Calibri'
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# --- TITOLO ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "Il cliente che chiede se hai usato l\u2019AI: "
    "la risposta che la Legge 132 ti impone di avere"
)
r.font.name = 'Calibri'
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Dal 2025 informare il cliente sull\u2019uso di strumenti di intelligenza artificiale "
    "non \u00e8 pi\u00f9 una scelta comunicativa: \u00e8 un obbligo di legge, con conseguenze precise."
)
r.font.name = 'Calibri'
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  \u2022  17 aprile 2026")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# --- CORPO ---
para(doc,
    "Hai ricevuto una telefonata da un cliente che voleva sapere se l\u2019analisi inviatale "
    "era stata preparata con l\u2019aiuto di un\u2019intelligenza artificiale. Non era una critica, "
    "era curiosit\u00e0. Probabilmente hai risposto in modo vago, perch\u00e9 non eri sicuro di "
    "come inquadrare la questione. Dal secondo semestre del 2025, per\u00f2, quella risposta "
    "vaga non \u00e8 pi\u00f9 un\u2019opzione: la Legge 132 del 2025 stabilisce un obbligo esplicito "
    "di informativa per tutti i professionisti intellettuali, commercialisti inclusi."
)

para(doc,
    "L\u2019articolo 13 della legge stabilisce che chi esercita una professione intellettuale "
    "ha l\u2019obbligo di comunicare al cliente, con linguaggio chiaro, semplice ed esaustivo, "
    "quali sistemi di intelligenza artificiale vengono impiegati nell\u2019esecuzione "
    "dell\u2019incarico e per quali finalit\u00e0. L\u2019obbligo non si applica soltanto ai casi in "
    "cui l\u2019AI produce il testo finale: vale anche per l\u2019analisi dei dati, la redazione "
    "di bozze, la verifica automatica di adempimenti, l\u2019elaborazione di previsioni "
    "finanziarie. Se uno strumento AI ha contribuito all\u2019incarico in modo sostanziale, "
    "il cliente deve saperlo."
)

para(doc,
    "Il Consiglio Nazionale dei Dottori Commercialisti ha gi\u00e0 elaborato una clausola "
    "contrattuale tipo, disponibile sul sito istituzionale, che i professionisti possono "
    "integrare nel mandato professionale. La clausola descrive la tipologia di strumenti "
    "utilizzati, le categorie di dati trattati e i limiti della supervisione umana. "
    "Non \u00e8 un adempimento burocratico da far firmare in coda al contratto: \u00e8 il punto "
    "di partenza per un dialogo con il cliente sul perimetro dell\u2019AI nella relazione "
    "professionale."
)

para(doc,
    "Una precisazione utile: la legge non impone al professionista di ottenere un consenso "
    "separato per l\u2019uso dell\u2019AI, n\u00e9 di chiedere approvazione caso per caso. Impone "
    "trasparenza. Il cliente deve sapere che l\u2019AI fa parte degli strumenti dello studio, "
    "non deve approvare ogni singola query. Questa distinzione \u00e8 importante perch\u00e9 evita "
    "di trasformare l\u2019obbligo in un freno operativo: si informa all\u2019inizio del rapporto, "
    "si aggiorna se cambiano le pratiche, si risponde con chiarezza se il cliente chiede."
)

h2(doc, "L\u2019intreccio con il GDPR e la privacy dei dati")

para(doc,
    "La Legge 132 si sovrappone, in parte, alle disposizioni gi\u00e0 esistenti del GDPR. "
    "Quando si usano strumenti AI che elaborano dati personali del cliente, come i software "
    "di contabilit\u00e0 con intelligenza artificiale integrata, lo studio \u00e8 tenuto a verificare "
    "che il fornitore del servizio abbia sottoscritto un accordo di responsabile del "
    "trattamento (il cosiddetto DPA) e che i dati non vengano trasmessi a server "
    "extraeuropei senza adeguate garanzie. Molti professionisti usano oggi strumenti AI "
    "generalisti per redigere analisi o comunicazioni che contengono, anche indirettamente, "
    "dati identificativi dei clienti. Su questo punto la Legge 132 introduce una "
    "responsabilit\u00e0 aggiuntiva rispetto a quanto gi\u00e0 previsto dalla normativa sulla "
    "protezione dei dati."
)

para(doc,
    "Nella pratica, l\u2019obbligo informativo si gestisce in modo abbastanza lineare. Chi usa "
    "strumenti AI nella propria attivit\u00e0 professionale aggiorna il mandato con la clausola "
    "tipo, specifica quali categorie di strumenti utilizza (software specializzati, "
    "assistenti generalisti, entrambi) e mantiene questa informazione aggiornata se cambia "
    "le proprie abitudini operative. Non \u00e8 necessario comunicare al cliente ogni singola "
    "interazione con uno strumento AI: \u00e8 sufficiente che il perimetro generale sia chiaro "
    "all\u2019inizio del rapporto professionale."
)

h2(doc, "L\u2019obbligo come occasione")

para(doc,
    "Molti professionisti hanno reagito alla Legge 132 come a un ulteriore peso normativo. "
    "La prospettiva \u00e8 comprensibile, ma riduttiva. I clienti che capiscono come lavora il "
    "loro professionista, compresi gli strumenti che usa, tendono a fidarsi di pi\u00f9. La "
    "trasparenza sull\u2019uso dell\u2019AI pu\u00f2 diventare un elemento di posizionamento: lo studio "
    "che comunica con chiarezza quali strumenti usa, per quali funzioni e con quale "
    "supervisione si distingue dagli studi che tacciono, spesso perch\u00e9 non si sono ancora "
    "posti la domanda."
)

para(doc,
    "La scadenza di agosto 2026 per i sistemi AI ad alto rischio, prevista dall\u2019AI Act "
    "europeo, render\u00e0 questo tema ancora pi\u00f9 urgente. Il professionista che ha gi\u00e0 "
    "aggiornato il proprio mandato e ha discusso con il cliente il ruolo dell\u2019AI nel "
    "rapporto professionale si trover\u00e0 in una posizione molto pi\u00f9 solida rispetto a chi "
    "dovr\u00e0 fare tutto insieme, sotto pressione normativa. Iniziare adesso costa poco. "
    "Aspettare agosto potrebbe costare di pi\u00f9."
)

sep(doc)

# --- FONTI ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Fonti e riferimenti")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Legge 132/2025, art. 13 \u2014 Obblighi di trasparenza per l\u2019uso dell\u2019AI da parte "
    "dei professionisti intellettuali",
    "Consiglio Nazionale dei Dottori Commercialisti e degli Esperti Contabili \u2014 "
    "Clausola contrattuale tipo per l\u2019informativa AI ai clienti (2025)",
    "EC News \u2014 AI Act e Legge n. 132/2025: cosa cambia per i commercialisti",
    "Agenzie Wolters Kluwer Italia \u2014 Legge 132/2025 sull\u2019AI: gli obblighi che gi\u00e0 "
    "oggi gravano su commercialisti, avvocati e CDL",
    "Tom's Hardware Business \u2014 La doppia compliance che inchioda le imprese italiane "
    "all\u2019AI Act (2026)",
]
for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"\u2022 {s}")
    r.font.name = 'Calibri'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte")
r.font.name = 'Calibri'
r.font.size = Pt(8)
r.italic = True
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

doc.save(OUTPUT)
print(f"Salvato: {OUTPUT}")
