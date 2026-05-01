"""
Articolo 1 — Il modello che usa il computer
Ratio/articoli/2026-05-01_gpt54-computer-use-studio-professionale.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-05-01_gpt54-computer-use-studio-professionale.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Maggio 2026  |  Strumenti e Modelli AI")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Il modello che usa il computer:\ncosa cambia quando l’AI controlla lo schermo")
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "GPT-5.4 segna il passaggio da modello che risponde a modello che agisce. "
    "Per chi lavora con documenti, fogli e gestionali, le implicazioni sono concrete."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  1 maggio 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Immagina di tornare in studio dopo una riunione di due ore e trovare il riepilogo mensile "
    "delle fatture già compilato. Le anomalie IVA sono state segnalate in una nota separata, "
    "tre fatture problematiche sono state messe in evidenza con una spiegazione sintetica. "
    "Nessun collaboratore era in ufficio. Lo ha fatto il software, aprendo applicazioni, "
    "navigando tra finestre, copiando dati da un gestionale all’altro come farebbe un "
    "operatore alla scrivania."
)

para(doc,
    "Questo non è un esempio ipotetico. Il 5 marzo 2026, OpenAI ha rilasciato GPT-5.4, "
    "il primo modello general-purpose della casa con computer use nativo integrato: la "
    "capacità di controllare il mouse, aprire programmi, compilare campi, scaricare file e "
    "lavorare su più applicazioni in sequenza senza alcun intervento umano tra un passo e "
    "l’altro. Nei test di riferimento, il modello raggiunge il 75% di successo su "
    "OSWorld-Verified, il benchmark internazionale per la misurazione del controllo "
    "informatico, superando la media degli operatori umani attestata al 72,4%."
)

heading(doc, "Cosa significa “usare il computer” per un modello AI")

para(doc,
    "La distinzione con i modelli precedenti vale la pena di capirla per valutare l’impatto "
    "pratico. Gli assistenti AI che molti professionisti usano già da mesi (ChatGPT, Claude, "
    "Gemini) operano in modalità conversazionale: ricevono testo, producono testo. L’utente "
    "deve poi trasferire quell’output nel suo gestionale, nella sua email, nel suo foglio "
    "Excel, e questo passaggio richiede tempo e introduce la possibilità di errori. Il "
    "modello risponde, ma non agisce."
)

para(doc,
    "GPT-5.4 cambia questo schema. Può aprire il browser, cercare dati su un portale "
    "esterno, scaricare un documento, aprire Excel, inserire i valori nelle celle corrette "
    "e chiudere la sessione. Il termine tecnico è “computer use” o “agentic computer "
    "control”, e segnala una trasformazione qualitativa nel modo in cui il modello interagisce "
    "con l’ambiente digitale di lavoro. Per i professionisti che gestiscono flussi "
    "documentali strutturati (inserimento dati, riconciliazione, compilazione di moduli "
    "standard) questo significa che alcune attività oggi gestite da personale amministrativo "
    "diventano automatizzabili con una supervisione umana molto più leggera."
)

heading(doc, "I limiti che contano")

para(doc,
    "Il 75% di successo su benchmark significa anche il 25% di insuccesso o errore. Su "
    "un compito come la compilazione di una dichiarazione IVA o la riconciliazione di "
    "estratti bancari, un errore al 25% non è accettabile senza revisione umana. Il "
    "computer use funziona bene su task strutturati e ripetitivi, molto meno su eccezioni, "
    "interfacce non standard o decisioni che richiedono giudizio contestuale."
)

para(doc,
    "Questo è il punto su cui i professionisti devono prestare la massima attenzione. "
    "La tentazione, con uno strumento così potente, è di delegare troppo. Un agente che "
    "usa il computer produce risultati visibili, credibili, ben formattati, che sembrano "
    "opera di un operatore competente. Ma la plausibilità dell’output non è sinonimo di "
    "correttezza. La firma sul documento rimane del professionista, e la responsabilità "
    "verso il cliente non si trasferisce al modello che ha compilato la bozza."
)

heading(doc, "Come si usa in pratica nello studio professionale")

para(doc,
    "Le applicazioni più mature per chi lavora in ambito contabile o legale sono, al "
    "momento, quelle che operano su processi altamente strutturati: riconciliazione di "
    "estratti bancari con la prima nota, verifica di scadenze e posizioni debitorie su "
    "portali esterni, compilazione di preventivi e lettere di incarico da template "
    "predefiniti, raccolta e sintesi di dati da più fonti documentali. Sono tutti compiti "
    "che hanno in comune un input prevedibile e una sequenza di passi definita."
)

para(doc,
    "La configurazione richiede comunque una fase di setup: definire le credenziali che "
    "il modello può usare, i limiti di accesso alle applicazioni, le regole di escalation "
    "verso l’operatore umano quando il compito esce dallo schema previsto. Questo lavoro "
    "non è tecnicamente complesso, ma richiede una riflessione approfondita sui propri "
    "processi, non solo sul software da adottare."
)

para(doc,
    "Il consiglio più utile che emerge da chi ha già adottato questi strumenti in contesti "
    "professionali strutturati è uno: non partire dai compiti più visibili o importanti. "
    "Iniziare da un processo che già funziona bene, di cui si conoscono tutte le eccezioni "
    "possibili, e osservare dove il modello diverge dall’aspettativa. Solo così si capisce "
    "dove serve la supervisione umana e dove il modello è davvero affidabile senza "
    "controllo continuo."
)

heading(doc, "La governance del processo, non solo dell’output")

para(doc,
    "La disponibilità del computer use a livello general-purpose non cambia solo "
    "l’efficienza di certi compiti. Cambia la natura della supervisione professionale. "
    "Chi usa questi strumenti non deve più solo verificare un testo generato, ma capire "
    "cosa il modello ha fatto concretamente sul proprio sistema: quali dati ha letto, "
    "quali ha modificato, quali decisioni operative ha preso in autonomia. Il controllo "
    "si sposta dalla revisione dell’output alla governance del processo."
)

para(doc,
    "Studi e imprese che si organizzeranno in questa direzione avranno uno strumento "
    "molto potente. Quelli che si limiteranno ad attivare il computer use senza ripensare "
    "la supervisione si troveranno a gestire errori difficili da intercettare in tempo, "
    "proprio perché il modello agisce in modo autonomo e produce risultati che sembrano "
    "corretti anche quando non lo sono. La prossima competenza critica per il professionista "
    "non è saper usare l’AI: è saper definire cosa l’AI può fare senza essere guardato."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "OpenAI — GPT-5.4 Technical Report (marzo 2026)",
    "OSWorld-Verified Benchmark — Computer Use Evaluation (2026)",
    "AI4Business — “GPT-5.4: la nuova generazione dell’AI per il lavoro professionale” (2026)",
    "SmartWorld — “OpenAI lancia GPT-5.4: ora l’IA sa davvero lavorare al posto nostro” (2026)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
