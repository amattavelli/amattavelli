"""
Articolo 3 — Genya Dichiarativi con AI e la responsabilità che resta allo studio
Ratio/articoli/2026-04-24_genya-dichiarativi-ia-responsabilita-studio.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-24_genya-dichiarativi-ia-responsabilita-studio.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'; run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'; run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Strumenti e Gestionale")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "Quando il gestionale compila la dichiarazione:\n"
    "Genya, l'AI e la responsabilità che resta allo studio"
)
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Wolters Kluwer ha integrato l'AI nei dichiarativi. "
    "Cosa cambia nel lavoro quotidiano dello studio, cosa non cambia, "
    "e dove si nasconde il rischio maggiore."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  24 aprile 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Aprile 2026, stagione delle dichiarazioni IVA. In diversi studi professionali italiani, "
    "il software gestionale ha già pre-compilato le bozze delle dichiarazioni analizzando "
    "automaticamente i movimenti contabili dell'anno. Il commercialista apre il file, trova "
    "percentuali e righe già popolate, controlla, corregge dove serve e firma. Lo strumento "
    "si chiama Genya Dichiarativi con Intelligenza Artificiale, sviluppato da Wolters Kluwer, "
    "ed è disponibile dal 2026 per i clienti della piattaforma. Qualcosa sta cambiando nel "
    "modo concreto in cui lavora uno studio professionale. Vale la pena capire bene cosa."
)

heading(doc, "Cosa fa il sistema: assistenza, non sostituzione")

para(doc,
    "Genya Dichiarativi integra un sistema AI che analizza i dati contabili caricati nel "
    "gestionale e propone la compilazione dei modelli dichiarativi, a partire dall'IVA 2026. "
    "Il sistema suggerisce, il professionista verifica, approva o modifica. "
    "L'architettura del prodotto è pensata esplicitamente come assistenza: la firma resta "
    "al commercialista, e con essa la responsabilità. Non è un dettaglio tecnico: è la "
    "premessa su cui si costruisce tutto l'uso corretto dello strumento."
)

para(doc,
    "Per capire cosa cambia davvero in uno studio che adotta questo tipo di integrazione, "
    "occorre distinguere tre livelli di impatto. Il primo è il risparmio di tempo sulle "
    "fasi meccaniche della compilazione: l'inserimento manuale di dati già presenti nel "
    "sistema, la verifica delle quadrature, la formattazione delle righe. Questo lavoro, "
    "spesso delegato a collaboratori junior, viene velocizzato in modo significativo. "
    "Il secondo livello riguarda la qualità del controllo: un sistema che ha già elaborato "
    "i dati segnala in automatico le anomalie (righe incongruenti, percentuali fuori norma "
    "rispetto agli anni precedenti, dati mancanti), riducendo il rischio di errori di "
    "distrazione su dichiarazioni complesse. Il terzo livello è strategico: il commercialista "
    "libera ore da attività di inserimento e le reindirizza verso la relazione con il cliente, "
    "l'analisi dell'ottimizzazione fiscale, la consulenza a valore aggiunto."
)

heading(doc, "Cosa il sistema non fa: i limiti da conoscere prima di usarlo")

para(doc,
    "Ciò che lo strumento non fa è altrettanto rilevante da capire. Un sistema AI che "
    "pre-compila dati contabili non interpreta situazioni straordinarie, come operazioni "
    "infragruppo, ristrutturazioni o variazioni di regime fiscale che richiedono una "
    "valutazione caso per caso. Non gestisce la normativa più recente se non è stato "
    "aggiornato con essa, e non sostituisce il giudizio professionale su operazioni non "
    "standard. La qualità dell'output dipende interamente dalla qualità dei dati in ingresso: "
    "se la contabilità ha errori di classificazione, il sistema li replica e, in alcuni "
    "casi, li amplifica propagandoli su più righe del modello."
)

para(doc,
    "Questo introduce una questione pratica che gli studi devono gestire consapevolmente: "
    "il rischio della fiducia eccessiva nello strumento. Uno studio che usa l'AI per la "
    "compilazione e riduce i controlli manuali perché \"il sistema ci pensa\" espone il "
    "professionista a un rischio maggiore, non minore. La pre-compilazione automatica "
    "non è una garanzia di correttezza: è una proposta che richiede verifica. "
    "Trattarla come un punto di arrivo anziché come un punto di partenza è l'errore "
    "più comune nella prima fase di adozione di questi strumenti."
)

heading(doc, "L'obbligo di informativa e il valore della trasparenza con il cliente")

para(doc,
    "La Legge 132/2025, in vigore dal 10 ottobre 2025, aggiunge un elemento che molti studi "
    "non hanno ancora considerato: il professionista che usa strumenti AI nello svolgimento "
    "dell'incarico deve informare il cliente in modo chiaro e semplice. Questo vale anche "
    "per l'uso di Genya o di qualsiasi altro gestionale con AI integrata. L'adempimento "
    "non è una formalità burocratica aggiuntiva: è un'apertura del rapporto professionale "
    "che, affrontata con il giusto approccio, può diventare un'opportunità."
)

para(doc,
    "Comunicare al cliente che la bozza della dichiarazione è stata elaborata con il "
    "supporto di un sistema AI e poi verificata dallo studio significa anche spiegare "
    "il valore del controllo professionale che rimane centrale nel processo. Il cliente "
    "capisce che l'AI non sostituisce il commercialista, ma lo aiuta a fare meglio e più "
    "velocemente ciò che faceva già. Questa conversazione, se condotta bene, rafforza il "
    "rapporto di fiducia anziché indebolirlo."
)

heading(doc, "Come strutturare il controllo sull'output del sistema")

para(doc,
    "Per adottare Genya Dichiarativi con AI in modo professionale, lo studio ha bisogno "
    "di definire una procedura interna di revisione dell'output. Non si tratta di fare "
    "tutto il lavoro manualmente come prima: si tratta di identificare i punti critici "
    "dove il sistema ha più probabilità di sbagliare e concentrare lì l'attenzione del "
    "professionista. Le operazioni straordinarie dell'anno, le variazioni di regime, i "
    "clienti con situazioni atipiche: questi sono i casi da esaminare con la stessa "
    "attenzione di sempre, indipendentemente da quello che il sistema ha proposto."
)

para(doc,
    "Adottare un gestionale con AI integrata richiede quindi due cose in parallelo: "
    "fiducia nel sistema per le parti meccaniche e standardizzate, e presidio professionale "
    "invariato per le parti che richiedono giudizio. Uno studio che sa usare bene Genya "
    "sa anche dove il sistema si ferma e dove comincia il suo lavoro. La differenza tra i "
    "due momenti non è tecnica: è professionale, ed è quella per cui il cliente paga."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "Wolters Kluwer — Comunicato lancio Genya Dichiarativi con IA (aprile 2026)",
    "datamanager.it — \"Wolters Kluwer Tax and Accounting, l'IA ridisegna il ruolo del commercialista\" (aprile 2026)",
    "Legge 132/2025 — Disposizioni nazionali in materia di intelligenza artificiale, art. 3 (obbligo informativa)",
    "Consiglio Nazionale dei Dottori Commercialisti — \"L'aiuto intelligente al commercialista\", ver. 2.0 (ottobre 2025)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
