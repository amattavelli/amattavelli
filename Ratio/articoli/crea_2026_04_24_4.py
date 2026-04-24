"""
Articolo 4 — Cassazione e responsabilità professionale con l'AI nel mezzo
Ratio/articoli/2026-04-24_cassazione-firma-commercialista-responsabilita-ai.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-24_cassazione-firma-commercialista-responsabilita-ai.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'; run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'; run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Responsabilità Professionale")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "Cosa firma il commercialista\nquando firma con l'AI nel mezzo"
)
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "La Cassazione nel 2026 ha chiarito che trasmettere una dichiarazione "
    "equivale a verificarne il contenuto. Con l'AI che redige le bozze, "
    "questo principio acquisisce un peso nuovo."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  24 aprile 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "L'ordinanza della Corte di Cassazione n. 5635 del 2026 ha confermato una sanzione "
    "a carico di un commercialista che si era limitato a trasmettere telematicamente una "
    "dichiarazione fiscale senza averla redatta personalmente. La Corte ha stabilito che "
    "la firma del professionista implica la verifica nel merito del contenuto, non solo la "
    "trasmissione formale del documento. La sentenza riguardava un caso precedente all'AI, "
    "ma il principio che afferma vale con forza ancora maggiore oggi, quando i sistemi AI "
    "possono redigere bozze, compilare modelli e proporre elaborazioni che il professionista, "
    "a volte, approva con un clic."
)

heading(doc, "L'automation bias: il rischio che non si vede")

para(doc,
    "Quando un sistema AI genera una bozza di dichiarazione o redige un testo professionale, "
    "il rischio concreto è che il professionista lo tratti come se fosse già controllato, "
    "riducendo involontariamente la propria soglia di attenzione. Questo fenomeno, documentato "
    "in diversi studi sul comportamento umano nell'interazione con sistemi automatizzati, "
    "prende il nome di automation bias: tendiamo a fidarci di più dell'output di un sistema "
    "automatico rispetto a quello di un collega umano, anche quando non abbiamo basi concrete "
    "per farlo. La presentazione ordinata e formalmente corretta di un output AI tende a "
    "trasmettere un'impressione di affidabilità che non corrisponde necessariamente alla "
    "sostanza del contenuto."
)

para(doc,
    "Nel contesto professionale italiano, l'automation bias si traduce in un rischio "
    "deontologico preciso. Se il commercialista usa un sistema AI per pre-compilare una "
    "dichiarazione e poi la firma senza una verifica adeguata, sta mettendo la propria "
    "responsabilità in mano a un sistema che non risponde degli errori che produce. "
    "I sistemi AI sbagliano: commettono errori di calcolo in condizioni particolari, "
    "non gestiscono bene le eccezioni normative, possono presentare dati in modo "
    "formalmente corretto ma sostanzialmente fuorviante. La responsabilità professionale "
    "non si trasferisce al software con la pre-compilazione automatica."
)

heading(doc, "L'obbligo di informativa: trasparenza che diventa protezione")

para(doc,
    "La Legge 132/2025 affronta il problema da una prospettiva diversa ma complementare. "
    "L'obbligo di informativa al cliente sull'uso di sistemi AI nello svolgimento "
    "dell'incarico non è una formalità: è il riconoscimento che il cliente ha il diritto "
    "di sapere se la prestazione professionale che sta ricevendo è stata elaborata, anche "
    "in parte, da un sistema automatico. Questo obbligo crea un incentivo pratico alla "
    "trasparenza: un professionista che comunica al cliente di aver usato un sistema AI "
    "per la bozza e di averla poi verificata si vincola anche a quella verifica. "
    "La trasparenza, in questo caso, non espone: protegge."
)

para(doc,
    "Dal punto di vista deontologico, comunicare l'uso dell'AI non equivale ad ammettere "
    "una diminuzione della qualità del servizio. Significa descrivere il processo con "
    "cui il servizio è stato prodotto. Un cliente informato sa che la sua dichiarazione "
    "non è stata redatta interamente a mano, ma sa anche che è stata revisionata da un "
    "professionista qualificato che ne risponde. Questo è esattamente il modello che "
    "distingue un uso professionale dell'AI da un uso irresponsabile."
)

heading(doc, "Le polizze professionali: cosa verificare prima di usare l'AI sistematicamente")

para(doc,
    "Alcune polizze di responsabilità civile professionale stanno aggiornando le proprie "
    "condizioni per coprire esplicitamente i danni derivanti da errori di sistemi AI usati "
    "nello studio. Alcune coperture, però, prevedono esclusioni se il professionista non "
    "ha mantenuto una supervisione adeguata sull'output del sistema. Prima di adottare "
    "strumenti AI in modo sistematico nello studio, verificare la propria polizza con il "
    "broker di riferimento è un atto di prudenza concreta, non un eccesso di cautela. "
    "Se la polizza non copre esplicitamente i danni da AI, occorre capire se è necessario "
    "un'estensione o una modifica delle condizioni."
)

heading(doc, "Come costruire procedure interne di controllo")

para(doc,
    "Il perimetro di responsabilità professionale non si riduce con l'AI: si ridefinisce. "
    "Il commercialista non è responsabile di ciò che fa il software, ma lo è di quello "
    "che firma. Costruire procedure interne di controllo sull'output AI significa "
    "identificare i punti critici del processo (le operazioni straordinarie, le situazioni "
    "atipiche, le variazioni di regime), definire chi le controlla e come, e documentare "
    "le verifiche effettuate. Questa documentazione non è un onere aggiuntivo: è la prova "
    "che lo studio ha gestito lo strumento in modo professionale, utile in caso di "
    "contestazione da parte del cliente o dell'Amministrazione Finanziaria."
)

para(doc,
    "Formare i collaboratori a non trattare l'output automatico come già verificato è la "
    "parte più difficile di questo percorso, perché richiede un cambiamento di abitudine "
    "e non solo l'adozione di uno strumento. Uno studio che ha chiarito internamente il "
    "confine tra ciò che il sistema fa e ciò che il professionista deve fare è uno studio "
    "che usa l'AI in modo responsabile. La domanda giusta da porsi non è \"posso usare "
    "l'AI per questo?\", ma \"come struturo il mio processo affinché l'AI mi aiuti senza "
    "abbassare la qualità del mio controllo?\""
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "Corte di Cassazione, ordinanza n. 5635/2026 — Responsabilità del professionista per dichiarazione trasmessa",
    "Legge 132/2025 — Disposizioni nazionali in materia di intelligenza artificiale, art. 3",
    "ktsfinance.com — \"Assicurazione Professionale Commercialisti e IA: 3 Rischi del 730/2026\" (2026)",
    "costanzoeassociati.it — \"Responsabilità del commercialista: Cassazione 2026\" (2026)",
    "ecnews.it — \"Intelligenza Artificiale e professioni intellettuali: l'obbligo di trasparenza\" (2025)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
