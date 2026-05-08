"""
Articolo 1 — GPT-5.5 Instant e le allucinazioni
Ratio/articoli/2026-05-08_gpt55-instant-allucinazioni-professionisti.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-05-08_gpt55-instant-allucinazioni-professionisti.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Maggio 2026  |  Strumenti e Modelli AI")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Quando il modello sbaglia con più stile:\nGPT-5.5 e il controllo che non si può saltare")
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Il nuovo default di ChatGPT riduce del 52% le allucinazioni su prompt legali e fiscali. "
    "Ma un errore più raffinato è anche più difficile da intercettare."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  8 maggio 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Chi usa ChatGPT da almeno un anno per lavoro ha imparato a convivere con una sensazione "
    "scomoda: la risposta è ben scritta, cita fonti con sicurezza, sembra ragionata, ma poi "
    "cerchi il riferimento normativo e non esiste. Oppure esiste, ma dice esattamente il "
    "contrario di quanto riportato. Questo fenomeno si chiama allucinazione e non è un difetto "
    "del singolo modello: è la natura statistica di come funzionano i large language model. "
    "Quello che cambia, però, è la frequenza. Il 5 maggio 2026 OpenAI ha aggiornato il modello "
    "di default di ChatGPT a GPT-5.5 Instant, sostituendo il precedente GPT-5.3. "
    "Nell'annuncio ufficiale c'è un numero che merita attenzione: le affermazioni allucinanti "
    "su prompt ad alto rischio in campo medico, legale e finanziario sono scese del 52,5% "
    "rispetto alla versione precedente."
)

para(doc,
    "Il dato va letto con attenzione, perché dice due cose insieme. La prima è quella "
    "rassicurante: il modello sbaglia molto meno quando gli fai domande su normativa fiscale, "
    "su detrazioni, su circolari dell'Agenzia delle Entrate, su scadenze contributive. "
    "La seconda è quella che si dimentica troppo in fretta: meno della metà delle allucinazioni "
    "non vuol dire zero allucinazioni. Se su cento prompt ad alto rischio il modello precedente "
    "ne sbagliava venti, il nuovo ne sbaglia circa nove o dieci. Meglio, certamente. Ma se stai "
    "preparando una risposta per un cliente su una questione delicata, quel dieci per cento "
    "pesa ancora."
)

heading(doc, "Cosa cambia nel lavoro quotidiano")

para(doc,
    "Per uno studio professionale questo aggiornamento ha conseguenze pratiche immediate, "
    "perché GPT-5.5 Instant è il modello che si apre quando lanci ChatGPT senza selezionare "
    "nulla. Chiunque nel team lo stia usando per bozze di pareri, per controlli su scadenze "
    "fiscali, per risposte rapide su normativa IVA, si trova oggi davanti a uno strumento "
    "più affidabile di quello della settimana scorsa. Senza aver fatto nulla, senza aver "
    "cambiato abitudini."
)

para(doc,
    "Il cambiamento riguarda anche lo stile delle risposte. GPT-5.5 Instant è stato addestrato "
    "per rispondere in modo più conciso, con meno formattazione eccessiva. Se hai provato a "
    "chiedere a ChatGPT un parere su un'operazione societaria negli ultimi mesi, probabilmente "
    "hai ricevuto una risposta con cinque titoli in grassetto, dieci punti elenco e una "
    "lunghezza da report aziendale. Il nuovo modello tende a rispondere in prosa, più "
    "direttamente, con meno scaffolding artificiale. Per chi usa l'output di ChatGPT come "
    "punto di partenza per documenti da rifinire, questa è una differenza concreta: meno "
    "editing per togliere la formattazione generata in automatico."
)

heading(doc, "Il problema dell’errore più raffinato")

para(doc,
    "C'è un aspetto che merita una riflessione specifica. Quando il modello sbagliava in modo "
    "evidente, la citazione di una norma inesistente o un numero chiaramente fuori scala "
    "attivava il campanello d'allarme del professionista. L'errore era visibile, e la "
    "revisione diventava automatica. Quando il modello sbaglia con più raffinatezza, come "
    "citare correttamente un articolo di legge ma applicarlo a una fattispecie sbagliata, "
    "o indicare una percentuale reale ma riferita a un anno diverso, l'errore supera la "
    "prima lettura. La plausibilità dell'output non è sinonimo di correttezza."
)

para(doc,
    "Questo è il paradosso di un modello che migliora: la fiducia cresce, la soglia di "
    "attenzione si abbassa, e l'errore che passa diventa più costoso di quello che veniva "
    "intercettato subito. Per i professionisti che firmano documenti prodotti con il supporto "
    "di ChatGPT, un modello più preciso non è un invito a verificare meno: è un invito a "
    "verificare in modo più mirato, concentrandosi su ciò che il modello può ancora sbagliare "
    "con più discrezione."
)

heading(doc, "Il nodo dei dati e della privacy")

para(doc,
    "C'è un aspetto del nuovo modello che riguarda direttamente i professionisti che gestiscono "
    "più clienti: GPT-5.5 Instant può ora attingere al contesto delle conversazioni precedenti, "
    "ai file caricati in sessioni passate, e agli account Gmail collegati per calibrare le "
    "risposte. Questo amplia le capacità di contestualizzazione, ma pone anche una domanda "
    "che nessun aggiornamento tecnico risolve da solo: dove finiscono i documenti che carichi "
    "in chat?"
)

para(doc,
    "La Legge 132/2025 e le linee guida del Garante Privacy italiano chiedono ai professionisti "
    "di sapere con precisione dove vengono trattati i dati dei clienti. Caricare una pratica "
    "fiscale in ChatGPT, anche sulla versione Plus o Business, richiede di aver valutato se "
    "la policy di OpenAI è compatibile con le obbligazioni contrattuali verso il cliente. "
    "Il modello migliora, ma le responsabilità verso i clienti non diminuiscono con lui. "
    "GPT-5.5 Instant è disponibile anche per gli utenti del piano gratuito, il che significa "
    "che la riduzione delle allucinazioni raggiunge un'utenza molto più ampia di quella che "
    "ha mai riflettuto su questi temi. Per un'azienda che ha dato accesso a ChatGPT ai propri "
    "dipendenti senza una policy interna sull'uso dell'AI, questo mese è un buon momento "
    "per rivedere quella scelta."
)

heading(doc, "Verificare di più o verificare meglio")

para(doc,
    "Un modello più preciso non è un modello infallibile. La buona notizia è che GPT-5.5 "
    "Instant riduce il margine di errore su domande sensibili. La notizia che spesso si "
    "omette è che questo riduce anche la visibilità dell'errore. Verificare l'output "
    "dell'AI non è diventato meno necessario: è diventato più sottile. La verifica che "
    "conta non è quella che controlla se il testo ha senso: è quella che controlla se "
    "il riferimento normativo si applica davvero al caso specifico del cliente che hai "
    "davanti, con la sua storia, le sue scelte passate, le sue particolarità contrattuali. "
    "Questo è il lavoro del professionista. Il modello può avvicinarsi molto alla risposta "
    "giusta, ma non può conoscere il dettaglio che fa la differenza. Almeno per ora."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "OpenAI — Annuncio GPT-5.5 Instant, 5 maggio 2026",
    "TechCrunch — \"OpenAI releases GPT-5.5 Instant, a new default model for ChatGPT\" (5 maggio 2026)",
    "Punto Informatico — \"GPT-5.5 Instant ha meno allucinazioni e dà risposte più concise\" (2026)",
    "Zazoom — \"OpenAI: arriva GPT-5.5 Instant, calano del 52% le allucinazioni\" (6 maggio 2026)",
    "Legge 132/2025 — Disciplina organica sull'intelligenza artificiale in Italia",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
