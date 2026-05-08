"""
Articolo 3 — PMI come deployer nell'AI Act
Ratio/articoli/2026-05-08_pmi-chatgpt-deployer-ai-act-90-giorni.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-05-08_pmi-chatgpt-deployer-ai-act-90-giorni.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Maggio 2026  |  Normativa e Compliance AI")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("La PMI che usa ChatGPT ogni giorno\nsenza sapere di essere deployer")
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "A novanta giorni dall'AI Act, molte aziende italiane non sanno ancora "
    "che ruolo hanno nella catena normativa. La questione non riguarda le grandi imprese: "
    "riguarda chi usa ChatGPT, Copilot o un gestionale con AI integrata per lavorare."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  8 maggio 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Prendi un'azienda di quindici dipendenti che vende componenti industriali. "
    "Hanno un account ChatGPT Teams che usano per le email ai clienti stranieri, "
    "per bozze di offerte commerciali, per riassumere report di settore. Usano anche "
    "uno strumento di AI per analizzare le tendenze dei prezzi dei fornitori. "
    "Non hanno mai parlato di intelligenza artificiale in una riunione di direzione, "
    "nel senso normativo del termine. Tra novanta giorni, il 2 agosto 2026, dovranno "
    "rispondere a obblighi precisi previsti dall'AI Act europeo. "
    "La maggior parte di loro non lo sa."
)

para(doc,
    "L'AI Act europeo, pienamente applicabile dal 2 agosto 2026, distingue tre ruoli "
    "principali nella catena dell'intelligenza artificiale: il fornitore, cioè chi sviluppa "
    "e immette sul mercato il sistema AI; il distributore, che lo commercializza; e il "
    "deployer, cioè chi lo usa nella propria attività per uno scopo professionale. "
    "La stragrande maggioranza delle PMI italiane che usano ChatGPT, Microsoft Copilot, "
    "Gemini o qualsiasi strumento AI integrato in un software gestionale sono deployer. "
    "Non lo sanno, non lo dichiarano in nessun documento aziendale, e spesso nemmeno "
    "ci pensano."
)

heading(doc, "Cosa comporta essere deployer")

para(doc,
    "Per i deployer di sistemi a rischio limitato, come ChatGPT o Copilot in un contesto "
    "di uso generico, gli obblighi principali riguardano la trasparenza verso chi interagisce "
    "con il sistema e la garanzia che le persone che usano gli strumenti AI nell'organizzazione "
    "abbiano un livello sufficiente di competenza per farlo in modo critico. Questo secondo "
    "obbligo, spesso chiamato AI literacy, è entrato in vigore già il 2 febbraio 2025, "
    "sedici mesi prima della piena applicazione del Regolamento. Non riguarda una formazione "
    "certificata: riguarda la capacità, documentabile, di usare gli strumenti AI con "
    "consapevolezza critica."
)

para(doc,
    "Per una PMI, significa poter rispondere a una domanda come: chi ha usato l'AI in "
    "azienda, per quali scopi, e ha ricevuto un minimo di orientamento su come farlo in "
    "modo corretto? Se la risposta è che nessuno ci ha mai detto niente, si è già in "
    "zona di non conformità rispetto a un obbligo che esiste da oltre un anno. L'ACN, "
    "l'Agenzia per la Cybersicurezza Nazionale designata come autorità di vigilanza "
    "in Italia dalla Legge 132/2025, non ha ancora avviato controlli sistematici sulle PMI, "
    "ma la conformità non riguarda solo i controlli: riguarda la gestione del rischio "
    "interno e la responsabilità verso i propri dipendenti e clienti."
)

heading(doc, "Il rischio nascosto nei gestionali")

para(doc,
    "Molti dei software gestionali già in uso nelle aziende italiane integrano funzioni AI "
    "in modo quasi invisibile. Un ERP che suggerisce previsioni di magazzino, un CRM che "
    "assegna punteggi ai lead, un software di fatturazione che categorizza automaticamente "
    "le voci di costo: ognuno di questi è, tecnicamente, un sistema AI che l'azienda sta "
    "usando come deployer. La domanda da farsi entro agosto non è \"usiamo l'AI?\", "
    "ma \"sappiamo quali sistemi AI stiamo usando e a che scopo?\""
)

para(doc,
    "Questa distinzione diventa critica quando il sistema AI riguarda decisioni che toccano "
    "le persone. Un software che usa l'AI per valutare le candidature in fase di selezione "
    "del personale, o che assegna turni di lavoro sulla base di algoritmi di ottimizzazione, "
    "rientra nella categoria dei sistemi ad alto rischio secondo l'AI Act. Gli obblighi per "
    "questi sistemi sono molto più pesanti: documentazione tecnica, valutazione di conformità, "
    "registrazione nel database europeo. Molte PMI che usano questi strumenti non lo sanno "
    "e non si sono mai poste la domanda."
)

heading(doc, "Il censimento che nessuno ha ancora fatto")

para(doc,
    "Il primo passo concreto che nessun consulente o commercialista ha ancora proposto "
    "formalmente ai propri clienti è il censimento degli strumenti AI in uso in azienda. "
    "Non perché sia un compito difficile: perché nessuno l'ha ancora chiesto. Si tratta "
    "di raccogliere, in modo sistematico, quali strumenti digitali con funzioni AI vengono "
    "usati in azienda, da chi, per quali attività, e se il loro uso coinvolge dati personali "
    "di dipendenti, clienti o fornitori. Questo esercizio, che richiede alcune ore di lavoro "
    "e un foglio strutturato, produce due risultati utili: una mappa del rischio normativo "
    "e una base documentale per dimostrare la diligenza dell'organizzazione."
)

para(doc,
    "Per gli studi professionali che assistono PMI, questo è anche un'opportunità di servizio "
    "concreta. Aiutare un cliente a mappare i sistemi AI in uso, a valutare il loro livello "
    "di rischio e a predisporre una nota interna sull'uso responsabile degli strumenti non "
    "richiede una competenza legale specialistica: richiede metodo e conoscenza del contesto "
    "normativo. Molti studi stanno già offrendo questo servizio come estensione naturale "
    "della consulenza sulla privacy e sulla sicurezza informatica. Chi non lo ha ancora "
    "proposto ai propri clienti ha novanta giorni per farlo prima che diventi urgente."
)

heading(doc, "La scadenza come promemoria, non come punto di partenza")

para(doc,
    "Il 2 agosto 2026 non è una data che cambia tutto da un giorno all'altro. "
    "Le sanzioni per i deployer di sistemi a rischio limitato non sono quelle da prima "
    "pagina previste per i fornitori che violano i divieti fondamentali. Ma la logica "
    "della norma è chiara: chi usa l'AI in modo professionale ha una responsabilità verso "
    "chi è coinvolto in quelle decisioni. Per una PMI che ha dato in mano a ChatGPT "
    "il filtraggio delle candidature o la stesura delle condizioni di vendita, "
    "questa responsabilità è già attiva. Agosto è un promemoria. Il lavoro andava "
    "iniziato prima, e può iniziare adesso."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "Regolamento UE 2024/1689 — AI Act (Regolamento sull'Intelligenza Artificiale)",
    "Legge 132/2025 — Designazione ACN come autorità di vigilanza per l'AI Act in Italia",
    "FISCOeTASSE — \"Alfabetizzazione digitale e AI Act: nuovi obblighi per imprese e operatori\" (2026)",
    "Polaris AI — \"AI Act 2026: obblighi per PMI italiane\" (2026)",
    "ISTAT — Rilevazione sull'uso dell'AI nelle imprese italiane (2024)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
