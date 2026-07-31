"""
Quattro articoli Ratio -- 31 luglio 2026

1. 2026-07-31_ai-omnibus-luglio-2026-scadenze-cambiate.docx
2. 2026-07-31_budget-ai-184-milioni-governance-assente.docx
3. 2026-07-31_divario-ai-grandi-imprese-pmi-2026.docx
4. 2026-07-31_parita-prezzo-modelli-ai-20-euro-competenza.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# AI Omnibus in vigore dal 27 luglio 2026
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Normativa AI")
titolo(
    doc,
    "Il calendario dell’AI Act è cambiato ancora.",
    "Il 27 luglio 2026 è entrato in vigore il Digital Omnibus sull’AI. "
    "I sistemi ad alto rischio slittano a dicembre 2027. "
    "Gli obblighi di trasparenza del 2 agosto restano. "
    "Qualche semplificazione per le PMI, ma il perimetro non cambia.",
    "A cura della Redazione Ratio  •  31 luglio 2026"
)

para(doc,
    "Penelope tesseva di giorno e disfaceva di notte, tenendo i pretendenti a bada "
    "con la promessa di una risposta che non arrivava mai. Il calendario di compliance "
    "dell’AI Act ha preso ispirazione da quella strategia: ogni volta che le "
    "aziende prendono misura di una scadenza, il quadro normativo viene ricalibrato e "
    "il conto alla rovescia riparte. Il 27 luglio 2026 è entrato in vigore il "
    "Regolamento UE 2026/1744, che modifica in modo sostanziale le tempistiche "
    "applicative dell’AI Act originario, tre giorni prima della scadenza "
    "più attesa."
)

para(doc,
    "Il Regolamento UE 2026/1744, denominato Digital Omnibus sull’AI, è un "
    "intervento di semplificazione e ricalibrazione sull’AI Act. Il suo obiettivo "
    "dichiarato è ridurre gli oneri amministrativi per le imprese europee, in "
    "particolare per le piccole e medie, e adeguare le scadenze ai tempi reali di "
    "preparazione del mercato. Non introduce nuovi obblighi e non crea nuove categorie "
    "di sistemi vietati: interviene sulle tempistiche e sulla proporzionalità degli "
    "adempimenti già previsti. Per chi aveva imparato a memoria l’AI Act "
    "originale, alcune date importanti sono cambiate."
)

heading(doc, "Cosa è stato posticipato")

para(doc,
    "I sistemi ad alto rischio elencati nell’Allegato III dell’AI Act, quelli "
    "autonomi che non sono componenti di prodotti già regolamentati da altra "
    "normativa europea, slittano al 2 dicembre 2027. I sistemi ad alto rischio "
    "integrati come componenti di sicurezza in prodotti soggetti a normativa di "
    "armonizzazione dell’Unione (dispositivi medici, macchine industriali, veicoli) "
    "hanno tempo fino al 2 agosto 2028. Per le aziende che avevano questi sistemi nel "
    "mirino e non avevano ancora avviato il percorso di conformità, il Digital "
    "Omnibus ha prodotto l’effetto di un’amnistia temporanea: il percorso "
    "resta obbligatorio, la scadenza è stata allontanata."
)

heading(doc, "Cosa è rimasto fermo")

para(doc,
    "Gli obblighi di trasparenza dell’articolo 50 restano applicabili dal "
    "2 agosto 2026 come pianificato. I chatbot devono dichiarare la propria natura "
    "artificiale prima di ogni interazione con gli utenti, i contenuti sintetici "
    "devono essere marcati tecnicamente, i sistemi biometrici negli spazi accessibili "
    "al pubblico devono rendere visibile la propria presenza e funzione. Entrano in "
    "vigore nello stesso giorno le sanzioni per i fornitori di modelli di AI per uso "
    "generale e gli strumenti rafforzati di sorveglianza del mercato. Il Digital "
    "Omnibus ha lasciato intatto il livello di trasparenza visibile agli utenti "
    "finali, che rimane il perimetro degli adempimenti più immediati per chi "
    "usa sistemi AI nel proprio lavoro professionale."
)

heading(doc, "Le semplificazioni per le PMI")

para(doc,
    "Per le PMI e le imprese a media capitalizzazione, il Digital Omnibus introduce "
    "alleggerimenti documentali: requisiti di conformità proporzionati alle "
    "dimensioni, accesso facilitato alle risorse di orientamento degli organismi "
    "nazionali, e la possibilità di basarsi su modelli di compliance condivisi "
    "a livello settoriale anziché sviluppare documentazione completamente "
    "proprietaria. Queste semplificazioni incidono soprattutto sulla fase di "
    "valutazione del rischio e di documentazione tecnica, che per le imprese più "
    "piccole rappresentavano un costo sproporzionato rispetto alle attività "
    "effettivamente svolte con i sistemi AI."
)

para(doc,
    "Il Digital Omnibus sposta le scadenze ma non modifica il lavoro di base che ogni "
    "impresa deve compiere: censire i sistemi AI in uso, valutarne il profilo di "
    "rischio, formare il personale, documentare i processi decisionali. Ogni rinvio "
    "può essere letto come respiro aggiuntivo per le aziende che hanno già "
    "avviato questo percorso, perché hanno più tempo per completarlo con "
    "accuratezza. Per quelle che non hanno iniziato, il rinvio produce soprattutto "
    "l’impressione che ci sia ancora tempo per rimandare. Storicamente, le "
    "normative europee con scadenze multiple hanno avuto la caratteristica di "
    "concentrare l’ottanta per cento delle attività di adeguamento negli "
    "ultimi tre mesi utili."
)

para(doc,
    "Penelope alla fine finì la tela. Le scadenze normative, prima o poi, "
    "hanno la stessa tendenza."
)

riferimenti(doc, [
    "Federprivacy -- 'AI Omnibus in vigore dal 27 luglio 2026: cambiano le scadenze "
    "dell’AI Act, ma gli obblighi di trasparenza restano confermati' (27 luglio 2026)",
    "CityNext.it -- 'AI Omnibus, in vigore dal 27 luglio 2026' (27 luglio 2026)",
    "Altalex.com -- 'Il Digital Omnibus sull’Intelligenza Artificiale in GUUE: "
    "cosa cambia per le imprese' (27 luglio 2026)",
    "Rivista.ai -- 'AI Act, arriva il Digital Omnibus: meno burocrazia, nuove scadenze "
    "e qualche sorpresa. Pubblicato il Regolamento UE 2026/1744' (24 luglio 2026)",
    "LaborProject.it -- 'L’AI Act entra in una nuova fase: cosa cambia con il "
    "Digital Omnibus sull’AI' (28 luglio 2026)",
    "iSimply.it -- 'AI Act dal 2 agosto 2026: cosa diventa applicabile dopo il "
    "Digital Omnibus'",
    "CybersecItalia.it -- 'UE, in vigore l’AI Omnibus. Più semplificazione "
    "per le imprese'",
])
doc.save(BASE + "2026-07-31_ai-omnibus-luglio-2026-scadenze-cambiate.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# Budget AI 18,4 milioni, governance a zero
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Governance e Rischio AI")
titolo(
    doc,
    "18,4 milioni investiti. La governance AI può attendere.",
    "Le aziende italiane con oltre 500 addetti spenderanno 18,4 milioni in AI "
    "nel 2026. Il 40% non ha supervisione umana sui workflow agentici, il 25% "
    "non controlla gli accessi agli agenti, il 48% tiene un registro. "
    "La governance aspetta.",
    "A cura della Redazione Ratio  •  31 luglio 2026"
)

para(doc,
    "L’apprendista stregone della ballata di Goethe non aveva sbagliato "
    "l’incantesimo: aveva solo dimenticato di imparare quello per fermarlo. "
    "Il maestro era andato via, le scope portavano acqua a ritmo crescente, e il "
    "pavimento si allagava perché nessuno aveva previsto il momento in cui il "
    "sistema avrebbe smesso di essere una soluzione e sarebbe diventato un problema. "
    "Le aziende italiane che investono cifre crescenti in AI senza costruire i "
    "controlli corrispondenti stanno nella posizione di quell’apprendista: "
    "l’incantesimo funziona, il problema è che non si ferma quando vorrebbero."
)

para(doc,
    "Il dato è preciso e recente. Le aziende italiane con più di 500 addetti "
    "indicano per il 2026 una spesa media in AI di 18,4 milioni di dollari, con un "
    "incremento atteso del 45% nei due anni seguenti. Il ritorno economico previsto "
    "arriva al 38%, equivalente a 12,2 milioni di dollari. Il 73% dei dirigenti si "
    "dichiara soddisfatto del ROI corrente. Fin qui, il quadro è quello di "
    "un’adozione che accelera e produce risultati. Il problema emerge appena si "
    "guardano i numeri sul lato della governance."
)

heading(doc, "I numeri che i dirigenti non citano nelle presentazioni")

para(doc,
    "Lo stesso studio che ha prodotto quei numeri sull’investimento rivela che "
    "il 40% delle aziende italiane non ha processi con intervento umano per "
    "supervisionare i workflow agentici: agenti AI che prendono decisioni, inviano "
    "comunicazioni, modificano dati, senza che nessuna persona sia formalmente "
    "incaricata di verificarne l’output prima che diventi definitivo. Il 25% "
    "non ha sistemi di controllo degli accessi per gli agenti, il che significa che "
    "un agente può operare su sistemi aziendali senza che esista un perimetro "
    "preciso di ciò che può o non può fare. Solo il 48% mantiene un "
    "registro degli agenti in uso all’interno dell’organizzazione. La "
    "metà delle aziende italiane non sa con certezza quanti e quali agenti AI "
    "operano per suo conto."
)

para(doc,
    "Le fragilità non si fermano al livello tecnico. Meno della metà delle "
    "aziende ha un responsabile dedicato all’AI (45%), solo il 35% ha definito "
    "KPI per l’AI a livello manageriale, il 37% ha attivato percorsi di "
    "formazione sulle funzionalità e sui rischi degli strumenti adottati. "
    "In molte organizzazioni italiane l’AI è gestita come una spesa "
    "operativa anziché come un asset che richiede una funzione di presidio: "
    "si compra, si usa, si rinnova l’abbonamento. La governance arriva "
    "solo quando arriva il problema."
)

heading(doc, "Il rischio che scala con il budget")

para(doc,
    "La peculiarità del rischio AI è che cresce proporzionalmente "
    "all’investimento. Un’azienda che usa ChatGPT per redigere comunicazioni "
    "ha un’esposizione contenuta anche in assenza di governance formale. "
    "Un’azienda che ha agenti AI integrati nei processi di approvazione delle "
    "spese, nella comunicazione con i clienti e nella produzione di documenti "
    "contabili ha un profilo di rischio operativo e normativo molto più elevato, "
    "proporzionale all’autonomia che ha delegato ai sistemi. Il Digital Omnibus "
    "ha posticipato alcune scadenze per i sistemi ad alto rischio, ma non ha "
    "modificato la sostanza: chi usa agenti AI in processi decisionali rilevanti "
    "sta operando in un perimetro che l’AI Act regola già, e che le "
    "autorità di vigilanza inizieranno a ispezionare dal 2 agosto in poi."
)

para(doc,
    "Il maestro stregone di Goethe torna alla fine e sistema tutto con un incantesimo. "
    "Nel mondo reale, quello che torna non è il maestro: è l’ispettore."
)

riferimenti(doc, [
    "SbirciaNotizia.it -- 'AI, 18,4 milioni per azienda nel 2026: budget in crescita "
    "e controlli in ritardo' (19 luglio 2026)",
    "AdnKronos -- 'Le imprese italiane investiranno in AI 18,4 milioni di dollari "
    "nel 2026, +45% nei prossimi 2 anni' (16 luglio 2026)",
    "SpotAndWeb.it -- 'Le imprese italiane investiranno in AI 18,4 milioni di dollari "
    "nel 2026, +45% nei prossimi 2 anni'",
    "PuntoImpresaDigitale.camcom.it -- 'IA e cybersecurity nelle PMI: cresce "
    "l’adozione, ma la sfida è culturale'",
    "Regolamento UE 2024/1689 (AI Act), articoli 26-27 sugli obblighi dei deployer",
])
doc.save(BASE + "2026-07-31_budget-ai-184-milioni-governance-assente.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# Divario AI grandi imprese / PMI
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "PMI e Adozione AI")
titolo(
    doc,
    "Grandi imprese al 53%, PMI al 15%: il divario si allarga.",
    "Tra il 2023 e il 2026 l’adozione dell’AI nelle PMI italiane è "
    "triplicata. Le grandi imprese sono andate ancora più veloci. Il gap è "
    "passato da 25 a 37 punti. Nessuna norma, nemmeno l’AI Act, "
    "è pensata per chiuderlo.",
    "A cura della Redazione Ratio  •  31 luglio 2026"
)

para(doc,
    "Il Vangelo di Matteo, al capitolo 25, contiene un’osservazione "
    "sull’economia che molti studiosi hanno citato con più o meno pudore: "
    "a chi ha sarà dato, a chi non ha sarà tolto anche quello che ha. "
    "Robert Merton, negli anni Sessanta, aveva formalizzato il principio nel "
    "concetto di effetto Matteo: nei sistemi competitivi, il vantaggio iniziale "
    "tende ad autoamplificarsi, e le distanze tra i partecipanti crescono anche "
    "quando tutti migliorano in assoluto. I dati sull’adozione dell’AI "
    "nelle imprese italiane nel 2026 descrivono esattamente quello schema."
)

para(doc,
    "Tra il 2023 e il 2026, la quota di PMI italiane che dichiara di utilizzare "
    "almeno una tecnologia AI è triplicata: dal 5% circa al 15,7%. Una crescita "
    "reale e significativa, che molte analisi presentano come la svolta del mercato "
    "italiano. Il problema è che nello stesso periodo le grandi imprese sono "
    "passate dal 28% al 53,1% di adozione. Il gap tra le due fasce, che nel 2023 "
    "era di circa 25 punti percentuali, si è allargato a 37 punti nel 2026. "
    "Le PMI sono cresciute in assoluto; in termini relativi, si sono allontanate "
    "dalle grandi imprese più velocemente di quanto si stiano avvicinando."
)

heading(doc, "Perché il divario si allarga mentre tutti crescono")

para(doc,
    "Le grandi imprese dispongono di risorse da investire in infrastrutture dati, "
    "team dedicati, contratti enterprise con i fornitori di AI, e la capacità "
    "di attrarre profili specializzati. Le PMI, anche quando vogliono adottare "
    "l’AI in modo strutturato, spesso mancano della qualità dei dati "
    "necessaria per alimentarla, della figura professionale in grado di "
    "supervisionarne l’implementazione, e del tempo operativo per affrontare "
    "un cambiamento di processo mentre si gestisce il quotidiano. Il risultato è "
    "che le PMI adottano gli strumenti più semplici, quelli che richiedono il "
    "minore investimento iniziale e producono risultati visibili nel breve, "
    "lasciando alle grandi imprese le applicazioni più profonde e strutturali."
)

heading(doc, "Il paradosso normativo")

para(doc,
    "L’AI Act si applica con criteri sostanzialmente analoghi alle grandi "
    "imprese e alle PMI (con le semplificazioni introdotte dal Digital Omnibus per "
    "le seconde), senza però intervenire sulla causa del divario. La capacità "
    "di adottare l’AI in modo strutturato dipende da competenze, qualità "
    "dei dati e organizzazione interna, fattori su cui una norma europea di "
    "compliance non ha leva diretta. Il 76% delle PMI italiane non ha investito "
    "in AI e non prevede di farlo: per questa fascia, la questione della "
    "conformità all’AI Act è ancora lontana, e il problema a monte "
    "è più elementare. Riguarda la capacità di valutare se e come "
    "uno strumento digitale può cambiare un processo specifico."
)

heading(doc, "Il fattore che distingue chi recupera da chi resta indietro")

para(doc,
    "Le PMI che stanno effettivamente accorciando il divario condividono quasi "
    "sempre la stessa caratteristica: c’è qualcuno, interno o esterno, "
    "che ha le competenze per valutare dove l’AI può fare la differenza "
    "e come implementarla sul processo specifico. Commercialisti con formazione "
    "digitale, consulenti di gestione che conoscono gli strumenti, system integrator "
    "specializzati per settore: è questa figura, più che il budget o la "
    "tecnologia disponibile, a determinare se una PMI riesce a colmare il gap. "
    "Le politiche di supporto, gli incentivi fiscali, le piattaforme nazionali "
    "per le PMI riducono il costo dell’accesso. Ma il fattore che separa chi "
    "recupera da chi resta indietro è la presenza di qualcuno capace di usare "
    "quell’accesso."
)

para(doc,
    "Il Vangelo di Matteo non diceva come fermare l’effetto. "
    "Neanche i decreti attuativi italiani sull’AI."
)

riferimenti(doc, [
    "GazzettaDiRoma.it -- 'PMI italiane a due velocità: cresce l’intelligenza "
    "artificiale, ma resta forte il divario tra grandi e piccole imprese'",
    "ArenaDigitale.it -- 'Italia: cresce adozione AI ma non l’automazione' "
    "(21 luglio 2026)",
    "AgendaDigitale.eu -- 'AI nelle PMI italiane: competenze e dati frenano "
    "la svolta digitale'",
    "GazzettaDiMilano.it -- 'PMI italiane 2026: i dati tra AI, welfare e "
    "incertezza economica'",
    "Istat -- Rapporto sull’uso delle tecnologie ICT e AI nelle imprese "
    "italiane, 2026",
    "ESTE.it -- 'AI Act, le PMI italiane arrivano impreparate alla svolta'",
])
doc.save(BASE + "2026-07-31_divario-ai-grandi-imprese-pmi-2026.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Parità di prezzo dei modelli AI: tutti a 20 euro al mese
# ============================================================

doc = new_doc()
testata(doc, "Luglio 2026", "Strumenti e Modelli AI")
titolo(
    doc,
    "Tutti a 20 euro. Il costo ha smesso di essere l’alibi.",
    "Nel luglio 2026 ChatGPT, Claude e Gemini costano circa 20 euro al mese. "
    "La parità di prezzo cambia la domanda: non è più quanto costa "
    "ma chi sa usarlo. Per i professionisti, il vantaggio è passato "
    "dall’accesso alla competenza.",
    "A cura della Redazione Ratio  •  31 luglio 2026"
)

para(doc,
    "Nel 1900 George Eastman lancì la Kodak Brownie a un dollaro: per la prima "
    "volta nella storia chiunque poteva fare fotografie, senza bisogno "
    "dell’attrezzatura professionale. I grandi fotografi del Novecento non "
    "smisero di essere grandi fotografi. La fotocamera diventò universale; "
    "la capacità di vedere rimase rara. Nel luglio 2026, i piani professionali "
    "di ChatGPT, Claude e Gemini costano tutti intorno ai venti euro mensili. "
    "Lo strumento è diventato universale. Il fattore che fa la differenza "
    "è rimasto altrove."
)

para(doc,
    "La convergenza di prezzo tra i tre principali modelli AI è un fenomeno "
    "che dodici mesi fa sarebbe sembrato improbabile. La competizione tra Anthropic, "
    "OpenAI e Google ha compresso i margini sui piani consumer fino al punto in cui "
    "il costo di accesso alla fascia premium si è sostanzialmente allineato. "
    "Chi avesse aspettato che il prezzo dell’AI scendesse abbastanza da "
    "giustificarne l’adozione ha la risposta: è già successo. "
    "Il costo mensile di uno strumento AI professionale equivale oggi a quello di "
    "un abbonamento a una rivista specializzata, o a una cena con un cliente."
)

heading(doc, "Quando il prezzo smette di essere una barriera")

para(doc,
    "La parità di prezzo non è una notizia neutra: è un cambio di "
    "struttura del mercato. Finché l’AI costava significativamente di più, "
    "i professionisti che la adottavano erano quelli che avevano già giustificato "
    "l’investimento attraverso un caso d’uso preciso e misurabile. A venti "
    "euro al mese, la soglia si è abbassata al punto in cui molti la adottano "
    "senza aver chiarito cosa vogliono ottenere. Il risultato è che il mercato "
    "si divide sempre più chiaramente in due fasce: chi usa l’AI come "
    "strumento di lavoro con una metodologia, ottenendo risultati misurabili, e chi "
    "la usa come assistente generico per compiti che non richiedono elaborazione "
    "sofisticata. Tra i due gruppi, la distanza di risultati cresce, anche se il "
    "costo dello strumento è identico."
)

heading(doc, "La competenza che lo strumento non include")

para(doc,
    "Il 37% delle aziende italiane ha attivato percorsi di formazione sulle "
    "funzionalità e sui rischi degli strumenti AI adottati. Il 63% usa gli "
    "stessi strumenti senza formazione sistematica. Letto insieme alla parità "
    "di prezzo, questo dato produce una considerazione precisa: la barriera "
    "all’adozione dell’AI professionale non è più economica. "
    "È la capacità di formulare domande utili, di valutare criticamente "
    "le risposte, di integrare lo strumento in un processo di lavoro anziché "
    "aggiungerlo come passaggio extra. La competenza che distingue chi ottiene "
    "risultati da chi ottiene risposte è quella che potremmo chiamare, con "
    "una certa brutalità, saper lavorare con l’AI."
)

heading(doc, "Cosa cambiano i venti euro bene usati")

para(doc,
    "Un abbonamento da venti euro mensili a Claude, ChatGPT o Gemini produce "
    "risultati molto diversi a seconda di come viene usato. Per un commercialista "
    "che analizza contratti, lo strumento cambia qualitativamente se sa strutturare "
    "il prompt per ottenere una revisione critica anziché un riassunto. "
    "Per un consulente aziendale, fa differenza sapere quale modello è più "
    "adatto a ragionare su scenari complessi e quale ottimale per generare testi "
    "strutturati in tempi brevi. Per un avvocato, la differenza è tra uno "
    "strumento che accelera la ricerca di precedenti e uno che fornisce precedenti "
    "plausibili ma inesatti. La parità di prezzo ha reso questi venti euro "
    "accessibili a tutti. La competenza per usarli bene rimane distribuita "
    "in modo molto diseguale."
)

para(doc,
    "Eastman vendé milioni di Brownie. I grandi fotografi del Novecento "
    "le acquistarono anch’essi. E poi fecero cose molto diverse "
    "con lo stesso strumento."
)

riferimenti(doc, [
    "Jenova.ai -- 'GPT vs Claude vs Gemini: Confronto Completo dei Modelli AI "
    "per il 2026'",
    "Bleap.finance -- 'Claude vs GPT vs Gemini: Confronto tra i Migliori Modelli "
    "di IA del 2026'",
    "FelloAI.com -- 'Best AI Models in July 2026: ChatGPT, Claude, Gemini & Grok'",
    "SurePrompts.com -- 'Best AI Model in 2026: ChatGPT vs Claude vs Gemini Compared'",
    "SbirciaLaNotizia.it -- 'AI, 18,4 milioni per azienda nel 2026: budget in crescita "
    "e controlli in ritardo' (19 luglio 2026)",
])
doc.save(BASE + "2026-07-31_parita-prezzo-modelli-ai-20-euro-competenza.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
