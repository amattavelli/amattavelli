"""
Quattro articoli Ratio -- 14 agosto 2026

1. 2026-08-14_gpt56-prezzi-tagliati-scelta-non-semplificata.docx
2. 2026-08-14_ai-literacy-obbligo-scattato-chi-puo-dimostrarlo.docx
3. 2026-08-14_governance-ai-chi-risponde-quando-sbaglia-azienda.docx
4. 2026-08-14_sistemi-ai-produzione-senza-documentazione-pmi.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# I prezzi dei modelli AI si tagliano: la scelta si complica
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Strumenti e Modelli AI")
titolo(
    doc,
    "I prezzi dell'AI scendono. La scelta giusta non si semplifica.",
    "Il 30 luglio OpenAI ha tagliato dell'80% il costo di GPT-5.6 Luna e del 20% quello "
    "di Terra. In due anni i costi API sono calati di oltre il 90%. "
    "Con sette modelli di fascia alta disponibili simultaneamente, "
    "il problema non e' piu' il prezzo: e' sapere cosa scegliere.",
    "A cura della Redazione Ratio  •  14 agosto 2026"
)

para(doc,
    "Nei supermercati degli anni Novanta, la pasta si comprava in tre formati: "
    "la marca nazionale, la marca del distributore e quella di fascia bassa. "
    "Tre opzioni, criterio di scelta semplice. Oggi uno scaffale di pasta "
    "conta trenta referenze, con distinzioni per trafilatura al bronzo, "
    "semola integrale, pasta biologica, formati speciali, linee senza glutine "
    "e prezzi che vanno da 0,80 a 6 euro al chilo. Il prodotto e' lo stesso. "
    "La scelta e' diventata un lavoro. Il mercato dei modelli AI sta percorrendo "
    "la stessa traiettoria, con una velocita' che il settore alimentare "
    "ha impiegato trent'anni a raggiungere."
)

para(doc,
    "Il 30 luglio 2026, OpenAI ha tagliato i prezzi API di due dei tre livelli "
    "della lineup GPT-5.6: Luna ha subito una riduzione dell'80%, passando da 1 dollaro "
    "per milione di token in input a 0,20 dollari; Terra e' sceso del 20%, "
    "da 2,50 a 2 dollari in input. Sol, il modello di punta, rimane invariato "
    "a 5 dollari in input e 30 in output. Il taglio e' arrivato a meno di tre "
    "settimane dal lancio della lineup, il 9 luglio. Contemporaneamente, Anthropic "
    "ha confermato che il prezzo promozionale di Sonnet 5 (2 dollari in input) "
    "resterà attivo fino al 31 agosto, con un incremento poi a 3 dollari. "
    "In questo contesto, il mercato offre oggi almeno sette modelli di fascia alta "
    "da provider diversi, con caratteristiche sovrapposte e prezzi sempre piu' vicini."
)

heading(doc, "Dove e' finito il vantaggio del prezzo")

para(doc,
    "Per due anni, il principale argomento a favore dei modelli piu' economici "
    "e' stato il costo per token: chi faceva grandi volumi aveva un incentivo "
    "economico chiaro a usare modelli meno potenti per i compiti di routine "
    "e riservare i piu' costosi ai task complessi. Con Luna a 0,20 dollari "
    "per milione di token, quella logica si comprime ma non scompare. "
    "Il differenziale tra Luna e Sol rimane di 25 volte in input e 25 volte in output: "
    "per uno studio professionale che usa i modelli in modo diretto, "
    "non tramite API, questa differenza e' quasi irrilevante, "
    "perche' il costo marginale per interazione e' minimo. "
    "Per chi integra i modelli in flussi automatizzati ad alto volume, "
    "invece, il risparmio e' ancora reale e la scelta tra livelli ha senso economico."
)

heading(doc, "Tre modelli dentro la stessa interfaccia")

para(doc,
    "La particolarita' di GPT-5.6 e' che i tre livelli, Sol, Terra e Luna, "
    "sono accessibili dalla stessa interfaccia ChatGPT senza richiedere "
    "la conoscenza del nome tecnico del modello: l'utente sceglie tra "
    "'risposta rapida' e 'risposta approfondita' e il sistema seleziona il modello "
    "adeguato automaticamente nei piani standard. OpenAI dichiara Luna ottimale "
    "per compiti delimitati e ripetuti, Terra come base per il lavoro quotidiano "
    "e Sol per i casi in cui un errore ha un costo alto. "
    "Anthropic usa una logica simile con Sonnet 5 e Opus 5. "
    "Il risultato pratico e' che per la maggior parte degli utenti "
    "non tecnici, la scelta avviene senza piena consapevolezza di quale modello "
    "stia effettivamente elaborando la richiesta, e quindi senza una valutazione "
    "consapevole del rapporto tra qualita' attesa e costo."
)

heading(doc, "La competenza che conta di piu'")

para(doc,
    "Per uno studio professionale o un'impresa che usa AI in modo strutturato, "
    "il problema del 2026 non e' piu' il costo dei modelli: e' la capacita' "
    "di abbinare il modello al compito. Usare Sol per redigere un promemoria "
    "interno e' uno spreco; usare Luna per analizzare un contratto con clausole "
    "potenzialmente conflittuali e' un rischio. La moltiplicazione delle opzioni "
    "aumenta la probabilita' di scelte inadeguate in entrambe le direzioni: "
    "spese eccessive per compiti standard, qualita' insufficiente per compiti critici. "
    "Il taglio dei prezzi rende il mercato piu' accessibile, ma non risolve "
    "il problema centrale, che e' sapere dove la qualita' del modello "
    "ha un effetto diretto sul risultato. Quella valutazione richiede qualcuno "
    "che conosca sia il processo sia lo strumento."
)

para(doc,
    "Nello scaffale della pasta, il prezzo piu' basso non e' quasi mai la scelta giusta. "
    "Ma nemmeno il piu' alto lo e' per forza. "
    "La competenza sta nel sapere dove la trafilatura al bronzo cambia davvero il piatto."
)

riferimenti(doc, [
    "Hardware Upgrade – 'OpenAI taglia i prezzi di due modelli GPT-5.6: Luna costa l'80% in meno dopo 3 settimane dal lancio' (30 luglio 2026)",
    "Geek Club Italia – 'GPT-5.6 spiegato semplice: cosa sono Sol, Terra e Luna e cosa cambia davvero' (1 agosto 2026)",
    "QualeAI.it – 'GPT-5.6 prezzi 2026: Luna, Terra o Sol?' (2026)",
    "Bleap.finance – 'Recensione Claude AI 2026: Funzionalita', Prezzi, Vantaggi e Alternative' (2026)",
    "FelloAI.com – 'Best AI Models in August 2026: ChatGPT, Claude, Gemini & Grok' (agosto 2026)",
    "Anthropic – 'Introducing Claude Sonnet 5' (30 giugno 2026) e 'Introducing Claude Opus 5' (24 luglio 2026)",
])
doc.save(BASE + "2026-08-14_gpt56-prezzi-tagliati-scelta-non-semplificata.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# AI literacy: l'obbligo e' scattato. Chi puo' dimostrarlo?
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Normativa AI e Professionisti")
titolo(
    doc,
    "L'obbligo di AI literacy e' scattato. Pochi possono dimostrarlo.",
    "Dal 3 agosto 2026 le autorita' di vigilanza italiane possono aprire procedimenti "
    "e richiedere documentazione sulla formazione AI del personale. "
    "L'obbligo di legge era gia' attivo da febbraio 2025. "
    "La maggior parte delle aziende non ha carta da mostrare.",
    "A cura della Redazione Ratio  •  14 agosto 2026"
)

para(doc,
    "Nel diritto tributario esiste il concetto di obbligo documentale: "
    "non basta aver fatto la cosa giusta, bisogna poter dimostrare "
    "di averla fatta. Un'impresa che ha detratto correttamente l'IVA "
    "su acquisti legittimi ma non conserva le fatture e' esposta "
    "alla stessa contestazione di chi ha detratto l'IVA su acquisti inesistenti. "
    "La prova e' parte integrante dell'adempimento, non un accessorio. "
    "L'articolo 4 dell'AI Act, che impone agli utilizzatori di sistemi AI "
    "di garantire un livello adeguato di competenze al proprio personale, "
    "funziona con la stessa logica: l'obbligo esiste dal 2 febbraio 2025, "
    "ma senza documentazione e' come se non fosse mai stato rispettato."
)

para(doc,
    "Dal 3 agosto 2026 le autorita' di vigilanza italiane, ACN per la sicurezza "
    "e AgID per la notifica, hanno formalmente avviato le attivita' ispettive "
    "previste dall'AI Act. Possono aprire procedimenti, richiedere documentazione "
    "e comminare sanzioni. Per le aziende che usano sistemi AI nei propri processi, "
    "questo significa che il fascicolo della formazione AI del personale non e' piu' "
    "un documento futuro da preparare: e' un documento che avrebbe dovuto esistere "
    "gia' diciotto mesi fa, e che oggi potrebbe essere richiesto in caso "
    "di ispezione o di incidente."
)

heading(doc, "Cosa dice esattamente la norma")

para(doc,
    "L'articolo 4 del Regolamento UE 2024/1689 stabilisce che 'i fornitori e i deployer "
    "di sistemi di IA adottano misure per garantire, nella misura del possibile, "
    "un livello sufficiente di alfabetizzazione in materia di IA del loro personale "
    "e di altre persone che si occupano del funzionamento e dell'uso di sistemi di IA "
    "per loro conto'. La norma non impone un corso standard ne' una certificazione "
    "obbligatoria: impone che la formazione sia proporzionata al ruolo, documentata "
    "e aggiornata. Chi usa ChatGPT per redigere comunicazioni ai clienti "
    "deve aver ricevuto formazione su cosa il modello puo' e non puo' fare, "
    "sui rischi di allucinazioni, sulle politiche di privacy applicabili. "
    "Chi usa sistemi AI per selezionare candidati deve aver ricevuto formazione "
    "sul rischio di bias algoritmico. La specificita' del ruolo e' essenziale."
)

heading(doc, "Cosa serve in caso di ispezione")

para(doc,
    "Le indicazioni operative che emergono dalle linee guida applicative "
    "individuano cinque elementi documentali minimi: un programma formativo scritto "
    "con obiettivi e contenuti, la lista dei partecipanti con date e ore, "
    "i materiali usati o un riferimento al fornitore, una valutazione "
    "della comprensione anche semplice, e la data dell'aggiornamento previsto. "
    "Non e' un dossier complesso: e' un file che dimostra che qualcuno, "
    "in azienda, ha pensato a chi usa l'AI, per fare cosa, e a cosa dovrebbe sapere. "
    "Il problema e' che la maggior parte delle aziende italiane ha adottato "
    "strumenti AI senza un processo formativo strutturato: li ha messi a disposizione, "
    "ha lasciato che i dipendenti li usassero, e non ha documentato nulla."
)

heading(doc, "Il ruolo del professionista esterno")

para(doc,
    "Per le PMI, la figura piu' vicina a una funzione di compliance AI "
    "e' spesso il consulente del lavoro, il commercialista o il consulente "
    "di management che segue l'azienda. Questi professionisti sono in una "
    "posizione utile per segnalare la lacuna e aiutare a colmarla: "
    "non necessariamente erogando la formazione in prima persona, "
    "ma strutturando il processo, identificando i ruoli coinvolti "
    "e definendo il formato del fascicolo documentale. "
    "La Legge 132/2025, che ha recepito l'AI Act nell'ordinamento italiano, "
    "non esclude la responsabilita' del vertice aziendale in caso "
    "di mancato adempimento: chi ha delegato l'AI ai tecnici "
    "senza occuparsi della formazione risponde comunque della lacuna."
)

para(doc,
    "In molti studi fiscali, la conservazione sostitutiva delle fatture "
    "e' diventata una routine solo dopo che qualcuno aveva perso una causa "
    "per mancanza di documento. "
    "Aspettare la prima ispezione AI per costruire il fascicolo formativo "
    "e' la stessa logica, con lo stesso finale prevedibile."
)

riferimenti(doc, [
    "Archimedia.it – 'AI Act: scadenza 2 agosto 2026 per la formazione AI del personale' (2026)",
    "Antonio Sinibaldi – 'AI literacy per aziende: obbligo formativo AI Act e scadenza agosto 2026' (2026)",
    "Legal for Digital – 'AI literacy, obbligo formazione AI Act' (2026)",
    "Blog Unique.it – 'Formare i dipendenti sull'AI: cosa dice la legge e come farlo senza complicare tutto' (2026)",
    "Gianluca Girardi – 'Formazione AI obbligatoria in azienda: cosa prevede l'AI Act e come adeguarsi' (2026)",
    "Regolamento UE 2024/1689 (AI Act), articolo 4",
    "Legge 23 settembre 2025 n. 132 (recepimento AI Act in Italia)",
])
doc.save(BASE + "2026-08-14_ai-literacy-obbligo-scattato-chi-puo-dimostrarlo.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# Chi risponde quando l'AI sbaglia in azienda?
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Governance AI e Impresa")
titolo(
    doc,
    "Chi risponde quando l'AI sbaglia in azienda? Il 91% non lo sa.",
    "Il 91% delle grandi imprese italiane non ha un modello strutturato "
    "di governance dell'AI. Il 50% di quelle che gia' usano sistemi AI "
    "non ha definito modelli di responsabilita'. "
    "Con le sanzioni dell'AI Act operative, la domanda non e' piu' teorica.",
    "A cura della Redazione Ratio  •  14 agosto 2026"
)

para(doc,
    "In molte societa' per azioni, la delega alle funzioni e' uno strumento "
    "di governo rodato: il consiglio delega all'amministratore delegato, "
    "l'amministratore delegato delega ai direttori di funzione, "
    "i direttori delegano ai responsabili operativi. La catena funziona "
    "perche' ogni passaggio e' documentato, i confini della delega sono definiti "
    "e la responsabilita' risale in caso di superamento dei limiti. "
    "Nelle stesse aziende, l'AI e' spesso entrata senza che nessuno "
    "abbia costruito una catena analoga: qualcuno ha abilitato uno strumento, "
    "qualcun altro ha cominciato a usarlo, e la domanda 'chi risponde "
    "se qualcosa va storto' non e' mai stata posta."
)

para(doc,
    "I dati dell'Osservatorio Permanente sull'AI del 2026 fotografano "
    "una situazione che molti avevano intuito ma che ora ha numeri precisi: "
    "il 91% delle grandi imprese italiane dichiara di non disporre "
    "di un modello strutturato di governance dell'AI; il 50% delle organizzazioni "
    "che gia' utilizzano sistemi di intelligenza artificiale e' ancora in uno stadio "
    "iniziale nella definizione dei modelli di responsabilita'. "
    "Sono numeri che riguardano le grandi imprese, quelle con piu' risorse "
    "e strutture piu' articolate. Nelle PMI, la percentuale e' presumibilmente piu' alta."
)

heading(doc, "Cosa manca quando manca la governance")

para(doc,
    "La governance dell'AI in un'azienda e' un sistema di risposte "
    "a quattro domande operative: chi approva l'introduzione di un nuovo strumento AI, "
    "chi ha il potere di limitarne o sospenderne l'uso, "
    "chi supervisiona le decisioni prodotte dal sistema, "
    "e chi risponde verso l'esterno in caso di danno o contestazione. "
    "In assenza di queste risposte, i rischi non scompaiono: "
    "si distribuiscono in modo disordinato tra chi ha introdotto lo strumento, "
    "chi lo usa quotidianamente e chi firma i documenti prodotti anche "
    "con l'ausilio del sistema. Per un commercialista, un consulente del lavoro "
    "o un'impresa con clienti, questa distribuzione disordinata e' un problema "
    "concreto: la responsabilita' professionale e la responsabilita' contrattuale "
    "non ammettono vuoti di attribuzione."
)

heading(doc, "Cosa richiede l'AI Act")

para(doc,
    "Dal 2 agosto 2026, per i sistemi AI classificati ad alto rischio "
    "secondo l'Allegato III del Regolamento, le aziende devono avere "
    "pronti e documentati: la classificazione del rischio del sistema, "
    "le istruzioni d'uso seguite, i registri di supervisione umana, "
    "la valutazione di impatto e la formazione del personale coinvolto. "
    "Per i sistemi non classificati ad alto rischio ma comunque in uso "
    "in processi rilevanti, la Legge 132/2025 e le linee guida ACN "
    "indicano che il vertice aziendale risponde delle modalita' "
    "concrete di utilizzo e delle istruzioni impartite ai dipendenti. "
    "La responsabilita' non si trasferisce al fornitore del modello "
    "ne' al sistema stesso: rimane dove e' sempre rimasta, "
    "in capo a chi ha deciso di usarlo e a chi ne ha autorizzato l'uso."
)

heading(doc, "Da dove si comincia")

para(doc,
    "Per uno studio professionale o una PMI che voglia costruire "
    "un modello di governance AI senza una struttura enterprise, "
    "il punto di partenza e' un censimento: quali strumenti AI sono in uso, "
    "chi li usa, per quali processi, con quale frequenza. "
    "Da questo censimento si puo' derivare una mappa minima "
    "di responsabilita': chi approva l'uso di ciascuno strumento, "
    "chi verifica gli output prima che diventino definitivi, "
    "chi risponde verso il cliente se l'output e' errato. "
    "Il documento non deve essere complesso: deve essere scritto, "
    "aggiornato e noto alle persone coinvolte. "
    "Il 91% delle grandi imprese che non ha questo documento "
    "ha una struttura molto piu' elaborata della maggior parte "
    "degli studi professionali, eppure non ha trovato il tempo di farlo. "
    "Questo non e' un buon precedente da seguire."
)

para(doc,
    "La delega funziona quando e' scritta, firmata e nei cassetti giusti. "
    "L'AI che nessuno ha formalmente autorizzato opera fuori da qualsiasi cassetto, "
    "e quando sbaglia non c'e' un cassetto da aprire per trovare chi risponde."
)

riferimenti(doc, [
    "Digitalic.it – 'Rapporto annuale dell'Osservatorio permanente sull'AI 2026' (2026)",
    "Agenda Digitale – 'Governance dell'AI nelle imprese: la strategia resta umana' (2026)",
    "ChannelTech.it – 'L'IA non e' solo tecnologia, ma una questione di governo d'impresa' (29 giugno 2026)",
    "Dgroove.it – 'AI Act 2026 Italia: guida pratica per aziende e IT Manager' (2026)",
    "Studio Legale CLF – 'Governance dell'IA in azienda: chi risponde?' (2026)",
    "FiscoeTasse.com – 'Regolamento intelligenza artificiale in Italia 2026: approvati i decreti' (2026)",
    "Regolamento UE 2024/1689 (AI Act), Allegato III e artt. 26-27",
    "Legge 23 settembre 2025 n. 132",
])
doc.save(BASE + "2026-08-14_governance-ai-chi-risponde-quando-sbaglia-azienda.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Sistemi AI in produzione senza documentazione: il rischio
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Compliance AI e PMI")
titolo(
    doc,
    "Sistemi AI in produzione, documentazione a zero: il rischio che le PMI ignorano.",
    "Audit di compliance AI Act su PMI italiane tra 20 e 300 dipendenti "
    "mostrano un pattern ricorrente: sistemi gia' in produzione classificabili "
    "ad alto rischio, nessuna documentazione, nessun registro, nessuna valutazione di impatto. "
    "Le polizze cyber non coprono ancora questo rischio.",
    "A cura della Redazione Ratio  •  14 agosto 2026"
)

para(doc,
    "Negli anni Novanta, molte piccole imprese italiane usavano "
    "software gestionali non licenziati. Non perche' fossero disoneste, "
    "ma perche' il mercato aveva corso piu' veloce della consapevolezza normativa "
    "e qualcuno aveva installato il programma senza che nessuno si chiedesse "
    "se ci fosse una licenza valida. Quando arrivo' la prima ondata "
    "di verifiche della Guardia di Finanza, il problema non era la malafede: "
    "era che nessuno aveva mai fatto la domanda giusta. "
    "Con i sistemi AI, nel 2026, si sta seguendo lo stesso schema: "
    "lo strumento e' gia' in produzione, nessuno ha chiesto se rientra "
    "tra i sistemi ad alto rischio, e la documentazione non esiste."
)

para(doc,
    "I primi audit di compliance AI Act condotti su PMI italiane con un organico "
    "tra venti e trecento dipendenti mostrano un pattern consistente. "
    "Le aziende hanno adottato strumenti di AI in processi che l'Allegato III "
    "del Regolamento UE 2024/1689 classifica come ad alto rischio: "
    "sistemi usati per valutare candidati in fase di assunzione, "
    "strumenti che contribuiscono a decisioni di credito o solvibilita', "
    "software che classificano clienti per profilo di rischio. "
    "In quasi tutti i casi rilevati: zero documentazione tecnica, "
    "zero registro delle operazioni, zero valutazione di impatto, "
    "zero procedura di supervisione umana documentata. "
    "Il sistema e' operativo. Gli obblighi normativi non sono stati soddisfatti."
)

heading(doc, "Cosa dice l'Allegato III")

para(doc,
    "L'Allegato III dell'AI Act elenca le categorie di sistemi ad alto rischio "
    "per i quali gli obblighi di documentazione, registrazione e supervisione "
    "umana sono i piu' stringenti. Rientrano in questa categoria, tra gli altri: "
    "i sistemi usati per la valutazione e selezione di candidati al lavoro "
    "o per promozioni e licenziamenti, i sistemi che determinano l'accesso "
    "o il punteggio creditizio, quelli usati per la classificazione di clienti "
    "in base a profili di rischio finanziario o assicurativo, "
    "i sistemi che influenzano decisioni in materia di prestazioni sociali. "
    "Molte PMI italiane usano strumenti che rientrano in queste categorie "
    "senza saperlo, perche' i fornitori non hanno comunicato la classificazione "
    "o perche' lo strumento e' stato introdotto come 'assistente' "
    "senza che nessuno ne abbia analizzato la funzione concreta nel processo."
)

heading(doc, "Il buco nelle polizze cyber")

para(doc,
    "A complicare il quadro, la maggior parte delle polizze di responsabilita' "
    "civile e cyber assicurazione disponibili sul mercato italiano "
    "non copre esplicitamente i rischi legati all'uso di agenti AI autonomi "
    "o di sistemi classificati ad alto rischio dall'AI Act. "
    "Alcune compagnie assicurative stanno introducendo prodotti dedicati, "
    "ma i premi sono ancora alti per via dell'assenza di dati storici "
    "sui sinistri correlati all'AI. Per una PMI che subisce un danno "
    "a terzi derivante da una decisione parzialmente assistita da un sistema AI "
    "non documentato, la polizza cyber standard potrebbe non rispondere, "
    "e la responsabilita' ricadrebbe direttamente sul soggetto giuridico."
)

heading(doc, "Come impostare un censimento minimo")

para(doc,
    "Il punto di partenza pratico, per una PMI o uno studio che voglia "
    "verificare la propria posizione, e' un censimento degli strumenti AI in uso. "
    "Non tutti i sistemi AI sono ad alto rischio: l'obiettivo del censimento "
    "e' identificare quelli che influenzano decisioni rilevanti su persone "
    "o su risorse significative. Per ciascuno di questi, la domanda e' "
    "se il fornitore ha fornito la classificazione di rischio, se esiste "
    "una procedura documentata di supervisione umana sugli output, "
    "e se il personale che usa il sistema ha ricevuto formazione adeguata. "
    "Nella maggior parte dei casi, la risposta a tutte e tre le domande sara' negativa. "
    "Saperlo e' il primo passo per non trovarsi nella stessa posizione "
    "di chi installava software senza licenza trent'anni fa: "
    "non in malafede, ma con lo stesso risultato."
)

para(doc,
    "La Guardia di Finanza degli anni Novanta non cercava malintenzionati: "
    "cercava chi non aveva fatto le domande giuste al momento giusto. "
    "Le autorita' di vigilanza AI del 2026 stanno iniziando lo stesso giro."
)

riferimenti(doc, [
    "AI4Business.it – 'AI Agentica: rischi operativi e nuove strategie di governance' (2026)",
    "Agenda Digitale – 'Agenti AI in azienda: i 10 controlli per proteggersi dagli attacchi' (2026)",
    "Manager.it – 'Sicurezza degli Agenti AI in Azienda: la Guida Completa 2026' (2026)",
    "Blog Unique.it – 'AI Act italiano: cosa devono sapere le PMI nel 2026 per non essere fuori norma' (2026)",
    "TuxWeb.it – 'AI Act 2026 per PMI italiane' (2026)",
    "AscenSys.it – 'AI Act 2026 PMI: obblighi, scadenze e sanzioni' (2026)",
    "Regolamento UE 2024/1689 (AI Act), Allegato III",
    "Legge 23 settembre 2025 n. 132; D.Lgs. attuativi approvati dal CdM il 10 giugno 2026",
])
doc.save(BASE + "2026-08-14_sistemi-ai-produzione-senza-documentazione-pmi.docx")
print("Salvato: articolo 4")


print("\nTutti e 4 gli articoli generati in:", BASE)
