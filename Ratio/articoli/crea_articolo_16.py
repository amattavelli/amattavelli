"""
Articolo 16: Claude Managed Agents — gli agenti AI entrano in azienda
Ratio/articoli/2026-04-10_claude-managed-agents-aziende-italiane.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-10_claude-managed-agents-aziende-italiane.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_paragraph(doc, text, size=11, italic=False, color=None,
                  space_after=8, space_before=0,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


# ---- TESTATA ----
p_rivista = doc.add_paragraph()
p_rivista.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = p_rivista.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
run_r.font.name = 'Calibri'
run_r.font.size = Pt(9)
run_r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run_r.font.bold = True
run_r.font.all_caps = True

add_separator(doc)

p_data = doc.add_paragraph()
p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_d = p_data.add_run("Aprile 2026  |  Strumenti e Innovazione")
run_d.font.name = 'Calibri'
run_d.font.size = Pt(8.5)
run_d.italic = True
run_d.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_after = Pt(6)
run_t = p_title.add_run(
    "Claude Managed Agents: Anthropic porta gli agenti AI direttamente in azienda"
)
run_t.font.name = 'Calibri'
run_t.font.size = Pt(24)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_after = Pt(10)
run_s = p_sub.add_run(
    "Con il lancio dei Managed Agents, Anthropic abbassa la soglia tecnica per chi vuole automatizzare "
    "processi aziendali con l\u2019AI. Non serve pi\u00f9 un team di sviluppo. Basta sapere cosa si vuole "
    "far fare all\u2019agente."
)
run_s.font.name = 'Calibri'
run_s.font.size = Pt(13)
run_s.italic = True
run_s.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_separator(doc)

p_autore = doc.add_paragraph()
p_autore.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_autore.paragraph_format.space_after = Pt(16)
run_a = p_autore.add_run("A cura della Redazione Ratio  \u2022  10 aprile 2026")
run_a.font.name = 'Calibri'
run_a.font.size = Pt(9)
run_a.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- CORPO ----

add_paragraph(doc,
    "Il 9 aprile 2026, Anthropic ha annunciato il lancio dei Claude Managed Agents, una piattaforma "
    "che consente alle aziende di creare, distribuire e gestire agenti AI autonomi senza dover scrivere "
    "codice. \u00c8 una data che vale la pena segnare, non perch\u00e9 sia la prima soluzione di questo "
    "tipo sul mercato, ma perch\u00e9 segna un passaggio importante: la costruzione di agenti AI autonomi "
    "smette di essere un progetto da software house e diventa uno strumento accessibile a un responsabile "
    "di processo, a un controller, a un office manager con un obiettivo chiaro e la pazienza di configurare "
    "qualcosa di nuovo."
)

add_paragraph(doc,
    "Un agente AI, nella sua forma pi\u00f9 semplice, \u00e8 un sistema che riceve un obiettivo e lo "
    "persegue in autonomia: raccoglie informazioni, le elabora, prende decisioni intermedie, produce "
    "un output o avvia un\u2019azione. Non risponde a una domanda: esegue un compito. La differenza con "
    "un chatbot \u00e8 sostanziale. Un agente pu\u00f2 connettersi a un gestionale, leggere i dati di "
    "un cliente, confrontarli con una scadenza fiscale, compilare un documento e inviarlo per la revisione "
    "finale al professionista \u2014 tutto in sequenza, senza intervento umano nelle fasi intermedie."
)

add_heading2(doc, "Come funzionano i Managed Agents in pratica")

add_paragraph(doc,
    "La proposta di Anthropic si articola su tre livelli. Il primo \u00e8 la creazione dell\u2019agente: "
    "l\u2019utente definisce in linguaggio naturale l\u2019obiettivo dell\u2019agente, i dati a cui "
    "pu\u00f2 accedere, i sistemi con cui pu\u00f2 interagire e i limiti entro cui deve operare. "
    "Il secondo livello \u00e8 la distribuzione: l\u2019agente viene reso disponibile a uno o pi\u00f9 "
    "utenti nell\u2019organizzazione, con controlli di accesso e un log delle azioni eseguite. "
    "Il terzo livello \u00e8 la gestione: la piattaforma tiene traccia di ogni operazione compiuta "
    "dall\u2019agente, consente di modificarne le istruzioni e di intervenire in caso di comportamenti "
    "anomali."
)

add_paragraph(doc,
    "In termini pratici, questo significa che uno studio professionale pu\u00f2 configurare un agente "
    "dedicato al monitoraggio delle scadenze: l\u2019agente accede al calendario dei clienti, verifica "
    "la documentazione caricata sul portale, segnala le situazioni in cui mancano elementi e predispone "
    "una bozza di comunicazione al cliente. Un altro agente pu\u00f2 occuparsi della riconciliazione dei "
    "movimenti bancari, confrontando estratti conto e partitari e segnalando le discrepanze che richiedono "
    "l\u2019occhio del professionista. Attivit\u00e0 che oggi costano ore di lavoro a basso valore aggiunto."
)

add_heading2(doc, "Il tema della fiducia: quando lasciare agire l\u2019agente")

add_paragraph(doc,
    "La domanda che molti professionisti si pongono \u00e8 legittima: fino a che punto ci si pu\u00f2 "
    "fidare di un agente che opera in autonomia? La risposta onesta \u00e8 che dipende dal compito. "
    "Per attivit\u00e0 strutturate, ripetitive e con output verificabili \u2014 come il monitoraggio "
    "delle scadenze o la riconciliazione dei dati \u2014 il livello di affidabilit\u00e0 dei modelli "
    "attuali \u00e8 gi\u00e0 sufficiente per ridurre significativamente il lavoro umano, a patto di "
    "mantenere un punto di controllo finale. Per attivit\u00e0 che richiedono giudizio professionale "
    "\u2014 una valutazione fiscale, una raccomandazione strategica \u2014 l\u2019agente pu\u00f2 "
    "preparare il materiale, ma la decisione resta al professionista."
)

add_paragraph(doc,
    "Anthropic ha costruito nei Managed Agents un sistema di audit trail obbligatorio: ogni azione "
    "compiuta dall\u2019agente viene registrata con timestamp, fonte dei dati usati e motivazione "
    "sintetica della scelta. Questo non \u00e8 solo un requisito di compliance \u2014 \u00e8 anche lo "
    "strumento che permette al professionista di capire cosa ha fatto l\u2019agente e perch\u00e9, "
    "costruendo progressivamente la fiducia necessaria per ampliare l\u2019autonomia del sistema."
)

add_heading2(doc, "Il mercato italiano: dove siamo")

add_paragraph(doc,
    "In Italia, secondo i dati dell\u2019Osservatorio Intelligenza Artificiale del Politecnico di Milano, "
    "il mercato dell\u2019AI agentica vale oggi circa 280 milioni di euro e cresce a un ritmo del 67% "
    "annuo. La componente enterprise \u00e8 predominante, ma il segmento PMI inizia a muoversi "
    "concretamente, soprattutto nei settori manifatturiero, logistico e dei servizi professionali. "
    "I fornitori di software gestionali italiani \u2014 da TeamSystem a Zucchetti \u2014 stanno "
    "integrando capacit\u00e0 agentiche nei loro prodotti, il che significa che molte aziende "
    "troveranno queste funzionalit\u00e0 gi\u00e0 dentro i sistemi che usano quotidianamente, "
    "senza dover scegliere una piattaforma separata."
)

add_paragraph(doc,
    "La soglia di ingresso si sta abbassando rapidamente. Un anno fa, costruire un agente autonomo "
    "richiedeva competenze di programmazione e accesso alle API. Oggi, piattaforme come Claude "
    "Managed Agents, Microsoft Copilot Studio e Make rendono possibile configurare automazioni "
    "agentiche senza scrivere una riga di codice. Il passo successivo non \u00e8 tecnologico: "
    "\u00e8 organizzativo. Richiede di identificare i processi ripetitivi ad alto volume, "
    "definire le regole di ingaggio per l\u2019agente e decidere dove mantenere la supervisione "
    "umana. Questo \u00e8 esattamente il tipo di lavoro in cui il consulente aziendale o il "
    "commercialista pu\u00f2 portare valore reale al proprio cliente."
)

add_heading2(doc, "La domanda che vale la pena farsi")

add_paragraph(doc,
    "La domanda non \u00e8 \u201cpossiamo permetterci un agente AI?\u201d \u2014 i costi sono "
    "scesi al punto che questa non \u00e8 pi\u00f9 la barriera. La domanda \u00e8 \u201cquali sono "
    "i tre processi nel nostro studio o nella nostra azienda che consumano pi\u00f9 tempo a basso "
    "valore aggiunto?\u201d Qualunque sia la risposta, \u00e8 l\u00ec che vale la pena cominciare. "
    "Non con una visione strategica triennale sull\u2019AI, ma con un agente che fa una cosa, "
    "la fa bene, e libera tempo per il lavoro che richiede davvero un essere umano."
)

add_separator(doc)

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(10)
run_n = p_note.add_run("Fonti e riferimenti")
run_n.font.name = 'Calibri'
run_n.font.size = Pt(9)
run_n.font.bold = True
run_n.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Anthropic \u2014 Lancio Claude Managed Agents, comunicato stampa, 9 aprile 2026",
    "TecnoAndroid \u2014 Claude Managed Agents: Anthropic lancia gli agenti IA per le aziende (aprile 2026)",
    "Osservatorio Intelligenza Artificiale \u2014 Politecnico di Milano, Report 2026",
    "AI4Business \u2014 AI 2026: l\u2019anno dell\u2019adozione sistemica nelle aziende",
    "Nordest24 \u2014 IA Agentica 2026: trasformazione digitale nelle aziende italiane",
]
for s in sources:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_after = Pt(2)
    run_s = p_s.add_run(f"\u2022 {s}")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(8.5)
    run_s.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(16)
run_f = p_footer.add_run(
    "\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte"
)
run_f.font.name = 'Calibri'
run_f.font.size = Pt(8)
run_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_f.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
