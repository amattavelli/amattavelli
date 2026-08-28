"""
Quattro articoli Ratio -- 28 agosto 2026

1. 2026-08-28_digital-omnibus-alto-rischio-pmi-dicembre-2027.docx
2. 2026-08-28_agenti-ai-produzione-governance-pmi.docx
3. 2026-08-28_opus5-gpt56-gemini37-venti-euro-quale-scegliere.docx
4. 2026-08-28_mercato-ai-18-miliardi-pmi-ferme-8-percento.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO    = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(data_autore)
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# Digital Omnibus: l'alto rischio spostato al 2027
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Normativa AI e Compliance")
titolo(
    doc,
    "L’alto rischio aspetta il 2027. Ma non tutto.",
    "Il Regolamento UE 2026/1744 — il cosiddetto Digital Omnibus — ha spostato "
    "gli obblighi più severi dell’AI Act al dicembre 2027. "
    "Per PMI e studi professionali, cambia il calendario, non la direzione. "
    "E alcune scadenze sono già scattate.",
    "A cura della Redazione Ratio  •  28 agosto 2026"
)

para(doc,
    "Nel 1814 il Congresso di Vienna rimandò di qualche mese la ridefinizione "
    "dei confini europei. I diplomatici speravano che il tempo addolcisse le posizioni. "
    "In parte funzionò. Ma chi si era già preparato per la nuova mappa "
    "non perse niente: chi aveva aspettato il rinvio per cominciare "
    "a ragionare si ritrovò comunque indietro. "
    "Il Regolamento UE 2026/1744, entrato in vigore il 27 luglio 2026 "
    "e noto come Digital Omnibus on AI, ha modificato alcuni articoli dell’AI Act "
    "spostando le scadenze più impegnative per i sistemi ad alto rischio. "
    "Gli obblighi che sarebbero scattati il 2 agosto 2026 per questa categoria "
    "sono stati rinviati: al 2 dicembre 2027 per alcune tipologie, "
    "al 2 agosto 2028 per le più complesse. "
    "La notizia ha circolato come un sospiro di sollievo. "
    "Ma il sospiro andrebbe modulato con attenzione."
)

para(doc,
    "Il Digital Omnibus non ha toccato gli obblighi già in vigore dal 2 agosto 2026. "
    "L’obbligo di trasparenza per i sistemi AI che interagiscono con le persone "
    "— chatbot, assistenti vocali, risponditori automatici — è pienamente operativo. "
    "L’obbligo di AI literacy, che impone alle imprese di garantire "
    "che il personale utilizzi l’AI con competenza adeguata, è in vigore. "
    "La governance dei modelli di uso generale — GPT, Claude, Gemini "
    "— è soggetta ai nuovi obblighi documentali. "
    "Quello che il rinvio ha spostato sono gli obblighi di valutazione della conformità, "
    "registrazione e sorveglianza post-immissione sul mercato "
    "per i sistemi classificati ad alto rischio: "
    "sistemi di selezione del personale, scoring creditizio, "
    "valutazione degli studenti, gestione delle infrastrutture critiche."
)

heading(doc, "Chi è davvero ad alto rischio")

para(doc,
    "La categoria ‘alto rischio’ dell’AI Act è definita dall’Allegato III del Regolamento. "
    "Per la grande maggioranza delle PMI e degli studi professionali italiani, "
    "i sistemi in uso — chatbot di front-office, strumenti di redazione, "
    "assistenti per la contabilità, analisi documentale — "
    "non rientrano in questa categoria. "
    "Rientrano invece in alto rischio i sistemi che incidono su decisioni "
    "con effetti significativi sulle persone fisiche: "
    "ammissione al credito, selezione e valutazione del personale, "
    "accesso a servizi pubblici essenziali, sistemi biometrici. "
    "Se uno studio professionale o una PMI usa un software di selezione del personale "
    "con componenti AI — screening automatico dei curricula, scoring dei candidati — "
    "è molto probabile che quel sistema ricada nell’alto rischio. "
    "In quel caso, il rinvio al 2027 offre tempo per prepararsi, "
    "non per ignorare il tema."
)

heading(doc, "La mappa degli obblighi aggiornata")

para(doc,
    "Con il Digital Omnibus, il calendario degli obblighi si articola su tre livelli. "
    "Primo livello — già in vigore dal 2 agosto 2026: trasparenza sui sistemi AI "
    "che interagiscono con persone, marcatura dei contenuti generati da AI, "
    "AI literacy del personale, governance dei modelli di uso generale. "
    "Secondo livello — in vigore dal 2 dicembre 2027: obblighi per i sistemi "
    "ad alto rischio elencati nell’Allegato III che non erano già coperti "
    "dalla normativa di settore preesistente. "
    "Terzo livello — in vigore dal 2 agosto 2028: obblighi per i sistemi "
    "ad alto rischio incorporati in prodotti già soggetti a normative "
    "di sicurezza europee (macchinari, dispositivi medici, veicoli). "
    "Le sanzioni per le violazioni degli obblighi di trasparenza restano invariate: "
    "fino al 3% del fatturato mondiale o 15 milioni di euro. "
    "Per la non conformità dei sistemi ad alto rischio, fino al 7% o 35 milioni."
)

heading(doc, "Cosa fare adesso")

para(doc,
    "Per PMI e studi professionali, l’azione pratica post-Digital Omnibus "
    "è una verifica in due tempi. "
    "Nel breve periodo — settembre-ottobre 2026 — completare la compliance "
    "sugli obblighi già in vigore: verifica dei sistemi AI in uso, "
    "configurazione delle disclosure per i sistemi rivolti ai clienti, "
    "documentazione della formazione ricevuta dal personale. "
    "Nel medio periodo — entro il primo trimestre 2027 — verificare "
    "se uno o più sistemi in uso o in valutazione ricadono nell’alto rischio "
    "e avviare la valutazione della conformità con i requisiti "
    "dell’Allegato III. Chi aspetta il 2027 per cominciare a ragionarci "
    "si ritroverà nella stessa posizione di chi aspettava il 2 agosto "
    "per occuparsi degli obblighi di trasparenza."
)

para(doc,
    "Il Congresso di Vienna si concluse nel giugno 1815, "
    "poco prima di Waterloo. I confini furono ridisegnati, "
    "ma chi aveva usato bene i mesi del negoziato "
    "arribò al tavolo con una posizione più solida. "
    "Il rinvio è un’opportunità di preparazione, non di attesa."
)

riferimenti(doc, [
    "Regolamento UE 2026/1744 (Digital Omnibus on AI) — entrato in vigore il 27 luglio 2026",
    "Regolamento UE 2024/1689 (AI Act), Allegato III — Sistemi AI ad alto rischio",
    "Meta Communications — 'Regolamentazione dell’AI. Cosa cambia da agosto 2026' (luglio 2026)",
    "UniverseIT — 'AI Act 2 agosto 2026: cosa cambia per le imprese'",
    "Legaledigitale.com — 'AI Act 2026: i nuovi obblighi dal 2 agosto spiegati a imprese e professionisti'",
    "Gruppo Informatica — 'AI Act 2026: obblighi per PMI, studi e sanità dal 2 agosto'",
    "PMI.it — 'AI Act dal 2 agosto 2026, obblighi in vigore e rinvii'",
])
doc.save(BASE + "2026-08-28_digital-omnibus-alto-rischio-pmi-dicembre-2027.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# Agenti AI in produzione: il 22% delle grandi imprese ci è già
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Strumenti AI per Professionisti")
titolo(
    doc,
    "Il 22% delle grandi imprese ha già agenti AI in produzione.",
    "Gli agenti autonomi non sono più un esperimento di laboratorio. "
    "Una quota significativa delle aziende italiane leader li usa già su processi reali. "
    "Per le PMI, il divario si apre in tempo reale. "
    "La domanda non è se adottarli: è come farlo senza perdere il controllo.",
    "A cura della Redazione Ratio  •  28 agosto 2026"
)

para(doc,
    "Nel 1903, i fratelli Wright volarono per 12 secondi a Kitty Hawk. "
    "La maggior parte degli ingegneri aeronautici dell’epoca non era presente, "
    "e molti di quelli che sentirono la notizia la ritennero esagerata o irrilevante. "
    "Il volo commerciale era ancora due decenni lontano. "
    "Eppure chi ignorava ciò che era successo quel mattino "
    "si ritrovò tagliato fuori da una curva tecnologica "
    "che aveva già cambiato direzione. "
    "Nel 2026, il 22% delle grandi imprese italiane ha già implementato "
    "agenti AI autonomi su workflow multi-step senza supervisione costante. "
    "Non sono prototipi: sono sistemi in produzione, "
    "che leggono documenti, interrogano database, inviano comunicazioni, "
    "prenotano risorse e producono report senza che un operatore umano "
    "approvi ogni singolo passaggio. "
    "Il volo è già avvenuto. La domanda è cosa si fa adesso."
)

para(doc,
    "Un agente AI è un sistema che non si limita a rispondere a una domanda "
    "ma esegue una sequenza di azioni per raggiungere un obiettivo: "
    "legge le email in arrivo, identifica quelle che richiedono una risposta urgente, "
    "recupera le informazioni rilevanti dal CRM o dal gestionale, "
    "redige una bozza di risposta contestualizzata e la presenta al responsabile "
    "per l’approvazione finale — o la invia direttamente se rientra "
    "in un insieme di regole predefinite. "
    "La differenza rispetto a un chatbot tradizionale è la catena di azioni: "
    "non una risposta, ma un processo. "
    "I principali sistemi agentic disponibili nel 2026 — "
    "Claude Opus 5 con computer use, GPT-5.6 con Operator, "
    "Gemini 3.7 con Project Mariner — "
    "possono operare su interfacce web, leggere email, compilare moduli, "
    "navigare gestionali e produrre output strutturati."
)

heading(doc, "I casi d’uso che stanno entrando nelle aziende italiane")

para(doc,
    "Nelle imprese italiane che hanno già adottato agenti AI, "
    "i processi coperti rientrano in quattro categorie principali. "
    "La prima è la gestione documentale: classificazione automatica "
    "di contratti, fatture e corrispondenza; estrazione strutturata "
    "di dati da documenti eterogenei; instradamento ai responsabili "
    "in base al contenuto. "
    "La seconda è il supporto alle vendite: analisi del portafoglio clienti, "
    "prioritizzazione dei lead, redazione di offerte su template approvati. "
    "La terza è la reportistica interna: aggregazione di dati da fonti multiple, "
    "produzione di dashboard settimanali, allert su anomalie nei KPI. "
    "La quarta è il servizio clienti di primo livello: "
    "risposta a richieste frequenti, instradamento dei casi complessi "
    "a operatori umani, aggiornamento dello stato delle pratiche. "
    "Nessuno di questi casi richiede un sistema completamente autonomo: "
    "tutti prevedono un punto di supervisione umana "
    "prima delle azioni più rilevanti."
)

heading(doc, "Il rischio della delega senza governance")

para(doc,
    "Il principale rischio degli agenti AI non è tecnico: è organizzativo. "
    "Un agente che opera su processi reali può commettere errori "
    "con conseguenze concrete: inviare una comunicazione sbagliata a un cliente, "
    "classificare un documento in modo errato, "
    "aggiornare un record con dati imprecisi. "
    "La velocità di esecuzione — che è uno dei principali vantaggi degli agenti — "
    "amplifica anche gli errori: un agente che lavora su cento pratiche al giorno "
    "può replicare un errore sistematico cento volte prima che venga rilevato. "
    "Le aziende che stanno ottenendo risultati positivi dagli agenti AI "
    "hanno in comune una caratteristica: "
    "hanno definito in anticipo i perimetri di autonomia del sistema, "
    "i punti di approvazione umana obbligatoria e le procedure di escalation "
    "quando il sistema incontra casi non previsti. "
    "La governance non è un freno agli agenti AI: è la condizione "
    "che ne rende sostenibile l’uso nel tempo."
)

heading(doc, "Per le PMI: da dove iniziare")

para(doc,
    "Per una PMI o uno studio professionale che non ha ancora sperimentato "
    "agenti AI, il punto di ingresso più sicuro è un processo ad alto volume, "
    "bassa variabilità e conseguenze reversibili in caso di errore. "
    "La classificazione della posta in arrivo, la sintesi di documenti interni, "
    "la generazione di bozze di risposta standard: "
    "sono processi dove l’agente produce un output "
    "che un operatore umano verifica prima di agire. "
    "Il passaggio successivo — delegare all’agente la comunicazione diretta — "
    "richiede che il perimetro dei casi gestibili sia ben definito "
    "e che il sistema abbia dimostrato un tasso di errore accettabile "
    "nelle fasi di supervisione. "
    "Il 22% delle grandi imprese ci è già arrivato. "
    "Per l’altro 78%, e per le PMI, la curva è ancora tutta davanti."
)

para(doc,
    "I fratelli Wright impiegarono sei anni da Kitty Hawk "
    "a produrre un aereo commercialmente affidabile. "
    "Sei anni, nell’economia dell’AI del 2026, "
    "equivalgono probabilmente a diciotto mesi. "
    "Non c’è molto tempo per restare a guardare."
)

riferimenti(doc, [
    "AI4Business — 'Guida agenti AI 2026: cosa sono, come funzionano, perché ora' (2026)",
    "Impesud — 'Agentic AI in Italia 2026: Dalla Teoria all’Azione' (2026)",
    "Everest Innovation — 'Agenti AI Autonomi: il 2026 è l’Anno della Rivoluzione Intelligente' (2026)",
    "MM-One — 'AI 2026: come cambiano davvero le aziende con gli agenti autonomi' (2026)",
    "Corriere Nazionale — 'La Nuova Era dell’Efficienza: Come l’AI Agentica sta Trasformando il Lavoro' (aprile 2026)",
    "Osservatorio Artificial Intelligence, Politecnico di Milano — Rapporto mercato AI 2025-2026",
    "Regolamento UE 2024/1689 (AI Act), artt. 4, 6, Allegato III",
])
doc.save(BASE + "2026-08-28_agenti-ai-produzione-governance-pmi.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# Opus 5, GPT-5.6, Gemini 3.7: tutti a 20 euro, quale scegliere
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Strumenti AI per Professionisti")
titolo(
    doc,
    "Tutti a 20 euro al mese. Ma non sono uguali.",
    "Claude Opus 5, GPT-5.6 Sol e Gemini 3.7 Pro costano la stessa cifra. "
    "La parità di prezzo non significa parità di prestazioni. "
    "Per un professionista o un’azienda, la scelta giusta dipende da cosa si fa, "
    "non da quanto si spende.",
    "A cura della Redazione Ratio  •  28 agosto 2026"
)

para(doc,
    "Nel 1970, tutte le auto di fascia media costavano circa la stessa cifra. "
    "Una Fiat 128 e una BMW 2002 avevano prezzi simili, "
    "ma erano progettate per guidatori diversi, con abitudini diverse, "
    "su strade diverse. Il confronto non aveva senso in astratto: "
    "aveva senso solo sapendo dove si abitava, quanto si guidava "
    "e cosa si chiedeva all’auto. "
    "Nel 2026, i tre principali assistenti AI per professionisti "
    "— Claude Opus 5 di Anthropic, GPT-5.6 Sol di OpenAI, "
    "Gemini 3.7 Pro di Google — "
    "costano tutti circa 20 euro al mese per l’abbonamento individuale. "
    "La parità di prezzo, come quella delle auto del 1970, "
    "non dice niente su quale sia quello giusto per il proprio lavoro."
)

para(doc,
    "Claude Opus 5, rilasciato il 24 luglio 2026, "
    "rappresenta attualmente il vertice delle prestazioni "
    "nei compiti che richiedono ragionamento prolungato, "
    "coerenza su documenti lunghi e qualità della scrittura in italiano. "
    "Introduce l’‘Effort Dial’: la possibilità di modulare "
    "la profondità del ragionamento interno in base alla complessità del compito, "
    "con effetti diretti sulla velocità di risposta e sulla qualità dell’output. "
    "Su compiti di analisi documentale, redazione di pareri, "
    "sintesi di contratti complessi e risposta a quesiti normativi articolati "
    "è il modello con le prestazioni più costanti. "
    "GPT-5.6 Sol, lanciato il 9 luglio come primo tier della famiglia GPT-5.6, "
    "eccelle per la velocità di risposta, la gestione di task paralleli "
    "e l’integrazione con l’ecosistema Office e Teams. "
    "Gemini 3.7 Pro, rilasciato il 13 agosto, "
    "domina per chi lavora in Google Workspace "
    "e ha bisogno di informazioni aggiornate in tempo reale "
    "senza uscire dall’ambiente di lavoro."
)

heading(doc, "Il test che conta: il proprio lavoro")

para(doc,
    "Il modo più affidabile per scegliere tra i tre modelli "
    "è testarne le prestazioni sui propri casi d’uso reali, "
    "non sui benchmark generali. "
    "I benchmark misurano prestazioni medie su insiemi eterogenei di compiti: "
    "un modello che vince su un benchmark accademico "
    "può perdere su un documento fiscale in italiano del 2025. "
    "La procedura pratica è semplice: identificare tre o quattro compiti "
    "che si svolgono con frequenza nel proprio lavoro — "
    "redazione di una lettera al cliente, analisi di una visura camerale, "
    "risposta a un quesito normativo, sintesi di un contratto — "
    "eseguirli su tutti e tre i modelli con lo stesso prompt "
    "e valutare l’output su tre criteri: accuratezza, leggibilità e tempo. "
    "Il modello che vince sui propri casi d’uso specifici "
    "è il modello giusto per la propria realtà, "
    "indipendentemente da cosa dicono i report generalisti."
)

heading(doc, "Il fattore ecosistema")

para(doc,
    "Oltre alle prestazioni sui singoli task, la scelta del modello "
    "è influenzata dall’ecosistema di strumenti già in uso. "
    "Uno studio professionale che lavora prevalentemente su Microsoft 365 "
    "— Outlook, Word, Excel, Teams — troverà nel piano M365 Copilot "
    "con GPT-5.6 sottostante un’integrazione nativa "
    "difficilmente replicabile con gli altri modelli. "
    "Un’azienda che usa Google Workspace in modo intensivo "
    "— Gmail, Drive, Docs, Meet — troverà la stessa integrazione in Gemini. "
    "Chi lavora principalmente da browser o da strumenti non integrati "
    "con i grandi ecosystem può scegliere il modello più adatto al task "
    "senza vincoli di ecosistema. "
    "La scelta dell’assistente AI principale non è solo una scelta di modello: "
    "è sempre più una scelta di piattaforma."
)

heading(doc, "Il rischio della dipendenza da un unico fornitore")

para(doc,
    "La parità di prezzo tra i tre grandi modelli nasconde "
    "un rischio strutturale che molte aziende non stanno ancora considerando: "
    "la dipendenza operativa da un’unica piattaforma AI. "
    "Se uno studio professionale o un’azienda costruisce "
    "i propri flussi di lavoro interamente su un modello specifico, "
    "un cambio di policy commerciale, un’interruzione del servizio "
    "o un aumento di prezzo da parte del fornitore "
    "diventano un problema operativo immediato. "
    "L’approccio più robusto è usare il modello principale per l’80% dei casi d’uso "
    "e mantenere la familiarità con un modello alternativo per i rimanenti. "
    "Non è un problema tecnologico: è una scelta di gestione del rischio."
)

para(doc,
    "La Fiat 128 e la BMW 2002 avevano entrambe quattro ruote "
    "e un volante. Chi scelse bene guardò alle strade su cui guidava, "
    "non alla scheda tecnica. "
    "La scelta del modello AI funziona esattamente nello stesso modo."
)

riferimenti(doc, [
    "Oreate AI Guides — 'Claude 5, GPT-5.6, and Gemini 3.7: The State of AI Model Releases in August 2026'",
    "Bleap Finance — 'Claude vs GPT vs Gemini: Confronto tra i Migliori Modelli di IA del 2026'",
    "Pasquale Pillitteri — 'ChatGPT vs Claude vs Gemini: Quale Scegliere nel 2026'",
    "Jenova AI — 'GPT vs Claude vs Gemini: Confronto Completo dei Modelli AI per il 2026'",
    "Medium — 'I Tested Every Major AI Model in August 2026. Here’s the Winner' (agosto 2026)",
    "Anthropic — Note di rilascio Claude Opus 5 (24 luglio 2026)",
    "OpenAI — Note di rilascio GPT-5.6 (9 luglio 2026)",
])
doc.save(BASE + "2026-08-28_opus5-gpt56-gemini37-venti-euro-quale-scegliere.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# Il mercato AI italiano raddoppia. Le PMI restano all'8%
# ============================================================

doc = new_doc()
testata(doc, "Agosto 2026", "Mercato AI e Strategia Aziendale")
titolo(
    doc,
    "Il mercato AI italiano a 1,8 miliardi. Le PMI restano all’8%.",
    "Il mercato dell’intelligenza artificiale in Italia è cresciuto del 50% nel 2025 "
    "e ha superato 1,8 miliardi di euro. "
    "Ma l’adozione nelle PMI — che rappresentano il 99% del tessuto produttivo italiano "
    "— è ancora ferma all’otto per cento. "
    "Il divario si allarga. E non è solo un problema di tecnologia.",
    "A cura della Redazione Ratio  •  28 agosto 2026"
)

para(doc,
    "Nel 1844, Samuel Morse inviò il primo messaggio telegrafico "
    "tra Washington e Baltimora: ‘What hath God wrought’. "
    "Nei dieci anni successivi, il telegrafo connesse le principali città americane. "
    "Ma le piccole imprese rurali, i commercianti locali, "
    "i professionisti nelle città secondarie "
    "impiegarono ancora vent’anni per incorporarlo nei propri processi. "
    "Il telegrafo era disponibile. Non era accessibile. "
    "I dati sull’adozione dell’AI nelle imprese italiane nel 2026 "
    "raccontano una storia simile. "
    "Il mercato AI italiano ha chiuso il 2025 a 1,8 miliardi di euro, "
    "con una crescita del 50% in un anno, "
    "secondo il rapporto dell’Osservatorio Artificial Intelligence "
    "del Politecnico di Milano. "
    "La quota di imprese italiane con almeno dieci addetti "
    "che usa tecnologie AI è passata dal 5% del 2023 al 16,4% del 2026. "
    "Numeri che descrivono un cambiamento reale. "
    "Ma che nascondono una frattura altrettanto reale: "
    "nelle PMI con meno di cinquanta addetti, "
    "la quota di adozione si ferma all’8-11%."
)

para(doc,
    "Il divario non sorprende chi lavora vicino al tessuto produttivo italiano. "
    "Le grandi imprese hanno reparti IT, budget dedicati all’innovazione, "
    "figure come il Chief AI Officer o il Chief Digital Officer. "
    "Possono permettersi investimenti in piattaforme enterprise, "
    "progetti pilota, team di data scientist. "
    "Le PMI — che sono il 99,9% delle imprese italiane "
    "e producono il 70% del PIL manifatturiero — "
    "non hanno nessuna di queste risorse. "
    "Hanno un titolare che segue quattro telefoni contemporaneamente, "
    "un responsabile amministrativo che gestisce anche la logistica "
    "e un consulente esterno che viene una volta al mese. "
    "L’AI è tecnicamente accessibile a tutti. "
    "Ma l’adozione strutturata non lo è ancora."
)

heading(doc, "Perché la tecnologia non basta")

para(doc,
    "Il principale ostacolo all’adozione AI nelle PMI italiane "
    "non è il costo degli strumenti: "
    "20 euro al mese per un assistente AI di fascia alta "
    "è una cifra alla portata di qualsiasi attività con più di tre dipendenti. "
    "Il principale ostacolo è la combinazione di tre fattori. "
    "Primo: la mancanza di una figura interna "
    "che abbia tempo e competenza per impostare e mantenere l’uso degli strumenti. "
    "Secondo: l’assenza di casi d’uso chiari e misurabili "
    "sui quali valutare se l’investimento vale. "
    "Terzo: la percezione — spesso fondata — "
    "che gli strumenti disponibili richiedano un adattamento significativo "
    "ai processi specifici dell’azienda prima di diventare utili. "
    "Questi tre ostacoli non si risolvono con la formazione generica sull’AI. "
    "Si risolvono con un accompagnamento contestualizzato: "
    "qualcuno che conosce l’azienda, ne capisce i processi "
    "e può tradurre le capacità degli strumenti "
    "in applicazioni concrete e misurabili."
)

heading(doc, "Il ruolo dei professionisti nel colmare il gap")

para(doc,
    "Per i professionisti che affiancano PMI — commercialisti, consulenti del lavoro, "
    "consulenti di direzione — il divario di adozione "
    "è una delle opportunità più concrete del decennio. "
    "Le aziende che non trovano competenze AI interne "
    "e non hanno budget per consulenti specializzati "
    "si rivolgono alle figure di fiducia già presenti nel loro ecosistema. "
    "Il commercialista che ha sviluppato competenze AI "
    "non deve trasformarsi in un consulente tecnologico: "
    "deve essere in grado di rispondere a domande concrete "
    "come ‘questo strumento è adatto al mio processo di fatturazione?’, "
    "‘come posso usare l’AI per analizzare il mio portafoglio crediti?’, "
    "‘cosa devo documentare per essere conforme all’AI Act?’. "
    "La risposta a queste domande vale consulenza, "
    "e può essere offerta in modo naturale nel contesto "
    "di una relazione professionale già esistente."
)

heading(doc, "Gli incentivi disponibili")

para(doc,
    "Il gap di adozione delle PMI italiane è anche al centro delle politiche industriali. "
    "Il credito d’imposta per investimenti in beni strumentali digitali "
    "— l’iperammortamento al 180% per investimenti in software AI — "
    "è operativo per gli investimenti effettuati nel 2026. "
    "I bandi del PNRR per la digitalizzazione delle PMI "
    "includono esplicitamente strumenti e piattaforme AI. "
    "Il programma SME AI Accelerator, sviluppato da OpenAI con Confartigianato, "
    "offre percorsi di adozione guidata per le piccole imprese. "
    "Non mancano le risorse: manca spesso la figura "
    "che aiuta l’imprenditore a capire quali risorse esistono "
    "e come accedervi in modo efficace."
)

para(doc,
    "Morse aspettò dieci anni prima che il telegrafo diventasse "
    "una infrastruttura diffusa. "
    "Nell’economia dell’AI, dieci anni sono diventati tre. "
    "Le PMI che iniziano adesso non sono in ritardo: "
    "sono esattamente dove erano i commercianti del 1854 "
    "quando la prima linea raggiunse la loro città. "
    "Chi le aiuta a salire sul treno adesso "
    "è esattamente dove voleva essere."
)

riferimenti(doc, [
    "Osservatorio Artificial Intelligence, Politecnico di Milano — Rapporto mercato AI 2025-2026",
    "ISTAT — 'ICT nelle imprese: quota di imprese con almeno 10 addetti che usa AI' (dicembre 2025)",
    "AI4Business — 'Trend 2026: AI in accelerazione — non solo hype, il 2026 sarà l’anno del valore concreto' (Google Cloud)",
    "Etalentum — IV Report ‘L’impatto dell’AI generativa su PMI e Governance’ (2026)",
    "Edizioni Este — 'AI Act, le PMI italiane arrivano impreparate alla svolta' (2026)",
    "Incentivimpresa.it — 'AI e Finanza Agevolata 2026: Guida Completa per Imprese'",
    "SME AI Accelerator — Programma OpenAI / Confartigianato (giugno 2026)",
    "Agenzia delle Entrate — Credito d’imposta beni strumentali digitali — Iperammortamento 180% (2026)",
])
doc.save(BASE + "2026-08-28_mercato-ai-18-miliardi-pmi-ferme-8-percento.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
