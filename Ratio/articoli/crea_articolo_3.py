"""
Articolo 3: Strumenti AI per commercialisti - Genya Expert AI
Ratio/articoli/2026-04-03_ai-gestionali-commercialisti-genya.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-03_ai-gestionali-commercialisti-genya.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_paragraph(doc, text, size=11, italic=False, color=None,
                  space_after=8, space_before=0,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

# ---- TESTATA ----
p_rivista = doc.add_paragraph()
p_rivista.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = p_rivista.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
run_r.font.name = 'Calibri'
run_r.font.size = Pt(9)
run_r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run_r.font.bold = True
run_r.font.all_caps = True

add_separator(doc)

p_data = doc.add_paragraph()
p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_d = p_data.add_run("Aprile 2026  |  Strumenti per il Professionista")
run_d.font.name = 'Calibri'
run_d.font.size = Pt(8.5)
run_d.italic = True
run_d.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_after = Pt(6)
run_t = p_title.add_run(
    "Il software che risponde al rigo VJ6: l\u2019AI entra nei gestionali dei commercialisti"
)
run_t.font.name = 'Calibri'
run_t.font.size = Pt(24)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_after = Pt(10)
run_s = p_sub.add_run(
    "Con Genya Expert AI, Wolters Kluwer porta l\u2019intelligenza artificiale direttamente "
    "nella dichiarazione IVA. Il mercato degli strumenti AI per gli studi \u00e8 in rapida trasformazione."
)
run_s.font.name = 'Calibri'
run_s.font.size = Pt(13)
run_s.italic = True
run_s.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_separator(doc)

p_autore = doc.add_paragraph()
p_autore.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_autore.paragraph_format.space_after = Pt(16)
run_a = p_autore.add_run("A cura della Redazione Ratio  \u2022  3 aprile 2026")
run_a.font.name = 'Calibri'
run_a.font.size = Pt(9)
run_a.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- CORPO ----

add_paragraph(doc,
    "Il 23 marzo 2026, durante la compilazione della dichiarazione IVA annuale, un commercialista "
    "che usa Genya \u2014 la piattaforma cloud di Wolters Kluwer Tax & Accounting Italia, adottata "
    "da decine di migliaia di studi nel paese \u2014 pu\u00f2 scrivere in linguaggio naturale "
    "\u201ccome si compila il rigo VJ6?\u201d e ricevere una risposta contestualizzata al cliente "
    "aperto in quel momento. Non si apre una nuova scheda su Google. Non si cerca sul sito "
    "dell\u2019Agenzia delle Entrate. La risposta arriva dentro il software, con riferimento "
    "alle istruzioni ministeriali e ai dati specifici del soggetto in lavorazione."
)

add_paragraph(doc,
    "Sembra un dettaglio. Non lo \u00e8. Significa che l\u2019AI \u00e8 entrata nel flusso di lavoro "
    "ordinario dello studio, non pi\u00f9 come strumento separato da aprire in un\u2019altra finestra, "
    "ma come parte del gestionale che il professionista usa gi\u00e0 ogni giorno. Il confine tra "
    "software professionale e assistente AI si \u00e8 dissolto."
)

add_heading2(doc, "Come funziona Genya Expert AI")

add_paragraph(doc,
    "La scelta di design pi\u00f9 rilevante di Genya Expert AI non \u00e8 tecnica ma editoriale: il sistema "
    "attinge esclusivamente da fonti ufficiali, principalmente le istruzioni ministeriali e le circolari "
    "dell\u2019Agenzia delle Entrate. Nessuna fonte esterna, nessun contenuto generato liberamente. "
    "Questo approccio, che Wolters Kluwer chiama \u201cprincipii di AI responsabile\u201d, riduce "
    "drasticamente il rischio di allucinazioni su materia normativa \u2014 il problema pi\u00f9 "
    "pericoloso per chi opera in contesti fiscali e contabili."
)

add_paragraph(doc,
    "Il sistema fa anche qualcosa di pi\u00f9: segnala proattivamente le incongruenze nei dati inseriti. "
    "Se un valore nel modello non \u00e8 coerente con quanto dichiarato in un altro rigo, l\u2019AI "
    "lo evidenzia prima che il professionista proceda. \u00c8 una funzione che non sostituisce il "
    "controllo umano ma lo supporta, rendendo pi\u00f9 difficile che un errore di compilazione passi "
    "inosservato nella fase di caricamento dei dati."
)

add_heading2(doc, "Il mercato si allarga: gli altri strumenti gi\u00e0 operativi")

add_paragraph(doc,
    "Genya non \u00e8 un caso isolato. Il mercato italiano degli strumenti AI specificamente progettati "
    "per professionisti \u00e8 in rapida espansione. Normo.ai, una soluzione italiana, offre un "
    "assistente AI per fiscalit\u00e0, diritto del lavoro e normativa civilistica, con supporto alla "
    "compilazione di modelli dichiarativi come 730, CU, ISA e dichiarazione IVA. Aptus.AI \u00e8 "
    "orientata ai professionisti legali e agli avvocati. Il Sole 24 Ore Professionale ha lanciato "
    "soluzioni AI dedicate per commercialisti e aziende. Non si tratta pi\u00f9 di strumenti generici "
    "adattati all\u2019uso professionale: sono prodotti costruiti specificamente per il contesto "
    "normativo e operativo italiano."
)

add_paragraph(doc,
    "Secondo un\u2019indagine condotta dalla Fondazione Nazionale di Ricerca dei Commercialisti "
    "tra luglio e settembre 2025, il 34% dei commercialisti italiani usa gi\u00e0 strumenti AI "
    "nell\u2019attivit\u00e0 professionale. La stessa ricerca stima che questa percentuale possa "
    "salire al 72% nei prossimi tre anni, con una visione favorevole dell\u2019AI condivisa "
    "dall\u201985% degli intervistati. Il mercato mondiale del software contabile-fiscale \u00e8 "
    "atteso in crescita del 15,5% annuo fino al 2035, con l\u2019AI come principale motore."
)

add_heading2(doc, "Cosa cambia nel lavoro dello studio")

add_paragraph(doc,
    "Il cambiamento pi\u00f9 profondo non riguarda la velocit\u00e0 di compilazione. Riguarda dove "
    "si concentra l\u2019attenzione del professionista. Quando la ricerca normativa e il controllo "
    "formale dei dati vengono supportati dall\u2019AI, il commercialista viene liberato da una parte "
    "significativa del lavoro che assorbiva tempo senza generare valore percepito dal cliente. "
    "Il tempo ricuperato pu\u00f2 essere reinvestito nella consulenza proattiva: pianificazione "
    "fiscale, analisi della situazione finanziaria del cliente, supporto alle scelte gestionali."
)

add_paragraph(doc,
    "Questo \u00e8 anche il nodo pi\u00f9 delicato. Gli studi che adotteranno questi strumenti senza "
    "ripensare il proprio modello di servizio rischiano di abbattere i costi operativi senza "
    "aumentare il valore offerto. Quelli che useranno il tempo recuperato per costruire un "
    "rapporto con il cliente pi\u00f9 orientato all\u2019anticipazione dei problemi \u2014 invece "
    "che alla produzione di adempimenti \u2014 troveranno in questi strumenti un vantaggio reale."
)

add_heading2(doc, "La responsabilit\u00e0 resta dove \u00e8 sempre stata")

add_paragraph(doc,
    "La Legge 132/2025 e le norme deontologiche degli ordini professionali ribadiscono un principio "
    "che non \u00e8 cambiato: la responsabilit\u00e0 del lavoro professionale resta interamente al "
    "professionista che lo firma. L\u2019AI \u00e8 uno strumento di supporto. Se il software suggerisce "
    "una compilazione errata e il commercialista la trasmette senza verificarla, la responsabilit\u00e0"
    " \u00e8 del professionista. Questo non \u00e8 un limite dello strumento: \u00e8 la condizione "
    "che consente al professionista di mantenere il proprio ruolo, la propria competenza e la "
    "propria relazione di fiducia con il cliente."
)

add_paragraph(doc,
    "Chi lavora con questi strumenti non fa meno il commercialista. Lo fa altrove. Meno nel rigo VJ6, "
    "pi\u00f9 nella domanda che il cliente non sa ancora di dover fare."
)

add_separator(doc)

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(10)
run_n = p_note.add_run("Fonti e riferimenti")
run_n.font.name = 'Calibri'
run_n.font.size = Pt(9)
run_n.font.bold = True
run_n.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Wolters Kluwer Tax & Accounting Italia \u2014 Annuncio Genya Dichiarativi Expert AI (23 marzo 2026)",
    "ANSA \u2014 Wolters Kluwer presenta Genya Dichiarativi con IA integrata (23 marzo 2026)",
    "Fondazione Nazionale di Ricerca dei Commercialisti \u2014 Indagine sull\u2019uso dell\u2019AI, luglio-settembre 2025",
    "Normo.ai \u2014 Documentazione del prodotto (2026)",
    "DATALOG \u2014 Digitalizzazione dei commercialisti: evoluzioni e trend 2026",
    "Companeo \u2014 I migliori software di contabilit\u00e0 AI per il 2026",
]
for s in sources:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_after = Pt(2)
    run_s = p_s.add_run(f"\u2022 {s}")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(8.5)
    run_s.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(16)
run_f = p_footer.add_run(
    "\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte"
)
run_f.font.name = 'Calibri'
run_f.font.size = Pt(8)
run_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_f.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
