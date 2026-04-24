"""
Articolo 2 — PMI e AI: uso vs integrazione
Ratio/articoli/2026-04-24_pmi-ai-uso-integrazione-misurare.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-24_pmi-ai-uso-integrazione-misurare.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'; run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'; run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Aprile 2026  |  Adozione AI nelle PMI")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("L'81% delle PMI dice di usare l'AI.\nMeno di una su quattro sa misurare il risultato.")
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Tra chi usa ChatGPT per scrivere email e chi ha automatizzato "
    "la riconciliazione contabile c'è una distanza enorme. "
    "E il mercato sta già premiando chi ha fatto il salto."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  24 aprile 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Una ricerca pubblicata nel primo trimestre del 2026 ha rilevato che l'81% delle "
    "piccole e medie imprese italiane dichiara di usare strumenti di intelligenza artificiale. "
    "Il dato sembra promettente, finché non si scopre che per “usare l’AI” la maggior parte "
    "intende chiedere a ChatGPT come scrivere un'email o generare qualche immagine per i "
    "social. Meno del 25% ha integrato l'AI in un processo aziendale misurabile. "
    "La distanza tra le due situazioni è enorme, e il modo in cui si sta colmando "
    "determinerà buona parte del vantaggio competitivo dei prossimi anni."
)

heading(doc, "Il problema non è la tecnologia: è la domanda sbagliata")

para(doc,
    "Le PMI che ottengono risultati concreti dall'AI partono da un processo specifico, "
    "spesso noioso, ripetitivo e ad alto costo di tempo, e chiedono: possiamo automatizzarlo "
    "o accelerarlo in modo misurabile? Quelle che si perdono partono invece dalla tecnologia "
    "e chiedono: dove potremmo usare l'AI? La seconda domanda sembra più aperta e "
    "strategica, ma in pratica produce sperimentazioni senza metodo, entusiasmo iniziale "
    "e poi abbandono quando non si vede un ritorno chiaro."
)

para(doc,
    "I casi d'uso che funzionano nelle PMI italiane hanno caratteristiche comuni. "
    "Riguardano attività standardizzabili, con output verificabili e con un costo del "
    "lavoro manuale documentabile. L'analisi e l'archiviazione di fatture in ingresso, "
    "la riconciliazione di estratti conto bancari con il gestionale, la redazione di "
    "risposte standard a richieste di informazioni commerciali, la sintesi di contratti "
    "prima di una riunione negoziale: questi sono i campi dove l'AI produce risparmi "
    "reali, nell'ordine del 50-70% del tempo precedentemente impiegato. Chi ha fatto "
    "questa scelta e l'ha misurata ha i numeri per giustificare l'investimento successivo."
)

heading(doc, "Il vero freno: la mancanza di un referente interno")

para(doc,
    "Ciò che frena l'adozione non è quasi mai la tecnologia disponibile. La soglia di "
    "accesso economica si è abbassata considerevolmente: gli strumenti SaaS che permettono "
    "di costruire flussi automatizzati (come Make o n8n) costano tra i 20 e i 200 euro al "
    "mese e permettono di collegare sorgenti dati aziendali a modelli AI senza scrivere "
    "una riga di codice. Una PMI con dieci dipendenti può automatizzare la gestione delle "
    "richieste di informazioni via email con un investimento iniziale dell'ordine di qualche "
    "giornata di setup e una spesa mensile equivalente a poche ore di lavoro impiegatizio."
)

para(doc,
    "Il vero collo di bottiglia è la mancanza di un referente interno che conosca sia il "
    "processo da automatizzare sia le basi del funzionamento degli strumenti. Nelle PMI "
    "senza una funzione IT dedicata, questa figura manca strutturalmente. La conseguenza "
    "è che si rimane bloccati a un uso superficiale degli strumenti, senza mai arrivare "
    "all'integrazione. La soluzione non è assumere un esperto AI, che sarebbe sproporzionato "
    "per la maggior parte delle dimensioni in gioco: è formare una persona già in azienda "
    "sulle basi dell'automazione e affiancarla per il primo progetto con supporto esterno."
)

heading(doc, "Il nodo della riservatezza dei dati")

para(doc,
    "Un secondo freno ricorrente riguarda la riservatezza dei dati. Molti titolari di PMI "
    "sono giustamente cauti nell'inserire dati aziendali su piattaforme cloud esterne. "
    "La risposta a questa preoccupazione non è evitare l'AI, ma scegliere con attenzione "
    "la piattaforma: esistono soluzioni con hosting europeo e trattamento dei dati conforme "
    "al GDPR, oppure soluzioni ibride che processano i dati localmente e usano il modello "
    "cloud solo per l'elaborazione del linguaggio. Per le attività più sensibili esistono "
    "anche modelli linguistici open source (come Mistral o Llama 3) installabili su server "
    "aziendali, eliminando il problema alla radice. La scelta della piattaforma giusta "
    "richiede un'ora di analisi, non mesi di attesa."
)

heading(doc, "Come misurare: il prima e il dopo")

para(doc,
    "Il problema della misurazione è più semplice di quanto sembri, ma richiede disciplina. "
    "Prima di avviare qualsiasi progetto AI, occorre documentare il tempo attualmente impiegato "
    "nel processo che si vuole automatizzare, la frequenza degli errori, il costo orario delle "
    "risorse coinvolte. Dopo l'implementazione, gli stessi indicatori permettono di calcolare "
    "il risparmio effettivo. Senza questo passaggio, l'AI rimane un'impressione positiva, "
    "non un dato di bilancio. Le aziende che hanno questa disciplina riescono ad estendere "
    "il progetto pilota ad altri processi perché hanno i numeri per giustificare l'investimento "
    "al management o ai soci."
)

para(doc,
    "Il divario tra le imprese che stanno costruendo vantaggio competitivo reale con l'AI "
    "e quelle che lo dichiarano senza averlo è destinato ad ampliarsi nei prossimi diciotto "
    "mesi. La differenza non è nella capacità di spesa ma nel metodo: scegliere un processo, "
    "misurare il prima, implementare, misurare il dopo. Chi ha fatto questo esercizio una volta "
    "lo fa di nuovo. Chi non lo ha mai fatto continuerà a rispondere positivamente ai sondaggi "
    "sull'uso dell'AI senza che nulla cambi nel proprio conto economico."
)

para(doc,
    "La domanda utile da porre nella prossima riunione di staff è semplice: quale attività "
    "ci costa più tempo e produce il risultato meno differenziato rispetto ai concorrenti? "
    "La risposta a quella domanda è il primo progetto AI da fare."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "Osservatorio Artificial Intelligence — Politecnico di Milano, Rapporto 2025/2026",
    "pmi.it — \"AI nelle PMI, l'81% la usa ma solo 1 su 4 l'ha integrata\" (marzo 2026)",
    "kinetikon.com — \"Agenti IA nelle PMI italiane: adozione e casi d'uso\" (2026)",
    "managementcue.it — \"Gli agenti AI in azienda: cosa cambierà davvero nel 2026\" (marzo 2026)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
