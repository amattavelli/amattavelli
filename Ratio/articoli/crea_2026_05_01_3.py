"""
Articolo 3 — Claude 4.7, GPT-5.5, Gemini: come scegliere senza restare intrappolati
Ratio/articoli/2026-05-01_scegliere-modello-ai-2026-professionisti.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-05-01_scegliere-modello-ai-2026-professionisti.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


# ---- TESTATA ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.bold = True; r.font.all_caps = True

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Maggio 2026  |  Strumenti e Modelli AI")
r.font.name = 'Calibri'; r.font.size = Pt(8.5)
r.font.italic = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Claude 4.7, GPT-5.5, Gemini:\ncome scegliere senza restare intrappolati")
r.font.name = 'Calibri'; r.font.size = Pt(24)
r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Il confronto tra modelli AI è diventato uno sport popolare. "
    "Per un professionista o un'azienda, però, la domanda giusta non è quale vince i benchmark."
)
r.font.name = 'Calibri'; r.font.size = Pt(13)
r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sep(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A cura della Redazione Ratio  •  1 maggio 2026")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- TESTO ----
para(doc,
    "Uno studio legale milanese ha cambiato tre volte strumento AI negli ultimi diciotto "
    "mesi. Prima ChatGPT, poi Claude perché \"scrive meglio in italiano\", poi di nuovo "
    "ChatGPT con il nuovo modello, poi una soluzione enterprise di Microsoft. Ogni cambio "
    "ha richiesto una fase di adattamento: nuovi prompt, nuove abitudini, nuovi template. "
    "Il tempo perso nei transiti è stato stimato internamente in quattro settimane di "
    "produttività ridotta. Nessuno dei modelli era sbagliato. La scelta, ogni volta, era "
    "stata fatta senza criteri."
)

para(doc,
    "La proliferazione di modelli nel 2026 ha reso questo problema molto più comune. "
    "Anthropic ha rilasciato Claude Opus 4.7 nell'aprile 2026, con capacità agenziali "
    "avanzate e un contesto di 200.000 token. OpenAI ha risposto a fine aprile con GPT-5.5, "
    "con computer use migliorato e 78,7% su OSWorld-Verified. Google continua ad aggiornare "
    "la famiglia Gemini con integrazioni sempre più profonde in Google Workspace. Il "
    "confronto tecnico tra modelli è reale, ma per un professionista che deve scegliere "
    "uno strumento da usare nel lavoro quotidiano, i benchmark sono spesso irrilevanti."
)

heading(doc, "I criteri che contano davvero")

para(doc,
    "La prima domanda da farsi non è \"quale modello scrive meglio?\" ma \"in quale "
    "sistema di lavoro si inserisce questo strumento?\". Un modello eccellente per la "
    "redazione di testi giuridici, ma senza integrazione nativa con il gestionale dello "
    "studio, genera più frizione di un modello tecnicamente meno potente ma già integrato "
    "nella piattaforma usata ogni giorno. La scelta del modello è, in buona parte, una "
    "scelta di architettura del flusso di lavoro."
)

para(doc,
    "Per i professionisti e le PMI italiane, i criteri pratici più rilevanti nel 2026 "
    "sono quattro. Il primo è la residenza dei dati: le informazioni professionali riservate "
    "(dati fiscali, contenuto di pratiche, dati personali dei clienti) non possono "
    "transitare su server extra-UE senza adeguate garanzie contrattuali. Microsoft Copilot, "
    "Google Workspace con Gemini e la versione enterprise di Claude offrono oggi opzioni "
    "con data residency in Europa; ChatGPT Plus standard non garantisce questa localizzazione."
)

para(doc,
    "Il secondo criterio è l'integrazione. I modelli che si collegano agli strumenti già "
    "in uso (Microsoft 365, Google Workspace, gestionali come TeamSystem o Zucchetti) "
    "eliminano il passaggio manuale di dati tra sistemi e riducono il rischio di errori "
    "di trascrizione. Il terzo è la stabilità dell'API: chi usa i modelli via API deve "
    "valutare la politica di deprecazione del fornitore, perché migrare da un'API all'altra "
    "ha costi non banali in termini di adattamento dei prompt e revisione dei workflow "
    "automatizzati già in produzione."
)

para(doc,
    "Il quarto criterio, spesso il più trascurato, è il costo per caso d'uso reale. I "
    "listini pubblici indicano prezzi per milione di token, ma la spesa effettiva dipende "
    "da quanti token consuma ogni richiesta tipica dello studio. Chi usa l'AI per analizzare "
    "contratti lunghi consuma token molto diversamente da chi la usa per rispondere a email "
    "brevi. Senza una stima del consumo effettivo, il confronto economico tra modelli è "
    "puramente teorico."
)

heading(doc, "Le differenze pratiche tra Claude, GPT e Gemini nel 2026")

para(doc,
    "Claude Opus 4.7 si distingue, nelle valutazioni dei professionisti che lo usano "
    "quotidianamente, per la qualità del ragionamento su testi lunghi e complessi e per "
    "la capacità di seguire istruzioni precise in contesti multi-step. La finestra di "
    "contesto estesa a 200.000 token lo rende la scelta più naturale per chi deve analizzare "
    "contratti articolati, bilanci consolidati o fascicoli di pratiche di grandi dimensioni "
    "senza perdere il filo del ragionamento."
)

para(doc,
    "GPT-5.5, con il computer use e la forte integrazione nell'ecosistema OpenAI, è la "
    "scelta più performante per i workflow che richiedono automazione operativa: compilare "
    "moduli, navigare portali, estrarre dati da siti web, operare su fogli di calcolo in "
    "modo autonomo. Gemini Pro, nell'ecosistema Google, è la scelta più naturale per chi "
    "ha già Google Workspace come piattaforma principale e vuole un'integrazione profonda "
    "con Gmail, Drive e Meet senza infrastrutture aggiuntive."
)

heading(doc, "Il rischio del lock-in")

para(doc,
    "Il rischio più concreto nel 2026 non è scegliere il modello sbagliato, ma costruire "
    "una dipendenza troppo profonda da un singolo fornitore. I workflow costruiti intorno "
    "a un modello specifico (prompt ottimizzati per quel modello, template, integrazioni "
    "API, automazioni) sono difficili da trasferire quando il fornitore cambia i prezzi, "
    "riduce le capacità del tier in uso, o depreca una funzione su cui è stata costruita "
    "un'automazione critica. Questo è già accaduto più volte negli ultimi due anni, e "
    "accadrà ancora."
)

para(doc,
    "La strategia più robusta per uno studio o una PMI di dimensioni medie è quella di "
    "costruire workflow portabili: testo dei prompt documentato e versionato, logiche di "
    "automazione che non dipendano da caratteristiche esclusive di un singolo modello, e "
    "una competenza interna sufficiente per valutare le alternative quando si rende "
    "necessario un cambio. Non significa usare tutti i modelli contemporaneamente, ma "
    "non affidarsi a uno solo come se fosse infrastruttura critica senza alternative."
)

heading(doc, "La scelta che dura")

para(doc,
    "Il panorama dei modelli AI cambierà ancora nei prossimi mesi, come ha fatto nei "
    "diciotto precedenti. La scelta giusta non è quella del modello più potente in "
    "assoluto, ma quella che si integra meglio nel modo di lavorare dello studio o "
    "dell'azienda, con il minore rischio di lock-in e il maggiore controllo sui dati "
    "dei clienti. Chi sceglie con questi criteri ha un vantaggio duraturo rispetto a chi "
    "insegue i benchmark: non cambia strumento ogni tre mesi e riesce a consolidare "
    "competenze e workflow che rimangono utili anche quando arriva il modello successivo."
)

sep(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Riferimenti")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for s in [
    "Anthropic — Claude Opus 4.7 Release Notes (aprile 2026)",
    "OpenAI — GPT-5.5 for ChatGPT and Codex (aprile 2026)",
    "Jenova.ai — GPT vs Claude vs Gemini: Confronto Completo dei Modelli AI per il 2026",
    "Luca Mastella — Claude Opus 4.7 + Claude Design: la guida pratica per il tuo lavoro (2026)",
    "InfoData / Il Sole 24 Ore — GPT-5.4, il modello che usa il computer (marzo 2026)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {s}")
    r.font.name = 'Calibri'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run(
    "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
    "Riproduzione consentita con citazione della fonte"
)
r.font.name = 'Calibri'; r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.font.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
