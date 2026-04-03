"""
Articolo 2: AI Act + L.132/2025 - Scadenze agosto 2026
Ratio/articoli/2026-04-03_quattro-mesi-agosto-2026-obblighi-ai.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/home/user/amattavelli/Ratio/articoli/2026-04-03_quattro-mesi-agosto-2026-obblighi-ai.docx"

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_paragraph(doc, text, size=11, italic=False, color=None,
                  space_after=8, space_before=0,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

# ---- TESTATA ----
p_rivista = doc.add_paragraph()
p_rivista.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = p_rivista.add_run("RATIO  \u2022  Approfondimenti per Professionisti e Imprese")
run_r.font.name = 'Calibri'
run_r.font.size = Pt(9)
run_r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run_r.font.bold = True
run_r.font.all_caps = True

add_separator(doc)

p_data = doc.add_paragraph()
p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_d = p_data.add_run("Aprile 2026  |  Normativa e Compliance")
run_d.font.name = 'Calibri'
run_d.font.size = Pt(8.5)
run_d.italic = True
run_d.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ---- TITOLO ----
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_after = Pt(6)
run_t = p_title.add_run(
    "Quattro mesi all\u2019agosto 2026: la verifica che ogni impresa italiana deve fare sull\u2019AI"
)
run_t.font.name = 'Calibri'
run_t.font.size = Pt(24)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_after = Pt(10)
run_s = p_sub.add_run(
    "Il 2 agosto 2026 entrano in piena applicazione le norme AI Act sui sistemi ad alto rischio. "
    "Chi usa AI in HR, credito o gestione operativa deve fare una verifica urgente."
)
run_s.font.name = 'Calibri'
run_s.font.size = Pt(13)
run_s.italic = True
run_s.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_separator(doc)

p_autore = doc.add_paragraph()
p_autore.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_autore.paragraph_format.space_after = Pt(16)
run_a = p_autore.add_run("A cura della Redazione Ratio  \u2022  3 aprile 2026")
run_a.font.name = 'Calibri'
run_a.font.size = Pt(9)
run_a.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---- CORPO ----

add_paragraph(doc,
    "Il 2 agosto 2026 scattano le disposizioni pi\u00f9 operative del Regolamento europeo sull\u2019intelligenza "
    "artificiale (AI Act, Reg. UE 2024/1689) per i sistemi classificati ad alto rischio. Mancano quattro mesi. "
    "Per buona parte delle aziende italiane, questa \u00e8 ancora una data astratta \u2014 qualcosa che "
    "riguarda i grandi gruppi, i fornitori di software, chi sviluppa tecnologia. Non \u00e8 cos\u00ec. "
    "Qualsiasi impresa che usa un sistema AI per prendere decisioni su persone \u2014 dipendenti, clienti, "
    "fornitori \u2014 potrebbe gi\u00e0 essere dentro il perimetro regolamentato."
)

add_paragraph(doc,
    "L\u2019AI Act classifica i sistemi AI in quattro categorie di rischio. Quelli a rischio inaccettabile "
    "sono vietati dal febbraio 2025. Quelli a rischio alto devono soddisfare requisiti tecnici e "
    "documentali stringenti prima di poter essere messi in uso o continuare a operare. Quelli a rischio "
    "limitato hanno solo obblighi di trasparenza. La maggior parte degli strumenti AI che uno studio "
    "professionale o una PMI usa quotidianamente \u2014 per scrivere testi, analizzare dati, rispondere "
    "a email \u2014 ricade nella categoria a rischio limitato o minimo, e non subir\u00e0 cambiamenti "
    "significativi. Il nodo critico riguarda i sistemi che influenzano decisioni ad alto impatto sulle persone."
)

add_heading2(doc, "Cosa rende un sistema AI \u201cad alto rischio\u201d")

add_paragraph(doc,
    "L\u2019allegato III dell\u2019AI Act elenca le categorie di utilizzo che determinano la classificazione "
    "ad alto rischio. Tra quelle pi\u00f9 rilevanti per aziende e studi professionali: i sistemi usati "
    "nella selezione e valutazione del personale (screening di CV, valutazione delle performance, decisioni "
    "su avanzamenti o licenziamenti), i sistemi usati in ambito creditizio per valutare l\u2019affidabilit\u00e0 "
    "di un soggetto, e quelli usati in processi amministrativi o giudiziari. Molti software HR moderni "
    "che integrano AI per il recruiting o la gestione delle performance rientrano in questa categoria."
)

add_paragraph(doc,
    "Per i sistemi ad alto rischio, l\u2019AI Act impone obblighi precisi. Il fornitore del sistema deve "
    "garantire la conformit\u00e0 tecnica e documentarla. Ma anche l\u2019utilizzatore \u2014 l\u2019azienda "
    "o lo studio che lo usa \u2014 ha responsabilit\u00e0 dirette: deve verificare che il sistema sia stato "
    "certificato, mantenere un registro degli eventi rilevanti (audit trail), assicurarsi che le decisioni "
    "prodotte dall\u2019AI possano essere spiegate in modo comprensibile (explainability), e nominare un "
    "responsabile interno per la supervisione. Le sanzioni per le violazioni arrivano fino al 3% del "
    "fatturato mondiale annuo; per le infrazioni pi\u00f9 gravi, al 6%."
)

add_heading2(doc, "La Legge italiana 132/2025: il quadro nazionale")

add_paragraph(doc,
    "Al framework europeo si aggiunge la Legge 132 del 23 settembre 2025, la prima legge nazionale "
    "sull\u2019AI in Europa, in vigore dal 10 ottobre 2025. Il provvedimento ha stanziato un miliardo "
    "di euro per supportare startup e PMI nei settori AI e tecnologie correlate, ma ha anche introdotto "
    "norme operative che si sommano agli obblighi europei. Tra le disposizioni pi\u00f9 rilevanti per le "
    "imprese: l\u2019obbligo di valutazione preventiva dell\u2019impatto sui diritti dei lavoratori per "
    "i sistemi AI usati in HR e gestione operativa, con procedure di consultazione sindacale in alcuni casi."
)

add_paragraph(doc,
    "I dati dell\u2019ISTAT certificati a dicembre 2025 mostrano che il 16,4% delle imprese italiane con "
    "almeno dieci addetti usa gi\u00e0 almeno una tecnologia AI \u2014 un dato raddoppiato rispetto all\u2019"
    "8,2% del 2024. Significa che circa una impresa su sei \u00e8 gi\u00e0 dentro il perimetro "
    "regolamentato, spesso senza averlo verificato sistematicamente."
)

add_heading2(doc, "Il ruolo dei consulenti e dei commercialisti")

add_paragraph(doc,
    "Per i consulenti d\u2019impresa, gli avvocati e i commercialisti, questo scenario apre un fronte "
    "di lavoro significativo. La compliance AI non \u00e8 una questione puramente informatica: richiede "
    "la capacit\u00e0 di leggere un processo aziendale, identificare dove vengono usati sistemi AI, "
    "classificarli rispetto alle categorie di rischio del Regolamento, e valutare se la documentazione "
    "tecnica prodotta dal fornitore del software \u00e8 sufficiente. \u00c8 un percorso analogo a quello "
    "che molti studi hanno fatto con il GDPR tra il 2017 e il 2018: chi si \u00e8 attrezzato prima ha "
    "potuto costruire una competenza spendibile sul mercato."
)

add_paragraph(doc,
    "Sul piano pratico, il primo passo per uno studio \u00e8 mappare i sistemi AI usati dai propri "
    "clienti PMI: software HR, strumenti di credit scoring, sistemi di valutazione delle performance. "
    "Per ciascuno, verificare se il fornitore ha gi\u00e0 avviato il percorso di conformit\u00e0 AI Act "
    "e se esiste documentazione tecnica disponibile. In molti casi, questa verifica porter\u00e0 a "
    "scoprire che il fornitore del software non ha ancora affrontato il tema, il che \u00e8 gi\u00e0 "
    "un\u2019informazione importante da trasferire al cliente."
)

add_heading2(doc, "Cosa fare nei prossimi quattro mesi")

add_paragraph(doc,
    "Quattro mesi sono pochi per costruire un sistema di compliance completo da zero. Sono sufficienti "
    "per fare le cose pi\u00f9 urgenti. La prima \u00e8 un censimento interno: quali sistemi AI vengono "
    "usati in azienda, per quali decisioni, su quali soggetti. La seconda \u00e8 verificare, per ciascun "
    "sistema, se il fornitore ha avviato il processo di certificazione AI Act e quali documenti sono "
    "gi\u00e0 disponibili. La terza \u00e8 identificare un responsabile interno che conosca il tema e "
    "possa fare da interlocutore con consulenti esterni e autorit\u00e0 di vigilanza."
)

add_paragraph(doc,
    "In Italia, le autorit\u00e0 nazionali competenti per la vigilanza sull\u2019AI Act sono l\u2019Agenzia "
    "per l\u2019Italia Digitale (AGID) e l\u2019Agenzia per la Cybersicurezza Nazionale (ACN), mentre "
    "il Garante Privacy mantiene le proprie competenze sui profili di trattamento dei dati personali. "
    "Un\u2019impresa che affronta oggi il censimento dei propri sistemi AI, anche in modo informale, "
    "si mette in una posizione molto pi\u00f9 difendibile rispetto ad agosto 2026 di una che aspetta."
)

add_separator(doc)

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(10)
run_n = p_note.add_run("Fonti e riferimenti")
run_n.font.name = 'Calibri'
run_n.font.size = Pt(9)
run_n.font.bold = True
run_n.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

sources = [
    "Regolamento (UE) 2024/1689 \u2014 AI Act (GU UE, 12 luglio 2024)",
    "Legge 23 settembre 2025, n. 132 \u2014 Disposizioni in materia di intelligenza artificiale",
    "ISTAT \u2014 Report sull\u2019adozione AI nelle imprese italiane, dicembre 2025",
    "AGID / ACN \u2014 Comunicazione sulle autorit\u00e0 nazionali competenti AI Act",
    "Tom\u2019s Hardware Business \u2014 Agenti AI e nuove regolamentazioni, 2026",
    "Agenda Digitale \u2014 AI Act 2026: obblighi per le imprese",
]
for s in sources:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_after = Pt(2)
    run_s = p_s.add_run(f"\u2022 {s}")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(8.5)
    run_s.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(16)
run_f = p_footer.add_run(
    "\u00a9 2026 Ratio  \u2022  Riproduzione consentita con citazione della fonte"
)
run_f.font.name = 'Calibri'
run_f.font.size = Pt(8)
run_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_f.italic = True

doc.save(OUTPUT_FILE)
print(f"Salvato: {OUTPUT_FILE}")
