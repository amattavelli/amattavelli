"""
Quattro articoli Ratio — 29 maggio 2026

1. 2026-05-29_digital-omnibus-scadenze-agosto-2026.docx
2. 2026-05-29_dlgs-47-2026-governance-digitale-imprese.docx
3. 2026-05-29_gap-ai-pmi-italiane-partire-da-dove.docx
4. 2026-05-29_ai-literacy-obbligo-formazione-imprese.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/amattavelli/Ratio/articoli/"

BLU_SCURO  = RGBColor(0x1F, 0x49, 0x7D)
BLU_MEDIO  = RGBColor(0x2E, 0x74, 0xB5)
GRIGIO     = RGBColor(0x60, 0x60, 0x60)
GRIGIO_CH  = RGBColor(0x80, 0x80, 0x80)
GRIGIO_PI  = RGBColor(0xA0, 0xA0, 0xA0)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(3)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    sn = doc.styles["Normal"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bot   = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = BLU_SCURO


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def testata(doc, mese_anno, categoria):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = BLU_SCURO
    r.font.bold = True; r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{mese_anno}  |  {categoria}")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.font.italic = True; r.font.color.rgb = GRIGIO
    doc.add_paragraph()


def titolo(doc, titolo_testo, occhiello):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo_testo)
    r.font.name = "Calibri"; r.font.size = Pt(24)
    r.font.bold = True; r.font.color.rgb = BLU_SCURO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(occhiello)
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.italic = True; r.font.color.rgb = BLU_MEDIO
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("A cura della Redazione Ratio  •  29 maggio 2026")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = GRIGIO_CH


def riferimenti(doc, fonti):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = GRIGIO
    for s in fonti:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = "Calibri"; r.font.size = Pt(8.5)
        r.font.color.rgb = GRIGIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = "Calibri"; r.font.size = Pt(8)
    r.font.color.rgb = GRIGIO_PI; r.font.italic = True


# ============================================================
# ARTICOLO 1
# Il Digital Omnibus ha slittato le scadenze. Non tutte.
# ============================================================

doc = new_doc()
testata(doc, "Maggio 2026", "Normativa AI")
titolo(
    doc,
    "Il Digital Omnibus ha slittato\nle scadenze. Non tutte.",
    "L'accordo del 7 maggio 2026 sposta al 2027 le regole sui sistemi ad alto rischio. "
    "Trasparenza, divieti e AI Literacy restano fermi ad agosto."
)

para(doc,
    "Nelle ultime settimane si è diffusa nelle imprese italiane una lettura parzialmente sbagliata "
    "del pacchetto normativo che l'Unione Europea ha chiamato Digital Omnibus. La notizia vera, "
    "quella che ha circolato di più, è che il 7 maggio 2026 il Parlamento europeo e il Consiglio "
    "hanno raggiunto un accordo provvisorio per spostare in avanti le scadenze più impegnative "
    "dell'AI Act. La notizia che si è persa per strada è che non tutte le scadenze hanno slittato. "
    "Alcune restano esattamente dove erano."
)

para(doc,
    "Il rinvio riguarda le norme sui sistemi di intelligenza artificiale classificati ad alto rischio: "
    "quelle regole, tra cui la documentazione tecnica, la supervisione umana obbligatoria e la "
    "registrazione nei database europei, vengono spostate al 2 dicembre 2027 e al 2 agosto 2028. "
    "Per le PMI italiane che non sviluppano sistemi AI ma li utilizzano come utenti finali, "
    "questo rinvio è concreto e significativo. Ma non è l'unica data in calendario."
)

heading(doc, "Cosa resta fermo al 2 agosto 2026")

para(doc,
    "Il Digital Omnibus non modifica tre categorie di obblighi che erano già entrati nel quadro "
    "normativo e che rimangono applicabili dalla data originaria. Il primo è il divieto assoluto "
    "di determinati usi dell'intelligenza artificiale, quelli che l'AI Act classifica come sistemi "
    "a rischio inaccettabile: la profilazione comportamentale di massa, il riconoscimento delle "
    "emozioni nei luoghi di lavoro, la manipolazione cognitiva attraverso tecniche subliminali. "
    "Questi divieti sono in vigore dal 2 febbraio 2025."
)

para(doc,
    "Il secondo gruppo riguarda gli obblighi di trasparenza dell'articolo 50. Dal 2 agosto 2026, "
    "qualsiasi sistema che interagisce con le persone in forma testuale o vocale deve dichiarare "
    "in modo chiaro che si tratta di un sistema artificiale. Un chatbot sul sito di uno studio "
    "professionale, un assistente virtuale per la gestione delle prenotazioni, un sistema di "
    "risposta automatica alle email dei clienti: tutti devono comunicare in modo inequivocabile "
    "la propria natura. Parallelamente, i contenuti generati con strumenti AI e destinati a "
    "circolare come documenti o comunicazioni devono essere etichettati come tali. Non è una "
    "formalità: l'ACN, l'Agenzia per la Cybersicurezza Nazionale designata dalla Legge 132/2025 "
    "come autorità di vigilanza italiana, può applicare sanzioni significative."
)

para(doc,
    "Il terzo obbligo che resta intatto è quello sull'AI Literacy, cioè sull'alfabetizzazione "
    "minima di chiunque utilizzi strumenti AI nell'ambito lavorativo. Su questo punto si tornerà "
    "in dettaglio, perché merita un ragionamento separato. Qui basta notare che l'accordo "
    "Digital Omnibus non ha toccato questo requisito."
)

heading(doc, "L'equivoco del rinvio")

para(doc,
    "L'equivoco nasce da un dato di fatto: la parte più complessa dell'AI Act, quella che richiedeva "
    "analisi del rischio dettagliate, documentazione tecnica e adeguamenti organizzativi profondi, "
    "è quella che ha slittato. È comprensibile che molte imprese abbiano interpretato questo come "
    "un segnale di allentamento generale. Il messaggio implicito che ha circolato nei mesi scorsi, "
    "quello per cui \"l'AI Act è stato rinviato\", non è falso ma è incompleto e rischia di "
    "diventare un problema per chi lo prende alla lettera."
)

para(doc,
    "Un'azienda che ha sul sito un chatbot non etichettato, che produce comunicazioni AI-generated "
    "senza marcatura, o che non ha fatto nulla per garantire una consapevolezza minima ai propri "
    "dipendenti sull'uso degli strumenti AI, non è coperta dal rinvio Digital Omnibus. Questi tre "
    "elementi erano nella parte dell'AI Act che non è stata modificata, e la scadenza rimane il "
    "2 agosto 2026. A questo si aggiunge la Legge 132 del 10 ottobre 2025, la normativa italiana "
    "che ha introdotto obblighi aggiuntivi per i professionisti e le imprese che operano nel Paese, "
    "a prescindere dall'evoluzione del quadro europeo."
)

heading(doc, "Cosa fare da qui ad agosto")

para(doc,
    "Per uno studio professionale o una PMI che vuole essere in regola entro agosto, il percorso "
    "minimo ha tre passi. Il primo è verificare se nell'attività quotidiana vengono usati strumenti "
    "AI che interagiscono con i clienti: chatbot, risponditori automatici, sistemi di analisi "
    "documentale che producono output verso l'esterno. Se sì, va valutato se sono dichiarati "
    "correttamente. Il secondo è capire se i documenti prodotti con strumenti AI, pareri, report, "
    "comunicazioni, vengono inviati ai clienti senza alcuna indicazione sulla loro origine. "
    "Il terzo è capire quale livello di consapevolezza hanno i dipendenti e i collaboratori "
    "sugli strumenti che già usano, o che potrebbero usare, nel loro lavoro."
)

para(doc,
    "Il Digital Omnibus ha guadagnato tempo sulla parte più strutturata degli adempimenti. "
    "Ma ha lasciato intatte quelle che, per la maggior parte delle PMI e degli studi professionali "
    "italiani, sono le prime cose da sistemare. E agosto è tra nove settimane."
)

riferimenti(doc, [
    "Parlamento europeo e Consiglio UE — Accordo provvisorio Digital Omnibus, 7 maggio 2026",
    "Regolamento UE 2024/1689 (AI Act) — artt. 4, 5, 50",
    "Legge 10 ottobre 2025, n. 132 — Disciplina organica dell'intelligenza artificiale in Italia",
    "Agenda Digitale — \"Digital Omnibus e AI Act: cosa può slittare e cosa va fatto prima di agosto 2026\"",
    "ACN — Agenzia per la Cybersicurezza Nazionale, autorità nazionale di vigilanza AI",
])
doc.save(BASE + "2026-05-29_digital-omnibus-scadenze-agosto-2026.docx")
print("Salvato: articolo 1")


# ============================================================
# ARTICOLO 2
# D.Lgs. 47/2026: quando la governance digitale entra in bilancio
# ============================================================

doc = new_doc()
testata(doc, "Maggio 2026", "Normativa e Governance")
titolo(
    doc,
    "Il D.Lgs. 47/2026 mette\nla cybersecurity nell'organigramma",
    "Dal 29 aprile 2026 AI, privacy e sicurezza informatica sono materia degli amministratori, "
    "non dell'IT. Per chi assiste le imprese, cambia il perimetro da valutare."
)

para(doc,
    "Dal 29 aprile 2026 è in vigore il D.Lgs. 27 marzo 2026, n. 47. Per molti studi professionali "
    "è ancora un testo sullo sfondo, uno di quei provvedimenti che si rimanda a quando arriva la "
    "circolare. Sarebbe un errore, perché questo decreto cambia in modo diretto il perimetro di "
    "responsabilità degli organi societari su tre temi che fino a ieri erano considerati materia "
    "tecnica: cybersecurity, intelligenza artificiale e protezione dei dati. E quando cambia "
    "il perimetro di responsabilità degli amministratori, cambia anche il perimetro di valutazione "
    "del commercialista o del consulente che affianca l'impresa."
)

para(doc,
    "Il decreto aggiorna il codice civile e il Testo Unico della Finanza per portare questi tre "
    "temi all'interno degli assetti organizzativi che le società di capitali sono tenute ad avere "
    "ai sensi dell'articolo 2086. Fino ad oggi, la verifica di adeguatezza di quegli assetti "
    "riguardava principalmente l'organizzazione contabile e finanziaria, la struttura decisionale, "
    "il sistema di controllo interno. Da fine aprile, quell'assetto deve includere anche la gestione "
    "del rischio digitale, dove per digitale si intende l'insieme di cybersecurity, intelligenza "
    "artificiale e protezione dei dati personali."
)

heading(doc, "Il collegio sindacale come quarta linea di difesa")

para(doc,
    "Una delle novità più rilevanti del D.Lgs. 47/2026 riguarda il collegio sindacale. "
    "Il decreto gli assegna esplicitamente un ruolo di presidio sulla conformità alle norme "
    "NIS2 e GDPR, affiancandolo ai meccanismi di controllo interno già esistenti come audit, "
    "risk management e compliance. In alcune letture che stanno circolando tra gli addetti ai "
    "lavori si parla di \"quarta linea di difesa\". La terminologia è discutibile, ma il "
    "messaggio sostanziale è chiaro: il collegio non può più limitarsi ai conti. Se l'impresa "
    "usa sistemi di intelligenza artificiale in processi rilevanti, il collegio è tenuto a "
    "verificare che esista una governance di quel rischio."
)

para(doc,
    "Per un sindaco di una società di medie dimensioni, questo significa che la verifica "
    "periodica deve ora includere domande che non erano nell'agenda tradizionale. L'azienda "
    "ha mappato quali strumenti AI utilizza nei propri processi? C'è una policy interna? "
    "Qualcuno ha verificato se quegli strumenti rientrano nelle categorie regolamentate dall'AI "
    "Act? Il risparmio di tempo ottenuto con l'automazione vale il rischio di non aver fatto "
    "questa valutazione? Queste non sono domande da lasciare al reparto IT."
)

heading(doc, "Cosa cambia per chi assiste le imprese")

para(doc,
    "Per il commercialista o il consulente che segue una PMI o una società di capitali, "
    "il D.Lgs. 47/2026 apre un tema concreto. La verifica dell'adeguatezza degli assetti "
    "organizzativi, quella che si fa tipicamente in occasione della revisione del sistema "
    "di controllo interno o della due diligence su operazioni straordinarie, ha ora un "
    "capitolo in più. Non necessariamente lungo, ma necessariamente presente."
)

para(doc,
    "Il primo passo pratico è capire in quali aree dell'impresa vengono già utilizzati "
    "strumenti di intelligenza artificiale. Spesso la risposta sorprende: software gestionali "
    "aggiornati negli ultimi due anni incorporano funzioni AI per la categorizzazione delle "
    "fatture, per la previsione di flussi di cassa, per la generazione automatica di report. "
    "Il fatto che non siano stati introdotti con una scelta esplicita non li esclude dall'ambito "
    "del decreto. Anzi, il caso più rischioso è proprio quello in cui gli strumenti AI sono "
    "entrati dall'aggiornamento software senza che nessuno in azienda se ne sia accorto."
)

para(doc,
    "Il secondo passo è verificare se l'impresa ha adottato una policy anche minima sull'uso "
    "dell'intelligenza artificiale. La Legge 132/2025 e il D.Lgs. 47/2026 non richiedono "
    "documenti di centinaia di pagine. Richiedono che l'organizzazione abbia consapevolezza "
    "degli strumenti che usa e che abbia identificato chi è responsabile delle decisioni che "
    "derivano da quell'uso. Una pagina ben scritta vale più di un manuale che nessuno ha letto."
)

heading(doc, "Il passaggio dalla compliance alla governance")

para(doc,
    "Il cambiamento più profondo che questo decreto introduce non è nella lista degli adempimenti "
    "formali. È nel cambio di prospettiva che richiede. Fino ad oggi, cybersecurity e protezione "
    "dei dati erano argomenti da delegare a un consulente specializzato, da trattare come "
    "compliance separata dalla gestione ordinaria dell'impresa. Il D.Lgs. 47/2026 dice che "
    "questa separazione non regge più. Se la società usa strumenti AI per prendere decisioni "
    "che impattano sul cliente, sul dipendente o sull'organizzazione, quella scelta è materia "
    "degli amministratori, non di un fornitore esterno."
)

para(doc,
    "Per chi affianca le imprese come consulente, il momento più utile per sollevare questo "
    "tema non è la scadenza normativa, ma il prossimo incontro operativo con il cliente. "
    "La domanda da fare non è se l'impresa usa l'AI, perché quasi certamente la risposta è già "
    "sì anche senza saperlo. La domanda è chi, in azienda, ha la responsabilità di sapere cosa "
    "sta usando e perché."
)

riferimenti(doc, [
    "D.Lgs. 27 marzo 2026, n. 47 — in vigore dal 29 aprile 2026",
    "Federprivacy — \"Dalla compliance alla governance: il Dlgs 47/2026 e la trasformazione di cybersecurity, AI e protezione dei dati negli assetti societari\"",
    "Agenda Digitale — \"D.Lgs. 47/2026, perché privacy e cybersecurity entrano nella governance delle imprese\"",
    "Actainfo — \"Il Dlgs 47/2026: un cambio di paradigma per la governance digitale delle imprese\"",
    "Art. 2086 c.c. — Gestione dell'impresa e adeguatezza degli assetti organizzativi",
])
doc.save(BASE + "2026-05-29_dlgs-47-2026-governance-digitale-imprese.docx")
print("Salvato: articolo 2")


# ============================================================
# ARTICOLO 3
# L'azienda accanto a te usa l'AI. La tua probabilmente no.
# ============================================================

doc = new_doc()
testata(doc, "Maggio 2026", "Adozione AI nelle PMI")
titolo(
    doc,
    "L'azienda accanto a te\nusa l'AI. La tua probabilmente no.",
    "Il 71% delle grandi imprese ha avviato almeno un progetto AI. "
    "Meno del 10% delle PMI italiane ha fatto lo stesso. Il divario si allarga ogni mese."
)

para(doc,
    "Immagina di gestire un'azienda che produce componenti industriali. Un tuo cliente ti chiede "
    "una disponibilità su un ordine urgente. Ricevi la risposta del tuo concorrente diretto in "
    "meno di due ore, con preventivo, stima di consegna e proposta di alternativa tecnica inclusa. "
    "La tua risposta arriva il giorno dopo, perché il responsabile commerciale era in visita a un "
    "altro cliente e hai dovuto aspettare che rientrasse per elaborare i dati. Non hai perso "
    "quell'ordine per una questione di prezzo o di qualità. L'hai perso per una questione di tempo. "
    "Il concorrente che ti ha battuto ha un agente AI collegato al suo gestionale."
)

para(doc,
    "Questo scenario non è ipotetico. Lo raccontano imprenditori di tutta Italia che, in questo "
    "primo semestre del 2026, stanno scoprendo che il divario nell'adozione dell'intelligenza "
    "artificiale tra le grandi imprese e le PMI non riguarda solo la produttività interna. "
    "Inizia a riguardare la competitività commerciale. Secondo i dati disponibili, il 71% delle "
    "grandi imprese italiane ha avviato almeno un progetto di intelligenza artificiale. La quota "
    "scende sotto il 10% per le PMI. ISTAT rilevava nel 2025 che solo il 16% delle imprese con "
    "almeno dieci dipendenti usava soluzioni AI, e appena il 7% tra quelle più piccole aveva "
    "avviato qualcosa di concreto."
)

heading(doc, "Perché le PMI aspettano")

para(doc,
    "Le ragioni del ritardo sono diverse e spesso si sommano. La prima è la percezione del costo: "
    "molti imprenditori associano l'intelligenza artificiale a investimenti infrastrutturali "
    "importanti, a consulenze specialistiche costose, a progetti pilota che richiedono mesi. "
    "Questa percezione è rimasta ancorata a una realtà che è cambiata. Oggi esistono strumenti "
    "AI accessibili con abbonamenti mensili nell'ordine delle decine di euro, integrabili con "
    "i software gestionali già in uso senza richiedere implementazioni complesse."
)

para(doc,
    "La seconda ragione è la mancanza di una persona interna che si assuma la responsabilità "
    "di introdurre lo strumento. Nelle grandi aziende c'è spesso un responsabile digitale, "
    "un CTO, qualcuno che ha questo nel mandato. In una PMI con venti dipendenti, nessuno ha "
    "formalmente il compito di capire cosa fa la concorrenza con l'AI e di tradurlo in un "
    "processo interno. La responsabilità diffusa diventa responsabilità di nessuno."
)

para(doc,
    "La terza ragione è l'incertezza normativa. Molti imprenditori hanno sentito parlare "
    "dell'AI Act, delle sanzioni, della compliance, e hanno dedotto che aspettare fosse la "
    "scelta più prudente. L'effetto paradossale è che aspettare di capire le regole ha "
    "lasciato campo libero a chi ha deciso di muoversi comunque, magari con meno attenzione "
    "alla conformità ma con un vantaggio operativo reale."
)

heading(doc, "Cosa stanno facendo le PMI che si sono mosse")

para(doc,
    "Le piccole imprese che hanno avviato progetti AI non sono partite da grandi trasformazioni "
    "digitali. Sono partite da un processo ripetitivo che consumava tempo. Il risparmio di tempo "
    "è la metrica più immediata e la più convincente per chi deve giustificare l'investimento. "
    "Il Politecnico di Milano ha stimato che il 40% del tempo del reparto amministrativo di una "
    "PMI media è occupato da attività che potrebbero essere automatizzate con strumenti già "
    "disponibili: categorizzazione delle fatture, riconciliazioni bancarie, preparazione di "
    "report periodici, risposta a email standardizzate."
)

para(doc,
    "Un caso tipico è quello di uno studio di consulenza che ha introdotto un sistema AI per "
    "la gestione delle richieste iniziali dei clienti. Il sistema legge la email, classifica "
    "il tipo di richiesta, estrae le informazioni rilevanti e prepara una bozza di risposta "
    "che il consulente rivede e invia. Il tempo medio per gestire una richiesta standard è "
    "passato da venti minuti a tre. Non è stato necessario cambiare software: il sistema si "
    "integrava con il client di posta già in uso. Il costo mensile era inferiore a quello "
    "di un abbonamento a un quotidiano professionale."
)

heading(doc, "Da dove partire")

para(doc,
    "Per una PMI o uno studio professionale che vuole capire da dove iniziare, il punto di "
    "partenza non è scegliere lo strumento. Il punto di partenza è identificare il processo "
    "più costoso in termini di tempo manuale, ripetitivo, a basso valore decisionale. "
    "Quel processo è il candidato naturale per il primo esperimento con l'AI. Non un progetto "
    "di trasformazione, non un cambio di piattaforma: un singolo processo, uno strumento, "
    "tre mesi di misura. Se funziona si allarga. Se non funziona si cambia approccio senza "
    "aver perso molto."
)

para(doc,
    "Il divario tra grandi imprese e PMI nell'adozione dell'AI non si chiude con un convegno "
    "o con una guida normativa. Si chiude quando un imprenditore smette di chiedersi se l'AI "
    "funzioni e inizia a chiedersi quale specifica attività della sua azienda smette di farla "
    "fare a una persona."
)

riferimenti(doc, [
    "ISTAT — Indagine sull'uso dell'intelligenza artificiale nelle imprese italiane, 2025",
    "Osservatorio Digital Innovation, Politecnico di Milano — Automazione dei processi amministrativi nelle PMI",
    "Best Tech Partner — \"Automazione lavoro IA: liberare tempo e valore nelle PMI italiane\" (27 maggio 2026)",
    "CUENEWS / ManagementCUE — \"Gli agenti AI in azienda: cosa cambierà davvero nel 2026 per le PMI italiane\"",
    "Ivemind — \"Agenti AI per PMI: Guida Pratica 2026\"",
])
doc.save(BASE + "2026-05-29_gap-ai-pmi-italiane-partire-da-dove.docx")
print("Salvato: articolo 3")


# ============================================================
# ARTICOLO 4
# L'AI Literacy non è un corso. È un obbligo che scade ad agosto.
# ============================================================

doc = new_doc()
testata(doc, "Maggio 2026", "Formazione e Compliance")
titolo(
    doc,
    "L'AI Literacy non è un corso.\nScade ad agosto.",
    "L'articolo 4 dell'AI Act impone a qualsiasi impresa che usa strumenti AI di garantire "
    "competenze adeguate al proprio personale. Il termine è il 2 agosto 2026. "
    "La maggior parte delle imprese non ha ancora fatto nulla."
)

para(doc,
    "Proviamo a partire da una situazione concreta. Un'impresa di quindici persone usa "
    "ChatGPT per redigere offerte commerciali, rispondere alle email dei clienti e riassumere "
    "i verbali delle riunioni. Lo usano quattro collaboratori, senza una formazione specifica, "
    "senza una policy interna, senza che nessuno abbia verificato se quello che producono "
    "sia corretto prima di inviarlo. Il titolare sa che \"usano l'AI\" e ritiene di aver fatto "
    "la cosa giusta portando lo strumento in azienda. Dal 2 agosto 2026, quella situazione "
    "configura una violazione dell'articolo 4 dell'AI Act, il quale prevede che chiunque "
    "dispieghi sistemi AI nel proprio contesto lavorativo assicuri che il personale abbia "
    "un livello adeguato di AI Literacy."
)

para(doc,
    "La data non è negoziabile e non è stata toccata dal Digital Omnibus. Il rinvio che "
    "l'accordo del 7 maggio ha introdotto riguarda i sistemi ad alto rischio: documentazione "
    "tecnica, supervisione obbligatoria, registrazione europea. L'obbligo di alfabetizzazione "
    "minima era in una parte diversa del regolamento e non è stato modificato. La Legge "
    "132/2025 ha poi introdotto obblighi aggiuntivi a livello nazionale, collegando la "
    "formazione sull'AI anche alle normative sulla sicurezza sul lavoro e sui diritti dei "
    "lavoratori nell'uso degli strumenti algoritmici."
)

heading(doc, "Cosa significa AI Literacy nella pratica")

para(doc,
    "Uno degli equivoci più diffusi sull'obbligo di AI Literacy è che si tratti di formare "
    "tutti i dipendenti come esperti di machine learning. La norma non dice questo. Dice che "
    "chi usa strumenti AI nel proprio lavoro deve avere consapevolezza sufficiente a usarli "
    "in modo responsabile. Il livello richiesto dipende dal ruolo: chi usa un chatbot per "
    "rispondere ai clienti deve capire che il sistema può sbagliare, che non va usato per "
    "comunicazioni che richiedono una verifica critica, che i dati dei clienti non vanno "
    "inseriti in strumenti non valutati dal punto di vista della privacy. Chi firma documenti "
    "prodotti con il supporto dell'AI deve sapere cosa sta firmando e perché ne è responsabile."
)

para(doc,
    "Per un'impresa con cinque o dieci dipendenti, questo può tradursi in una sessione di "
    "due ore con qualcuno che conosce gli strumenti usati dall'azienda, seguita da un "
    "documento sintetico con le regole d'uso. Per uno studio professionale con personale che "
    "usa strumenti AI per la gestione dei clienti, il livello atteso è più alto: capire la "
    "differenza tra un output da usare come bozza e uno da firmare direttamente, riconoscere "
    "i segnali di una risposta plausibile ma sbagliata, sapere quando è obbligatorio verificare "
    "la fonte. Non è un corso universitario. È la stessa responsabilità che si chiede a chi "
    "usa un software gestionale: sapere cosa fa e sapere dove può sbagliare."
)

heading(doc, "La documentazione che serve")

para(doc,
    "La Legge 132/2025 e le indicazioni attuative della Commissione europea convergono su un "
    "punto: l'obbligo di AI Literacy non si assolve con una dichiarazione interna. Serve "
    "documentazione tracciabile. Questo significa registri di formazione, datati e firmati, "
    "con indicazione degli strumenti AI oggetto della formazione e del contenuto trattato. "
    "Significa anche che la formazione non può essere generica: \"corso sull'intelligenza "
    "artificiale\" non basta se l'azienda usa specificamente uno strumento per la gestione "
    "delle comunicazioni commerciali. La formazione deve essere adeguata all'uso concreto."
)

para(doc,
    "Per le imprese con dipendenti, questo entra anche nel quadro dei diritti dei lavoratori. "
    "Il D.Lgs. 47/2026 ha rafforzato il collegamento tra AI governance e struttura organizzativa, "
    "il che significa che un dipendente che subisce conseguenze lavorative da decisioni prese "
    "con il supporto di sistemi AI ha diritto a sapere che quello strumento era in uso, "
    "con quali criteri funzionava e quali verifiche erano previste. L'obbligo di trasparenza "
    "interna non è separato dall'obbligo di formazione: sono due facce dello stesso adempimento."
)

heading(doc, "Il rischio di arrivare ad agosto senza averlo fatto")

para(doc,
    "L'ACN può avviare procedure di verifica su segnalazione o d'ufficio. Le sanzioni per "
    "la violazione degli obblighi di trasparenza e formazione non raggiungono le cifre "
    "previste per i sistemi ad alto rischio, ma non sono simboliche. Per una PMI con meno "
    "di cinquanta dipendenti, una sanzione anche minima ha un impatto reputazionale e "
    "operativo che non si risolve in fretta."
)

para(doc,
    "C'è però un rischio più immediato di quello sanzionatorio. Un collaboratore che usa "
    "ChatGPT per preparare una comunicazione commerciale e non sa che il modello può "
    "generare numeri plausibili ma sbagliati ha già creato un problema, indipendentemente "
    "da agosto. La formazione obbligatoria non è pensata per proteggere l'azienda "
    "dall'ispezione: è pensata per far sì che gli strumenti AI producano valore invece "
    "di rischi. Trattarla come un adempimento burocratico da archiviare è il modo più "
    "sicuro per sprecare sia il tempo della formazione che i vantaggi dello strumento."
)

riferimenti(doc, [
    "Regolamento UE 2024/1689 (AI Act) — art. 4: Alfabetizzazione in materia di AI",
    "Legge 10 ottobre 2025, n. 132 — obblighi di formazione e trasparenza algoritmica",
    "D.Lgs. 47/2026 — governance digitale e diritti dei lavoratori",
    "Commissione europea — FAQ su AI Literacy: chiarimenti attuativi 2026",
    "Antonio Sinibaldi — \"AI literacy e obbligo formativo: cosa fare entro agosto 2026\"",
    "Randstad — \"AI literacy e obbligo di formazione nelle aziende: come muoversi?\"",
])
doc.save(BASE + "2026-05-29_ai-literacy-obbligo-formazione-imprese.docx")
print("Salvato: articolo 4")

print("\nTutti e 4 gli articoli generati in:", BASE)
