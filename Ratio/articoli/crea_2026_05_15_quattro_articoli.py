"""
Crea 4 articoli Ratio — 15 maggio 2026
Temi: Digital Omnibus / AI Act, GPT-5.5, AI verticale commercialisti, gap adozione PMI
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ────────────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.left_margin = Cm(3)
    s.right_margin = Cm(3)
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    return doc


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)


def testata(doc, sezione, data_label, titolo, sottotitolo, data_autore):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATIO  •  Approfondimenti per Professionisti e Imprese")
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r.font.bold = True
    r.font.all_caps = True
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{data_label}  |  {sezione}")
    r.font.name = 'Calibri'
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titolo)
    r.font.name = 'Calibri'
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(sottotitolo)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    sep(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(f"A cura della Redazione Ratio  •  {data_autore}")
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def footer(doc, refs):
    sep(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Riferimenti")
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    for s in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {s}")
        r.font.name = 'Calibri'
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "© 2026 Mattavelli Amodeo — Commercialisti Associati  •  "
        "Riproduzione consentita con citazione della fonte"
    )
    r.font.name = 'Calibri'
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
    r.font.italic = True


# ═══════════════════════════════════════════════════════════════════════════
# ARTICOLO 1 — Digital Omnibus e AI Act
# ═══════════════════════════════════════════════════════════════════════════

def crea_articolo_1():
    doc = new_doc()
    testata(
        doc,
        sezione="Normativa e Compliance",
        data_label="Maggio 2026",
        titolo=(
            "Il rinvio dell’AI Act non è un salvacondotto:\n"
            "cosa resta in piedi da agosto 2026"
        ),
        sottotitolo=(
            "L’accordo del 7 maggio tra Parlamento e Consiglio UE ha spostato le scadenze "
            "per i sistemi ad alto rischio al 2027 e al 2028. Ma l’obbligo di formazione "
            "resta invariato, e agosto arriva comunque."
        ),
        data_autore="15 maggio 2026",
    )

    para(doc,
        "Nelle ultime settimane molte aziende italiane hanno rallentato il lavoro di adeguamento "
        "all’AI Act aspettando notizie da Bruxelles. Il 7 maggio, nella notte tra martedì "
        "e mercoledì, Parlamento europeo e Consiglio dell’Unione europea hanno raggiunto "
        "un accordo provvisorio nel quadro del pacchetto Digital Omnibus: le scadenze per i sistemi "
        "di intelligenza artificiale classificati ad alto rischio vengono spostate in avanti di oltre "
        "un anno. Gli obblighi per i sistemi autonomi slittano dal 2 agosto 2026 al 2 dicembre 2027; "
        "quelli per i sistemi integrati in prodotti soggetti a normativa settoriale si spostano al "
        "2 agosto 2028. In molti hanno letto la notizia come una boccata d’ossigeno. "
        "Conviene però capire bene cosa si è spostato e cosa è rimasto fermo."
    )

    para(doc,
        "Il rinvio riguarda una parte specifica del Regolamento: gli obblighi che si applicano ai "
        "sistemi di IA classificati ad alto rischio nell’Allegato III (biometria, infrastrutture "
        "critiche, selezione del personale, merito creditizio, istruzione, settore legale) e "
        "nell’Allegato I (prodotti soggetti a normativa di settore, come dispositivi medici o "
        "macchinari industriali). Sono sistemi rilevanti per certi comparti, ma la maggior parte degli "
        "studi professionali e delle PMI italiane non li utilizza nella forma tecnica che il Regolamento "
        "intende con quel termine. Per loro, il rinvio cambia poco o nulla nella situazione di agosto 2026."
    )

    para(doc,
        "Quello che rimane in vigore ad agosto 2026 riguarda la formazione, i divieti assoluti già "
        "operativi dal 2 febbraio 2025, e la marcatura dei contenuti generati dall’AI. L’articolo "
        "4 del Regolamento, che impone a chiunque utilizzi sistemi di intelligenza artificiale in un "
        "contesto professionale di garantire un livello sufficiente di comprensione delle loro "
        "caratteristiche, dei loro limiti e dei rischi che comportano, non è stato toccato dal "
        "Digital Omnibus. Entra in vigore il 2 agosto 2026, senza rinvii."
    )

    heading(doc, "Cosa cambia davvero con il Digital Omnibus")

    para(doc,
        "L’accordo del 7 maggio introduce alcune novità concrete che vanno oltre il semplice "
        "slittamento di date. Le esenzioni per le PMI vengono estese anche alle mid-cap (imprese fino a "
        "750 dipendenti e 150 milioni di fatturato), riducendo il perimetro degli obblighi più "
        "gravosi per una fascia più ampia di imprese. Viene introdotto un divieto esplicito, con "
        "scadenza dicembre 2026, per i sistemi progettati per generare contenuti sessuali non consensuali "
        "o materiale che coinvolge minori, una misura attesa da più parti. Il periodo di tolleranza "
        "per la marcatura dei contenuti generati dall’AI (immagini, audio, video) viene accorciato "
        "da sei a tre mesi, con scadenza al 2 dicembre 2026."
    )

    para(doc,
        "Per le aziende che si erano già strutturate in vista di agosto 2026, il rinvio sugli "
        "Allegati I e III lascia tempo per completare i percorsi di adeguamento con meno pressione. "
        "Per quelle che non avevano ancora iniziato, il rischio concreto è che il rinvio diventi "
        "una giustificazione per non fare niente, e che si arrivi alle nuove scadenze del 2027 e del "
        "2028 nelle stesse condizioni di oggi."
    )

    heading(doc, "L’obbligo che nessun rinvio tocca")

    para(doc,
        "La AI literacy è l’obbligo che più tocca direttamente la quotidianità "
        "di uno studio professionale o di un ufficio amministrativo. Il Regolamento non richiede "
        "certificazioni formali, ma chiede che chiunque lavori con sistemi di IA abbia una comprensione "
        "proporzionata alla propria funzione: un responsabile che valida l’output di un sistema "
        "automatizzato ha bisogno di una comprensione diversa da quella di chi usa un assistente di "
        "scrittura per redigere comunicazioni commerciali. La differenza non è banale, e l’assenza "
        "di percorsi formativi documentati può diventare una vulnerabilità in caso di controllo."
    )

    para(doc,
        "In Italia, le autorità competenti designate ai sensi dell’AI Act sono l’Agenzia "
        "per l’Italia Digitale (AGID) e l’Agenzia per la Cybersicurezza Nazionale (ACN), "
        "affiancate dal Garante Privacy per i profili connessi al trattamento dei dati personali. "
        "La Legge 132/2025, il testo italiano di recepimento entrato in vigore il 10 ottobre 2025, "
        "ha introdotto disposizioni specifiche sulla responsabilità professionale che si sommano "
        "agli obblighi del Regolamento europeo. Avere documentato l’attività formativa svolta "
        "e le policy interne sull’uso degli strumenti AI è già una forma concreta di "
        "compliance, indipendentemente dalla classificazione del sistema usato."
    )

    heading(doc, "Il vantaggio di chi continua a prepararsi")

    para(doc,
        "Le aziende che continuano il lavoro di adeguamento pur avendo guadagnato tempo sui sistemi "
        "ad alto rischio si trovano in una posizione migliore su più fronti. Consolidano la "
        "governance interna sull’uso degli strumenti AI, che è utile indipendentemente dalle "
        "scadenze normative. Costruiscono competenze nel team che rendono l’adozione di nuovi "
        "strumenti più rapida e sicura. Documentano i processi in modo che sarà più "
        "facile aggiornare quando le scadenze del 2027 e del 2028 si avvicineranno."
    )

    para(doc,
        "Il Digital Omnibus ha dato più tempo su una parte degli obblighi. Non ha tolto la "
        "necessità di avere una posizione chiara su come l’intelligenza artificiale entra "
        "nel lavoro quotidiano, chi la controlla, e cosa succede quando sbaglia. Queste domande non "
        "hanno una scadenza normativa: hanno una rilevanza pratica che precede qualsiasi obbligo di legge."
    )

    footer(doc, [
        "Regolamento (UE) 2024/1689 — AI Act, come modificato dal Digital Omnibus",
        "Accordo provvisorio Parlamento UE / Consiglio dell’UE, 7 maggio 2026",
        "agendadigitale.eu — “Accordo su Digital Omnibus, così l’Europa rinvia gli obblighi AI Act” (7 maggio 2026)",
        "avvera.it — “AI Act Omnibus: l’accordo del 7 maggio tra Parlamento e Consiglio” (maggio 2026)",
        "Legge 132/2025 — Disciplina organica sull’intelligenza artificiale in Italia",
    ])

    out = "/home/user/amattavelli/Ratio/articoli/2026-05-15_digital-omnibus-ai-act-rinvio-agosto-2026.docx"
    doc.save(out)
    print(f"Salvato: {out}")


# ═══════════════════════════════════════════════════════════════════════════
# ARTICOLO 2 — GPT-5.5 agentivo
# ═══════════════════════════════════════════════════════════════════════════

def crea_articolo_2():
    doc = new_doc()
    testata(
        doc,
        sezione="Strumenti e Modelli AI",
        data_label="Maggio 2026",
        titolo=(
            "GPT-5.5 sa usare il computer:\ncosa cambia quando il modello inizia ad agire"
        ),
        sottotitolo=(
            "Lanciato il 23 aprile 2026, il nuovo modello di OpenAI gestisce flussi di lavoro "
            "complessi in autonomia. Per studi e uffici, la variabile critica non è la "
            "precisione dell’output ma la governance del processo."
        ),
        data_autore="15 maggio 2026",
    )

    para(doc,
        "Immagina di avere in studio un collaboratore a cui puoi dire: analizza le ultime sei "
        "relazioni di bilancio dei nostri clienti del settore manifatturiero, confronta i principali "
        "indici con i benchmark di settore e prepara una sintesi con i casi che meritano attenzione. "
        "Il collaboratore apre i file, fa i calcoli, consulta i dati di riferimento, scrive il testo. "
        "Non ti chiede come fare ogni singolo passaggio: lo sa già. Il 23 aprile 2026 OpenAI ha "
        "rilasciato GPT-5.5, un modello costruito esattamente per fare questo: ricevere un obiettivo "
        "articolato e portarlo a termine attraverso una sequenza di azioni autonome, usando strumenti, "
        "navigando interfacce, correggendo gli errori in corso d’opera."
    )

    para(doc,
        "La novità rispetto alle versioni precedenti non è nella qualità delle risposte "
        "singole, che era già elevata con GPT-5.4 rilasciato a marzo 2026. La novità è "
        "nella capacità di concatenare azioni senza bisogno di un input umano a ogni passo. "
        "GPT-5.5 è ottimizzato per quattro ambiti specifici: programmazione avanzata, lavoro "
        "d’ufficio (analisi dati, ricerca, redazione di documenti strutturati), uso pratico del "
        "computer (navigare interfacce, inserire dati, compilare moduli), e ricerca professionale su "
        "larga scala. Per chi lavora in uno studio o in un ufficio amministrativo, queste quattro aree "
        "coprono una fetta significativa delle attività quotidiane."
    )

    heading(doc, "Cosa riesce a fare nella pratica")

    para(doc,
        "A differenza dei chatbot che rispondono a domande, un agente come GPT-5.5 pianifica. "
        "Riceve un obiettivo, lo scompone in sotto-compiti, usa gli strumenti disponibili (file, "
        "database, motori di ricerca, API) e produce un output strutturato. Nel lavoro contabile, "
        "questo si traduce in operazioni come la riconciliazione automatica di un estratto conto con "
        "i movimenti del gestionale, la compilazione di bozze di documenti a partire da dati "
        "strutturati, o la sintesi di un fascicolo documentale in vista di una scadenza. Nel lavoro "
        "legale, copre attività come l’analisi comparativa di clausole contrattuali o la "
        "ricerca di precedenti giurisprudenziali su un tema specifico."
    )

    para(doc,
        "La distribuzione iniziale riguarda gli utenti di Codex e ChatGPT, con estensione prevista "
        "alle API nelle settimane successive. Chi usa già ChatGPT in abbonamento trova la nuova "
        "versione disponibile senza aggiornamenti particolari. Chi invece vuole integrare le capacità "
        "agentiche nei propri sistemi o gestionali dovrà attendere l’accesso API, che apre la "
        "possibilità di costruire flussi di lavoro personalizzati. Diversi provider di software "
        "gestionali per studi professionali stanno già lavorando a integrazioni dirette, sull’onda "
        "del percorso avviato da Anthropic con i suoi finance agents rilasciati il mese scorso."
    )

    heading(doc, "Il punto che i comunicati stampa non affrontano")

    para(doc,
        "Ogni lancio di un nuovo modello porta con sé una quantità di casi d’uso positivi "
        "e quasi nessuna discussione sui meccanismi di controllo. Per un professionista che usa un agente "
        "per la chiusura mensile o per la preparazione di un’analisi da presentare a un cliente, "
        "la domanda critica non è “questo modello è abbastanza bravo?”, ma “chi "
        "nel mio studio capisce abbastanza il processo da accorgersi quando l’agente ha fatto "
        "qualcosa di sbagliato?”."
    )

    para(doc,
        "GPT-5.5 è progettato per correggere i propri errori in itinere, il che riduce le "
        "problematiche più grossolane. L’errore silenzioso però, quello che produce un "
        "output plausibile ma scorretto, è più insidioso di quello evidente. "
        "Un’analisi finanziaria che usa un indice calcolato in modo leggermente diverso da come "
        "lo calcola il settore, o una bozza di lettera che omette un passaggio rilevante, può "
        "passare il controllo di chi non conosce bene il processo sottostante. Per questo, nel contesto "
        "di un lavoro professionale, l’automazione agentiva aumenta il valore della competenza "
        "tecnica, non la sostituisce."
    )

    heading(doc, "Il profilo di utilizzo più sostenibile")

    para(doc,
        "Per studi professionali e uffici che vogliono avvicinarsi a GPT-5.5 in modo strutturato, "
        "il profilo di utilizzo più ragionevole nella fase attuale è quello della prima "
        "bozza supervisionata: l’agente esegue la parte meccanica, un professionista verifica "
        "l’output prima che entri in un documento ufficiale o in una comunicazione al cliente. "
        "Questo consente di costruire confidenza con lo strumento, di capirne i limiti specifici nel "
        "proprio contesto di lavoro, e di formare il team in modo progressivo."
    )

    para(doc,
        "L’integrazione con Microsoft 365, resa disponibile in parallelo al lancio di GPT-5.5, "
        "riduce una delle barriere pratiche più citate: il dover copiare e incollare tra ambienti "
        "diversi. L’agente può ora lavorare direttamente sui file Word, Excel e Outlook, "
        "senza uscire dall’ambiente di lavoro familiare. Per uno studio che usa già "
        "Microsoft 365 come infrastruttura, questa integrazione abbassa il costo di adozione in modo "
        "significativo e rende l’inizio più accessibile di quanto fosse sei mesi fa."
    )

    para(doc,
        "GPT-5.5 arriva in un momento in cui la curva di adozione degli strumenti AI in molti studi "
        "italiani si sta appiattendo dopo la fase di esplorazione iniziale. Il modello offre capacità "
        "concrete e già disponibili. Il lavoro che rimane da fare non è tecnologico: è "
        "capire quali processi vale la pena automatizzare, chi nel team ha la competenza per supervisionare "
        "l’output, e come documentare il fatto che la responsabilità professionale finale "
        "è rimasta in capo al professionista, anche quando ad eseguire il lavoro è stato un agente."
    )

    footer(doc, [
        "OpenAI — Annuncio GPT-5.5 (23 aprile 2026)",
        "Sky TG24 — “OpenAI, annunciato GPT 5.5: ecco le principali novità” (24 aprile 2026)",
        "techbusiness.it — “OpenAI GPT-5.5: il nuovo modello per rivoluzionare il lavoro”",
        "ninja.it — “OpenAI lancia GPT-5.5 e introduce agenti AI autonomi” (aprile 2026)",
        "Legge 132/2025 — Disciplina organica sull’intelligenza artificiale in Italia",
    ])

    out = "/home/user/amattavelli/Ratio/articoli/2026-05-15_gpt55-agenti-computer-use-professionisti.docx"
    doc.save(out)
    print(f"Salvato: {out}")


# ═══════════════════════════════════════════════════════════════════════════
# ARTICOLO 3 — AI verticale per commercialisti
# ═══════════════════════════════════════════════════════════════════════════

def crea_articolo_3():
    doc = new_doc()
    testata(
        doc,
        sezione="Professione e Strumenti",
        data_label="Maggio 2026",
        titolo=(
            "Quando lo strumento sa la risposta ma non sa perché:\n"
            "i rischi dell’AI verticale negli studi professionali"
        ),
        sottotitolo=(
            "Il 65% dei commercialisti italiani usa già strumenti AI. "
            "La questione non è l’adozione, ma quanto il professionista "
            "che firma conosce la logica di ciò che ha delegato."
        ),
        data_autore="15 maggio 2026",
    )

    para(doc,
        "Sei un commercialista. Hai sottoscritto un abbonamento a uno strumento di intelligenza "
        "artificiale verticale, uno di quelli specializzati su normativa fiscale e contabilità "
        "italiana. Inserisci una domanda su come trattare una fattispecie IVA complessa, e in meno "
        "di dieci secondi ottieni una risposta strutturata, con riferimenti normativi e circolari "
        "dell’Agenzia delle Entrate. La risposta è quasi sempre corretta. Quasi. "
        "Il problema inizia quando sei così abituato alla correttezza che smetti di "
        "verificare i “quasi”."
    )

    para(doc,
        "A marzo 2026, l’Osservatorio Artificial Intelligence del Politecnico di Milano ha "
        "pubblicato dati che indicano che il 65% dei commercialisti italiani utilizza già "
        "strumenti di intelligenza artificiale nella propria attività professionale, con "
        "l’84% degli studi che percepisce un impatto positivo. Sono numeri che segnalano "
        "una transizione già in corso. Gli strumenti verticali per il settore, come Normo.ai "
        "o le suite sviluppate da provider specializzati con centinaia di studi clienti in tutta "
        "Italia, hanno trovato una base utente in tempi più rapidi di quanto ci si aspettasse."
    )

    heading(doc, "Il problema della specializzazione senza trasparenza")

    para(doc,
        "Gli strumenti di AI verticale hanno un vantaggio evidente rispetto ai modelli generalisti: "
        "sono addestrati o ottimizzati su corpora specifici (normativa fiscale, circolari, prassi "
        "contabile italiana) e producono risposte più coerenti con il contesto professionale "
        "del paese. Hanno però un limite strutturale che vale la pena tenere presente: "
        "rispondono con sicurezza anche quando la fattispecie è al limite della loro base di "
        "conoscenza. Un modello generalista tende a esprimere incertezza di fronte a una domanda "
        "che non sa rispondere bene. Alcuni modelli verticali, ottimizzati per sembrare competenti "
        "nel loro dominio, possono produrre risposte plausibili su casi che non rientrano nei loro "
        "dati di addestramento."
    )

    para(doc,
        "Il commercialista che usa questi strumenti come supporto alla ricerca e alla prima bozza "
        "di risposta è in una posizione diversa da chi li usa come oracolo a cui affidare la "
        "risposta definitiva. La distinzione sembra ovvia, ma nella pratica quotidiana di uno studio "
        "sotto pressione di scadenze, la linea tra i due modi di usare uno strumento si assottiglia "
        "velocemente. Quando il tempo stringe e lo strumento ha dato ottimi risultati nelle ultime "
        "dieci domande, controllare l’undicesima con la stessa attenzione richiede uno sforzo "
        "consapevole che non tutti i team hanno la disciplina di mantenere."
    )

    heading(doc, "La firma resta in capo al professionista")

    para(doc,
        "La Legge 132/2025, il testo italiano di disciplina dell’intelligenza artificiale "
        "entrato in vigore il 10 ottobre 2025, è esplicita sul punto della responsabilità "
        "professionale. Il documento firmato dal commercialista, dall’avvocato o dal consulente "
        "è quello del professionista, non dello strumento che ha aiutato a redigerlo. "
        "L’uso di un sistema AI non modifica l’imputabilità dell’errore, "
        "non costituisce un’attenuante riconosciuta, e non trasferisce la responsabilità "
        "al fornitore dello strumento, salvo ipotesi molto specifiche di difetto del prodotto."
    )

    para(doc,
        "Questo non significa che gli strumenti verticali siano da evitare: significa che l’uso "
        "professionale richiede una comprensione sufficiente della logica con cui questi strumenti "
        "ragionano. Sapere, per esempio, che un modello è stato aggiornato a una certa data e "
        "non incorpora le modifiche normative successive è già un’informazione che "
        "cambia il modo in cui si usa l’output. Sapere che su certi tipi di fattispecie il "
        "modello tende a citare la norma base senza considerare la prassi interpretativa più "
        "recente è un altro elemento che guida la revisione critica."
    )

    heading(doc, "Come lavorarci in modo professionale")

    para(doc,
        "Gli studi che stanno usando questi strumenti con maggiore efficacia hanno in comune alcune "
        "pratiche. Definiscono in modo esplicito per quali tipi di quesiti usano lo strumento come "
        "supporto e per quali la risposta definitiva richiede comunque una verifica manuale sulle "
        "fonti primarie. Tengono traccia delle aree in cui lo strumento ha prodotto output da "
        "correggere, costruendo nel tempo una mappa dei suoi limiti specifici nel proprio contesto "
        "di lavoro. Formano i collaboratori più giovani non solo su come usare lo strumento, "
        "ma su come riconoscere quando dubitarne."
    )

    para(doc,
        "Il mercato degli strumenti verticali per studi professionali italiani si sta consolidando "
        "rapidamente. Diversi provider stanno sviluppando integrazioni con i principali software "
        "gestionali del settore, e alcuni offrono già funzionalità di audit trail che "
        "permettono di documentare le consultazioni effettuate con lo strumento. Questa "
        "documentazione non è solo utile per eventuali verifiche: è anche un modo per "
        "costruire consapevolezza interna su quanto e come il team si affida all’AI nel lavoro "
        "quotidiano."
    )

    para(doc,
        "Il commercialista che esce vincitore dalla transizione verso l’AI non è quello "
        "che usa più strumenti, ma quello che conosce abbastanza ogni strumento da sapere quando "
        "smettere di fidarsi di esso. Questa competenza non arriva dall’uso passivo: arriva "
        "dall’uso critico. Ogni volta che verifichi una risposta e la trovi corretta, stai anche "
        "validando il tuo modello mentale dello strumento. Ogni volta che la trovi scorretta, stai "
        "imparando qualcosa che nessun corso di formazione può darti allo stesso modo."
    )

    footer(doc, [
        "Osservatorio Artificial Intelligence, Politecnico di Milano — Rapporto 2026",
        "italiaoggi.it — “Intelligenza artificiale per commercialisti: perché la specializzazione verticale”",
        "fisco7.it — “Intelligenza artificiale negli studi professionali: riflessi fiscali e nuove responsabilità” (maggio 2026)",
        "normo.ai — Documentazione ufficiale del servizio",
        "Legge 132/2025 — Disciplina organica sull’intelligenza artificiale in Italia",
    ])

    out = "/home/user/amattavelli/Ratio/articoli/2026-05-15_ai-verticale-commercialista-delegare-senza-capire.docx"
    doc.save(out)
    print(f"Salvato: {out}")


# ═══════════════════════════════════════════════════════════════════════════
# ARTICOLO 4 — Gap adozione AI nelle PMI italiane
# ═══════════════════════════════════════════════════════════════════════════

def crea_articolo_4():
    doc = new_doc()
    testata(
        doc,
        sezione="PMI e Strategia",
        data_label="Maggio 2026",
        titolo=(
            "Il mercato AI supera 1,8 miliardi, ma nelle piccole\n"
            "imprese italiane l’adozione è ancora ferma al 7%"
        ),
        sottotitolo=(
            "I dati del Politecnico di Milano segnalano una crescita esplosiva del mercato AI. "
            "Per le PMI più piccole, il problema non è l’accesso agli strumenti "
            "ma sapere da quale processo iniziare."
        ),
        data_autore="15 maggio 2026",
    )

    para(doc,
        "Il mercato italiano dell’intelligenza artificiale ha raggiunto 1,8 miliardi di euro "
        "nel 2025, con una crescita del 50% sull’anno precedente, secondo i dati dell’Osservatorio "
        "Artificial Intelligence del Politecnico di Milano. Nei convegni e nelle presentazioni questo "
        "numero viene citato come segnale di un’accelerazione strutturale che non si può "
        "ignorare. Nel frattempo, nelle piccole imprese italiane sotto i 50 dipendenti, che "
        "rappresentano la spina dorsale del tessuto produttivo del paese, il tasso di adozione "
        "effettiva di strumenti AI si ferma al 7%. Il gap tra il mercato che cresce e le imprese "
        "che restano ferme non si chiude da solo."
    )

    para(doc,
        "La domanda che vale la pena farsi non è “perché le PMI non adottano "
        "l’AI?” come se si trattasse di un problema culturale da correggere con campagne "
        "di sensibilizzazione. Vale la pena chiedersi cosa hanno in comune le imprese che hanno "
        "fatto il passo e quelle che non l’hanno ancora fatto, e cosa è cambiato "
        "concretamente per le prime. I dati del Politecnico indicano che tra le medie imprese "
        "(50-250 dipendenti) il tasso di adozione sale al 15%, e che il 20% di quelle che ancora "
        "non hanno iniziato sta valutando di farlo nel breve periodo. Il tema non è la "
        "consapevolezza: è la capacità di tradurla in un’azione concreta con "
        "risorse limitate."
    )

    heading(doc, "Cosa frena l’adozione nelle imprese più piccole")

    para(doc,
        "Chi gestisce una PMI con un team amministrativo di tre o quattro persone si trova di fronte "
        "a una difficoltà concreta: non ha una funzione IT interna che valuti gli strumenti, "
        "non ha un budget dedicato alla sperimentazione, e il tempo dei responsabili è "
        "completamente assorbito dall’operatività. Gli strumenti di AI generativa sono "
        "economicamente accessibili (i principali costano poche decine di euro al mese per persona), "
        "ma l’accesso allo strumento è solo il primo passo. Il secondo, più difficile, "
        "è capire su quale processo usarlo, come formare chi lo userà, e come verificare "
        "che stia funzionando."
    )

    para(doc,
        "A questo si aggiunge un problema di misurazione. Una grande azienda con processi "
        "documentati può costruire un caso business per l’adozione AI: misura il tempo "
        "risparmiato su un’attività specifica, calcola il costo, confronta con il "
        "beneficio. Una PMI con processi ancora in buona parte informali fatica a costruire lo "
        "stesso ragionamento, perché non ha il punto di partenza da cui misurare. "
        "Questo non significa che il beneficio non ci sia: significa che non è "
        "immediatamente quantificabile, e che chi deve prendere la decisione si trova a farlo "
        "senza dati."
    )

    heading(doc, "L’approccio collaborativo che cambia il problema")

    para(doc,
        "Negli ultimi mesi si è diffuso nell’uso aziendale un modello di utilizzo che "
        "vale la pena considerare: quello dell’AI collaborativa, dove lo strumento non "
        "sostituisce un ruolo ma lavora affianco a chi quel ruolo lo svolge, su compiti specifici "
        "e delimitati. Un responsabile amministrativo che usa un agente AI per la prima passata di "
        "riconciliazione dei conti, un responsabile commerciale che usa un modello generativo per "
        "costruire le bozze di offerta personalizzate, un titolare che usa un assistente AI per "
        "preparare la documentazione per una pratica bancaria. In tutti questi casi lo strumento "
        "non cambia la struttura del lavoro: riduce la quota di tempo dedicata alla parte meccanica "
        "di ogni attività."
    )

    para(doc,
        "Questo approccio è più accessibile per una PMI di un progetto di automazione "
        "complessa perché non richiede integrazioni tecniche elaborate, non modifica i processi "
        "esistenti in modo radicale, e consente di iniziare con un investimento limitato. Richiede "
        "però qualcosa che non è tecnologico: la capacità di scegliere il processo "
        "giusto su cui cominciare, quello dove il guadagno di tempo è visibile in settimane e "
        "non in mesi, e dove l’eventuale errore dello strumento è verificabile da chi "
        "lo usa."
    )

    heading(doc, "Da dove iniziare con risorse limitate")

    para(doc,
        "Per una piccola impresa o uno studio professionale che non ha ancora un progetto AI "
        "strutturato, il percorso più efficace nella fase attuale è quello per aree "
        "funzionali. Si sceglie un’area (gestione dei documenti, comunicazione con i clienti, "
        "analisi dei dati), si identifica il compito più ripetitivo che assorbe tempo utile, "
        "si prova uno strumento per quella specifica attività per quattro settimane, si misura "
        "il delta. Se funziona, si estende. Se non funziona, si cambia strumento o si cambia compito. "
        "Questo approccio non richiede un piano strategico quinquennale: richiede un’ora di "
        "analisi e la disponibilità di una persona a provare."
    )

    para(doc,
        "Il 20% delle piccole imprese che il Politecnico indica come “in valutazione” "
        "ha già superato la fase della curiosità. Quello che serve, in molti casi, non "
        "è un’altra ricerca di mercato o un convegno sull’AI: è avere a portata "
        "di mano un esempio concreto di cosa ha fatto un’altra impresa simile, in quanto tempo, "
        "con quali strumenti e con quale risultato. Questa informazione è la più rara e la "
        "più utile, ed è quella che circola meno nei canali formali."
    )

    para(doc,
        "Il mercato AI in Italia continuerà a crescere indipendentemente da quello che faranno "
        "le singole PMI. La questione rilevante non è restare al passo con un mercato astratto: "
        "è capire se c’è un’attività specifica nel proprio lavoro che "
        "potrebbe essere fatta meglio, o in meno tempo, con uno strumento che esiste già, "
        "costa poco e non richiede competenze tecniche per iniziare. Se la risposta è sì, "
        "il momento giusto per provare era ieri. Adesso va bene lo stesso."
    )

    footer(doc, [
        "Osservatorio Artificial Intelligence, Politecnico di Milano — Rapporto 2025 (osservatori.net)",
        "besttechpartner.ai — “Intelligenza artificiale collaborativa: il salto strategico per la produttività dei team nel 2026” (13 maggio 2026)",
        "besttechpartner.ai — “IA e lavoro: scenari, impatto e opportunità per il mercato italiano” (10 maggio 2026)",
        "deepelse.com — “AI per PMI italiane: guida completa 2026”",
        "Legge 132/2025 — Disciplina organica sull’intelligenza artificiale in Italia",
    ])

    out = "/home/user/amattavelli/Ratio/articoli/2026-05-15_mercato-ai-pmi-gap-adozione-dove-iniziare.docx"
    doc.save(out)
    print(f"Salvato: {out}")


# ── MAIN ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    crea_articolo_1()
    crea_articolo_2()
    crea_articolo_3()
    crea_articolo_4()
    print("\nTutti e 4 gli articoli creati con successo.")
